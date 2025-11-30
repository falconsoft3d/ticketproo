from django.db import models
from django.contrib.auth.models import User
from django.utils import timezone
from django.db import transaction
from django.db.models.signals import post_save
from django.dispatch import receiver
from datetime import timedelta
import os
import uuid


class Category(models.Model):
    name = models.CharField(max_length=100, unique=True, verbose_name='Nombre')
    description = models.TextField(blank=True, verbose_name='Descripción')
    color = models.CharField(
        max_length=7, 
        default='#007bff',
        help_text='Color en formato hexadecimal (ej: #007bff)',
        verbose_name='Color'
    )
    is_active = models.BooleanField(default=True, verbose_name='Activa')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    
    class Meta:
        ordering = ['name']
        verbose_name = 'Categoría'
        verbose_name_plural = 'Categorías'
    
    def __str__(self):
        return self.name


class Company(models.Model):
    """Modelo para gestionar empresas/clientes"""
    
    name = models.CharField(
        max_length=200, 
        unique=True, 
        verbose_name='Nombre de la Empresa'
    )
    description = models.TextField(
        blank=True, 
        verbose_name='Descripción',
        help_text='Descripción opcional de la empresa'
    )
    address = models.CharField(
        max_length=300, 
        blank=True, 
        verbose_name='Dirección'
    )
    phone = models.CharField(
        max_length=20, 
        blank=True, 
        verbose_name='Teléfono'
    )
    email = models.EmailField(
        blank=True, 
        verbose_name='Email de contacto'
    )
    website = models.URLField(
        blank=True, 
        verbose_name='Sitio web'
    )
    color = models.CharField(
        max_length=7, 
        default='#28a745',
        help_text='Color en formato hexadecimal (ej: #28a745)',
        verbose_name='Color identificativo'
    )
    logo = models.ImageField(
        upload_to='company_logos/',
        blank=True,
        null=True,
        verbose_name='Logo de la empresa',
        help_text='Logo de la empresa (opcional)'
    )
    is_active = models.BooleanField(
        default=True, 
        verbose_name='Empresa activa'
    )
    created_at = models.DateTimeField(
        default=timezone.now, 
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True, 
        verbose_name='Última actualización'
    )
    public_token = models.UUIDField(
        null=True,
        blank=True,
        editable=True,
        verbose_name='Token público',
        help_text='Token único para acceso público a estadísticas'
    )
    client_requests_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token de Solicitudes',
        help_text='Token único para que los clientes accedan a sus solicitudes'
    )
    uuid = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        null=True,
        blank=True,
        verbose_name='UUID público',
        help_text='Identificador único para URLs públicas'
    )
    business_objectives = models.TextField(
        blank=True,
        verbose_name='Objetivos Empresariales',
        help_text='Objetivos, metas y KPIs estratégicos de la empresa. Estos serán considerados en los resúmenes ejecutivos generados por IA.'
    )
    
    # Información fiscal y de ubicación
    tax_id = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='NIF/CIF/Tax ID',
        help_text='Número de identificación fiscal'
    )
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad',
        help_text='Ciudad de la empresa'
    )
    state = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Estado/Provincia',
        help_text='Estado o provincia de la empresa'
    )
    postal_code = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Código Postal',
        help_text='Código postal'
    )
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País',
        help_text='País de la empresa'
    )
    
    # Información bancaria
    bank_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre del Banco',
        help_text='Nombre del banco'
    )
    bank_url = models.URLField(
        blank=True,
        verbose_name='URL del Banco',
        help_text='Sitio web del banco'
    )
    bank_account_holder = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Titular de la Cuenta',
        help_text='Nombre del titular de la cuenta bancaria'
    )
    bank_account = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Número de Cuenta',
        help_text='IBAN o número de cuenta bancaria'
    )
    bank_swift = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='SWIFT/BIC',
        help_text='Código SWIFT o BIC del banco'
    )
    
    # Condiciones de pago
    payment_terms = models.TextField(
        blank=True,
        verbose_name='Condiciones de Pago',
        help_text='Términos y condiciones de pago (se mostrarán en las cotizaciones)'
    )
    payment_terms_description = models.TextField(
        blank=True,
        verbose_name='Descripción de Condiciones de Pago',
        help_text='Descripción detallada de las condiciones de pago para las cotizaciones'
    )
    quotation_general_description = models.TextField(
        blank=True,
        verbose_name='Descripción General de Cotizaciones',
        help_text='Descripción general que aparecerá en todas las cotizaciones de esta empresa'
    )
    
    class Meta:
        ordering = ['name']
        verbose_name = 'Empresa'
        verbose_name_plural = 'Empresas'
    
    def __str__(self):
        return self.name
    
    def get_active_tickets_count(self):
        """Obtiene el número de tickets activos de la empresa"""
        return self.tickets.exclude(status='closed').count()
    
    def get_total_tickets_count(self):
        """Obtiene el total de tickets de la empresa"""
        return self.tickets.count()
    
    def get_users_count(self):
        """Obtiene el número de usuarios asociados a la empresa"""
        return self.users.count()
    
    def get_public_stats(self):
        """Obtiene estadísticas públicas de la empresa"""
        from django.db.models import Count, Q
        
        tickets = self.tickets.all()
        
        return {
            'total_tickets': tickets.count(),
            'open_tickets': tickets.filter(status='open').count(),
            'in_progress_tickets': tickets.filter(status='in_progress').count(),
            'closed_tickets': tickets.filter(status='closed').count(),
            'high_priority': tickets.filter(priority='high').count(),
            'medium_priority': tickets.filter(priority='medium').count(),
            'low_priority': tickets.filter(priority='low').count(),
            'users_count': self.get_users_count(),
            'last_ticket_date': tickets.order_by('-created_at').first().created_at if tickets.exists() else None,
        }
    
    def regenerate_public_token(self):
        """Regenera el token público para acceso a estadísticas"""
        self.public_token = uuid.uuid4()
        self.save()
        return self.public_token
    
    def get_client_requests_url(self):
        """Obtiene la URL pública para que el cliente vea sus solicitudes"""
        from django.urls import reverse
        return reverse('public_client_requests', kwargs={'token': self.client_requests_token})
    
    def regenerate_client_requests_token(self):
        """Regenera el token para solicitudes de cliente"""
        self.client_requests_token = uuid.uuid4()
        self.save()
        return self.client_requests_token


class Project(models.Model):
    """Modelo para gestionar proyectos de trabajo"""
    
    STATUS_CHOICES = [
        ('planning', 'Planificación'),
        ('active', 'Activo'),
        ('on_hold', 'En Pausa'),
        ('completed', 'Completado'),
        ('cancelled', 'Cancelado'),
    ]
    
    name = models.CharField(max_length=200, unique=True, verbose_name='Nombre del Proyecto')
    description = models.TextField(blank=True, verbose_name='Descripción')
    color = models.CharField(
        max_length=7, 
        default='#28a745',
        help_text='Color en formato hexadecimal (ej: #28a745)',
        verbose_name='Color'
    )
    status = models.CharField(
        max_length=20, 
        choices=STATUS_CHOICES, 
        default='active', 
        verbose_name='Estado'
    )
    start_date = models.DateField(null=True, blank=True, verbose_name='Fecha de Inicio')
    end_date = models.DateField(null=True, blank=True, verbose_name='Fecha de Fin')
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE, 
        related_name='created_projects',
        verbose_name='Creado por'
    )
    assigned_users = models.ManyToManyField(
        User,
        blank=True,
        related_name='assigned_projects',
        verbose_name='Usuarios Asignados',
        help_text='Usuarios que tienen acceso a este proyecto'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    public_share_token = models.CharField(
        max_length=64,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token de Compartir Público',
        help_text='Token único para compartir información del proyecto públicamente'
    )
    
    class Meta:
        ordering = ['-created_at', 'name']
        verbose_name = 'Proyecto'
        verbose_name_plural = 'Proyectos'
    
    def __str__(self):
        return self.name
    
    def get_total_hours(self):
        """Obtiene el total de horas trabajadas en este proyecto"""
        from django.db.models import F, Sum, Case, When, Value
        from django.db.models.functions import Extract
        
        # Calcular la duración en segundos para cada entrada
        total_seconds = 0
        
        for entry in self.time_entries.all():
            if entry.fecha_salida:
                # Entrada completada
                duration = (entry.fecha_salida - entry.fecha_entrada).total_seconds()
            else:
                # Entrada activa
                duration = (timezone.now() - entry.fecha_entrada).total_seconds()
            total_seconds += duration
        
        # Convertir a horas y redondear
        return round(total_seconds / 3600, 2)
    
    def get_active_workers_count(self):
        """Obtiene el número de trabajadores actualmente trabajando en este proyecto"""
        return self.time_entries.filter(fecha_salida__isnull=True).count()
    
    def can_edit(self, user):
        """Verifica si el usuario puede editar este proyecto"""
        from .utils import is_agent
        return is_agent(user) or user == self.created_by
    
    def can_delete(self, user):
        """Verifica si el usuario puede eliminar este proyecto"""
        from .utils import is_agent
        return is_agent(user) or user == self.created_by
    
    def generate_public_share_token(self):
        """Genera un token unico para compartir el proyecto publicamente"""
        if not self.public_share_token:
            import secrets
            self.public_share_token = secrets.token_urlsafe(32)
            self.save()
        return self.public_share_token
    
    def get_public_share_url(self):
        if not self.public_share_token:
            self.generate_public_share_token()
        return "/public/project/" + str(self.public_share_token) + "/"
    
    def get_tickets_count(self):
        """Obtiene el total de tickets relacionados con este proyecto"""
        from django.db.models import Q
        
        # Usar Q objects para obtener tickets únicos (directamente asignados O en time_entries)
        tickets_ids = set()
        
        # Tickets directamente asignados al proyecto
        direct_tickets = self.tickets.values_list('id', flat=True)
        tickets_ids.update(direct_tickets)
        
        # Tickets de time_entries
        time_entry_tickets = self.time_entries.filter(
            ticket__isnull=False
        ).values_list('ticket_id', flat=True)
        tickets_ids.update(time_entry_tickets)
        
        return len(tickets_ids)
    
    def get_resolved_tickets_count(self):
        """Obtiene el total de tickets resueltos relacionados con este proyecto"""
        from django.db.models import Q
        
        # Usar Q objects para obtener tickets únicos resueltos
        resolved_tickets_ids = set()
        
        # Tickets directamente asignados y resueltos
        direct_resolved = self.tickets.filter(
            status__in=['resolved', 'closed']
        ).values_list('id', flat=True)
        resolved_tickets_ids.update(direct_resolved)
        
        # Tickets resueltos de time_entries
        time_entry_resolved = self.time_entries.filter(
            ticket__isnull=False,
            ticket__status__in=['resolved', 'closed']
        ).values_list('ticket_id', flat=True)
        resolved_tickets_ids.update(time_entry_resolved)
        
        return len(resolved_tickets_ids)
    
    def get_completed_work_orders_count(self):
        """Obtiene el total de órdenes de trabajo completadas relacionadas con este proyecto"""
        # Contar órdenes directamente asignadas y completadas
        direct_completed = self.work_orders.filter(status='finished').count()
        
        # Contar órdenes completadas de time_entries
        time_entry_completed = self.time_entries.filter(
            work_order__isnull=False,
            work_order__status='finished'
        ).values('work_order').distinct().count()
        
        # Usar set para eliminar duplicados
        completed_orders_ids = set()
        
        # Órdenes directamente asignadas y completadas
        direct_orders = self.work_orders.filter(
            status='finished'
        ).values_list('id', flat=True)
        completed_orders_ids.update(direct_orders)
        
        # Órdenes completadas de time_entries
        time_entry_orders = self.time_entries.filter(
            work_order__isnull=False,
            work_order__status='finished'
        ).values_list('work_order_id', flat=True)
        completed_orders_ids.update(time_entry_orders)
        
        return len(completed_orders_ids)
    
    def get_project_workers(self):
        """Obtiene los usuarios que han trabajado en este proyecto con sus datos de contacto"""
        from django.db.models import Q
        workers = User.objects.filter(
            time_entries__project=self
        ).distinct().select_related('profile')
        
        return workers
    
    def get_total_cost(self):
        """Calcula el coste total del proyecto basado en las horas trabajadas y el coste por hora de cada empleado"""
        from decimal import Decimal
        total_cost = Decimal('0.00')
        
        for entry in self.time_entries.all():
            hours_worked = entry.horas_trabajadas
            if hours_worked and entry.user and hasattr(entry.user, 'profile'):
                coste_hora = entry.user.profile.coste_hora
                total_cost += Decimal(str(hours_worked)) * coste_hora
        
        return total_cost
    
    def get_total_revenue(self):
        """Calcula el total de venta del proyecto basado en las horas trabajadas y el precio por hora de cada empleado"""
        from decimal import Decimal
        total_revenue = Decimal('0.00')
        
        for entry in self.time_entries.all():
            hours_worked = entry.horas_trabajadas
            if hours_worked and entry.user and hasattr(entry.user, 'profile'):
                precio_hora = entry.user.profile.precio_hora
                total_revenue += Decimal(str(hours_worked)) * precio_hora
        
        return total_revenue
    
    def get_profit(self):
        """Calcula el beneficio del proyecto (venta - coste)"""
        return self.get_total_revenue() - self.get_total_cost()
    
    def get_profit_margin(self):
        """Calcula el margen de beneficio como porcentaje"""
        revenue = self.get_total_revenue()
        if revenue > 0:
            return (self.get_profit() / revenue) * 100
        return 0


class Ticket(models.Model):
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('urgent', 'Urgente'),
    ]
    
    STATUS_CHOICES = [
        ('open', 'Abierto'),
        ('in_progress', 'En Progreso'),
        ('resolved', 'Resuelto'),
        ('closed', 'Cerrado'),
    ]
    
    TYPE_CHOICES = [
        ('desarrollo', 'Desarrollo'),
        ('error', 'Error'),
    ]
    
    title = models.CharField(max_length=200, verbose_name='Título')
    ticket_number = models.CharField(
        max_length=10, 
        unique=True,
        blank=True,
        verbose_name='Número de Ticket',
        help_text='Formato: XX-XXX (ID Usuario-Secuencia)'
    )
    description = models.TextField(verbose_name='Descripción')
    category = models.ForeignKey(
        Category,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='tickets',
        verbose_name='Categoría'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='tickets',
        verbose_name='Empresa/Cliente',
        help_text='Empresa o cliente asociado al ticket'
    )
    project = models.ForeignKey(
        Project,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='tickets',
        verbose_name='Proyecto',
        help_text='Proyecto asociado al ticket (opcional)'
    )
    priority = models.CharField(
        max_length=10, 
        choices=PRIORITY_CHOICES, 
        default='medium',
        verbose_name='Prioridad'
    )
    status = models.CharField(
        max_length=15, 
        choices=STATUS_CHOICES, 
        default='open',
        verbose_name='Estado'
    )
    ticket_type = models.CharField(
        max_length=15,
        choices=TYPE_CHOICES,
        default='desarrollo',
        verbose_name='Tipo'
    )
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE, 
        related_name='created_tickets',
        verbose_name='Creado por'
    )
    assigned_to = models.ForeignKey(
        User, 
        on_delete=models.SET_NULL, 
        null=True, 
        blank=True,
        related_name='assigned_tickets',
        verbose_name='Asignado a'
    )
    hours = models.DecimalField(
        max_digits=6,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Horas estimadas/trabajadas',
        help_text='Cantidad de horas estimadas o trabajadas en el ticket'
    )
    # Campo para enlace único público
    public_share_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de compartir',
        help_text='Token único para compartir el ticket públicamente'
    )
    is_public_shareable = models.BooleanField(
        default=False,
        verbose_name='Compartir públicamente',
        help_text='Permite que este ticket sea visible mediante un enlace público'
    )
    is_approved = models.BooleanField(
        default=False,
        verbose_name='Aprobado',
        help_text='Indica si el ticket ha sido aprobado por el cliente'
    )
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='approved_tickets',
        verbose_name='Aprobado por',
        help_text='Usuario que aprobó el ticket'
    )
    approved_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de aprobación',
        help_text='Fecha y hora en que fue aprobado el ticket'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Ticket'
        verbose_name_plural = 'Tickets'
    
    def generate_ticket_number(self):
        """Genera un numero de ticket unico con formato XX-XXX"""
        if self.ticket_number:
            return self.ticket_number
            
        # Formatear ID del usuario con ceros a la izquierda (2 digitos)
        user_id_formatted = "{:02d}".format(self.created_by.id)
        
        # Buscar el ultimo ticket creado por este usuario
        last_ticket = Ticket.objects.filter(
            created_by=self.created_by,
            ticket_number__startswith=user_id_formatted + "-"
        ).exclude(id=self.id).order_by('-ticket_number').first()
        
        if last_ticket and last_ticket.ticket_number:
            # Extraer el número de secuencia del último ticket
            try:
                last_sequence = int(last_ticket.ticket_number.split('-')[1])
                new_sequence = last_sequence + 1
            except (IndexError, ValueError):
                new_sequence = 1
        else:
            new_sequence = 1
        
        # Formatear secuencia con ceros a la izquierda (3 dígitos)
        sequence_formatted = f"{new_sequence:03d}"
        
        return f"{user_id_formatted}-{sequence_formatted}"
    
    def save(self, *args, **kwargs):
        """Override del método save para generar número de ticket automáticamente"""
        is_new_ticket = self.pk is None
        
        if not self.ticket_number:
            # Usar transacción para evitar condiciones de carrera
            with transaction.atomic():
                self.ticket_number = self.generate_ticket_number()
                # Verificar que el número sea único (por si acaso)
                counter = 1
                original_number = self.ticket_number
                while Ticket.objects.filter(ticket_number=self.ticket_number).exists():
                    user_id_formatted = f"{self.created_by.id:02d}"
                    sequence_num = int(original_number.split('-')[1]) + counter
                    self.ticket_number = f"{user_id_formatted}-{sequence_num:03d}"
                    counter += 1
        
        super().save(*args, **kwargs)
        
        # Enviar notificación de Telegram solo para tickets nuevos
        if is_new_ticket:
            try:
                import logging
                logger = logging.getLogger(__name__)
                logger.info(f"Ticket nuevo creado: {self.ticket_number}, iniciando notificación Telegram")
                
                from .telegram_utils import notify_ticket_created
                result = notify_ticket_created(self)
                logger.info(f"Resultado notificación Telegram: {result}")
            except Exception as e:
                # No queremos que falle la creación del ticket si falla la notificación
                import logging
                logger = logging.getLogger(__name__)
                logger.error(f"Error al enviar notificación de Telegram para ticket {self.ticket_number}: {e}")
                import traceback
                logger.error(f"Traceback completo: {traceback.format_exc()}")
    
    def __str__(self):
        if self.ticket_number:
            return f"{self.ticket_number} - {self.title}"
        return f"#{self.id} - {self.title}"
    
    def get_priority_badge_class(self):
        """Retorna la clase CSS para el badge de prioridad"""
        priority_classes = {
            'low': 'bg-success',
            'medium': 'bg-warning',
            'high': 'bg-danger',
            'urgent': 'bg-dark',
        }
        return priority_classes.get(self.priority, 'bg-secondary')
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado"""
        status_classes = {
            'open': 'bg-primary',
            'in_progress': 'bg-warning',
            'resolved': 'bg-success',
            'closed': 'bg-secondary',
        }
        return status_classes.get(self.status, 'bg-secondary')
    
    def get_type_badge_class(self):
        """Retorna la clase CSS para el badge de tipo"""
        type_classes = {
            'desarrollo': 'bg-info',
            'error': 'bg-danger',
        }
        return type_classes.get(self.ticket_type, 'bg-secondary')
    
    def get_age_in_hours(self):
        """Retorna la antigüedad del ticket en horas"""
        from django.utils import timezone
        now = timezone.now()
        time_diff = now - self.created_at
        return time_diff.total_seconds() / 3600
    
    def get_age_badge_class(self):
        """Retorna la clase CSS para el badge de antigüedad según las horas"""
        hours = self.get_age_in_hours()
        if hours < 72:  # Menos de 72 horas - Verde
            return 'bg-success'
        elif hours <= 140:  # Entre 72 y 140 horas - Naranja
            return 'bg-warning'
        else:  # Más de 140 horas - Rojo
            return 'bg-danger'
    
    def get_age_display(self):
        """Retorna la antigüedad formateada para mostrar"""
        hours = self.get_age_in_hours()
        if hours < 24:
            return f"{int(hours)}h"
        elif hours < 72:
            days = int(hours / 24)
            remaining_hours = int(hours % 24)
            if remaining_hours > 0:
                return f"{days}d {remaining_hours}h"
            else:
                return f"{days}d"
        else:
            days = int(hours / 24)
            return f"{days}d"
    
    def get_public_url(self):
        """Retorna la URL pública del ticket si está habilitado para compartir"""
        if self.is_public_shareable:
            from django.urls import reverse
            return reverse('tickets:public_ticket', kwargs={'token': self.public_share_token})
        return None
    
    def regenerate_public_token(self):
        """Regenera el token público del ticket"""
        self.public_share_token = uuid.uuid4()
        self.save(update_fields=['public_share_token'])
    
    def approve_ticket(self, user):
        """Aprobar el ticket por un usuario"""
        from django.utils import timezone
        self.is_approved = True
        self.approved_by = user
        self.approved_at = timezone.now()
        self.save(update_fields=['is_approved', 'approved_by', 'approved_at'])
    
    def unapprove_ticket(self):
        """Quitar la aprobación del ticket"""
        self.is_approved = False
        self.approved_by = None
        self.approved_at = None
        self.save(update_fields=['is_approved', 'approved_by', 'approved_at'])


def ticket_attachment_upload_path(instance, filename):
    """Función para determinar dónde subir los adjuntos"""
    return f'ticket_attachments/ticket_{instance.ticket.id}/{filename}'


class TicketAttachment(models.Model):
    ticket = models.ForeignKey(
        Ticket, 
        on_delete=models.CASCADE, 
        related_name='attachments',
        verbose_name='Ticket'
    )
    file = models.FileField(
        upload_to=ticket_attachment_upload_path,
        verbose_name='Archivo'
    )
    original_filename = models.CharField(
        max_length=255,
        verbose_name='Nombre original'
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Subido por'
    )
    uploaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de subida'
    )
    file_size = models.PositiveIntegerField(
        default=0,
        verbose_name='Tamaño del archivo (bytes)'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Adjunto'
        verbose_name_plural = 'Adjuntos'
    
    def __str__(self):
        return f"{self.original_filename} - Ticket #{self.ticket.id}"
    
    def get_file_size_display(self):
        """Retorna el tamaño del archivo en formato legible"""
        if self.file_size < 1024:
            return f"{self.file_size} bytes"
        elif self.file_size < 1024 * 1024:
            return f"{self.file_size / 1024:.1f} KB"
        else:
            return f"{self.file_size / (1024 * 1024):.1f} MB"
    
    def get_file_extension(self):
        """Retorna la extensión del archivo"""
        return os.path.splitext(self.original_filename)[1].lower()
    
    def delete(self, *args, **kwargs):
        """Elimina el archivo físico cuando se elimina el registro"""
        if self.file:
            if os.path.isfile(self.file.path):
                os.remove(self.file.path)
        super().delete(*args, **kwargs)


class TicketComment(models.Model):
    """Modelo para comentarios en tickets"""
    ticket = models.ForeignKey(
        Ticket,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Ticket'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    content = models.TextField(
        verbose_name='Comentario',
        help_text='Escribe tu comentario aquí'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Comentario'
        verbose_name_plural = 'Comentarios'
    
    def __str__(self):
        return f"Comentario de {self.user.username} en Ticket #{self.ticket.ticket_number}"
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar este comentario"""
        # Solo el autor del comentario o un agente puede editarlo
        from .utils import is_agent
        return self.user == user or is_agent(user)
    
    def can_delete(self, user):
        """Verifica si un usuario puede eliminar este comentario"""
        # Solo el autor del comentario o un agente puede eliminarlo
        from .utils import is_agent
        return self.user == user or is_agent(user)


class UserProfile(models.Model):
    """Modelo para extender la información del usuario"""
    user = models.OneToOneField(
        User,
        on_delete=models.CASCADE,
        related_name='profile',
        verbose_name='Usuario'
    )
    phone = models.CharField(
        max_length=20,
        blank=True,
        null=True,
        verbose_name='Teléfono',
        help_text='Número de teléfono de contacto'
    )
    bio = models.TextField(
        max_length=500,
        blank=True,
        verbose_name='Biografía',
        help_text='Descripción corta del usuario'
    )
    cargo = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Cargo',
        help_text='Cargo o posición del usuario en la empresa'
    )
    descripcion_cargo = models.TextField(
        max_length=300,
        blank=True,
        null=True,
        verbose_name='Descripción del Cargo',
        help_text='Descripción detallada de las responsabilidades del cargo'
    )
    notification_preferences = models.JSONField(
        default=dict,
        blank=True,
        verbose_name='Preferencias de notificación',
        help_text='Configuración de notificaciones del usuario'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='users',
        verbose_name='Empresa',
        help_text='Empresa a la que pertenece el usuario'
    )
    
    # Campos de facturación
    precio_hora = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0.00,
        verbose_name='Precio por hora',
        help_text='Precio de venta por hora de trabajo (€)'
    )
    coste_hora = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0.00,
        verbose_name='Coste por hora',
        help_text='Coste interno por hora de trabajo (€)'
    )
    
    # Token para acceso público a tareas
    public_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        verbose_name='Token público',
        help_text='Token único para acceso público a tareas'
    )
    
    # Formulario de contacto público
    enable_public_contact_form = models.BooleanField(
        default=False,
        verbose_name='Activar formulario de contacto público',
        help_text='Permite que visitantes envíen formularios de contacto para crear empresas'
    )
    
    # Configuración de ubicación para el clima
    city = models.CharField(
        max_length=100,
        default='Valencia',
        verbose_name='Ciudad',
        help_text='Ciudad para mostrar información del clima'
    )
    country = models.CharField(
        max_length=100,
        default='España',
        verbose_name='País',
        help_text='País para mostrar información del clima'
    )
    country_code = models.CharField(
        max_length=2,
        default='ES',
        verbose_name='Código de país',
        help_text='Código ISO de 2 letras del país (ej: ES, US, FR)'
    )
    
    # Fecha de cumpleaños
    birth_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de cumpleaños',
        help_text='Fecha de nacimiento del usuario'
    )
    
    # Términos y condiciones para cotizaciones
    quote_terms_conditions = models.TextField(
        blank=True,
        verbose_name='Términos y Condiciones de Cotizaciones',
        help_text='Términos y condiciones que se mostrarán en las cotizaciones rápidas'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    def get_total_hours(self):
        """Retorna el total de horas trabajadas por el usuario"""
        from datetime import timedelta
        total_minutes = 0
        for entry in self.user.time_entries.filter(fecha_salida__isnull=False):
            total_minutes += entry.duracion_trabajada
        return round(total_minutes / 60, 1)
    
    def get_monthly_hours(self):
        """Retorna las horas trabajadas en el mes actual"""
        from datetime import datetime, timedelta
        now = timezone.now()
        start_month = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        total_minutes = 0
        entries = self.user.time_entries.filter(
            fecha_entrada__gte=start_month,
            fecha_salida__isnull=False
        )
        for entry in entries:
            total_minutes += entry.duracion_trabajada
        return round(total_minutes / 60, 1)
    
    def get_weekly_hours(self):
        """Retorna las horas trabajadas en esta semana"""
        from datetime import datetime, timedelta
        now = timezone.now()
        # Calcular el inicio de la semana (lunes)
        start_week = now - timedelta(days=now.weekday())
        start_week = start_week.replace(hour=0, minute=0, second=0, microsecond=0)
        
        total_minutes = 0
        entries = self.user.time_entries.filter(
            fecha_entrada__gte=start_week,
            fecha_salida__isnull=False
        )
        for entry in entries:
            total_minutes += entry.duracion_trabajada
        return round(total_minutes / 60, 1)
    
    def get_daily_hours(self):
        """Retorna las horas trabajadas hoy"""
        from datetime import datetime
        today = timezone.now().date()
        
        total_minutes = 0
        entries = self.user.time_entries.filter(
            fecha_entrada__date=today,
            fecha_salida__isnull=False
        )
        for entry in entries:
            total_minutes += entry.duracion_trabajada
            
        # También incluir sesión activa si existe
        active_entry = self.user.time_entries.filter(
            fecha_entrada__date=today,
            fecha_salida__isnull=True
        ).first()
        
        if active_entry:
            total_minutes += active_entry.duracion_trabajada
            
        return round(total_minutes / 60, 1)
    
    class Meta:
        verbose_name = 'Perfil de Usuario'
        verbose_name_plural = 'Perfiles de Usuario'
    
    def __str__(self):
        return f"Perfil de {self.user.username}"
    
    def get_full_name(self):
        """Retorna el nombre completo del usuario"""
        if self.user.first_name and self.user.last_name:
            return f"{self.user.first_name} {self.user.last_name}"
        elif self.user.first_name:
            return self.user.first_name
        else:
            return self.user.username

    def regenerate_token(self):
        """Regenera el token público"""
        self.public_token = uuid.uuid4()
        self.save()
        return self.public_token


@receiver(post_save, sender=User)
def create_user_profile(sender, instance, created, **kwargs):
    """Crea automáticamente un perfil cuando se crea un usuario"""
    if created:
        UserProfile.objects.create(user=instance)


@receiver(post_save, sender=User)
def save_user_profile(sender, instance, **kwargs):
    """Guarda el perfil cuando se guarda el usuario"""
    if hasattr(instance, 'profile'):
        instance.profile.save()
    else:
        UserProfile.objects.create(user=instance)


class UserNote(models.Model):
    """Modelo para notas internas asociadas a usuarios"""
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo de la nota'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Contenido detallado de la nota'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='notes',
        verbose_name='Usuario asociado',
        help_text='Usuario al que está asociada esta nota'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_notes',
        verbose_name='Creado por',
        help_text='Agente que creó la nota'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Empresa',
        help_text='Empresa asociada a esta nota'
    )
    tickets = models.ManyToManyField(
        Ticket,
        blank=True,
        related_name='notes',
        verbose_name='Tickets relacionados',
        help_text='Tickets asociados a esta nota'
    )
    is_private = models.BooleanField(
        default=False,
        verbose_name='Nota privada',
        help_text='Si está marcada, solo tú podrás ver, editar y eliminar esta nota'
    )
    share_token = models.CharField(
        max_length=50,
        blank=True,
        null=True,
        unique=True,
        verbose_name='Token de compartir',
        help_text='Token único para compartir la nota públicamente'
    )
    is_shareable = models.BooleanField(
        default=False,
        verbose_name='Compartible públicamente',
        help_text='Si la nota puede ser compartida mediante un enlace público'
    )
    view_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de visualizaciones',
        help_text='Veces que la nota ha sido vista mediante enlace público'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Nota de Usuario'
        verbose_name_plural = 'Notas de Usuario'
    
    def __str__(self):
        return f"{self.title} - {self.user.username}"
    
    def can_view(self, user):
        """Verifica si un usuario puede ver esta nota"""
        from .utils import is_agent
        
        # Si la nota es privada, solo el creador puede verla
        if self.is_private and self.created_by != user:
            return False
            
        # Los agentes pueden ver notas públicas
        if is_agent(user):
            return True
            
        # El creador puede ver sus propias notas (privadas y públicas)
        if self.created_by == user:
            return True
            
        # Si la nota no es privada, el usuario asociado puede verla
        if not self.is_private and self.user == user:
            return True
            
        return False
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar esta nota"""
        from .utils import is_agent
        
        # Si la nota es privada, solo el creador puede editarla
        if self.is_private and self.created_by != user:
            return False
            
        # Solo los agentes pueden crear/editar notas (para notas públicas)
        # O el creador puede editar sus propias notas privadas
        return is_agent(user) or (self.is_private and self.created_by == user)
    
    def can_delete(self, user):
        """Verifica si un usuario puede eliminar esta nota"""
        from .utils import is_agent
        
        # Si la nota es privada, solo el creador puede eliminarla
        if self.is_private and self.created_by != user:
            return False
            
        # Solo los agentes pueden eliminar notas (para notas públicas)
        # O el creador puede eliminar sus propias notas privadas
        return is_agent(user) or (self.is_private and self.created_by == user)
    
    def generate_share_token(self):
        """Genera un token único para compartir la nota"""
        import uuid
        self.share_token = str(uuid.uuid4())[:12]
        self.is_shareable = True
        self.save()
        return self.share_token
    
    def get_public_url(self, request=None):
        """Obtiene la URL pública de la nota si es compartible"""
        if not self.is_shareable or not self.share_token:
            return None
        
        if request:
            protocol = 'https' if request.is_secure() else 'http'
            domain = request.get_host()
            return f"{protocol}://{domain}/shared/note/{self.share_token}/"
        else:
            return f"/shared/note/{self.share_token}/"
    
    def disable_sharing(self):
        """Deshabilita el compartir público de la nota"""
        self.is_shareable = False
        self.save()
    
    def increment_view_count(self):
        """Incrementa el contador de visualizaciones"""
        self.view_count += 1
        self.save(update_fields=['view_count'])
    
    def get_related_tickets_count(self):
        """Retorna el número de tickets relacionados"""
        return self.tickets.count()
    
    def get_tickets_summary(self):
        """Retorna un resumen de los tickets relacionados"""
        tickets = self.tickets.all()[:3]  # Mostrar máximo 3 tickets
        if tickets:
            summary = ', '.join([f"#{ticket.ticket_number}" for ticket in tickets])
            remaining = self.tickets.count() - 3
            if remaining > 0:
                summary += f" y {remaining} más"
            return summary
        return "Sin tickets asociados"


class TimeEntry(models.Model):
    """Modelo para registrar entrada y salida de trabajo de empleados (agentes)"""
    
    END_REASON_CHOICES = [
        ('comer', 'Comer'),
        ('fin_jornada', 'Fin de Jornada'),
        ('otro', 'Otro'),
    ]
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Empleado'
    )
    project = models.ForeignKey(
        Project,
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Proyecto',
        null=True,
        blank=True
    )
    ticket = models.ForeignKey(
        'Ticket',
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Ticket',
        null=True,
        blank=True,
        help_text='Ticket en el que se trabajará (opcional)'
    )
    work_order = models.ForeignKey(
        'WorkOrder',
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Orden de Trabajo',
        null=True,
        blank=True,
        help_text='Orden de trabajo en la que se trabajará (opcional)'
    )
    task = models.ForeignKey(
        'Task',
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Tarea',
        null=True,
        blank=True,
        help_text='Tarea en la que se trabajará (opcional)'
    )
    fecha_entrada = models.DateTimeField(
        verbose_name='Fecha y hora de entrada'
    )
    fecha_salida = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha y hora de salida'
    )
    end_reason = models.CharField(
        max_length=20,
        choices=END_REASON_CHOICES,
        null=True,
        blank=True,
        verbose_name='Causa de finalización',
        help_text='Motivo por el cual se finaliza la jornada'
    )
    notas = models.TextField(
        blank=True,
        max_length=500,
        verbose_name='Notas',
        help_text='Comentarios opcionales sobre la jornada'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Registrado el'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-fecha_entrada']
        verbose_name = 'Registro de Horario'
        verbose_name_plural = 'Registros de Horario'
        unique_together = ['user', 'fecha_entrada']
    
    def __str__(self):
        fecha = self.fecha_entrada.strftime('%d/%m/%Y')
        estado = "Activo" if not self.fecha_salida else "Completado"
        return f"{self.user.username} - {fecha} ({estado})"
    
    @property
    def is_active(self):
        """Verifica si el registro está activo (sin fecha de salida)"""
        return self.fecha_salida is None
    
    @property
    def duracion_trabajada(self):
        """Calcula la duración trabajada en minutos"""
        if not self.fecha_salida:
            # Si aún está trabajando, calcular desde entrada hasta ahora
            return int((timezone.now() - self.fecha_entrada).total_seconds() / 60)
        return int((self.fecha_salida - self.fecha_entrada).total_seconds() / 60)
    
    @property
    def duracion_formateada(self):
        """Retorna la duración en formato HH:MM"""
        minutos = self.duracion_trabajada
        horas = minutos // 60
        mins = minutos % 60
        return f"{horas:02d}:{mins:02d}"
    
    @property
    def horas_trabajadas(self):
        """Retorna las horas trabajadas como decimal"""
        return round(self.duracion_trabajada / 60, 2)
    
    def can_view(self, user):
        """Verifica si un usuario puede ver este registro"""
        from .utils import is_agent
        # Solo los agentes pueden ver registros de horario
        if not is_agent(user):
            return False
        # Los agentes pueden ver solo sus propios registros
        return self.user == user
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar este registro"""
        from .utils import is_agent
        # Solo el propietario del registro puede editarlo
        return is_agent(user) and self.user == user
    
    def finalizar_jornada(self, notas_salida=None):
        """Marca la salida de la jornada laboral"""
        if self.fecha_salida:
            raise ValueError("Esta jornada ya ha sido finalizada")
        
        self.fecha_salida = timezone.now()
        if notas_salida:
            self.notas = notas_salida
        self.save()
        return self
    
    @classmethod
    def get_active_entry(cls, user):
        """Obtiene el registro activo (sin salida) de un usuario"""
        return cls.objects.filter(user=user, fecha_salida__isnull=True).first()
    
    @classmethod
    def get_active_entry_for_agent(cls, user):
        """Obtiene el registro activo (sin salida) de un agente"""
        from .utils import is_agent
        if not is_agent(user):
            return None
        return cls.objects.filter(user=user, fecha_salida__isnull=True).first()
    
    @classmethod
    def create_entry(cls, user, project=None, ticket=None, work_order=None, task=None, notas_entrada=None):
        """Crea un nuevo registro de entrada (para agentes)"""
        from .utils import is_agent
        if not is_agent(user):
            raise ValueError("Solo los agentes pueden registrar horarios")
        
        # Verificar que no haya una entrada activa
        active_entry = cls.get_active_entry(user)
        if active_entry:
            raise ValueError("Ya tienes una jornada activa. Finalízala antes de crear una nueva.")
        
        return cls.objects.create(
            user=user,
            project=project,
            ticket=ticket,
            work_order=work_order,
            task=task,
            fecha_entrada=timezone.now(),
            notas=notas_entrada or ''
        )
    
    @classmethod
    def create_public_entry(cls, user, notas_entrada=None):
        """Crea un nuevo registro de entrada para acceso público (sin restricción de agente)"""
        # Verificar que no haya una entrada activa
        active_entry = cls.get_active_entry(user)
        if active_entry:
            raise ValueError("Ya tienes una jornada activa. Finalízala antes de crear una nueva.")
        
        return cls.objects.create(
            user=user,
            project=None,
            ticket=None,
            work_order=None,
            task=None,
            fecha_entrada=timezone.now(),
            notas=notas_entrada or ''
        )


class TimeEntryAuditLog(models.Model):
    """Modelo para auditar cambios en registros de tiempo"""
    
    time_entry = models.ForeignKey(
        TimeEntry,
        on_delete=models.CASCADE,
        related_name='audit_logs',
        verbose_name='Registro de Tiempo'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario que modificó'
    )
    timestamp = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha y hora de modificación'
    )
    field_name = models.CharField(
        max_length=50,
        verbose_name='Campo modificado'
    )
    old_value = models.TextField(
        null=True,
        blank=True,
        verbose_name='Valor anterior'
    )
    new_value = models.TextField(
        null=True,
        blank=True,
        verbose_name='Valor nuevo'
    )
    change_reason = models.CharField(
        max_length=200,
        null=True,
        blank=True,
        verbose_name='Razón del cambio'
    )
    
    class Meta:
        ordering = ['-timestamp']
        verbose_name = 'Log de Auditoría de Tiempo'
        verbose_name_plural = 'Logs de Auditoría de Tiempo'
    
    def __str__(self):
        return f"{self.time_entry} - {self.field_name} - {self.timestamp.strftime('%d/%m/%Y %H:%M')}"


class PublicTimeAccess(models.Model):
    """Modelo para gestionar acceso público al control de horario"""
    
    user = models.OneToOneField(
        User,
        on_delete=models.CASCADE,
        related_name='public_time_access',
        verbose_name='Usuario'
    )
    token = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Token de Acceso',
        help_text='Token único para acceso público'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está activo, el usuario puede usar el acceso público'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Creado el'
    )
    last_used = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Último uso'
    )
    
    # Configuraciones adicionales
    require_location = models.BooleanField(
        default=False,
        verbose_name='Requiere ubicación',
        help_text='Si está activado, registrará la ubicación GPS al marcar entrada/salida'
    )
    allowed_ip_addresses = models.TextField(
        blank=True,
        verbose_name='IPs permitidas',
        help_text='Lista de IPs permitidas separadas por comas (opcional)'
    )
    
    class Meta:
        verbose_name = 'Acceso Público de Horario'
        verbose_name_plural = 'Accesos Públicos de Horario'
        ordering = ['-created_at']
    
    def __str__(self):
        status = "Activo" if self.is_active else "Inactivo"
        return f"{self.user.get_full_name() or self.user.username} - {status}"
    
    def save(self, *args, **kwargs):
        if not self.token:
            import secrets
            self.token = secrets.token_urlsafe(32)
        super().save(*args, **kwargs)
    
    @property
    def public_url(self):
        """Genera la URL pública para el acceso al control de horario"""
        from django.urls import reverse
        return reverse('public_time_clock', kwargs={'token': self.token})
    
    def is_ip_allowed(self, ip_address):
        """Verifica si una IP está permitida"""
        if not self.allowed_ip_addresses:
            return True
        
        allowed_ips = [ip.strip() for ip in self.allowed_ip_addresses.split(',')]
        return ip_address in allowed_ips
    
    def update_last_used(self):
        """Actualiza la fecha de último uso"""
        self.last_used = timezone.now()
        self.save(update_fields=['last_used'])


class SystemConfiguration(models.Model):
    """Modelo para configuraciones generales del sistema"""
    
    # Configuración de registro de usuarios
    allow_user_registration = models.BooleanField(
        default=True,
        verbose_name='Permitir registro de usuarios',
        help_text='Permite que nuevos usuarios se registren en el sistema'
    )
    
    # Configuraciones adicionales que se pueden agregar en el futuro
    site_name = models.CharField(
        max_length=100,
        default='TicketPro',
        verbose_name='Nombre del sitio',
        help_text='Nombre que aparece en el encabezado del sistema'
    )
    
    # Configuraciones de tickets
    default_ticket_priority = models.CharField(
        max_length=20,
        choices=[
            ('low', 'Baja'),
            ('medium', 'Media'),
            ('high', 'Alta'),
            ('urgent', 'Urgente')
        ],
        default='medium',
        verbose_name='Prioridad por defecto de tickets',
        help_text='Prioridad asignada automáticamente a nuevos tickets'
    )
    
    # Configuración de moneda
    default_currency = models.CharField(
        max_length=10,
        choices=[
            ('EUR', 'Euro (€)'),
            ('USD', 'Dólar Americano ($)'),
            ('GBP', 'Libra Esterlina (£)'),
            ('JPY', 'Yen Japonés (¥)'),
            ('CAD', 'Dólar Canadiense (C$)'),
            ('AUD', 'Dólar Australiano (A$)'),
            ('CHF', 'Franco Suizo (CHF)'),
            ('CNY', 'Yuan Chino (¥)'),
            ('MXN', 'Peso Mexicano ($)'),
            ('COP', 'Peso Colombiano ($)'),
        ],
        default='EUR',
        verbose_name='Moneda por defecto',
        help_text='Moneda utilizada para mostrar valores en el sistema'
    )
    
    # Configuración de Chat IA
    openai_api_key = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='API Key de OpenAI',
        help_text='Clave de API para conectar con ChatGPT'
    )
    
    openai_model = models.CharField(
        max_length=50,
        default='gpt-4o',
        verbose_name='Modelo de OpenAI',
        help_text='Modelo de IA a utilizar por defecto (ej: gpt-4o, gpt-3.5-turbo)'
    )
    
    ai_chat_enabled = models.BooleanField(
        default=False,
        verbose_name='Chat IA habilitado',
        help_text='Habilita o deshabilita la funcionalidad de chat con IA'
    )
    
    # Configuración de análisis de empleados con IA
    ai_employee_analysis_prompt = models.TextField(
        default='Analiza el siguiente perfil de candidato y dame un resumen profesional y una puntuación del 1 al 10 para el cargo solicitado. Considera experiencia, formación, habilidades y adecuación al puesto.\n\nDatos del candidato:\n{datos}\n\nPor favor, responde en este formato:\nResumen: [Aquí tu análisis profesional del candidato]\nPuntuación: [Número del 1 al 10]',
        verbose_name='Prompt para análisis de empleados',
        help_text='Plantilla del prompt que se enviará a la IA para analizar candidatos. Usa {datos} donde se insertarán los datos del candidato.'
    )
    
    # Configuración de Telegram
    enable_telegram_notifications = models.BooleanField(
        default=False,
        verbose_name='Activar notificaciones de Telegram',
        help_text='Envía notificaciones a un grupo de Telegram cuando se crean tickets'
    )
    telegram_bot_token = models.CharField(
        max_length=200,
        blank=True,
        null=True,
        verbose_name='Token del Bot de Telegram',
        help_text='Token del bot de Telegram proporcionado por @BotFather'
    )
    telegram_chat_id = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='ID del Chat/Grupo de Telegram',
        help_text='ID del grupo o chat donde enviar las notificaciones (ej: -100123456789)'
    )
    
    # Configuración de Email
    enable_email_notifications = models.BooleanField(
        default=False,
        verbose_name='Activar notificaciones por email',
        help_text='Envía notificaciones por email cuando se reciben contactos desde la web'
    )
    email_host = models.CharField(
        max_length=255,
        blank=True,
        default='smtp.gmail.com',
        verbose_name='Servidor SMTP',
        help_text='Servidor de correo saliente (ej: smtp.gmail.com, mail.tu-dominio.com)'
    )
    email_port = models.IntegerField(
        blank=True,
        null=True,
        default=587,
        verbose_name='Puerto SMTP',
        help_text='Puerto del servidor SMTP (587 para TLS, 465 para SSL, 25 para sin cifrado)'
    )
    email_host_user = models.EmailField(
        blank=True,
        verbose_name='Usuario SMTP',
        help_text='Dirección de email para autenticación SMTP'
    )
    email_host_password = models.CharField(
        max_length=255,
        blank=True,
        verbose_name='Contraseña SMTP',
        help_text='Contraseña o token de aplicación para SMTP'
    )
    email_use_tls = models.BooleanField(
        default=True,
        verbose_name='Usar TLS',
        help_text='Activar cifrado TLS (recomendado para puerto 587)'
    )
    email_use_ssl = models.BooleanField(
        default=False,
        verbose_name='Usar SSL',
        help_text='Activar cifrado SSL (para puerto 465)'
    )
    email_from = models.EmailField(
        blank=True,
        verbose_name='Email remitente',
        help_text='Dirección de email que aparecerá como remitente'
    )
    notification_emails = models.TextField(
        blank=True,
        verbose_name='Emails de notificación',
        help_text='Direcciones de email donde enviar notificaciones (una por línea)'
    )
    
    # Configuración de PayPal
    paypal_enabled = models.BooleanField(
        default=False,
        verbose_name='PayPal habilitado',
        help_text='Habilita o deshabilita los pagos con PayPal'
    )
    paypal_mode = models.CharField(
        max_length=20,
        choices=[
            ('sandbox', 'Sandbox (Pruebas)'),
            ('live', 'Live (Producción)')
        ],
        default='sandbox',
        verbose_name='Modo de PayPal',
        help_text='Modo de operación de PayPal'
    )
    paypal_client_id = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='PayPal Client ID',
        help_text='Client ID de la aplicación PayPal'
    )
    paypal_client_secret = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='PayPal Client Secret',
        help_text='Client Secret de la aplicación PayPal'
    )
    paypal_webhook_id = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='PayPal Webhook ID',
        help_text='ID del webhook configurado en PayPal para recibir notificaciones de pagos'
    )
    
    # Token público para catálogo de productos
    products_catalog_token = models.UUIDField(
        null=True,
        blank=True,
        unique=True,
        verbose_name='Token del Catálogo Público',
        help_text='Token único para acceso público al catálogo de productos'
    )
    
    # Metadatos
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        verbose_name = 'Configuración del Sistema'
        verbose_name_plural = 'Configuración del Sistema'
    
    def __str__(self):
        return f'Configuración del Sistema - {self.site_name}'
    
    @classmethod
    def get_config(cls):
        """Obtener la configuración del sistema (singleton)"""
        config, created = cls.objects.get_or_create(pk=1)
        return config
    
    def get_currency_symbol(self):
        """Obtener el símbolo de la moneda configurada"""
        currency_symbols = {
            'EUR': '€',
            'USD': '$',
            'GBP': '£',
            'JPY': '¥',
            'CAD': 'C$',
            'AUD': 'A$',
            'CHF': 'CHF',
            'CNY': '¥',
            'MXN': '$',
            'COP': '$',
        }
        return currency_symbols.get(self.default_currency, '€')
    
    def format_currency(self, amount):
        """Formatear un monto con la moneda configurada"""
        if self.default_currency in ['EUR', 'GBP']:
            return f"{self.get_currency_symbol()}{amount:,.2f}"
        else:
            return f"{amount:,.2f} {self.get_currency_symbol()}"
    
    def save(self, *args, **kwargs):
        # Asegurar que solo exista una instancia de configuración
        self.pk = 1
        super().save(*args, **kwargs)


class Document(models.Model):
    """Modelo para archivos de documentación compartibles públicamente"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre descriptivo del documento'
    )
    description = models.TextField(
        max_length=500,
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional del contenido del documento'
    )
    file = models.FileField(
        upload_to='documents/%Y/%m/',
        verbose_name='Archivo',
        help_text='Archivo de documentación (PDF, DOC, DOCX, TXT, etc.)'
    )
    file_size = models.PositiveIntegerField(
        verbose_name='Tamaño del archivo (bytes)',
        editable=False
    )
    file_type = models.CharField(
        max_length=100,
        verbose_name='Tipo de archivo',
        editable=False
    )
    
    # Sistema de compartir público (como en tickets)
    public_share_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de compartir público',
        help_text='Token único para compartir el documento públicamente'
    )
    is_public = models.BooleanField(
        default=True,
        verbose_name='Disponible públicamente',
        help_text='Si está marcado, el documento puede ser accedido sin autenticación'
    )
    download_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de descargas',
        help_text='Contador de veces que se ha descargado el documento'
    )
    
    # Metadatos del documento
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_documents',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='documents',
        verbose_name='Empresa/Cliente',
        help_text='Empresa o cliente asociado al documento'
    )
    tags = models.CharField(
        max_length=300,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas para categorizar el documento'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Documento'
        verbose_name_plural = 'Documentos'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        if self.file:
            self.file_size = self.file.size
            self.file_type = self.file.name.split('.')[-1].upper() if '.' in self.file.name else 'UNKNOWN'
        super().save(*args, **kwargs)
    
    @property
    def file_size_formatted(self):
        """Retorna el tamaño del archivo en formato legible"""
        if self.file_size < 1024:
            return f"{self.file_size} B"
        elif self.file_size < 1024 * 1024:
            return f"{self.file_size / 1024:.1f} KB"
        elif self.file_size < 1024 * 1024 * 1024:
            return f"{self.file_size / (1024 * 1024):.1f} MB"
        else:
            return f"{self.file_size / (1024 * 1024 * 1024):.1f} GB"
    
    @property
    def public_url(self):
        """URL pública para compartir el documento"""
        from django.urls import reverse
        return reverse('document_public', kwargs={'token': self.public_share_token})
    
    def increment_download_count(self):
        """Incrementa el contador de descargas"""
        self.download_count += 1
        self.save(update_fields=['download_count'])
    
    def get_tags_list(self):
        """Retorna las etiquetas como lista"""
        if self.tags:
            return [tag.strip() for tag in self.tags.split(',') if tag.strip()]
        return []


class UrlManager(models.Model):
    """Modelo para gestionar URLs con credenciales encriptadas"""
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre descriptivo para la URL'
    )
    url = models.URLField(
        max_length=500,
        verbose_name='URL',
        help_text='Dirección web completa (incluye http:// o https://)'
    )
    username = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Usuario',
        help_text='Nombre de usuario para acceder'
    )
    encrypted_password = models.TextField(
        blank=True,
        null=True,
        verbose_name='Contraseña encriptada',
        help_text='Contraseña almacenada de forma segura'
    )
    description = models.TextField(
        blank=True,
        null=True,
        verbose_name='Descripción',
        help_text='Información adicional sobre esta URL'
    )
    category = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Categoría',
        help_text='Ej: Servidor, Base de datos, Servicio web, etc.'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Indica si la URL está activa o no'
    )
    is_principal = models.BooleanField(
        default=False,
        verbose_name='URL Principal',
        help_text='Marcar como URL principal para mostrar en el dashboard'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_urls',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    last_accessed = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Último acceso'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Gestor de URL'
        verbose_name_plural = 'Gestores de URLs'
        indexes = [
            models.Index(fields=['category', 'is_active']),
            models.Index(fields=['created_by', 'is_active']),
            models.Index(fields=['is_principal', 'is_active']),
        ]
    
    def __str__(self):
        return f"{self.title} - {self.url}"
    
    def set_password(self, raw_password):
        """Encripta y guarda la contraseña"""
        if not raw_password:
            self.encrypted_password = None
            return
            
        from cryptography.fernet import Fernet
        from django.conf import settings
        import base64
        
        # Generar una clave de encriptación basada en SECRET_KEY
        key = base64.urlsafe_b64encode(settings.SECRET_KEY[:32].encode().ljust(32, b'0'))
        fernet = Fernet(key)
        
        # Encriptar la contraseña
        encrypted_password = fernet.encrypt(raw_password.encode())
        self.encrypted_password = base64.urlsafe_b64encode(encrypted_password).decode()
    
    def get_password(self):
        """Desencripta y retorna la contraseña"""
        if not self.encrypted_password:
            return ""
            
        from cryptography.fernet import Fernet
        from django.conf import settings
        import base64
        
        try:
            # Generar la misma clave de encriptación
            key = base64.urlsafe_b64encode(settings.SECRET_KEY[:32].encode().ljust(32, b'0'))
            fernet = Fernet(key)
            
            # Desencriptar la contraseña
            encrypted_data = base64.urlsafe_b64decode(self.encrypted_password.encode())
            decrypted_password = fernet.decrypt(encrypted_data)
            return decrypted_password.decode()
        except Exception as e:
            return f"Error al desencriptar: {str(e)}"
    
    def can_view(self, user):
        """Verifica si un usuario puede ver esta URL"""
        from .utils import is_agent
        # Solo los agentes pueden ver las URLs
        return is_agent(user)
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar esta URL"""
        from .utils import is_agent
        # Solo los agentes pueden editar URLs
        return is_agent(user)
    
    def mark_accessed(self):
        """Marca la URL como accedida recientemente"""
        self.last_accessed = timezone.now()
        self.save(update_fields=['last_accessed'])
    
    @property
    def domain(self):
        """Extrae el dominio de la URL"""
        from urllib.parse import urlparse
        try:
            return urlparse(self.url).netloc
        except:
            return 'URL inválida'
    
    @property
    def status_display(self):
        """Retorna el estado en formato display"""
        return "Activo" if self.is_active else "Inactivo"


class WorkOrder(models.Model):
    """Modelo para gestionar órdenes de trabajo"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('accepted', 'Aceptado'),
        ('finished', 'Terminado'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo de la orden de trabajo'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del trabajo a realizar'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='work_orders',
        verbose_name='Empresa',
        help_text='Empresa a la que se dirige la orden de trabajo'
    )
    project = models.ForeignKey(
        Project,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='work_orders',
        verbose_name='Proyecto',
        help_text='Proyecto asociado a la orden de trabajo (opcional)'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado',
        help_text='Estado actual de la orden de trabajo'
    )
    due_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de entrega',
        help_text='Fecha límite para completar la orden (opcional)'
    )
    is_public = models.BooleanField(
        default=True,
        verbose_name='Publicar públicamente',
        help_text='Permite que esta orden sea visible públicamente'
    )
    public_share_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de compartir público',
        help_text='Token único para compartir la orden públicamente'
    )
    estimated_hours = models.DecimalField(
        max_digits=6,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Horas estimadas',
        help_text='Tiempo estimado para completar la orden'
    )
    amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Importe',
        help_text='Precio o costo de la orden de trabajo'
    )
    priority = models.CharField(
        max_length=10,
        choices=[
            ('low', 'Baja'),
            ('medium', 'Media'),
            ('high', 'Alta'),
            ('urgent', 'Urgente'),
        ],
        default='medium',
        verbose_name='Prioridad'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_work_orders',
        verbose_name='Creado por'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='assigned_work_orders',
        verbose_name='Asignado a',
        help_text='Usuario responsable de la orden'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de finalización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Orden de Trabajo'
        verbose_name_plural = 'Órdenes de Trabajo'
        indexes = [
            models.Index(fields=['status', 'company']),
            models.Index(fields=['due_date', 'status']),
            models.Index(fields=['created_by', 'status']),
        ]
    
    def __str__(self):
        return f"OT-{self.id:04d} - {self.title}"
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado"""
        status_classes = {
            'draft': 'bg-secondary',
            'accepted': 'bg-primary',
            'finished': 'bg-success',
        }
        return status_classes.get(self.status, 'bg-secondary')
    
    def get_priority_badge_class(self):
        """Retorna la clase CSS para el badge de prioridad"""
        priority_classes = {
            'low': 'bg-success',
            'medium': 'bg-warning',
            'high': 'bg-danger',
            'urgent': 'bg-dark',
        }
        return priority_classes.get(self.priority, 'bg-secondary')
    
    def get_public_url(self):
        """Retorna la URL pública de la orden si está habilitada para compartir"""
        if self.is_public:
            from django.urls import reverse
            return reverse('public_work_order', kwargs={'token': self.public_share_token})
        return None
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar esta orden"""
        from .utils import is_agent
        return is_agent(user) or user == self.created_by
    
    def can_delete(self, user):
        """Verifica si un usuario puede eliminar esta orden"""
        from .utils import is_agent
        # Solo se puede eliminar si está en borrador
        return (is_agent(user) or user == self.created_by) and self.status == 'draft'
    
    def mark_as_finished(self):
        """Marca la orden como terminada"""
        self.status = 'finished'
        self.completed_at = timezone.now()
        self.save(update_fields=['status', 'completed_at'])
    
    def is_overdue(self):
        """Verifica si la orden está vencida"""
        if self.status == 'finished' or not self.due_date:
            return False
        return timezone.now().date() > self.due_date
    
    def days_until_due(self):
        """Retorna los días hasta la fecha de entrega"""
        if self.status == 'finished' or not self.due_date:
            return None
        delta = self.due_date - timezone.now().date()
        return delta.days
    
    def get_amount_display(self):
        """Retorna el importe formateado"""
        if self.amount:
            return f"${self.amount:,.2f}"
        return "No especificado"
    
    @property
    def order_number(self):
        """Retorna el número de orden formateado"""
        return f"OT-{self.id:04d}"
    
    def get_tasks_summary(self):
        """Retorna un resumen de las tareas"""
        tasks = self.tasks.all()
        total_tasks = tasks.count()
        pending_tasks = tasks.filter(status='pending').count()
        in_progress_tasks = tasks.filter(status='in_progress').count()
        completed_tasks = tasks.filter(status='completed').count()
        
        return {
            'total': total_tasks,
            'pending': pending_tasks,
            'in_progress': in_progress_tasks,
            'completed': completed_tasks,
            'progress_percentage': (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
        }
    
    def get_total_actual_hours(self):
        """Retorna el total de horas reales trabajadas en todas las tareas"""
        return self.tasks.aggregate(
            total=models.Sum('actual_hours')
        )['total'] or 0
    
    def get_total_estimated_hours(self):
        """Retorna el total de horas estimadas de todas las tareas"""
        return self.tasks.aggregate(
            total=models.Sum('estimated_hours')
        )['total'] or 0
    
    def get_completion_percentage(self):
        """Retorna el porcentaje de completitud basado en tareas"""
        summary = self.get_tasks_summary()
        return summary['progress_percentage']
    
    def get_total_hours_logged(self):
        """Retorna el total de horas trabajadas en todas las tareas"""
        from django.db.models import Sum
        from decimal import Decimal
        
        total_hours = Decimal('0')
        task_details = []
        
        for task in self.tasks.all():
            task_total = task.time_entries.aggregate(
                total_hours=Sum('hours')
            )['total_hours']
            
            task_hours = task_total if task_total else Decimal('0')
            total_hours += task_hours
            task_details.append(f"Task {task.id} ({task.title}): {task_hours}h")
        
        # Debug: Imprimir detalles para troubleshooting
        print(f"WorkOrder {self.id} total hours calculation:")
        for detail in task_details:
            print(f"  - {detail}")
        print(f"  Total: {total_hours}h")
        
        return float(total_hours)
    
    def get_progress_percentage(self):
        """Retorna el porcentaje de progreso basado en tareas completadas"""
        total_tasks = self.get_total_tasks()
        if total_tasks == 0:
            return 0
        
        completed_tasks = self.get_completed_tasks()
        progress = round((completed_tasks / total_tasks) * 100, 1)
        
        # Debug: Imprimir los valores para troubleshooting
        print(f"WorkOrder {self.id}: total_tasks={total_tasks}, completed_tasks={completed_tasks}, progress={progress}")
        
        return progress
    
    def get_total_tasks(self):
        """Retorna el número total de tareas"""
        return self.tasks.count()
    
    def get_pending_tasks(self):
        """Retorna el número de tareas pendientes"""
        return self.tasks.filter(status='pending').count()
    
    def get_in_progress_tasks(self):
        """Retorna el número de tareas en progreso"""
        return self.tasks.filter(status='in_progress').count()
    
    def get_completed_tasks(self):
        """Retorna el número de tareas completadas"""
        return self.tasks.filter(status='completed').count()
    
    def update_actual_hours(self):
        """Actualiza las horas reales basándose en las entradas de tiempo de todas las tareas"""
        from decimal import Decimal
        
        total_hours = Decimal('0.00')
        for task in self.tasks.all():
            task.update_actual_hours()
            if task.actual_hours:
                total_hours += task.actual_hours
        
        self.actual_hours = total_hours
        self.save(update_fields=['actual_hours'])
        return total_hours


def work_order_attachment_upload_path(instance, filename):
    """Función para determinar dónde subir los adjuntos de órdenes de trabajo"""
    return f'work_order_attachments/work_order_{instance.work_order.id}/{filename}'


class WorkOrderAttachment(models.Model):
    """Modelo para adjuntos de órdenes de trabajo"""
    work_order = models.ForeignKey(
        WorkOrder, 
        on_delete=models.CASCADE, 
        related_name='attachments',
        verbose_name='Orden de Trabajo'
    )
    file = models.FileField(
        upload_to=work_order_attachment_upload_path,
        verbose_name='Archivo'
    )
    original_filename = models.CharField(
        max_length=255,
        verbose_name='Nombre original'
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Subido por'
    )
    uploaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de subida'
    )
    file_size = models.PositiveIntegerField(
        default=0,
        verbose_name='Tamaño del archivo (bytes)'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Adjunto de Orden de Trabajo'
        verbose_name_plural = 'Adjuntos de Órdenes de Trabajo'
    
    def __str__(self):
        return f'Adjunto: {self.original_filename} - OT: {self.work_order.title}'
    
    def get_file_size_display(self):
        """Retorna el tamaño del archivo en formato legible"""
        size = self.file_size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} TB"


class WorkOrderTask(models.Model):
    """Modelo para gestionar tareas específicas de órdenes de trabajo"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completada'),
    ]
    
    work_order = models.ForeignKey(
        WorkOrder,
        on_delete=models.CASCADE,
        related_name='tasks',
        verbose_name='Orden de Trabajo'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la Tarea'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción detallada de la tarea (opcional)'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición de la tarea'
    )
    estimated_hours = models.DecimalField(
        max_digits=6,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Horas Estimadas',
        help_text='Tiempo estimado para completar la tarea'
    )
    actual_hours = models.DecimalField(
        max_digits=6,
        decimal_places=2,
        default=0,
        verbose_name='Horas Reales',
        help_text='Tiempo real trabajado en la tarea'
    )
    started_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Iniciado en'
    )
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Completado en'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['order', 'id']
        verbose_name = 'Tarea de Orden de Trabajo'
        verbose_name_plural = 'Tareas de Órdenes de Trabajo'
        indexes = [
            models.Index(fields=['work_order', 'status']),
            models.Index(fields=['work_order', 'order']),
        ]
    
    def __str__(self):
        return f"{self.work_order.order_number} - {self.title}"
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado"""
        status_classes = {
            'pending': 'bg-secondary',
            'in_progress': 'bg-primary',
            'completed': 'bg-success',
        }
        return status_classes.get(self.status, 'bg-secondary')
    
    def get_status_icon(self):
        """Retorna el icono para el estado"""
        status_icons = {
            'pending': 'fas fa-clock',
            'in_progress': 'fas fa-play-circle',
            'completed': 'fas fa-check-circle',
        }
        return status_icons.get(self.status, 'fas fa-question')
    
    def start_task(self):
        """Marca la tarea como iniciada"""
        if self.status == 'pending':
            self.status = 'in_progress'
            self.started_at = timezone.now()
            self.save(update_fields=['status', 'started_at'])
    
    def complete_task(self):
        """Marca la tarea como completada"""
        if self.status == 'in_progress':
            self.status = 'completed'
            self.completed_at = timezone.now()
            self.save(update_fields=['status', 'completed_at'])
    
    def get_duration(self):
        """Retorna la duración de la tarea si está completada"""
        if self.started_at and self.completed_at:
            delta = self.completed_at - self.started_at
            hours = delta.total_seconds() / 3600
            return round(hours, 2)
        return None
    
    def update_actual_hours(self):
        """Actualiza las horas reales basado en los time entries"""
        total_hours = self.time_entries.aggregate(
            total=models.Sum('hours')
        )['total'] or 0
        
        self.actual_hours = total_hours
        self.save(update_fields=['actual_hours'])
        return total_hours
    
    def has_active_time_entry(self):
        """Verifica si la tarea tiene una sesión de tiempo activa"""
        return self.active_sessions.filter(is_active=True).exists()
    
    def get_active_time_session(self, user=None):
        """Obtiene la sesión de tiempo activa para un usuario"""
        if user:
            return self.active_sessions.filter(is_active=True, user=user).first()
        return self.active_sessions.filter(is_active=True).first()
    
    def get_total_time_logged(self):
        """Retorna el total de horas registradas en time entries"""
        total = self.time_entries.aggregate(
            total=models.Sum('hours')
        )['total'] or 0
        return float(total)
    
    def get_progress_percentage(self):
        """Calcula el porcentaje de progreso basado en horas estimadas vs trabajadas"""
        if not self.estimated_hours or self.estimated_hours == 0:
            return 0
        
        total_logged = self.get_total_time_logged()
        percentage = (total_logged / float(self.estimated_hours)) * 100
        return min(percentage, 100)  # No más del 100%


class WorkOrderTaskTimeEntry(models.Model):
    """Modelo para registrar tiempo trabajado en tareas de órdenes de trabajo"""
    
    task = models.ForeignKey(
        WorkOrderTask,
        on_delete=models.CASCADE,
        related_name='time_entries',
        verbose_name='Tarea'
    )
    hours = models.DecimalField(
        max_digits=6,
        decimal_places=2,
        verbose_name='Horas Trabajadas'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción del Trabajo',
        help_text='Descripción opcional del trabajo realizado'
    )
    date = models.DateField(
        default=timezone.now,
        verbose_name='Fecha'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Registrado en'
    )
    
    class Meta:
        ordering = ['-date', '-created_at']
        verbose_name = 'Registro de Tiempo de Tarea'
        verbose_name_plural = 'Registros de Tiempo de Tareas'
        indexes = [
            models.Index(fields=['task', 'date']),
        ]
    
    def __str__(self):
        return f"{self.task} - {self.hours}h ({self.date})"


class WorkOrderTaskTimeSession(models.Model):
    """Modelo para gestionar sesiones de tiempo activas"""
    
    task = models.ForeignKey(
        WorkOrderTask,
        on_delete=models.CASCADE,
        related_name='active_sessions',
        verbose_name='Tarea'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        verbose_name='Usuario',
        help_text='Usuario que inició la sesión (puede ser vacío para sesiones públicas)'
    )
    start_time = models.DateTimeField(
        default=timezone.now,
        verbose_name='Inicio'
    )
    end_time = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fin'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción del trabajo realizado'
    )
    
    class Meta:
        ordering = ['-start_time']
        verbose_name = 'Sesión de Tiempo de Tarea'
        verbose_name_plural = 'Sesiones de Tiempo de Tareas'
        indexes = [
            models.Index(fields=['task', 'is_active']),
            models.Index(fields=['user', 'is_active']),
        ]
    
    def __str__(self):
        status = "Activa" if self.is_active else "Finalizada"
        username = self.user.username if self.user else "Público"
        return f"{self.task} - {username} ({status})"
    
    def get_duration_hours(self):
        """Retorna la duración en horas"""
        if self.end_time:
            delta = self.end_time - self.start_time
        else:
            delta = timezone.now() - self.start_time
        
        hours = delta.total_seconds() / 3600
        return round(hours, 2)
    
    def stop_session(self, description=""):
        """Detiene la sesión y crea una entrada de tiempo"""
        if self.is_active:
            self.end_time = timezone.now()
            self.is_active = False
            self.description = description
            self.save()
            
            # Crear entrada de tiempo
            hours = self.get_duration_hours()
            time_entry = WorkOrderTaskTimeEntry.objects.create(
                task=self.task,
                hours=hours,
                description=description,
                date=self.start_time.date()
            )
            
            # Actualizar horas reales de la tarea
            self.task.update_actual_hours()
            
            return time_entry
        return None


class WorkOrderRating(models.Model):
    """Modelo para calificaciones de órdenes de trabajo desde clientes"""
    
    work_order = models.ForeignKey(
        WorkOrder,
        on_delete=models.CASCADE,
        related_name='ratings',
        verbose_name='Orden de Trabajo'
    )
    rating = models.IntegerField(
        verbose_name='Calificación',
        help_text='Calificación de 1 a 5 estrellas',
        choices=[(1, '1 estrella'), (2, '2 estrellas'), (3, '3 estrellas'), (4, '4 estrellas'), (5, '5 estrellas')]
    )
    comment = models.TextField(
        blank=True,
        verbose_name='Comentario',
        help_text='Comentario opcional sobre el servicio'
    )
    rated_by_name = models.CharField(
        max_length=200,
        verbose_name='Nombre',
        help_text='Nombre de quien califica'
    )
    rated_by_email = models.EmailField(
        verbose_name='Email',
        help_text='Email de quien califica'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de calificación'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Calificación de Orden de Trabajo'
        verbose_name_plural = 'Calificaciones de Órdenes de Trabajo'
    
    def __str__(self):
        return f"{self.work_order} - {self.rating} estrellas por {self.rated_by_name}"


class WorkOrderComment(models.Model):
    """Modelo para comentarios públicos en órdenes de trabajo"""
    
    work_order = models.ForeignKey(
        WorkOrder,
        on_delete=models.CASCADE,
        related_name='public_comments',
        verbose_name='Orden de Trabajo'
    )
    comment = models.TextField(
        verbose_name='Comentario'
    )
    author_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del autor'
    )
    author_email = models.EmailField(
        verbose_name='Email del autor'
    )
    is_public = models.BooleanField(
        default=True,
        verbose_name='Visible públicamente',
        help_text='Si está marcado, el comentario será visible en la vista pública'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Comentario de Orden de Trabajo'
        verbose_name_plural = 'Comentarios de Órdenes de Trabajo'
    
    def __str__(self):
        return f"{self.work_order} - {self.author_name} ({self.created_at.strftime('%d/%m/%Y')})"


class Task(models.Model):
    """Modelo para gestionar tareas que pueden ser asignadas a múltiples usuarios"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completada'),
        ('cancelled', 'Cancelada'),
    ]
    
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('urgent', 'Urgente'),
    ]
    
    task_number = models.CharField(
        max_length=20,
        unique=True,
        editable=False,
        verbose_name='Número de Tarea',
        null=True,
        blank=True
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    description = models.TextField(
        verbose_name='Descripción'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    priority = models.CharField(
        max_length=20,
        choices=PRIORITY_CHOICES,
        default='medium',
        verbose_name='Prioridad'
    )
    assigned_users = models.ManyToManyField(
        User,
        related_name='assigned_tasks',
        blank=True,
        verbose_name='Usuarios Asignados'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_tasks',
        verbose_name='Creado por'
    )
    due_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de vencimiento'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de finalización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Tarea'
        verbose_name_plural = 'Tareas'
    
    def save(self, *args, **kwargs):
        if not self.task_number:
            from django.db import transaction
            with transaction.atomic():
                # Obtener el último número de tarea con bloqueo
                last_task = Task.objects.select_for_update().order_by('-id').first()
                if last_task and last_task.task_number:
                    try:
                        last_number = int(last_task.task_number.split('-')[1])
                        new_number = last_number + 1
                    except (IndexError, ValueError):
                        new_number = 1
                else:
                    new_number = 1
                
                # Verificar que el número no exista
                while Task.objects.filter(task_number=f'TASK-{new_number:04d}').exists():
                    new_number += 1
                
                self.task_number = f'TASK-{new_number:04d}'
        
        super().save(*args, **kwargs)
    
    def __str__(self):
        if self.task_number:
            return f'{self.task_number} - {self.title}'
        return f'{self.title} - {self.get_status_display()}'
    
    def get_priority_color(self):
        """Retorna el color CSS según la prioridad"""
        colors = {
            'low': 'success',
            'medium': 'warning',
            'high': 'danger',
            'urgent': 'dark'
        }
        return colors.get(self.priority, 'secondary')
    
    def get_status_color(self):
        """Retorna el color CSS según el estado"""
        colors = {
            'pending': 'secondary',
            'in_progress': 'primary',
            'completed': 'success',
            'cancelled': 'danger'
        }
        return colors.get(self.status, 'secondary')
    
    def mark_as_completed(self):
        """Marca la tarea como completada"""
        self.status = 'completed'
        self.completed_at = timezone.now()
        self.save()
    
    def is_overdue(self):
        """Verifica si la tarea está vencida"""
        if self.due_date and self.status not in ['completed', 'cancelled']:
            return timezone.now() > self.due_date
        return False


class DailyTaskSession(models.Model):
    """Modelo para gestionar sesiones diarias de tareas por usuario"""
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='daily_task_sessions',
        verbose_name='Usuario'
    )
    date = models.DateField(
        verbose_name='Fecha de la sesión'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas del día',
        help_text='Notas generales sobre el plan del día'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Creado en'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Actualizado en'
    )
    
    class Meta:
        ordering = ['-created_at', '-date']
        verbose_name = 'Gestión de Tareas'
        verbose_name_plural = 'Gestiones de Tareas'
    
    def __str__(self):
        return f'{self.user.username} - {self.date} ({self.created_at.strftime("%H:%M")})'
    
    def get_total_tasks(self):
        """Retorna el total de tareas en esta sesión"""
        return self.daily_task_items.count()
    
    def get_completed_tasks(self):
        """Retorna las tareas completadas en esta sesión"""
        return self.daily_task_items.filter(completed=True).count()
    
    def get_pending_tasks(self):
        """Retorna las tareas pendientes en esta sesión"""
        return self.daily_task_items.filter(completed=False).count()
    
    def get_completion_percentage(self):
        """Retorna el porcentaje de completitud"""
        total = self.get_total_tasks()
        if total == 0:
            return 0
        return round((self.get_completed_tasks() / total) * 100, 1)


class DailyTaskItem(models.Model):
    """Modelo para items de tareas dentro de una sesión diaria"""
    
    session = models.ForeignKey(
        DailyTaskSession,
        on_delete=models.CASCADE,
        related_name='daily_task_items',
        verbose_name='Sesión diaria'
    )
    task = models.ForeignKey(
        Task,
        on_delete=models.CASCADE,
        related_name='daily_items',
        verbose_name='Tarea'
    )
    completed = models.BooleanField(
        default=False,
        verbose_name='Completada'
    )
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Completada en'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas específicas',
        help_text='Notas específicas sobre el trabajo en esta tarea'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden'
    )
    
    class Meta:
        ordering = ['order', 'task__priority']
        unique_together = ['session', 'task']
        verbose_name = 'Item de Tarea Diaria'
        verbose_name_plural = 'Items de Tareas Diarias'
    
    def __str__(self):
        status = "✓" if self.completed else "○"
        return f'{status} {self.task.title} - {self.session.date}'
    
    def mark_completed(self):
        """Marca el item como completado"""
        self.completed = True
        self.completed_at = timezone.now()
        self.save()
    
    def mark_pending(self):
        """Marca el item como pendiente"""
        self.completed = False
        self.completed_at = None
        self.save()


class ChatRoom(models.Model):
    """Modelo para las salas de chat entre usuarios"""
    name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre de la sala'
    )
    participants = models.ManyToManyField(
        User,
        related_name='chat_rooms',
        verbose_name='Participantes'
    )
    is_group = models.BooleanField(
        default=False,
        verbose_name='Es grupo',
        help_text='True para chats grupales, False para chats 1:1'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    last_activity = models.DateTimeField(
        default=timezone.now,
        verbose_name='Última actividad'
    )
    
    class Meta:
        ordering = ['-last_activity']
        verbose_name = 'Sala de Chat'
        verbose_name_plural = 'Salas de Chat'
    
    def __str__(self):
        if self.is_group and self.name:
            return self.name
        elif not self.is_group:
            participants = list(self.participants.all())
            if len(participants) == 2:
                return f"{participants[0].get_full_name() or participants[0].username} & {participants[1].get_full_name() or participants[1].username}"
        return f"Chat {self.id}"
    
    def get_other_participant(self, user):
        """Para chats 1:1, obtiene el otro participante"""
        if not self.is_group:
            return self.participants.exclude(id=user.id).first()
        return None
    
    def get_last_message(self):
        """Obtiene el último mensaje de la sala"""
        return self.messages.order_by('-created_at').first()
    
    def mark_as_read_for_user(self, user):
        """Marca todos los mensajes como leídos para un usuario"""
        unread_messages = self.messages.filter(read_by__isnull=True).exclude(sender=user)
        for message in unread_messages:
            message.read_by.add(user)
    
    def get_unread_count_for_user(self, user):
        """Obtiene el número de mensajes no leídos para un usuario"""
        return self.messages.exclude(sender=user).exclude(read_by=user).count()


class ChatMessage(models.Model):
    """Modelo para los mensajes del chat"""
    room = models.ForeignKey(
        ChatRoom,
        on_delete=models.CASCADE,
        related_name='messages',
        verbose_name='Sala de chat'
    )
    sender = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='chat_messages',
        verbose_name='Remitente'
    )
    message = models.TextField(
        blank=True,
        verbose_name='Mensaje'
    )
    attachment = models.FileField(
        upload_to='chat_attachments/%Y/%m/%d/',
        blank=True,
        null=True,
        verbose_name='Archivo adjunto'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de envío'
    )
    read_by = models.ManyToManyField(
        User,
        blank=True,
        related_name='read_messages',
        verbose_name='Leído por'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Mensaje de Chat'
        verbose_name_plural = 'Mensajes de Chat'
    
    def __str__(self):
        if self.message:
            preview = self.message[:50] + "..." if len(self.message) > 50 else self.message
            return f"{self.sender.username}: {preview}"
        elif self.attachment:
            return f"{self.sender.username}: [Archivo adjunto]"
        return f"Mensaje {self.id}"
    
    def is_read_by_user(self, user):
        """Verifica si el mensaje ha sido leído por un usuario específico"""
        return self.read_by.filter(id=user.id).exists()
    
    def get_attachment_name(self):
        """Obtiene el nombre del archivo adjunto"""
        if self.attachment:
            return os.path.basename(self.attachment.name)
        return None
    
    def get_attachment_size(self):
        """Obtiene el tamaño del archivo adjunto en formato legible"""
        if self.attachment:
            try:
                size = self.attachment.size
                for unit in ['B', 'KB', 'MB', 'GB']:
                    if size < 1024.0:
                        return f"{size:.1f} {unit}"
                    size /= 1024.0
                return f"{size:.1f} TB"
            except:
                return "Tamaño desconocido"
        return None


class Command(models.Model):
    """Modelo para gestionar biblioteca de comandos"""
    
    CATEGORY_CHOICES = [
        ('linux', 'Linux/Unix'),
        ('windows', 'Windows'),
        ('docker', 'Docker'),
        ('git', 'Git'),
        ('python', 'Python'),
        ('javascript', 'JavaScript'),
        ('database', 'Base de Datos'),
        ('network', 'Redes'),
        ('security', 'Seguridad'),
        ('devops', 'DevOps'),
        ('web', 'Desarrollo Web'),
        ('mobile', 'Desarrollo Móvil'),
        ('other', 'Otros'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo del comando'
    )
    command = models.TextField(
        verbose_name='Comando',
        help_text='El comando completo con sus parámetros'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada de lo que hace el comando'
    )
    category = models.CharField(
        max_length=20,
        choices=CATEGORY_CHOICES,
        default='other',
        verbose_name='Categoría'
    )
    tags = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas para facilitar la búsqueda'
    )
    example_usage = models.TextField(
        blank=True,
        verbose_name='Ejemplo de uso',
        help_text='Ejemplo práctico de cómo usar el comando'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas adicionales',
        help_text='Notas importantes, advertencias o tips'
    )
    is_dangerous = models.BooleanField(
        default=False,
        verbose_name='Comando peligroso',
        help_text='Marcar si el comando puede ser destructivo o peligroso'
    )
    is_favorite = models.BooleanField(
        default=False,
        verbose_name='Favorito',
        help_text='Marcar como comando favorito'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    usage_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Veces usado',
        help_text='Contador de cuántas veces se ha copiado/usado este comando'
    )
    
    class Meta:
        ordering = ['-is_favorite', '-usage_count', '-created_at']
        verbose_name = 'Comando'
        verbose_name_plural = 'Comandos'
        indexes = [
            models.Index(fields=['category']),
            models.Index(fields=['is_favorite']),
            models.Index(fields=['created_by']),
            models.Index(fields=['title']),
        ]
    
    def __str__(self):
        return self.title
    
    def get_category_display_with_icon(self):
        """Retorna la categoría con su icono correspondiente"""
        icons = {
            'linux': 'bi-terminal',
            'windows': 'bi-windows',
            'docker': 'bi-box',
            'git': 'bi-git',
            'python': 'bi-code-slash',
            'javascript': 'bi-braces',
            'database': 'bi-database',
            'network': 'bi-wifi',
            'security': 'bi-shield-check',
            'devops': 'bi-gear',
            'web': 'bi-globe',
            'mobile': 'bi-phone',
            'other': 'bi-three-dots',
        }
        return {
            'name': self.get_category_display(),
            'icon': icons.get(self.category, 'bi-three-dots')
        }
    
    def get_tags_list(self):
        """Retorna las etiquetas como lista"""
        if self.tags:
            return [tag.strip() for tag in self.tags.split(',') if tag.strip()]
        return []
    
    def increment_usage(self):
        """Incrementa el contador de uso"""
        self.usage_count = models.F('usage_count') + 1
        self.save(update_fields=['usage_count'])
        self.refresh_from_db()
    
    def get_short_description(self, max_length=100):
        """Retorna una versión corta de la descripción"""
        if len(self.description) <= max_length:
            return self.description
        return self.description[:max_length] + '...'


class ContactFormSubmission(models.Model):
    """Modelo para gestionar formularios de contacto público"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('approved', 'Aprobada'),
        ('rejected', 'Rechazada'),
        ('company_created', 'Empresa Creada'),
    ]
    
    # Información del formulario de contacto
    company_name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Empresa'
    )
    contact_name = models.CharField(
        max_length=100,
        verbose_name='Nombre del Contacto'
    )
    email = models.EmailField(
        verbose_name='Correo Electrónico'
    )
    phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Teléfono'
    )
    website = models.URLField(
        blank=True,
        verbose_name='Sitio Web'
    )
    address = models.CharField(
        max_length=300,
        blank=True,
        verbose_name='Dirección'
    )
    message = models.TextField(
        blank=True,
        verbose_name='Mensaje Adicional',
        help_text='Información adicional sobre la empresa o servicios requeridos'
    )
    
    # Metadatos del formulario
    submitted_by_user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='contact_form_owner',
        verbose_name='Propietario del formulario',
        help_text='Usuario que tiene activado el formulario de contacto'
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    # Información de seguimiento
    ip_address = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    # Empresa creada (si se aprueba)
    created_company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Empresa Creada'
    )
    
    # Fechas
    submitted_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Envío'
    )
    processed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Procesamiento'
    )
    processed_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='processed_contact_forms',
        verbose_name='Procesado por'
    )
    
    # Notas internas
    admin_notes = models.TextField(
        blank=True,
        verbose_name='Notas del Administrador',
        help_text='Notas internas para el seguimiento de la solicitud'
    )
    
    class Meta:
        ordering = ['-submitted_at']
        verbose_name = 'Formulario de Contacto'
        verbose_name_plural = 'Formularios de Contacto'
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['submitted_by_user']),
            models.Index(fields=['submitted_at']),
        ]
    
    def __str__(self):
        return f"{self.company_name} - {self.contact_name} ({self.get_status_display()})"
    
    def mark_as_processed(self, user, status, company=None):
        """Marca el formulario como procesado"""
        self.status = status
        self.processed_by = user
        self.processed_at = timezone.now()
        if company:
            self.created_company = company
        self.save()
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge del estado"""
        status_classes = {
            'pending': 'bg-warning',
            'approved': 'bg-success',
            'rejected': 'bg-danger',
            'company_created': 'bg-primary',
        }
        return status_classes.get(self.status, 'bg-secondary')


# ============= MODELOS CRM =============

class OpportunityStatus(models.Model):
    """Estados configurables para las oportunidades CRM"""
    
    name = models.CharField(
        max_length=100, 
        unique=True, 
        verbose_name='Nombre del Estado'
    )
    description = models.TextField(
        blank=True, 
        verbose_name='Descripción'
    )
    color = models.CharField(
        max_length=7, 
        default='#007bff',
        help_text='Color en formato hexadecimal (ej: #007bff)',
        verbose_name='Color'
    )
    is_final = models.BooleanField(
        default=False,
        verbose_name='Estado Final',
        help_text='Indica si este estado representa el cierre de la oportunidad'
    )
    is_won = models.BooleanField(
        default=False,
        verbose_name='Estado Ganado',
        help_text='Indica si este estado representa una venta exitosa'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en el pipeline'
    )
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    created_by = models.ForeignKey(
        User, 
        on_delete=models.SET_NULL, 
        null=True, 
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['order', 'name']
        verbose_name = 'Estado de Oportunidad'
        verbose_name_plural = 'Estados de Oportunidad'
    
    def __str__(self):
        return self.name


class Opportunity(models.Model):
    """Modelo para gestionar oportunidades CRM"""
    
    PROBABILITY_CHOICES = [
        (0, '0%'),
        (10, '10%'),
        (20, '20%'),
        (30, '30%'),
        (40, '40%'),
        (50, '50%'),
        (60, '60%'),
        (70, '70%'),
        (80, '80%'),
        (90, '90%'),
        (100, '100%'),
    ]
    
    name = models.CharField(
        max_length=200, 
        verbose_name='Nombre de la Oportunidad'
    )
    description = models.TextField(
        blank=True, 
        verbose_name='Descripción'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='opportunities',
        verbose_name='Empresa'
    )
    contact_name = models.CharField(
        max_length=200,
        verbose_name='Contacto Principal'
    )
    contact_position = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Cargo del Contacto',
        help_text='Posición o cargo del contacto en la empresa'
    )
    contact_email = models.EmailField(
        blank=True,
        verbose_name='Email de Contacto'
    )
    contact_phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Teléfono de Contacto'
    )
    value = models.DecimalField(
        max_digits=12, 
        decimal_places=2,
        verbose_name='Valor Estimado',
        help_text='Valor estimado de la oportunidad'
    )
    probability = models.IntegerField(
        choices=PROBABILITY_CHOICES,
        default=20,
        verbose_name='Probabilidad de Cierre (%)'
    )
    status = models.ForeignKey(
        OpportunityStatus,
        on_delete=models.PROTECT,
        related_name='opportunities',
        verbose_name='Estado'
    )
    expected_close_date = models.DateField(
        verbose_name='Fecha Estimada de Cierre'
    )
    actual_close_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Fecha Real de Cierre'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='assigned_opportunities',
        verbose_name='Asignado a'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_opportunities',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    # Campos de seguimiento
    last_contact_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Último Contacto'
    )
    next_follow_up = models.DateField(
        blank=True,
        null=True,
        verbose_name='Próximo Seguimiento'
    )
    source = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Fuente',
        help_text='¿Cómo se originó esta oportunidad?'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Oportunidad'
        verbose_name_plural = 'Oportunidades'
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['assigned_to']),
            models.Index(fields=['company']),
            models.Index(fields=['expected_close_date']),
        ]
    
    def __str__(self):
        return f"{self.name} - {self.company.name}"
    
    @property
    def expected_value(self):
        """Calcula el valor esperado basado en la probabilidad"""
        return (self.value * self.probability) / 100
    
    @property
    def is_overdue(self):
        """Verifica si la oportunidad está vencida"""
        if self.expected_close_date and not self.status.is_final:
            return self.expected_close_date < timezone.now().date()
        return False
    
    @property
    def days_to_close(self):
        """Calcula los días hasta el cierre esperado"""
        if self.expected_close_date:
            delta = self.expected_close_date - timezone.now().date()
            return delta.days
        return None
    
    @property
    def days_overdue(self):
        """Calcula los días vencidos (número positivo)"""
        if self.is_overdue and self.expected_close_date:
            delta = timezone.now().date() - self.expected_close_date
            return delta.days
        return 0
    
    def can_be_edited_by(self, user):
        """Verifica si un usuario puede editar esta oportunidad"""
        from .utils import is_agent
        return is_agent(user) or user == self.created_by or user == self.assigned_to


class OpportunityNote(models.Model):
    """Notas y seguimiento de oportunidades"""
    
    opportunity = models.ForeignKey(
        Opportunity,
        on_delete=models.CASCADE,
        related_name='notes',
        verbose_name='Oportunidad'
    )
    content = models.TextField(verbose_name='Contenido')
    is_important = models.BooleanField(
        default=False,
        verbose_name='Importante',
        help_text='Marcar como nota importante'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Nota de Oportunidad'
        verbose_name_plural = 'Notas de Oportunidad'
    
    def __str__(self):
        return f"Nota para {self.opportunity.name} - {self.created_at.strftime('%d/%m/%Y')}"


class OpportunityStatusHistory(models.Model):
    """Historial de cambios de estado de las oportunidades"""
    
    opportunity = models.ForeignKey(
        Opportunity,
        on_delete=models.CASCADE,
        related_name='status_history',
        verbose_name='Oportunidad'
    )
    previous_status = models.ForeignKey(
        OpportunityStatus,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='previous_changes',
        verbose_name='Estado Anterior'
    )
    new_status = models.ForeignKey(
        OpportunityStatus,
        on_delete=models.CASCADE,
        related_name='new_changes',
        verbose_name='Nuevo Estado'
    )
    changed_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Cambiado por'
    )
    changed_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha del cambio')
    comment = models.TextField(
        blank=True,
        verbose_name='Comentario',
        help_text='Razón del cambio de estado'
    )
    
    class Meta:
        ordering = ['-changed_at']
        verbose_name = 'Historial de Estado'
        verbose_name_plural = 'Historial de Estados'
    
    def __str__(self):
        return f"{self.opportunity.name}: {self.previous_status} → {self.new_status}"


class OpportunityActivity(models.Model):
    """Modelo para actividades de seguimiento de oportunidades"""
    
    ACTIVITY_TYPE_CHOICES = [
        ('call', 'Llamada'),
        ('meeting', 'Reunión'),
        ('email', 'Correo Electrónico'),
        ('task', 'Tarea'),
        ('demo', 'Demostración'),
        ('proposal', 'Propuesta'),
        ('follow_up', 'Seguimiento'),
        ('other', 'Otro'),
    ]
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completada'),
        ('cancelled', 'Cancelada'),
        ('overdue', 'Vencida'),
    ]
    
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('urgent', 'Urgente'),
    ]
    
    opportunity = models.ForeignKey(
        Opportunity,
        on_delete=models.CASCADE,
        related_name='activities',
        verbose_name='Oportunidad',
        null=True,
        blank=True
    )
    contact = models.ForeignKey(
        'Contact',
        on_delete=models.CASCADE,
        related_name='activities',
        verbose_name='Contacto',
        null=True,
        blank=True,
        help_text='Contacto relacionado con esta actividad'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la Actividad'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    activity_type = models.CharField(
        max_length=20,
        choices=ACTIVITY_TYPE_CHOICES,
        default='task',
        verbose_name='Tipo de Actividad'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    priority = models.CharField(
        max_length=10,
        choices=PRIORITY_CHOICES,
        default='medium',
        verbose_name='Prioridad'
    )
    scheduled_date = models.DateTimeField(
        verbose_name='Fecha Programada',
        help_text='Fecha y hora cuando debe realizarse la actividad'
    )
    completed_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Completado'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='assigned_activities',
        verbose_name='Asignado a'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_activities',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    # Campos adicionales
    result = models.TextField(
        blank=True,
        verbose_name='Resultado',
        help_text='Resultado o notas de la actividad completada'
    )
    duration_minutes = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Duración (minutos)',
        help_text='Duración estimada o real de la actividad'
    )
    location = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Ubicación',
        help_text='Lugar donde se realizará la actividad'
    )
    
    class Meta:
        ordering = ['scheduled_date', '-priority']
        verbose_name = 'Actividad de Oportunidad'
        verbose_name_plural = 'Actividades de Oportunidades'
        indexes = [
            models.Index(fields=['status']),
            models.Index(fields=['assigned_to']),
            models.Index(fields=['scheduled_date']),
            models.Index(fields=['activity_type']),
            models.Index(fields=['priority']),
        ]
    
    def __str__(self):
        return f"{self.title} - {self.opportunity.name}"
    
    @property
    def is_overdue(self):
        """Verifica si la actividad está vencida"""
        if self.status in ['pending', 'in_progress'] and self.scheduled_date:
            return self.scheduled_date < timezone.now()
        return False
    
    @property
    def days_until_due(self):
        """Calcula los días hasta el vencimiento"""
        if self.scheduled_date:
            delta = self.scheduled_date.date() - timezone.now().date()
            return delta.days
        return None
    
    @property
    def is_due_today(self):
        """Verifica si la actividad vence hoy"""
        if self.scheduled_date:
            return self.scheduled_date.date() == timezone.now().date()
        return False
    
    @property
    def is_due_soon(self):
        """Verifica si la actividad vence en los próximos 3 días"""
        days_until = self.days_until_due
        return days_until is not None and 0 <= days_until <= 3
    
    def mark_completed(self, result=None, user=None):
        """Marca la actividad como completada"""
        self.status = 'completed'
        self.completed_date = timezone.now()
        if result:
            self.result = result
        self.save()
    
    def get_priority_color(self):
        """Retorna el color para la prioridad"""
        colors = {
            'low': 'success',
            'medium': 'warning', 
            'high': 'danger',
            'urgent': 'dark'
        }
        return colors.get(self.priority, 'secondary')
    
    def get_status_color(self):
        """Retorna el color para el estado"""
        colors = {
            'pending': 'warning',
            'in_progress': 'info',
            'completed': 'success',
            'cancelled': 'secondary',
            'overdue': 'danger'
        }
        return colors.get(self.status, 'secondary')


class Meeting(models.Model):
    """Modelo para gestionar reuniones"""
    
    STATUS_CHOICES = [
        ('scheduled', 'Programada'),
        ('in_progress', 'En progreso'),
        ('finished', 'Finalizada'),
        ('cancelled', 'Cancelada'),
    ]
    
    title = models.CharField(max_length=200, verbose_name='Título de la reunión')
    description = models.TextField(blank=True, verbose_name='Descripción')
    spin_methodology = models.TextField(
        blank=True, 
        verbose_name='Metodología SPIN',
        help_text='Preguntas generadas automáticamente usando la metodología SPIN de ventas'
    )
    date = models.DateTimeField(verbose_name='Fecha y hora')
    duration = models.IntegerField(default=60, verbose_name='Duración (minutos)')
    location = models.CharField(max_length=200, blank=True, verbose_name='Ubicación')
    status = models.CharField(
        max_length=20, 
        choices=STATUS_CHOICES, 
        default='scheduled',
        verbose_name='Estado'
    )
    
    # Organizador
    organizer = models.ForeignKey(
        User, 
        on_delete=models.CASCADE, 
        related_name='organized_meetings',
        verbose_name='Organizador'
    )
    
    # Empresa (opcional para uso interno)
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Empresa',
        help_text='Empresa relacionada con la reunión (opcional)'
    )
    
    # Producto relacionado con la reunión
    product = models.ForeignKey(
        'Product',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Producto',
        help_text='Producto principal que se discutirá en la reunión'
    )
    
    # Link público único
    public_token = models.CharField(
        max_length=100, 
        unique=True, 
        verbose_name='Token público'
    )
    
    # Configuración
    allow_questions = models.BooleanField(default=True, verbose_name='Permitir preguntas')
    require_email = models.BooleanField(default=False, verbose_name='Requerir email')
    is_active = models.BooleanField(default=True, verbose_name='Activa')
    
    # Estadísticas
    public_views_count = models.IntegerField(default=0, verbose_name='Vistas del enlace público')
    
    # Metadata
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Creada el')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Actualizada el')
    
    class Meta:
        ordering = ['-date']
        verbose_name = 'Reunión'
        verbose_name_plural = 'Reuniones'
    
    def __str__(self):
        return f"{self.title} - {self.date.strftime('%d/%m/%Y %H:%M')}"
    
    def save(self, *args, **kwargs):
        if not self.public_token:
            import uuid
            self.public_token = str(uuid.uuid4())
        super().save(*args, **kwargs)
    
    def get_public_url(self):
        """Retorna la URL pública para registro de asistentes"""
        return f"/meetings/public/{self.public_token}/"
    
    def get_questions_url(self):
        """Retorna la URL pública para hacer preguntas"""
        return f"/meetings/public/{self.public_token}/questions/"
    
    def get_attendees_count(self):
        """Retorna el total de asistentes registrados"""
        return self.meetingattendee_set.count()
    
    def get_questions_count(self):
        """Retorna el total de preguntas"""
        return self.meetingquestion_set.count()


class MeetingAttendee(models.Model):
    """Modelo para asistentes a reuniones"""
    
    RESPONSE_CHOICES = [
        ('pending', 'Pendiente'),
        ('accepted', 'Aceptada'),
        ('declined', 'Rechazada'),
    ]
    
    meeting = models.ForeignKey(Meeting, on_delete=models.CASCADE, verbose_name='Reunión')
    name = models.CharField(max_length=100, verbose_name='Nombre')
    email = models.EmailField(blank=True, verbose_name='Email')
    company = models.CharField(max_length=100, blank=True, verbose_name='Empresa')
    registered_at = models.DateTimeField(default=timezone.now, verbose_name='Registrado el')
    ip_address = models.GenericIPAddressField(null=True, blank=True, verbose_name='IP')
    response_status = models.CharField(
        max_length=20, 
        choices=RESPONSE_CHOICES, 
        default='pending',
        verbose_name='Estado de respuesta'
    )
    response_date = models.DateTimeField(null=True, blank=True, verbose_name='Fecha de respuesta')
    
    class Meta:
        ordering = ['registered_at']
        verbose_name = 'Asistente'
        verbose_name_plural = 'Asistentes'
        # Removido unique_together para permitir múltiples registros
    
    def __str__(self):
        return f"{self.name} - {self.meeting.title}"


class MeetingQuestion(models.Model):
    """Modelo para preguntas en reuniones"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('answered', 'Respondida'),
        ('ignored', 'Ignorada'),
    ]
    
    meeting = models.ForeignKey(Meeting, on_delete=models.CASCADE, verbose_name='Reunión')
    attendee = models.ForeignKey(
        MeetingAttendee, 
        on_delete=models.CASCADE, 
        null=True, 
        blank=True,
        verbose_name='Asistente'
    )
    question = models.TextField(verbose_name='Pregunta')
    answer = models.TextField(blank=True, verbose_name='Respuesta')
    status = models.CharField(
        max_length=20, 
        choices=STATUS_CHOICES, 
        default='pending',
        verbose_name='Estado'
    )
    
    # Datos del preguntante (por si no está registrado como asistente)
    asker_name = models.CharField(max_length=100, blank=True, verbose_name='Nombre del preguntante')
    asker_email = models.EmailField(blank=True, verbose_name='Email del preguntante')
    
    # Metadata
    asked_at = models.DateTimeField(default=timezone.now, verbose_name='Preguntado el')
    answered_at = models.DateTimeField(null=True, blank=True, verbose_name='Respondido el')
    answered_by = models.ForeignKey(
        User, 
        on_delete=models.SET_NULL, 
        null=True, 
        blank=True,
        verbose_name='Respondido por'
    )
    ip_address = models.GenericIPAddressField(null=True, blank=True, verbose_name='IP')
    
    class Meta:
        ordering = ['asked_at']
        verbose_name = 'Pregunta de reunión'
        verbose_name_plural = 'Preguntas de reuniones'
    
    def __str__(self):
        return f"Pregunta en {self.meeting.title}: {self.question[:50]}..."
    
    def get_asker_display_name(self):
        """Retorna el nombre del preguntante"""
        if self.attendee:
            return self.attendee.name
        return self.asker_name or "Anónimo"


class MeetingRating(models.Model):
    """Modelo para calificaciones de reuniones con caras (triste, neutro, feliz)"""
    
    RATING_CHOICES = [
        ('sad', '😞 Triste'),
        ('neutral', '😐 Neutro'),
        ('happy', '😊 Feliz'),
    ]
    
    meeting = models.ForeignKey(Meeting, on_delete=models.CASCADE, verbose_name='Reunión', related_name='ratings')
    attendee = models.ForeignKey(
        MeetingAttendee, 
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Asistente'
    )
    rating = models.CharField(
        max_length=10,
        choices=RATING_CHOICES,
        verbose_name='Calificación'
    )
    comment = models.TextField(blank=True, verbose_name='Comentario')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de evaluación')
    ip_address = models.GenericIPAddressField(null=True, blank=True, verbose_name='IP')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Evaluación de reunión'
        verbose_name_plural = 'Evaluaciones de reuniones'
        # Un asistente puede evaluar solo una vez
        unique_together = ['meeting', 'attendee']
    
    def __str__(self):
        return f"Evaluación de {self.meeting.title}: {self.get_rating_display()}"


class MeetingAccessLog(models.Model):
    """Modelo para registrar estadísticas de acceso al enlace público de reuniones"""
    
    meeting = models.ForeignKey(Meeting, on_delete=models.CASCADE, verbose_name='Reunión', related_name='access_logs')
    ip_address = models.GenericIPAddressField(verbose_name='Dirección IP')
    user_agent = models.TextField(blank=True, verbose_name='User Agent')
    device_type = models.CharField(max_length=50, blank=True, verbose_name='Tipo de Dispositivo')  # mobile, tablet, desktop
    browser = models.CharField(max_length=100, blank=True, verbose_name='Navegador')
    os = models.CharField(max_length=100, blank=True, verbose_name='Sistema Operativo')
    country = models.CharField(max_length=100, blank=True, verbose_name='País')
    city = models.CharField(max_length=100, blank=True, verbose_name='Ciudad')
    accessed_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha y hora de acceso')
    
    class Meta:
        ordering = ['-accessed_at']
        verbose_name = 'Registro de acceso'
        verbose_name_plural = 'Registros de acceso'
        indexes = [
            models.Index(fields=['meeting', '-accessed_at']),
            models.Index(fields=['country']),
            models.Index(fields=['device_type']),
        ]
    
    def __str__(self):
        return f"{self.meeting.title} - {self.country} - {self.device_type} - {self.accessed_at.strftime('%d/%m/%Y %H:%M')}"


class MeetingLink(models.Model):
    """Modelo para gestionar enlaces relacionados con reuniones"""
    
    meeting = models.ForeignKey(Meeting, on_delete=models.CASCADE, verbose_name='Reunión', related_name='links')
    title = models.CharField(max_length=200, verbose_name='Título del enlace')
    url = models.URLField(max_length=500, verbose_name='URL')
    description = models.TextField(blank=True, verbose_name='Descripción')
    order = models.IntegerField(default=0, verbose_name='Orden')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Creado el')
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Enlace de reunión'
        verbose_name_plural = 'Enlaces de reunión'
    
    def __str__(self):
        return f"{self.meeting.title} - {self.title}"


class Course(models.Model):
    """Modelo para gestionar cursos de capacitación"""
    
    title = models.CharField(max_length=200, verbose_name='Título del Curso')
    description = models.TextField(blank=True, verbose_name='Descripción')
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE, 
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        verbose_name='Empresa',
        help_text='Si se selecciona una empresa, solo los usuarios de esa empresa podrán ver este curso'
    )
    public_token = models.UUIDField(
        null=True,
        blank=True,
        verbose_name='Token público',
        help_text='Token para acceso público al curso sin autenticación'
    )
    is_public_enabled = models.BooleanField(
        default=False,
        verbose_name='Acceso público habilitado',
        help_text='Permite que el curso sea accesible públicamente mediante un enlace'
    )
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Curso'
        verbose_name_plural = 'Cursos'
    
    def __str__(self):
        return self.title
    
    def get_classes_count(self):
        """Retorna el número de clases en el curso"""
        return self.classes.count()
    
    def generate_public_token(self):
        """Genera un nuevo token público para el curso"""
        import uuid
        self.public_token = uuid.uuid4()
        self.is_public_enabled = True
        self.save()
        return self.public_token
    
    def disable_public_access(self):
        """Deshabilita el acceso público al curso"""
        self.is_public_enabled = False
        self.save()
    
    def get_public_url(self, request=None):
        """Retorna la URL pública del curso si está habilitada"""
        if self.is_public_enabled and self.public_token:
            from django.urls import reverse
            if request:
                return request.build_absolute_uri(
                    reverse('course_public', kwargs={'token': self.public_token})
                )
            return reverse('course_public', kwargs={'token': self.public_token})
        return None
    
    def can_user_access(self, user):
        """Verifica si un usuario puede acceder a este curso"""
        if not user.is_authenticated:
            return False
            
        # Si el curso no tiene empresa asignada, todos pueden verlo
        if not self.company:
            return True
            
        # Si el curso tiene empresa, verificar que el usuario pertenezca a esa empresa
        try:
            # Intentar obtener o crear el UserProfile si no existe
            user_profile, created = UserProfile.objects.get_or_create(user=user)
            return user_profile.company == self.company
        except Exception as e:
            return False
    
    def is_public(self):
        """Retorna True si el curso es público (sin empresa asignada)"""
        return self.company is None


class CourseRegistrationToken(models.Model):
    """Modelo para gestionar tokens de registro público para cursos"""
    
    course = models.ForeignKey(
        Course,
        on_delete=models.CASCADE,
        related_name='registration_tokens',
        verbose_name='Curso'
    )
    token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración',
        help_text='Dejar vacío para que no expire'
    )
    max_registrations = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Máximo de registros',
        help_text='Número máximo de usuarios que pueden registrarse con este token'
    )
    registration_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Registros realizados'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Token de Registro de Curso'
        verbose_name_plural = 'Tokens de Registro de Curso'
    
    def __str__(self):
        return f"Token para {self.course.title} - {self.token}"
    
    def is_valid(self):
        """Verifica si el token es válido para registro"""
        if not self.is_active:
            return False, "El token no está activo"
        
        if self.expires_at and timezone.now() > self.expires_at:
            return False, "El token ha expirado"
        
        if self.max_registrations and self.registration_count >= self.max_registrations:
            return False, "Se ha alcanzado el límite máximo de registros"
        
        return True, "Token válido"
    
    def increment_registration_count(self):
        """Incrementa el contador de registros"""
        self.registration_count += 1
        self.save(update_fields=['registration_count'])
    
    def get_registration_url(self):
        """Retorna la URL completa para registro público"""
        from django.urls import reverse
        return reverse('course_public_register', kwargs={'token': self.token})


class CourseClass(models.Model):
    """Modelo para gestionar clases dentro de un curso"""
    
    course = models.ForeignKey(
        Course, 
        on_delete=models.CASCADE, 
        related_name='classes',
        verbose_name='Curso'
    )
    title = models.CharField(max_length=200, verbose_name='Título de la Clase')
    description = models.TextField(blank=True, verbose_name='Descripción')
    video_url = models.URLField(
        blank=True, 
        null=True,
        verbose_name='URL del Video',
        help_text='URL del video de la clase (YouTube, Vimeo, etc.)'
    )
    order = models.PositiveIntegerField(
        default=0, 
        verbose_name='Orden',
        help_text='Orden de la clase dentro del curso'
    )
    duration_minutes = models.PositiveIntegerField(
        blank=True, 
        null=True,
        verbose_name='Duración (minutos)',
        help_text='Duración estimada de la clase en minutos'
    )
    is_active = models.BooleanField(default=True, verbose_name='Activa')
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE, 
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        ordering = ['course', 'order', 'title']
        verbose_name = 'Clase'
        verbose_name_plural = 'Clases'
        unique_together = ['course', 'order']
    
    def __str__(self):
        return f"{self.course.title} - {self.title}"
    
    def has_video(self):
        """Verifica si la clase tiene video"""
        return bool(self.video_url)
    
    def get_view_count(self):
        """Retorna el número de usuarios que han visto esta clase"""
        return self.views.count()
    
    def is_viewed_by_user(self, user):
        """Verifica si un usuario específico ha visto esta clase"""
        if not user.is_authenticated:
            return False
        return self.views.filter(user=user).exists()


class CourseClassView(models.Model):
    """Modelo para rastrear qué usuarios han visto qué clases"""
    
    course_class = models.ForeignKey(
        CourseClass,
        on_delete=models.CASCADE,
        related_name='views',
        verbose_name='Clase'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    viewed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Visto el'
    )
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    
    class Meta:
        unique_together = ['course_class', 'user']
        ordering = ['-viewed_at']
        verbose_name = 'Visualización de Clase'
        verbose_name_plural = 'Visualizaciones de Clases'
    
    def __str__(self):
        return f"{self.user.username} vio {self.course_class.title}"


class Contact(models.Model):
    """Modelo para gestionar contactos de ventas"""
    
    STATUS_CHOICES = [
        ('positive', 'Positivo'),
        ('neutral', 'Neutral'),
        ('negative', 'Negativo'),
        ('not_now', 'No ahora'),
        ('do_not_contact', 'No Contactar'),
        ('won', 'Ganado'),
    ]
    
    COMPANY_SIZE_CHOICES = [
        ('small', 'Pequeña (Menos de 10 empleados)'),
        ('medium', 'Mediana (10 a 50 empleados)'),
        ('large', 'Grande (Más de 50 empleados)'),
    ]
    
    name = models.CharField(
        max_length=100,
        verbose_name='Nombre'
    )
    email = models.EmailField(
        blank=True,
        null=True,
        verbose_name='Email'
    )
    phone = models.CharField(
        max_length=20,
        blank=True,
        null=True,
        verbose_name='Teléfono'
    )
    position = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Cargo'
    )
    company = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Empresa'
    )
    country = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='País',
        help_text='País del contacto'
    )
    erp = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='ERP',
        help_text='Sistema ERP actual del contacto'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='negative',
        verbose_name='Estado'
    )
    source = models.CharField(
        max_length=200,
        blank=True,
        null=True,
        verbose_name='Fuente',
        help_text='¿De dónde obtuviste este contacto?'
    )
    company_size = models.CharField(
        max_length=10,
        choices=COMPANY_SIZE_CHOICES,
        default='small',
        verbose_name='Tamaño de Empresa',
        help_text='Tamaño de la empresa del contacto'
    )
    notes = models.TextField(
        blank=True,
        null=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el contacto'
    )
    
    # Seguimiento de contacto
    contacted_by_phone = models.BooleanField(
        default=False,
        verbose_name='Contactado por telefono',
        help_text='Marcar si se contacto al cliente por telefono'
    )
    contacted_by_web = models.BooleanField(
        default=False,
        verbose_name='Contactado por web/email',
        help_text='Marcar si se contacto al cliente por web o email'
    )
    had_meeting = models.BooleanField(
        default=False,
        verbose_name='Reunion realizada',
        help_text='Marcar si se tuvo una reunion con este contacto'
    )
    meeting_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Fecha de reunion',
        help_text='Fecha en la que se realizo la reunion'
    )
    contact_tracking_notes = models.TextField(
        blank=True,
        verbose_name='Notas de seguimiento',
        help_text='Notas sobre los contactos realizados'
    )
    last_contact_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha del ultimo contacto',
        help_text='Fecha del ultimo contacto realizado'
    )
    
    # Redes sociales y web
    website = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Sitio Web',
        help_text='Sitio web de la empresa o contacto'
    )
    facebook_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Facebook',
        help_text='URL del perfil de Facebook'
    )
    linkedin_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='LinkedIn',
        help_text='URL del perfil de LinkedIn'
    )
    twitter_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Twitter/X',
        help_text='URL del perfil de Twitter/X'
    )
    instagram_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Instagram',
        help_text='URL del perfil de Instagram'
    )
    youtube_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='YouTube',
        help_text='URL del canal de YouTube'
    )
    tiktok_url = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='TikTok',
        help_text='URL del perfil de TikTok'
    )
    
    contact_date = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Contacto'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='contacts_created',
        verbose_name='Creado por'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='contacts_assigned',
        verbose_name='Asignado a'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-contact_date']
        verbose_name = 'Contacto'
        verbose_name_plural = 'Contactos'
    
    def __str__(self):
        return f"{self.name} - {self.get_status_display()}"
    
    def get_status_color(self):
        """Retorna el color del estado"""
        if self.status == 'positive':
            return 'success'
        elif self.status == 'neutral':
            return 'primary'
        elif self.status == 'not_now':
            return 'warning'
        elif self.status == 'do_not_contact':
            return 'dark'
        elif self.status == 'won':
            return 'success'
        else:
            return 'danger'
    
    def get_status_icon(self):
        """Retorna el icono del estado"""
        if self.status == 'positive':
            return 'bi-check-circle'
        elif self.status == 'neutral':
            return 'bi-dash-circle'
        elif self.status == 'not_now':
            return 'bi-clock'
        elif self.status == 'do_not_contact':
            return 'bi-slash-circle'
        elif self.status == 'won':
            return 'bi-trophy-fill'
        else:
            return 'bi-x-circle'


def contact_attachment_upload_path(instance, filename):
    """Ruta de subida para adjuntos de contactos"""
    return f'contact_attachments/contact_{instance.contact.id}/{filename}'


class ContactComment(models.Model):
    """Modelo para comentarios en contactos"""
    contact = models.ForeignKey(
        Contact,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Contacto'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    content = models.TextField(
        verbose_name='Comentario',
        help_text='Escribe tu comentario aquí'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Comentario de Contacto'
        verbose_name_plural = 'Comentarios de Contacto'
    
    def __str__(self):
        return f"Comentario de {self.user.username} en Contacto {self.contact.name}"
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar este comentario"""
        from .utils import is_agent
        return self.user == user or is_agent(user)
    
    def can_delete(self, user):
        """Verifica si un usuario puede eliminar este comentario"""
        from .utils import is_agent
        return self.user == user or is_agent(user)


class ContactAttachment(models.Model):
    """Modelo para adjuntos en contactos"""
    contact = models.ForeignKey(
        Contact, 
        on_delete=models.CASCADE, 
        related_name='attachments',
        verbose_name='Contacto'
    )
    file = models.FileField(
        upload_to=contact_attachment_upload_path,
        verbose_name='Archivo'
    )
    original_filename = models.CharField(
        max_length=255,
        verbose_name='Nombre original'
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Subido por'
    )
    uploaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de subida'
    )
    file_size = models.PositiveIntegerField(
        default=0,
        verbose_name='Tamaño del archivo (bytes)'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional del adjunto'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Adjunto de Contacto'
        verbose_name_plural = 'Adjuntos de Contacto'
    
    def __str__(self):
        return f"{self.original_filename} - Contacto {self.contact.name}"
    
    def get_file_size_display(self):
        """Retorna el tamaño del archivo en formato legible"""
        if self.file_size < 1024:
            return f"{self.file_size} bytes"
        elif self.file_size < 1024 * 1024:
            return f"{self.file_size / 1024:.1f} KB"
        else:
            return f"{self.file_size / (1024 * 1024):.1f} MB"
    
    def get_file_extension(self):
        """Retorna la extensión del archivo"""
        return os.path.splitext(self.original_filename)[1].lower()
    
    def delete(self, *args, **kwargs):
        """Elimina el archivo físico cuando se elimina el registro"""
        if self.file:
            if os.path.isfile(self.file.path):
                os.remove(self.file.path)
        super().delete(*args, **kwargs)


class SalesPlan(models.Model):
    """Modelo para plan de ventas mensual de usuarios"""
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='sales_plans',
        verbose_name='Usuario'
    )
    monthly_contact_goal = models.PositiveIntegerField(
        default=0,
        verbose_name='Objetivo de contactos al mes',
        help_text='Número de contactos nuevos que debe crear al mes'
    )
    monthly_positive_contact_goal = models.PositiveIntegerField(
        default=0,
        verbose_name='Objetivo de contactos positivos al mes',
        help_text='Número de contactos positivos que debe conseguir al mes'
    )
    monthly_meeting_goal = models.PositiveIntegerField(
        default=0,
        verbose_name='Objetivo de reuniones mensuales',
        help_text='Número de reuniones que debe realizar al mes'
    )
    monthly_won_goal = models.PositiveIntegerField(
        default=0,
        verbose_name='Objetivo de ganados al mes',
        help_text='Número de contactos ganados que debe conseguir al mes'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Plan activo'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Plan de Venta'
        verbose_name_plural = 'Planes de Venta'
    
    def __str__(self):
        return f"Plan de {self.user.get_full_name() or self.user.username}"
    
    def get_monthly_progress(self):
        """Calcula el progreso del mes actual"""
        from django.db.models import Count, Q
        from datetime import timedelta
        from dateutil.relativedelta import relativedelta
        
        now = timezone.now()
        month_start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        
        # Calcular el mismo día del mes pasado
        last_month = now - relativedelta(months=1)
        last_month_start = last_month.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
        last_month_same_day = last_month.replace(hour=23, minute=59, second=59, microsecond=999999)
        
        # Contactos creados este mes
        contacts_created = Contact.objects.filter(
            created_by=self.user,
            created_at__gte=month_start,
            created_at__lte=now
        ).count()
        
        # Contactos del mismo día del mes pasado
        contacts_last_month = Contact.objects.filter(
            created_by=self.user,
            created_at__gte=last_month_start,
            created_at__lte=last_month_same_day
        ).count()
        
        # Contactos positivos este mes
        positive_contacts = Contact.objects.filter(
            created_by=self.user,
            status='positive',
            created_at__gte=month_start,
            created_at__lte=now
        ).count()
        
        # Positivos del mes pasado
        positive_last_month = Contact.objects.filter(
            created_by=self.user,
            status='positive',
            created_at__gte=last_month_start,
            created_at__lte=last_month_same_day
        ).count()
        
        # Reuniones este mes
        meetings_held = Contact.objects.filter(
            created_by=self.user,
            had_meeting=True,
            meeting_date__gte=month_start.date(),
            meeting_date__lte=now.date()
        ).count()
        
        # Reuniones del mes pasado
        meetings_last_month = Contact.objects.filter(
            created_by=self.user,
            had_meeting=True,
            meeting_date__gte=last_month_start.date(),
            meeting_date__lte=last_month_same_day.date()
        ).count()
        
        # Contactos ganados este mes
        won_contacts = Contact.objects.filter(
            created_by=self.user,
            status='won',
            created_at__gte=month_start,
            created_at__lte=now
        ).count()
        
        # Ganados del mes pasado
        won_last_month = Contact.objects.filter(
            created_by=self.user,
            status='won',
            created_at__gte=last_month_start,
            created_at__lte=last_month_same_day
        ).count()
        
        # Calcular diferencias
        contacts_diff = contacts_created - contacts_last_month
        positive_diff = positive_contacts - positive_last_month
        meetings_diff = meetings_held - meetings_last_month
        won_diff = won_contacts - won_last_month
        
        # Calcular porcentajes
        contact_percentage = (contacts_created / self.monthly_contact_goal * 100) if self.monthly_contact_goal > 0 else 0
        positive_percentage = (positive_contacts / self.monthly_positive_contact_goal * 100) if self.monthly_positive_contact_goal > 0 else 0
        meeting_percentage = (meetings_held / self.monthly_meeting_goal * 100) if self.monthly_meeting_goal > 0 else 0
        won_percentage = (won_contacts / self.monthly_won_goal * 100) if self.monthly_won_goal > 0 else 0
        
        # Promedio general
        overall_percentage = (contact_percentage + positive_percentage + meeting_percentage + won_percentage) / 4
        
        return {
            'contacts_created': contacts_created,
            'contacts_goal': self.monthly_contact_goal,
            'contact_percentage': round(contact_percentage, 1),
            'contacts_last_month': contacts_last_month,
            'contacts_diff': contacts_diff,
            'positive_contacts': positive_contacts,
            'positive_goal': self.monthly_positive_contact_goal,
            'positive_percentage': round(positive_percentage, 1),
            'positive_last_month': positive_last_month,
            'positive_diff': positive_diff,
            'meetings_held': meetings_held,
            'meeting_goal': self.monthly_meeting_goal,
            'meeting_percentage': round(meeting_percentage, 1),
            'meetings_last_month': meetings_last_month,
            'meetings_diff': meetings_diff,
            'won_contacts': won_contacts,
            'won_goal': self.monthly_won_goal,
            'won_percentage': round(won_percentage, 1),
            'won_last_month': won_last_month,
            'won_diff': won_diff,
            'overall_percentage': round(overall_percentage, 1)
        }


class BlogCategory(models.Model):
    """Categorías para los artículos del blog"""
    
    name = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Nombre'
    )
    slug = models.SlugField(
        max_length=100,
        unique=True,
        verbose_name='Slug URL'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    color = models.CharField(
        max_length=7,
        default='#007bff',
        verbose_name='Color',
        help_text='Color en formato hexadecimal (ej: #007bff)'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['name']
        verbose_name = 'Categoría de Blog'
        verbose_name_plural = 'Categorías de Blog'
    
    def __str__(self):
        return self.name


class BlogPost(models.Model):
    """Artículos del blog"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('published', 'Publicado'),
        ('archived', 'Archivado'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    slug = models.SlugField(
        max_length=200,
        unique=True,
        verbose_name='Slug URL'
    )
    excerpt = models.TextField(
        max_length=300,
        verbose_name='Resumen',
        help_text='Breve descripción del artículo (máx. 300 caracteres)'
    )
    content = models.TextField(
        verbose_name='Contenido'
    )
    featured_image = models.ImageField(
        upload_to='blog/images/',
        blank=True,
        null=True,
        verbose_name='Imagen Principal'
    )
    category = models.ForeignKey(
        BlogCategory,
        on_delete=models.SET_NULL,
        null=True,
        related_name='posts',
        verbose_name='Categoría'
    )
    status = models.CharField(
        max_length=10,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    tags = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas'
    )
    meta_description = models.CharField(
        max_length=160,
        blank=True,
        verbose_name='Meta Descripción',
        help_text='Descripción para SEO (máx. 160 caracteres)'
    )
    views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de visualizaciones'
    )
    is_featured = models.BooleanField(
        default=False,
        verbose_name='Artículo Destacado'
    )
    published_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de publicación'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='blog_posts',
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['-published_at', '-created_at']
        verbose_name = 'Artículo de Blog'
        verbose_name_plural = 'Artículos de Blog'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        # Auto-generar slug si no existe
        if not self.slug:
            from django.utils.text import slugify
            self.slug = slugify(self.title)
        
        # Establecer fecha de publicación cuando se publica
        if self.status == 'published' and not self.published_at:
            self.published_at = timezone.now()
        
        super().save(*args, **kwargs)
    
    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('blog_post_detail', kwargs={'slug': self.slug})
    
    def get_tags_list(self):
        """Retorna las etiquetas como lista"""
        if self.tags:
            return [tag.strip() for tag in self.tags.split(',')]
        return []
    
    def increment_views(self):
        """Incrementa el contador de visualizaciones"""
        self.views_count += 1
        self.save(update_fields=['views_count'])
    
    def get_pending_comments_count(self):
        """Retorna el número de comentarios pendientes de aprobación"""
        return self.comments.filter(is_approved=False).count()
    
    def get_approved_comments_count(self):
        """Retorna el número de comentarios aprobados"""
        return self.comments.filter(is_approved=True).count()


class BlogComment(models.Model):
    """Comentarios de los artículos del blog"""
    
    post = models.ForeignKey(
        BlogPost,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Artículo'
    )
    name = models.CharField(
        max_length=100,
        verbose_name='Nombre'
    )
    email = models.EmailField(
        verbose_name='Email'
    )
    website = models.URLField(
        blank=True,
        verbose_name='Sitio Web'
    )
    content = models.TextField(
        verbose_name='Comentario'
    )
    is_approved = models.BooleanField(
        default=False,
        verbose_name='Aprobado'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Comentario de Blog'
        verbose_name_plural = 'Comentarios de Blog'
    
    def __str__(self):
        return f"{self.name} en {self.post.title}"


class AIChatSession(models.Model):
    """Modelo para sesiones de chat con IA"""
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_chat_sessions',
        verbose_name='Usuario'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la conversación',
        help_text='Título descriptivo de la conversación'
    )
    ai_model = models.CharField(
        max_length=50,
        default='gpt-4o',
        verbose_name='Modelo de IA',
        help_text='Modelo utilizado para esta conversación'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Sesión activa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-updated_at']
        verbose_name = 'Sesión de Chat IA'
        verbose_name_plural = 'Sesiones de Chat IA'
    
    def __str__(self):
        return f"{self.title} - {self.user.username}"
    
    def get_message_count(self):
        """Retorna el número de mensajes en esta sesión"""
        return self.messages.count()
    
    def get_last_message(self):
        """Retorna el último mensaje de la sesión"""
        return self.messages.order_by('-created_at').first()


class AIChatMessage(models.Model):
    """Modelo para mensajes individuales del chat con IA"""
    
    ROLE_CHOICES = [
        ('user', 'Usuario'),
        ('assistant', 'IA'),
        ('system', 'Sistema'),
    ]
    
    session = models.ForeignKey(
        AIChatSession,
        on_delete=models.CASCADE,
        related_name='messages',
        verbose_name='Sesión'
    )
    role = models.CharField(
        max_length=10,
        choices=ROLE_CHOICES,
        verbose_name='Rol'
    )
    content = models.TextField(
        verbose_name='Contenido del mensaje'
    )
    tokens_used = models.PositiveIntegerField(
        default=0,
        verbose_name='Tokens utilizados',
        help_text='Número de tokens consumidos para este mensaje'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Mensaje de Chat IA'
        verbose_name_plural = 'Mensajes de Chat IA'
    
    def __str__(self):
        return f"{self.get_role_display()}: {self.content[:50]}..."
    
    def is_from_user(self):
        """Verifica si el mensaje es del usuario"""
        return self.role == 'user'
    
    def is_from_ai(self):
        """Verifica si el mensaje es de la IA"""
        return self.role == 'assistant'


class SharedAIChatMessage(models.Model):
    """DEPRECATED - Usar PublicAIMessageShare en su lugar"""
    
    message = models.ForeignKey(
        'AIChatMessage',
        on_delete=models.CASCADE,
        related_name='shared_instances_old',
        verbose_name='Mensaje'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='shared_ai_messages_old',
        verbose_name='Usuario que compartió'
    )
    share_token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token de compartición',
        help_text='Token único para acceder al mensaje compartido'
    )
    title = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Título público',
        help_text='Título opcional para el mensaje compartido'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Compartición activa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de compartición'
    )
    views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de vistas'
    )
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración',
        help_text='Fecha en que expira el enlace compartido (opcional)'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Mensaje IA Compartido (Deprecated)'
        verbose_name_plural = 'Mensajes IA Compartidos (Deprecated)'
        db_table = 'tickets_sharedaichatmessage_old'
    
    def __str__(self):
        title = self.title or f"Mensaje de {self.user.username}"
        return f"{title} - {self.share_token[:8]}..."
    
    def get_public_url(self):
        """Retorna la URL pública del mensaje compartido"""
        from django.urls import reverse
        return reverse('shared_ai_message_view', kwargs={'token': self.share_token})
    
    def is_expired(self):
        """Verifica si el enlace ha expirado"""
        if not self.expires_at:
            return False
        return timezone.now() > self.expires_at
    
    def increment_views(self):
        """Incrementa el contador de vistas"""
        self.views_count += 1
        self.save(update_fields=['views_count'])


class PublicAIMessageShare(models.Model):
    """Modelo nuevo para mensajes de chat IA compartidos públicamente"""
    
    chat_message = models.ForeignKey(
        'AIChatMessage',
        on_delete=models.CASCADE,
        related_name='public_shares',
        verbose_name='Mensaje de Chat'
    )
    shared_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_message_shares',
        verbose_name='Compartido por'
    )
    token = models.CharField(
        max_length=64,
        unique=True,
        db_index=True,
        verbose_name='Token de compartición',
        help_text='Token único para acceder al mensaje compartido'
    )
    public_title = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Título público',
        help_text='Título opcional para el mensaje compartido'
    )
    is_active = models.BooleanField(
        default=True,
        db_index=True,
        verbose_name='Compartición activa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        db_index=True,
        verbose_name='Fecha de compartición'
    )
    views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de vistas'
    )
    expiration_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración',
        help_text='Fecha en que expira el enlace compartido (opcional)'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Compartición de Mensaje IA'
        verbose_name_plural = 'Comparticiones de Mensajes IA'
        indexes = [
            models.Index(fields=['token']),
            models.Index(fields=['shared_by', '-created_at']),
            models.Index(fields=['is_active', '-created_at']),
        ]
    
    def __str__(self):
        title = self.public_title or f"Mensaje de {self.shared_by.username}"
        return f"{title} - {self.token[:8]}..."
    
    def get_public_url(self):
        """Retorna la URL pública del mensaje compartido"""
        from django.urls import reverse
        return reverse('shared_ai_message_view', kwargs={'token': self.token})
    
    def is_expired(self):
        """Verifica si el enlace ha expirado"""
        if not self.expiration_date:
            return False
        return timezone.now() > self.expiration_date
    
    def increment_views(self):
        """Incrementa el contador de vistas de forma atómica"""
        from django.db.models import F
        PublicAIMessageShare.objects.filter(pk=self.pk).update(views_count=F('views_count') + 1)
        self.refresh_from_db()


class Concept(models.Model):
    """Modelo para gestionar conceptos y definiciones que se muestran en el home"""
    
    term = models.CharField(
        max_length=200,
        unique=True,
        verbose_name='Término',
        help_text='El concepto o término a definir'
    )
    definition = models.TextField(
        verbose_name='Definición',
        help_text='Explicación del concepto'
    )
    category = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Categoría',
        help_text='Categoría del concepto (opcional)'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si el concepto debe mostrarse en el home'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición (menor número aparece primero)'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='concepts_created',
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['order', 'term']
        verbose_name = 'Concepto'
        verbose_name_plural = 'Conceptos'
    
    def __str__(self):
        return f"{self.term}: {self.definition[:50]}..."


class Exam(models.Model):
    """Modelo para los exámenes"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Examen'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='¿Es público?'
    )
    public_token = models.CharField(
        max_length=50,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token público'
    )
    passing_score = models.IntegerField(
        default=70,
        verbose_name='Puntuación mínima para aprobar (%)'
    )
    time_limit = models.IntegerField(
        blank=True,
        null=True,
        verbose_name='Tiempo límite (minutos)',
        help_text='Dejar vacío para sin límite de tiempo'
    )
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)
    
    class Meta:
        verbose_name = 'Examen'
        verbose_name_plural = 'Exámenes'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        if self.is_public and not self.public_token:
            import uuid
            self.public_token = str(uuid.uuid4())[:8]
        elif not self.is_public:
            self.public_token = None
        super().save(*args, **kwargs)
    
    def get_questions_count(self):
        return self.questions.count()
    
    def get_total_attempts(self):
        return self.attempts.count()


class ExamQuestion(models.Model):
    """Modelo para las preguntas del examen"""
    
    exam = models.ForeignKey(
        Exam,
        on_delete=models.CASCADE,
        related_name='questions',
        verbose_name='Examen'
    )
    question_text = models.TextField(
        verbose_name='Pregunta'
    )
    option_a = models.CharField(
        max_length=500,
        verbose_name='Opción A'
    )
    option_b = models.CharField(
        max_length=500,
        verbose_name='Opción B'
    )
    option_c = models.CharField(
        max_length=500,
        verbose_name='Opción C'
    )
    correct_option = models.CharField(
        max_length=1,
        choices=[('A', 'Opción A'), ('B', 'Opción B'), ('C', 'Opción C')],
        verbose_name='Opción correcta'
    )
    order = models.PositiveIntegerField(
        default=1,
        verbose_name='Orden'
    )
    created_at = models.DateTimeField(auto_now_add=True)
    
    class Meta:
        verbose_name = 'Pregunta del Examen'
        verbose_name_plural = 'Preguntas del Examen'
        ordering = ['order']
    
    def __str__(self):
        return f"{self.exam.title} - Pregunta {self.order}"


class ExamAttempt(models.Model):
    """Modelo para los intentos de examen"""
    
    exam = models.ForeignKey(
        Exam,
        on_delete=models.CASCADE,
        related_name='attempts',
        verbose_name='Examen'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Usuario'
    )
    participant_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del participante'
    )
    participant_email = models.EmailField(
        verbose_name='Email del participante'
    )
    score = models.FloatField(
        verbose_name='Puntuación (%)'
    )
    total_questions = models.IntegerField(
        verbose_name='Total de preguntas'
    )
    correct_answers = models.IntegerField(
        verbose_name='Respuestas correctas'
    )
    started_at = models.DateTimeField(
        verbose_name='Iniciado en'
    )
    completed_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Completado en'
    )
    time_taken = models.IntegerField(
        verbose_name='Tiempo tomado (segundos)'
    )
    passed = models.BooleanField(
        default=False,
        verbose_name='¿Aprobado?'
    )
    certificate_token = models.CharField(
        max_length=100,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token del certificado'
    )
    certificate_generated_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Certificado generado en'
    )
    
    class Meta:
        verbose_name = 'Intento de Examen'
        verbose_name_plural = 'Intentos de Examen'
        ordering = ['-completed_at']
    
    def __str__(self):
        return f"{self.participant_name} - {self.exam.title} ({self.score}%)"
    
    @property
    def time_taken_display(self):
        """Retorna el tiempo tomado en formato legible"""
        if not self.time_taken:
            return "No disponible"
        
        if self.time_taken < 60:
            return f"{self.time_taken} segundos"
        elif self.time_taken < 3600:
            minutes = self.time_taken // 60
            seconds = self.time_taken % 60
            return f"{minutes} min {seconds} seg"
        else:
            hours = self.time_taken // 3600
            minutes = (self.time_taken % 3600) // 60
            seconds = self.time_taken % 60
            return f"{hours}h {minutes}m {seconds}s"
    
    @property
    def incorrect_answers(self):
        """Retorna el número de respuestas incorrectas"""
        return self.total_questions - self.correct_answers
    
    def generate_certificate_token(self):
        """Genera un token único para el certificado"""
        if not self.certificate_token and self.passed:
            import uuid
            from django.utils import timezone
            self.certificate_token = str(uuid.uuid4())
            self.certificate_generated_at = timezone.now()
            self.save(update_fields=['certificate_token', 'certificate_generated_at'])
        return self.certificate_token
    
    def save(self, *args, **kwargs):
        # Generar token del certificado automáticamente si pasa el examen
        if self.passed and not self.certificate_token:
            import uuid
            self.certificate_token = str(uuid.uuid4())
            if not self.certificate_generated_at:
                from django.utils import timezone
                self.certificate_generated_at = timezone.now()
        super().save(*args, **kwargs)


class ExamAnswer(models.Model):
    """Modelo para las respuestas del examen"""
    
    attempt = models.ForeignKey(
        ExamAttempt,
        on_delete=models.CASCADE,
        related_name='answers',
        verbose_name='Intento'
    )
    question = models.ForeignKey(
        ExamQuestion,
        on_delete=models.CASCADE,
        verbose_name='Pregunta'
    )
    selected_option = models.CharField(
        max_length=1,
        choices=[('A', 'Opción A'), ('B', 'Opción B'), ('C', 'Opción C')],
        verbose_name='Opción seleccionada'
    )
    is_correct = models.BooleanField(
        default=False,
        verbose_name='¿Es correcta?'
    )
    
    class Meta:
        verbose_name = 'Respuesta del Examen'
        verbose_name_plural = 'Respuestas del Examen'
        unique_together = ['attempt', 'question']
    
    def __str__(self):
        return f"{self.attempt.participant_name} - {self.question.question_text[:50]}..."


class ContactoWeb(models.Model):
    """Modelo para gestionar contactos desde la web pública"""
    
    nombre = models.CharField(
        max_length=100,
        verbose_name='Nombre completo'
    )
    email = models.EmailField(
        verbose_name='Correo electrónico'
    )
    telefono = models.CharField(
        max_length=20,
        blank=True,
        null=True,
        verbose_name='Teléfono'
    )
    empresa = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Empresa'
    )
    asunto = models.CharField(
        max_length=200,
        verbose_name='Asunto'
    )
    mensaje = models.TextField(
        verbose_name='Mensaje'
    )
    fecha_creacion = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    leido = models.BooleanField(
        default=False,
        verbose_name='Leído'
    )
    respondido = models.BooleanField(
        default=False,
        verbose_name='Respondido'
    )
    ip_address = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        blank=True,
        null=True,
        verbose_name='User Agent'
    )
    
    class Meta:
        verbose_name = 'Contacto Web'
        verbose_name_plural = 'Contactos Web'
        ordering = ['-fecha_creacion']
    
    def __str__(self):
        return f"{self.nombre} - {self.asunto}"


class PublicDocumentUpload(models.Model):
    """Modelo para URLs públicas de subida de documentos"""
    
    # Token único para la URL pública
    upload_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de subida',
        help_text='Token único para la URL pública de subida'
    )
    
    # Configuración de la URL
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la subida',
        help_text='Descripción de qué tipo de documentos se esperan'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Instrucciones',
        help_text='Instrucciones para quien vaya a subir el documento'
    )
    
    # Usuario que creó la URL
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='public_upload_urls',
        verbose_name='Creado por'
    )
    
    # Empresa asociada (opcional)
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='public_upload_urls',
        verbose_name='Empresa/Cliente',
        help_text='Empresa o cliente para el cual se subirán los documentos'
    )
    
    # Configuración de expiración
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración',
        help_text='Opcional: fecha límite para usar esta URL'
    )
    
    # Estado
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si está desactivada, no se pueden subir más documentos'
    )
    
    # Configuración de límites
    max_uploads = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Máximo de subidas',
        help_text='Opcional: límite de documentos que se pueden subir'
    )
    upload_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Documentos subidos',
        help_text='Contador de documentos subidos usando esta URL'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    last_used_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Último uso'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'URL Pública de Subida'
        verbose_name_plural = 'URLs Públicas de Subida'
    
    def __str__(self):
        return f"{self.title} - {self.upload_token}"
    
    def is_valid(self):
        """Verifica si la URL está válida para usarse"""
        if not self.is_active:
            return False
        
        if self.expires_at and timezone.now() > self.expires_at:
            return False
        
        if self.max_uploads and self.upload_count >= self.max_uploads:
            return False
        
        return True
    
    def get_public_url(self):
        """Devuelve la URL pública para subir documentos"""
        from django.urls import reverse
        return reverse('public_document_upload', kwargs={'token': self.upload_token})


class Employee(models.Model):
    """Modelo para gestionar empleados de la empresa"""
    
    # Información personal básica
    first_name = models.CharField(
        max_length=100,
        verbose_name='Nombre'
    )
    last_name = models.CharField(
        max_length=100,
        verbose_name='Apellidos'
    )
    email = models.EmailField(
        unique=True,
        verbose_name='Correo electrónico'
    )
    phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Teléfono'
    )
    
    # Información profesional
    position = models.CharField(
        max_length=150,
        verbose_name='Cargo'
    )
    salary_euros = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Salario (EUR)'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción/Experiencia'
    )
    
    # Currículo
    resume_file = models.FileField(
        upload_to='resumes/',
        null=True,
        blank=True,
        verbose_name='Archivo de currículo'
    )
    
    # Análisis de IA
    ai_analysis = models.TextField(
        blank=True,
        verbose_name='Análisis de IA',
        help_text='Resumen automático generado por IA'
    )
    ai_score = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Puntuación IA (1-10)',
        help_text='Puntuación automática del candidato'
    )
    
    # Opinión de reunión de contratación
    hiring_meeting_opinion = models.TextField(
        blank=True,
        verbose_name='Opinión de reunión de contratación',
        help_text='Comentarios y evaluación de la reunión/entrevista de contratación'
    )
    
    # Relaciones
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='employees',
        verbose_name='Empresa'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_employees',
        verbose_name='Creado por'
    )
    
    # Estados
    STATUS_CHOICES = [
        ('candidate', 'Candidato'),
        ('in_process', 'En trámites'),
        ('employee', 'Empleado'),
        ('not_employee', 'No empleado'),
    ]
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='candidate',
        verbose_name='Estado'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de aplicación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Empleado'
        verbose_name_plural = 'Empleados'
    
    def __str__(self):
        return f"{self.first_name} {self.last_name} - {self.position}"
    
    def get_full_name(self):
        return f"{self.first_name} {self.last_name}"
    
    def is_candidate(self):
        return self.status == 'candidate'
    
    def is_in_process(self):
        return self.status == 'in_process'
    
    def is_employee(self):
        return self.status == 'employee'
    
    def is_not_employee(self):
        return self.status == 'not_employee'
    
    def get_status_color(self):
        """Devuelve el color para el badge según el estado"""
        colors = {
            'candidate': 'primary',
            'in_process': 'warning',
            'employee': 'success',
            'not_employee': 'danger',
        }
        return colors.get(self.status, 'secondary')


class EmployeePayroll(models.Model):
    """Modelo para gestionar nóminas de empleados"""
    
    # Relación con el empleado
    employee = models.ForeignKey(
        Employee,
        on_delete=models.CASCADE,
        related_name='payrolls',
        verbose_name='Empleado'
    )
    
    # Información de la nómina
    period_month = models.PositiveIntegerField(
        verbose_name='Mes',
        help_text='Mes de la nómina (1-12)'
    )
    period_year = models.PositiveIntegerField(
        verbose_name='Año',
        help_text='Año de la nómina'
    )
    
    # Archivos
    payroll_pdf = models.FileField(
        upload_to='payrolls/',
        verbose_name='Nómina (PDF)',
        help_text='Archivo PDF de la nómina'
    )
    payment_receipt = models.FileField(
        upload_to='payment_receipts/',
        null=True,
        blank=True,
        verbose_name='Comprobante de pago',
        help_text='Archivo del comprobante de pago (opcional)'
    )
    
    # Información adicional
    gross_salary = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Salario bruto'
    )
    net_salary = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Salario neto'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre esta nómina'
    )
    
    # Control
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_payrolls',
        verbose_name='Creado por'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de registro'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-period_year', '-period_month', '-created_at']
        verbose_name = 'Nómina de Empleado'
        verbose_name_plural = 'Nóminas de Empleados'
        unique_together = ['employee', 'period_month', 'period_year']
    
    def __str__(self):
        return f"{self.employee.get_full_name()} - {self.get_period_display()}"
    
    def get_period_display(self):
        """Devuelve el período en formato legible"""
        months = {
            1: 'Enero', 2: 'Febrero', 3: 'Marzo', 4: 'Abril',
            5: 'Mayo', 6: 'Junio', 7: 'Julio', 8: 'Agosto',
            9: 'Septiembre', 10: 'Octubre', 11: 'Noviembre', 12: 'Diciembre'
        }
        return f"{months.get(self.period_month, self.period_month)}/{self.period_year}"
    
    def has_payment_receipt(self):
        """Verifica si tiene comprobante de pago"""
        return bool(self.payment_receipt)


class JobApplicationToken(models.Model):
    """Modelo para gestionar tokens de aplicación pública a empleos"""
    
    # Token único
    application_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de aplicación'
    )
    
    # Información del puesto
    job_title = models.CharField(
        max_length=150,
        verbose_name='Título del puesto'
    )
    job_description = models.TextField(
        verbose_name='Descripción del puesto'
    )
    proposed_salary_euros = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Salario propuesto (EUR)',
        help_text='Salario que la empresa está dispuesta a pagar'
    )
    
    # Configuración
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='job_application_tokens',
        verbose_name='Empresa'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_job_tokens',
        verbose_name='Creado por'
    )
    
    # Estado y expiración
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración'
    )
    max_applications = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Máximo de aplicaciones'
    )
    application_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Aplicaciones recibidas'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Token de Aplicación de Empleo'
        verbose_name_plural = 'Tokens de Aplicación de Empleo'
    
    def __str__(self):
        return f"{self.job_title} - {self.application_token}"
    
    def is_valid(self):
        """Verifica si el token está válido para aplicaciones"""
        if not self.is_active:
            return False
        
        if self.expires_at and timezone.now() > self.expires_at:
            return False
        
        if self.max_applications and self.application_count >= self.max_applications:
            return False
        
        return True
    
    def get_public_url(self, request=None):
        """Devuelve la URL pública para aplicar al empleo"""
        from django.urls import reverse
        url = reverse('public_job_application', kwargs={'token': self.application_token})
        if request:
            return request.build_absolute_uri(url)
        return url
    
    def get_qr_data(self):
        """Datos para generar código QR"""
        from django.conf import settings
        base_url = getattr(settings, 'BASE_URL', 'http://localhost:8000')
        return f"{base_url}{self.get_public_url()}"


class Agreement(models.Model):
    """Modelo para gestionar acuerdos/contratos con firma digital"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('published', 'Publicado'),
        ('signed', 'Firmado'),
        ('cancelled', 'Cancelado'),
    ]
    
    # Información básica del acuerdo
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Acuerdo'
    )
    body = models.TextField(
        verbose_name='Cuerpo del Acuerdo',
        help_text='Contenido completo del acuerdo/contrato'
    )
    
    # Control de estado
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    
    # URLs y tokens
    public_token = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Token Público'
    )
    
    # Metadatos
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        verbose_name='Empresa'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_agreements',
        verbose_name='Creado por'
    )
    
    # Control de fechas
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    published_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de publicación'
    )
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración'
    )
    
    # Configuración
    max_signers = models.PositiveIntegerField(
        default=10,
        verbose_name='Máximo de firmantes',
        help_text='Número máximo de personas que pueden firmar este acuerdo'
    )
    requires_approval = models.BooleanField(
        default=False,
        verbose_name='Requiere aprobación',
        help_text='Las firmas requieren aprobación antes de ser válidas'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Acuerdo'
        verbose_name_plural = 'Acuerdos'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        if not self.public_token:
            import secrets
            self.public_token = secrets.token_urlsafe(32)
        super().save(*args, **kwargs)
    
    @property
    def is_editable(self):
        """Un acuerdo es editable si no tiene firmas válidas"""
        return not self.signatures.filter(is_valid=True).exists()
    
    @property
    def signature_count(self):
        """Número de firmas válidas"""
        return self.signatures.filter(is_valid=True).count()
    
    @property
    def is_expired(self):
        """Verifica si el acuerdo ha expirado"""
        if self.expires_at:
            return timezone.now() > self.expires_at
        return False
    
    @property
    def can_be_signed(self):
        """Verifica si el acuerdo puede ser firmado"""
        if self.status != 'published':
            return False
        if self.is_expired:
            return False
        if self.max_signers and self.signature_count >= self.max_signers:
            return False
        return True
    
    def get_public_url(self, request=None):
        """URL pública para firmar el acuerdo"""
        from django.urls import reverse
        url = reverse('public_agreement_sign', kwargs={'token': self.public_token})
        if request:
            return request.build_absolute_uri(url)
        return url
    
    def generate_ai_content(self, prompt_addition=""):
        """Generar contenido del acuerdo con IA basado en el título"""
        from tickets.models import SystemConfiguration
        
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            return None
            
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            base_prompt = f"""
Genera el contenido completo de un acuerdo/contrato profesional con el título: "{self.title}"

El acuerdo debe incluir:
1. Introducción y propósito
2. Términos y condiciones principales
3. Derechos y obligaciones de las partes
4. Cláusulas de cumplimiento
5. Términos de vigencia
6. Cláusulas de resolución de conflictos
7. Firmas y fecha

Haz que sea profesional, claro y legalmente sólido.
Usa un lenguaje formal pero comprensible.
Incluye espacios para nombres y firmas al final.

{prompt_addition}
"""
            
            response = client.chat.completions.create(
                model=config.openai_model or "gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un experto abogado especializado en redacción de contratos y acuerdos legales."},
                    {"role": "user", "content": base_prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            print(f"Error generando contenido IA: {e}")
            return None


class AgreementSignature(models.Model):
    """Modelo para las firmas de acuerdos"""
    
    # Relación con el acuerdo
    agreement = models.ForeignKey(
        Agreement,
        on_delete=models.CASCADE,
        related_name='signatures',
        verbose_name='Acuerdo'
    )
    
    # Información del firmante
    signer_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del firmante'
    )
    signer_email = models.EmailField(
        verbose_name='Email del firmante'
    )
    signature_data = models.TextField(
        verbose_name='Datos de la firma',
        help_text='Datos de la firma digital (base64)'
    )
    
    # Control de firma
    signed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de firma'
    )
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        null=True,
        blank=True,
        verbose_name='User Agent'
    )
    
    # Estado de la firma
    is_valid = models.BooleanField(
        default=True,
        verbose_name='Firma válida'
    )
    is_approved = models.BooleanField(
        default=True,
        verbose_name='Firma aprobada'
    )
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='approved_signatures',
        verbose_name='Aprobado por'
    )
    approved_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de aprobación'
    )
    
    # Descarga del PDF
    pdf_downloaded_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='PDF descargado en'
    )
    download_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Veces descargado'
    )
    
    class Meta:
        ordering = ['-signed_at']
        verbose_name = 'Firma de Acuerdo'
        verbose_name_plural = 'Firmas de Acuerdos'
        unique_together = ['agreement', 'signer_email']
    
    def __str__(self):
        return f"{self.signer_name} - {self.agreement.title}"
    
    def get_download_url(self, request=None):
        """URL para descargar el PDF firmado"""
        from django.urls import reverse
        url = reverse('download_signed_agreement', kwargs={
            'agreement_token': self.agreement.public_token,
            'signature_id': self.pk
        })
        if request:
            return request.build_absolute_uri(url)
        return url


class LandingPage(models.Model):
    """Modelo para Landing Pages de campañas de marketing"""
    
    # Información básica
    nombre_producto = models.CharField(
        max_length=200,
        verbose_name='Nombre del Producto'
    )
    descripcion = models.TextField(
        verbose_name='Descripción del Producto',
        help_text='Descripción que aparecerá en la landing page'
    )
    imagen = models.ImageField(
        upload_to='landing_pages/',
        blank=True,
        null=True,
        verbose_name='Imagen del Producto',
        help_text='Imagen principal de la landing page'
    )
    
    # Configuración de URL pública
    slug = models.SlugField(
        max_length=200,
        unique=True,
        verbose_name='URL de la Landing Page',
        help_text='URL única para acceder a la landing page (ej: mi-producto)'
    )
    
    # Token de API para seguridad
    api_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token de API',
        help_text='Token único para autenticar llamadas a la API de esta landing page'
    )
    
    # Estado
    is_active = models.BooleanField(
        default=True,
        verbose_name='Página Activa',
        help_text='Si está activa, la landing page será accesible públicamente'
    )
    
    # Configuración de diseño
    color_principal = models.CharField(
        max_length=7,
        default='#00A651',
        verbose_name='Color Principal',
        help_text='Color principal de la landing page (formato HEX: #00A651)'
    )
    color_secundario = models.CharField(
        max_length=7,
        default='#E8F5E8',
        verbose_name='Color Secundario',
        help_text='Color secundario para fondos (formato HEX: #E8F5E8)'
    )
    
    # Configuración del formulario
    titulo_formulario = models.CharField(
        max_length=200,
        default='VEA TODAS LAS DEMOSTRACIONES',
        verbose_name='Título del Formulario'
    )
    subtitulo_formulario = models.TextField(
        default='Salesforce acerca a las empresas con sus clientes. Vea nuestras demostraciones gratuitas y obtenga más información sobre cómo podemos ayudar a su negocio.',
        verbose_name='Subtítulo del Formulario'
    )
    
    # Configuración de empresa/campaña
    empresa_campana = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Empresa de la Campaña',
        help_text='Nombre de la empresa que ejecuta la campaña'
    )
    
    # Configuración de Telegram
    telegram_bot_token = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Bot Token de Telegram',
        help_text='Token del bot de Telegram para enviar notificaciones (ej: 123456789:ABCdefGHIjklMNOpqrsTUVwxyz)'
    )
    telegram_chat_id = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Chat ID de Telegram',
        help_text='ID del chat o canal donde enviar las notificaciones (ej: -1001234567890 o @mi_canal)'
    )
    
    # URL de reunión
    meeting_url = models.URLField(
        blank=True,
        verbose_name='URL de Reunión',
        help_text='URL para planificar reuniones (ej: Calendly, Google Meet, Zoom, etc.)'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Estadísticas
    views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de visitas'
    )
    submissions_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Formularios enviados'
    )
    meeting_button_clicks = models.PositiveIntegerField(
        default=0,
        verbose_name='Clics en "Planificar Reunión"'
    )
    contact_button_clicks = models.PositiveIntegerField(
        default=0,
        verbose_name='Clics en "Contactar"'
    )
    
    class Meta:
        verbose_name = 'Landing Page'
        verbose_name_plural = 'Landing Pages'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.nombre_producto} - {self.slug}"
    
    def get_public_url(self):
        """URL pública de la landing page"""
        from django.urls import reverse
        return reverse('landing_page_public', kwargs={'slug': self.slug})
    
    def get_absolute_url(self):
        """URL para el detalle de la landing page en admin"""
        from django.urls import reverse
        return reverse('landing_page_detail', kwargs={'pk': self.pk})
    
    def increment_views(self):
        """Incrementar contador de visitas"""
        self.views_count += 1
        self.save(update_fields=['views_count'])
    
    def increment_submissions(self):
        """Incrementar contador de envíos"""
        self.submissions_count += 1
        self.save(update_fields=['submissions_count'])
    
    @property
    def total_views(self):
        """Total de vistas de la landing page"""
        return self.views_count
    
    @property
    def total_submissions(self):
        """Total de envíos de la landing page"""
        return self.submissions_count
    
    @property
    def conversion_rate(self):
        """Tasa de conversión en porcentaje"""
        if self.views_count > 0:
            return round((self.submissions_count / self.views_count) * 100, 1)
        return 0


class LandingPageSubmission(models.Model):
    """Modelo para los envíos de formularios de landing pages"""
    
    landing_page = models.ForeignKey(
        LandingPage,
        on_delete=models.CASCADE,
        related_name='submissions',
        verbose_name='Landing Page'
    )
    
    # Datos del formulario
    nombre = models.CharField(
        max_length=100,
        verbose_name='Nombre'
    )
    apellido = models.CharField(
        max_length=100,
        verbose_name='Apellido'
    )
    email = models.EmailField(
        verbose_name='Email'
    )
    telefono = models.CharField(
        max_length=20,
        verbose_name='Teléfono'
    )
    empresa = models.CharField(
        max_length=200,
        verbose_name='Empresa'
    )
    cargo = models.CharField(
        max_length=100,
        verbose_name='Cargo'
    )
    mensaje = models.TextField(
        verbose_name='Mensaje o comentarios',
        help_text='Cuéntanos más sobre tus necesidades o proyecto'
    )
    
    # Preferencias de contacto
    preferred_contact_time = models.CharField(
        max_length=20,
        choices=[
            ('asap', 'Lo antes posible'),
            ('monday', 'El próximo lunes'),
            ('tuesday', 'El próximo martes'),
            ('wednesday', 'El próximo miércoles'),
            ('thursday', 'El próximo jueves'),
            ('friday', 'El próximo viernes'),
            ('next_week', 'La próxima semana'),
            ('flexible', 'Horario flexible'),
        ],
        default='asap',
        verbose_name='Cuándo prefieres que te contactemos'
    )
    
    # Metadatos de seguimiento
    ip_address = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        blank=True,
        null=True,
        verbose_name='User Agent'
    )
    utm_source = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Source'
    )
    utm_medium = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Medium'
    )
    utm_campaign = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Campaign'
    )
    
    # Estado
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de envío'
    )
    processed = models.BooleanField(
        default=False,
        verbose_name='Procesado'
    )
    processed_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de procesamiento'
    )
    processed_by = models.ForeignKey(
        'auth.User',
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        related_name='processed_submissions',
        verbose_name='Procesado por'
    )
    
    # Evaluación de IA
    ai_evaluated = models.BooleanField(
        default=False,
        verbose_name='Evaluado por IA'
    )
    ai_evaluation_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de evaluación IA'
    )
    ai_company_score = models.IntegerField(
        blank=True,
        null=True,
        verbose_name='Puntuación Empresa (1-10)',
        help_text='Evaluación de la calidad/seriedad de la empresa'
    )
    ai_contact_score = models.IntegerField(
        blank=True,
        null=True,
        verbose_name='Puntuación Contacto (1-10)',
        help_text='Evaluación de la calidad del contacto'
    )
    ai_project_score = models.IntegerField(
        blank=True,
        null=True,
        verbose_name='Puntuación Proyecto (1-10)',
        help_text='Evaluación del potencial del proyecto'
    )
    ai_overall_score = models.FloatField(
        blank=True,
        null=True,
        verbose_name='Puntuación General',
        help_text='Puntuación promedio general'
    )
    ai_evaluation_summary = models.TextField(
        blank=True,
        verbose_name='Resumen de Evaluación IA',
        help_text='Resumen detallado de la evaluación'
    )
    ai_recommendations = models.TextField(
        blank=True,
        verbose_name='Recomendaciones IA',
        help_text='Recomendaciones de seguimiento'
    )
    ai_priority_level = models.CharField(
        max_length=20,
        choices=[
            ('low', 'Baja'),
            ('medium', 'Media'),
            ('high', 'Alta'),
            ('urgent', 'Urgente'),
        ],
        blank=True,
        verbose_name='Nivel de Prioridad IA'
    )
    
    class Meta:
        verbose_name = 'Envío de Landing Page'
        verbose_name_plural = 'Envíos de Landing Pages'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.nombre} {self.apellido} - {self.landing_page.nombre_producto}"
    
    def evaluate_with_ai(self):
        """Evalúa el envío usando IA"""
        from .ai_landing_evaluator import evaluate_landing_submission
        
        try:
            result = evaluate_landing_submission(self)
            
            # Guardar resultados
            self.ai_evaluated = True
            self.ai_evaluation_date = timezone.now()
            self.ai_company_score = result.get('company_score')
            self.ai_contact_score = result.get('contact_score')
            self.ai_project_score = result.get('project_score')
            self.ai_overall_score = result.get('overall_score')
            self.ai_evaluation_summary = result.get('summary')
            self.ai_recommendations = result.get('recommendations')
            self.ai_priority_level = result.get('priority_level')
            
            self.save()
            return True
            
        except Exception as e:
            print(f"Error evaluando con IA: {e}")
            return False


class SharedFile(models.Model):
    """Modelo para gestionar archivos compartidos"""
    
    title = models.CharField(
        max_length=200, 
        verbose_name='Título del Archivo'
    )
    description = models.TextField(
        blank=True, 
        verbose_name='Descripción',
        help_text='Descripción opcional del archivo'
    )
    file = models.FileField(
        upload_to='shared_files/%Y/%m/',
        verbose_name='Archivo'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='shared_files',
        verbose_name='Empresa',
        null=True,
        blank=True
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='uploaded_files',
        verbose_name='Subido por',
        null=True,
        blank=True
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Si está marcado, todos los usuarios de la empresa pueden ver este archivo'
    )
    file_size = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Tamaño del archivo (bytes)'
    )
    file_type = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Tipo de archivo'
    )
    download_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de descargas'
    )
    created_at = models.DateTimeField(
        default=timezone.now, 
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True, 
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Archivo Compartido'
        verbose_name_plural = 'Archivos Compartidos'
    
    def __str__(self):
        return f"{self.title} - {self.company.name}"
    
    def get_file_size_display(self):
        """Convierte el tamaño del archivo a formato legible"""
        if not self.file_size:
            return "Desconocido"
        
        size = self.file_size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"
    
    def get_file_extension(self):
        """Obtiene la extensión del archivo"""
        if self.file:
            return os.path.splitext(self.file.name)[1][1:].upper()
        return ""
    
    def save(self, *args, **kwargs):
        if self.file:
            # Obtener tamaño del archivo
            self.file_size = self.file.size
            
            # Obtener tipo de archivo basado en la extensión
            ext = self.get_file_extension().lower()
            file_type_map = {
                'pdf': 'PDF Document',
                'doc': 'Word Document',
                'docx': 'Word Document',
                'xls': 'Excel Spreadsheet',
                'xlsx': 'Excel Spreadsheet',
                'ppt': 'PowerPoint Presentation',
                'pptx': 'PowerPoint Presentation',
                'txt': 'Text File',
                'jpg': 'Image',
                'jpeg': 'Image',
                'png': 'Image',
                'gif': 'Image',
                'zip': 'Archive',
                'rar': 'Archive',
                '7z': 'Archive',
            }
            self.file_type = file_type_map.get(ext, 'Archivo')
        
        super().save(*args, **kwargs)


class SharedFileDownload(models.Model):
    """Modelo para rastrear descargas de archivos compartidos"""
    
    shared_file = models.ForeignKey(
        SharedFile,
        on_delete=models.CASCADE,
        related_name='downloads',
        verbose_name='Archivo Compartido'
    )
    downloaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='file_downloads',
        verbose_name='Descargado por',
        null=True,
        blank=True
    )
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    downloaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de descarga'
    )
    
    class Meta:
        ordering = ['-downloaded_at']
        verbose_name = 'Descarga de Archivo'
        verbose_name_plural = 'Descargas de Archivos'
    
    def __str__(self):
        user = self.downloaded_by.username if self.downloaded_by else 'Anónimo'
        return f"{self.shared_file.title} - {user} - {self.downloaded_at}"


class Recording(models.Model):
    """Modelo para gestionar grabaciones de audio con transcripción IA"""
    
    TRANSCRIPTION_STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('processing', 'Procesando'),
        ('completed', 'Completada'),
        ('failed', 'Fallida'),
    ]
    
    title = models.CharField(
        max_length=200, 
        verbose_name='Título de la grabación'
    )
    description = models.TextField(
        blank=True, 
        verbose_name='Descripción',
        help_text='Descripción opcional de la grabación'
    )
    audio_file = models.FileField(
        upload_to='recordings/%Y/%m/',
        verbose_name='Archivo de audio'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='recordings',
        verbose_name='Empresa',
        null=True,
        blank=True
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='uploaded_recordings',
        verbose_name='Subido por',
        null=True,
        blank=True
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Si está marcado, todos los usuarios de la empresa pueden ver esta grabación'
    )
    file_size = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Tamaño del archivo (bytes)'
    )
    duration_seconds = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Duración en segundos'
    )
    
    # Campos de transcripción
    transcription_status = models.CharField(
        max_length=20,
        choices=TRANSCRIPTION_STATUS_CHOICES,
        default='pending',
        verbose_name='Estado de transcripción'
    )
    transcription_text = models.TextField(
        blank=True,
        verbose_name='Texto transcrito'
    )
    transcription_confidence = models.FloatField(
        null=True,
        blank=True,
        verbose_name='Confianza de transcripción (%)',
        help_text='Porcentaje de confianza de la transcripción'
    )
    transcription_language = models.CharField(
        max_length=10,
        default='es',
        verbose_name='Idioma detectado'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now, 
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True, 
        verbose_name='Última actualización'
    )
    transcribed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de transcripción'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Grabación'
        verbose_name_plural = 'Grabaciones'
    
    def __str__(self):
        company_name = self.company.name if self.company else 'Sin empresa'
        return f"{self.title} - {company_name}"
    
    def get_file_size_display(self):
        """Convierte el tamaño del archivo a formato legible"""
        if not self.file_size:
            return "Desconocido"
        
        size = self.file_size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"
    
    def get_duration_display(self):
        """Convierte la duración a formato legible"""
        if not self.duration_seconds:
            return "Desconocida"
        
        hours = self.duration_seconds // 3600
        minutes = (self.duration_seconds % 3600) // 60
        seconds = self.duration_seconds % 60
        
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        else:
            return f"{minutes:02d}:{seconds:02d}"
    
    def get_audio_extension(self):
        """Obtiene la extensión del archivo de audio"""
        if self.audio_file:
            return os.path.splitext(self.audio_file.name)[1][1:].upper()
        return ""
    
    def save(self, *args, **kwargs):
        if self.audio_file:
            # Obtener tamaño del archivo
            self.file_size = self.audio_file.size
        
        super().save(*args, **kwargs)


class RecordingPlayback(models.Model):
    """Modelo para rastrear reproducciones de grabaciones"""
    
    recording = models.ForeignKey(
        Recording,
        on_delete=models.CASCADE,
        related_name='playbacks',
        verbose_name='Grabación'
    )
    played_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='recording_playbacks',
        verbose_name='Reproducido por',
        null=True,
        blank=True
    )
    played_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de reproducción'
    )
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    class Meta:
        ordering = ['-played_at']
        verbose_name = 'Reproducción'
        verbose_name_plural = 'Reproducciones'
    
    def __str__(self):
        user = self.played_by.username if self.played_by else 'Anónimo'
        return f"{self.recording.title} - {user} - {self.played_at}"


class PageVisit(models.Model):
    """Modelo para registrar visitas a páginas públicas"""
    
    PAGE_CHOICES = [
        ('home', 'Página de Inicio'),
        ('blog', 'Blog'),
        ('blog_post', 'Artículo de Blog'),
        ('conceptos', 'Conceptos Públicos'),
        ('contact', 'Contacto'),
        ('landing', 'Landing Page'),
        ('course_public', 'Curso Público'),
        ('ticket_public', 'Ticket Público'),
        ('other', 'Otra'),
    ]
    
    # Información de la página visitada
    page_type = models.CharField(
        max_length=50,
        choices=PAGE_CHOICES,
        verbose_name='Tipo de Página'
    )
    page_url = models.URLField(
        max_length=500,
        verbose_name='URL Visitada'
    )
    page_title = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Título de la Página'
    )
    
    # Información del visitante
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP'
    )
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País'
    )
    country_code = models.CharField(
        max_length=2,
        blank=True,
        verbose_name='Código de País'
    )
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad'
    )
    region = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Región/Estado'
    )
    
    # Información técnica
    user_agent = models.TextField(
        verbose_name='User Agent'
    )
    browser = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Navegador'
    )
    browser_version = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Versión del Navegador'
    )
    operating_system = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Sistema Operativo'
    )
    device_type = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Tipo de Dispositivo'
    )
    is_mobile = models.BooleanField(
        default=False,
        verbose_name='Es Móvil'
    )
    is_bot = models.BooleanField(
        default=False,
        verbose_name='Es Bot'
    )
    
    # Información de referencia
    referrer = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Página de Referencia'
    )
    utm_source = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Source'
    )
    utm_medium = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Medium'
    )
    utm_campaign = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='UTM Campaign'
    )
    
    # Información de sesión
    session_id = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='ID de Sesión'
    )
    duration_seconds = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Duración en Segundos'
    )
    
    # Metadatos
    visited_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Visita'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    class Meta:
        ordering = ['-visited_at']
        verbose_name = 'Visita a Página'
        verbose_name_plural = 'Visitas a Páginas'
        indexes = [
            models.Index(fields=['page_type', 'visited_at']),
            models.Index(fields=['ip_address', 'visited_at']),
            models.Index(fields=['country_code']),
            models.Index(fields=['is_bot']),
        ]
    
    def __str__(self):
        return f"{self.get_page_type_display()} - {self.ip_address} - {self.visited_at}"
    
    @property
    def location(self):
        """Devuelve una representación legible de la ubicación"""
        parts = []
        if self.city:
            parts.append(self.city)
        if self.region:
            parts.append(self.region)
        if self.country:
            parts.append(self.country)
        return ', '.join(parts) if parts else 'Desconocida'
    
    @property
    def browser_info(self):
        """Devuelve información del navegador formateada"""
        if self.browser and self.browser_version:
            return f"{self.browser} {self.browser_version}"
        elif self.browser:
            return self.browser
        return 'Desconocido'


class AIBlogConfigurator(models.Model):
    """Configurador de IA para generar contenido de blog automáticamente"""
    
    # Información básica
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Configurador',
        help_text='Nombre descriptivo para identificar este configurador'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional de qué tipo de contenido genera'
    )
    
    # Configuración de palabras clave
    keywords = models.TextField(
        verbose_name='Palabras Clave',
        help_text='Palabras clave separadas por comas para generar contenido (ej: tecnología, desarrollo, python)'
    )
    
    # Configuración de generación
    topic_template = models.TextField(
        default='Escribe un artículo completo sobre {keyword} con ejemplos prácticos y consejos útiles.',
        verbose_name='Plantilla de Tema',
        help_text='Plantilla para generar el prompt. Usa {keyword} donde quieras insertar la palabra clave'
    )
    
    content_length = models.CharField(
        max_length=50,
        choices=[
            ('short', 'Corto (500-800 palabras)'),
            ('medium', 'Medio (800-1200 palabras)'),
            ('long', 'Largo (1200-2000 palabras)'),
        ],
        default='medium',
        verbose_name='Longitud del Contenido'
    )
    
    content_style = models.CharField(
        max_length=50,
        choices=[
            ('professional', 'Profesional'),
            ('casual', 'Casual'),
            ('technical', 'Técnico'),
            ('educational', 'Educativo'),
            ('news', 'Noticia'),
        ],
        default='professional',
        verbose_name='Estilo del Contenido'
    )
    
    # Configuración de programación
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está activo, el configurador puede generar contenido automáticamente'
    )
    
    schedule_time = models.TimeField(
        default='10:00',
        verbose_name='Hora de Ejecución',
        help_text='Hora del día para ejecutar la generación automática (formato 24h)'
    )
    
    frequency_days = models.PositiveIntegerField(
        default=1,
        verbose_name='Frecuencia (días)',
        help_text='Cada cuántos días ejecutar (1 = diario, 7 = semanal, etc.)'
    )
    
    max_posts_per_run = models.PositiveIntegerField(
        default=1,
        verbose_name='Máximo Posts por Ejecución',
        help_text='Número máximo de posts a generar en cada ejecución'
    )
    
    # Configuración de categoría
    default_category = models.ForeignKey(
        'BlogCategory',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Categoría por Defecto',
        help_text='Categoría a asignar a los posts generados'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    last_run = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última Ejecución'
    )
    total_posts_generated = models.PositiveIntegerField(
        default=0,
        verbose_name='Total Posts Generados'
    )
    
    class Meta:
        verbose_name = 'Configurador de IA para Blog'
        verbose_name_plural = 'Configuradores de IA para Blog'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.name} - {'Activo' if self.is_active else 'Inactivo'}"
    
    @property
    def keywords_list(self):
        """Devuelve las palabras clave como lista"""
        return [kw.strip() for kw in self.keywords.split(',') if kw.strip()]
    
    @property
    def next_run_time(self):
        """Calcula la próxima fecha de ejecución"""
        from datetime import datetime, timedelta
        if not self.last_run:
            # Si nunca se ha ejecutado, usar hoy
            today = datetime.now().date()
            return datetime.combine(today, self.schedule_time)
        
        # Calcular próxima ejecución basada en frecuencia
        next_date = self.last_run.date() + timedelta(days=self.frequency_days)
        return datetime.combine(next_date, self.schedule_time)
    
    def should_run(self):
        """Determina si el configurador debe ejecutarse ahora"""
        if not self.is_active:
            return False
        
        from datetime import datetime
        now = datetime.now()
        
        # Si nunca se ha ejecutado, puede ejecutarse
        if not self.last_run:
            return True
        
        # Verificar si ha pasado el tiempo suficiente según la frecuencia
        next_run = self.next_run_time
        return now >= next_run
    
    def should_run_now(self):
        """Verifica si debe ejecutarse ahora"""
        if not self.is_active:
            return False
        
        from django.utils import timezone
        now = timezone.now()
        next_run = self.next_run_time
        
        # Permitir un margen de 5 minutos
        return now >= next_run and now <= next_run + timedelta(minutes=5)


class AIBlogGenerationLog(models.Model):
    """Log de generaciones de contenido por IA"""
    
    configurator = models.ForeignKey(
        AIBlogConfigurator,
        on_delete=models.CASCADE,
        related_name='generation_logs',
        verbose_name='Configurador'
    )
    
    keyword_used = models.CharField(
        max_length=200,
        verbose_name='Palabra Clave Utilizada'
    )
    
    generated_post = models.ForeignKey(
        'BlogPost',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Post Generado'
    )
    
    prompt_used = models.TextField(
        verbose_name='Prompt Utilizado'
    )
    
    generation_status = models.CharField(
        max_length=20,
        choices=[
            ('success', 'Exitoso'),
            ('error', 'Error'),
            ('pending', 'Pendiente'),
        ],
        default='pending',
        verbose_name='Estado de Generación'
    )
    
    error_message = models.TextField(
        blank=True,
        verbose_name='Mensaje de Error'
    )
    
    execution_time_seconds = models.FloatField(
        null=True,
        blank=True,
        verbose_name='Tiempo de Ejecución (segundos)'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    class Meta:
        verbose_name = 'Log de Generación IA'
        verbose_name_plural = 'Logs de Generación IA'
        ordering = ['-created_at']
    
    def __str__(self):
        status_icon = '✅' if self.generation_status == 'success' else '❌' if self.generation_status == 'error' else '⏳'
        return f"{status_icon} {self.configurator.name} - {self.keyword_used} ({self.created_at.strftime('%d/%m/%Y %H:%M')})"


class MultipleDocumentation(models.Model):
    """Modelo para gestionar documentación múltiple en una sola URL pública"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la Documentación'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción de qué contiene esta documentación'
    )
    
    public_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        verbose_name='Token Público',
        help_text='Token único para acceso público'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está activo, se puede acceder públicamente'
    )
    
    # Protección con contraseña
    password_protected = models.BooleanField(
        default=False,
        verbose_name='Protegido con contraseña',
        help_text='Si está marcado, se requerirá una contraseña para acceder'
    )
    access_password = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Contraseña de acceso',
        help_text='Contraseña requerida para acceder a la documentación pública'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de Actualización'
    )
    
    class Meta:
        verbose_name = 'Documentación Múltiple'
        verbose_name_plural = 'Documentaciones Múltiples'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
    
    def get_public_url(self):
        """Obtiene la URL pública para compartir"""
        from django.urls import reverse
        return reverse('multiple_documentation_public', kwargs={'token': self.public_token})


class MultipleDocumentationItem(models.Model):
    """Modelo para cada archivo dentro de una documentación múltiple"""
    
    documentation = models.ForeignKey(
        MultipleDocumentation,
        on_delete=models.CASCADE,
        related_name='items',
        verbose_name='Documentación'
    )
    
    number = models.PositiveIntegerField(
        verbose_name='Número',
        help_text='Número de orden del documento'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Documento'
    )
    
    file = models.FileField(
        upload_to='multiple_documentation/',
        verbose_name='Archivo',
        help_text='Archivo a compartir'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción del archivo'
    )
    
    download_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Contador de Descargas',
        help_text='Número de veces que se ha descargado este archivo'
    )
    
    expiration_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Vencimiento',
        help_text='Fecha de vencimiento del documento (opcional)'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    class Meta:
        verbose_name = 'Item de Documentación'
        verbose_name_plural = 'Items de Documentación'
        ordering = ['number']
        unique_together = ['documentation', 'number']
    
    def __str__(self):
        return f"{self.number}. {self.name}"
    
    def get_file_size(self):
        """Obtiene el tamaño del archivo en formato legible"""
        try:
            size = self.file.size
            for unit in ['bytes', 'KB', 'MB', 'GB']:
                if size < 1024.0:
                    return f"{size:.1f} {unit}"
                size /= 1024.0
            return f"{size:.1f} TB"
        except:
            return "N/A"
    
    def get_file_extension(self):
        """Obtiene la extensión del archivo"""
        import os
        return os.path.splitext(self.file.name)[1].lower() if self.file else ""
    
    def is_expired(self):
        """Verifica si el documento está vencido"""
        if not self.expiration_date:
            return False
        from django.utils import timezone
        return timezone.now().date() > self.expiration_date
    
    def days_until_expiration(self):
        """Calcula los días hasta el vencimiento"""
        if not self.expiration_date:
            return None
        from django.utils import timezone
        delta = self.expiration_date - timezone.now().date()
        return delta.days


class MultipleDocumentationStats(models.Model):
    """Modelo para estadísticas de documentación múltiple"""
    
    documentation = models.ForeignKey(
        MultipleDocumentation,
        on_delete=models.CASCADE,
        related_name='stats',
        verbose_name='Documentación'
    )
    
    # Estadísticas de página
    page_views = models.PositiveIntegerField(
        default=0,
        verbose_name='Visualizaciones de Página',
        help_text='Número total de veces que se ha visitado la página pública'
    )
    
    unique_visitors = models.PositiveIntegerField(
        default=0,
        verbose_name='Visitantes Únicos',
        help_text='Número de visitantes únicos (basado en IP)'
    )
    
    # Estadísticas de descargas
    total_downloads = models.PositiveIntegerField(
        default=0,
        verbose_name='Descargas Totales',
        help_text='Número total de archivos descargados'
    )
    
    # Fechas de seguimiento
    first_view_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Primera Visualización'
    )
    
    last_view_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última Visualización'
    )
    
    last_download_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última Descarga'
    )
    
    # Timestamps
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de Actualización'
    )
    
    class Meta:
        verbose_name = 'Estadística de Documentación'
        verbose_name_plural = 'Estadísticas de Documentación'
        unique_together = ['documentation']
    
    def __str__(self):
        return f"Stats: {self.documentation.title}"
    
    def get_average_downloads_per_day(self):
        """Calcula el promedio de descargas por día"""
        if not self.first_view_date or self.total_downloads == 0:
            return 0
        
        days = (timezone.now() - self.first_view_date).days
        if days == 0:
            return self.total_downloads
        return round(self.total_downloads / days, 2)
    
    def get_conversion_rate(self):
        """Calcula la tasa de conversión (descargas/visualizaciones)"""
        if self.page_views == 0:
            return 0
        return round((self.total_downloads / self.page_views) * 100, 2)


class MultipleDocumentationItemStats(models.Model):
    """Modelo para estadísticas de archivos individuales"""
    
    item = models.OneToOneField(
        MultipleDocumentationItem,
        on_delete=models.CASCADE,
        related_name='stats',
        verbose_name='Archivo'
    )
    
    download_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de Descargas'
    )
    
    unique_downloaders = models.PositiveIntegerField(
        default=0,
        verbose_name='Descargadores Únicos',
        help_text='Número de IPs únicas que han descargado este archivo'
    )
    
    first_download_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Primera Descarga'
    )
    
    last_download_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última Descarga'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de Actualización'
    )
    
    class Meta:
        verbose_name = 'Estadística de Archivo'
        verbose_name_plural = 'Estadísticas de Archivos'
    
    def __str__(self):
        return f"Stats: {self.item.name} ({self.download_count} descargas)"
    
    def get_popularity_percentage(self):
        """Calcula el porcentaje de popularidad respecto a otros archivos de la misma documentación"""
        documentation_stats = self.item.documentation.stats.first()
        if not documentation_stats or documentation_stats.total_downloads == 0:
            return 0
        return round((self.download_count / documentation_stats.total_downloads) * 100, 2)


class MultipleDocumentationVisit(models.Model):
    """Modelo para trackear visitas individuales"""
    
    documentation = models.ForeignKey(
        MultipleDocumentation,
        on_delete=models.CASCADE,
        related_name='visits',
        verbose_name='Documentación'
    )
    
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP'
    )
    
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    referer = models.URLField(
        blank=True,
        null=True,
        verbose_name='Referente'
    )
    
    timestamp = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha y Hora'
    )
    
    class Meta:
        verbose_name = 'Visita a Documentación'
        verbose_name_plural = 'Visitas a Documentación'
        ordering = ['-timestamp']
    
    def __str__(self):
        return f"{self.ip_address} - {self.documentation.title} - {self.timestamp.strftime('%d/%m/%Y %H:%M')}"


class MultipleDocumentationDownload(models.Model):
    """Modelo para trackear descargas individuales"""
    
    item = models.ForeignKey(
        MultipleDocumentationItem,
        on_delete=models.CASCADE,
        related_name='downloads',
        verbose_name='Archivo'
    )
    
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP'
    )
    
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    referer = models.URLField(
        blank=True,
        null=True,
        verbose_name='Referente'
    )
    
    timestamp = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha y Hora'
    )
    
    class Meta:
        verbose_name = 'Descarga de Archivo'
        verbose_name_plural = 'Descargas de Archivos'
        ordering = ['-timestamp']
    
    def __str__(self):
        return f"{self.ip_address} - {self.item.name} - {self.timestamp.strftime('%d/%m/%Y %H:%M')}"


class TaskSchedule(models.Model):
    """Modelo para gestionar cronogramas de tareas con vista Gantt"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Cronograma'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='task_schedules',
        verbose_name='Empresa'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_schedules',
        verbose_name='Creado Por'
    )
    
    start_date = models.DateField(
        verbose_name='Fecha de Inicio'
    )
    
    end_date = models.DateField(
        verbose_name='Fecha de Fin'
    )
    
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Si está activo, el cronograma será accesible públicamente'
    )
    
    public_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token Público'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    
    class Meta:
        verbose_name = 'Cronograma de Tareas'
        verbose_name_plural = 'Cronogramas de Tareas'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
    
    def get_progress_percentage(self):
        """Calcula el porcentaje de progreso del cronograma"""
        total_tasks = self.tasks.count()
        if total_tasks == 0:
            return 0
        completed_tasks = self.tasks.filter(is_completed=True).count()
        return round((completed_tasks / total_tasks) * 100, 2)
    
    def get_total_tasks(self):
        """Retorna el total de tareas"""
        return self.tasks.count()
    
    def get_completed_tasks(self):
        """Retorna el total de tareas completadas"""
        return self.tasks.filter(is_completed=True).count()
    
    def get_pending_tasks(self):
        """Retorna el total de tareas pendientes"""
        return self.tasks.filter(is_completed=False).count()
    
    def get_overdue_tasks(self):
        """Retorna las tareas vencidas"""
        from django.utils import timezone
        today = timezone.now().date()
        return self.tasks.filter(
            is_completed=False,
            end_date__lt=today
        ).count()
    
    def get_on_time_completion_rate(self):
        """Calcula el porcentaje de tareas completadas a tiempo"""
        completed_tasks = self.tasks.filter(is_completed=True)
        if completed_tasks.count() == 0:
            return 0
        on_time_tasks = completed_tasks.filter(
            completed_at__lte=models.F('end_date')
        ).count()
        return round((on_time_tasks / completed_tasks.count()) * 100, 2)
    
    def is_overdue(self):
        """Verifica si el cronograma está vencido"""
        from django.utils import timezone
        today = timezone.now().date()
        return today > self.end_date and self.get_progress_percentage() < 100
    
    def calculate_auto_dates(self):
        """Calcula automáticamente las fechas de inicio y fin basándose en las tareas"""
        tasks = self.tasks.all()
        if tasks.exists():
            # Fecha de inicio: la fecha más temprana de todas las tareas
            earliest_start = tasks.order_by('start_date').first().start_date
            # Fecha de fin: la fecha más tardía de todas las tareas
            latest_end = tasks.order_by('-end_date').first().end_date
            return earliest_start, latest_end
        return None, None
    
    def update_auto_dates(self):
        """Actualiza automáticamente las fechas del cronograma basándose en las tareas"""
        start_date, end_date = self.calculate_auto_dates()
        if start_date and end_date:
            self.start_date = start_date
            self.end_date = end_date
            self.save(update_fields=['start_date', 'end_date'])
    
    def get_total_days(self):
        """Calcula la duración total del cronograma en días laborables (excluyendo sábados y domingos)"""
        if not self.start_date or not self.end_date:
            return 0
        
        # Importar datetime para cálculos de días laborables
        from datetime import timedelta
        
        current_date = self.start_date
        business_days = 0
        
        while current_date <= self.end_date:
            # 0=Lunes, 1=Martes, ..., 6=Domingo
            # Solo contar lunes a viernes (0-4)
            if current_date.weekday() < 5:
                business_days += 1
            current_date += timedelta(days=1)
        
        return business_days
    
    def get_total_task_days(self):
        """Calcula la suma total de días de todas las tareas"""
        return sum(task.get_duration_days() for task in self.tasks.all())
    
    def get_total_task_hours(self):
        """Calcula la suma total de horas de todas las tareas (8 horas por día)"""
        return self.get_total_task_days() * 8
    
    def get_total_work_weeks(self):
        """Calcula la suma total de semanas de trabajo (40 horas por semana)"""
        total_hours = self.get_total_task_hours()
        if total_hours == 0:
            return 0
        return round(total_hours / 40, 1)


class ScheduleTask(models.Model):
    """Modelo para las tareas individuales del cronograma"""
    
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('critical', 'Crítica'),
    ]
    
    schedule = models.ForeignKey(
        TaskSchedule,
        on_delete=models.CASCADE,
        related_name='tasks',
        verbose_name='Cronograma'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la Tarea'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    
    start_date = models.DateField(
        verbose_name='Fecha de Inicio'
    )
    
    end_date = models.DateField(
        verbose_name='Fecha de Fin'
    )
    
    priority = models.CharField(
        max_length=10,
        choices=PRIORITY_CHOICES,
        default='medium',
        verbose_name='Prioridad'
    )
    
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='schedule_tasks_assigned',
        verbose_name='Asignado A'
    )
    
    is_completed = models.BooleanField(
        default=False,
        verbose_name='Completada'
    )
    
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Completada el'
    )
    
    progress_percentage = models.IntegerField(
        default=0,
        verbose_name='Porcentaje de Progreso',
        help_text='Progreso de 0 a 100'
    )
    
    dependencies = models.ManyToManyField(
        'self',
        symmetrical=False,
        blank=True,
        related_name='dependent_tasks',
        verbose_name='Dependencias'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    
    class Meta:
        verbose_name = 'Tarea del Cronograma'
        verbose_name_plural = 'Tareas del Cronograma'
        ordering = ['start_date']
    
    def __str__(self):
        return f"{self.schedule.title} - {self.title}"
    
    def save(self, *args, **kwargs):
        # Si la tarea se marca como completada, establecer la fecha de completado
        if self.is_completed and not self.completed_at:
            self.completed_at = timezone.now()
            self.progress_percentage = 100
        elif not self.is_completed:
            self.completed_at = None
        
        # Guardar la tarea primero
        super().save(*args, **kwargs)
        
        # Actualizar las fechas del cronograma automáticamente
        self.schedule.update_auto_dates()
    
    def is_overdue(self):
        """Verifica si la tarea está vencida"""
        from django.utils import timezone
        today = timezone.now().date()
        return not self.is_completed and today > self.end_date
    
    def get_duration_days(self):
        """Calcula la duración de la tarea en días laborables (excluyendo sábados y domingos)"""
        if not self.start_date or not self.end_date:
            return 0
        
        # Importar datetime para cálculos de días laborables
        from datetime import timedelta
        
        current_date = self.start_date
        business_days = 0
        
        while current_date <= self.end_date:
            # 0=Lunes, 1=Martes, ..., 6=Domingo
            # Solo contar lunes a viernes (0-4)
            if current_date.weekday() < 5:
                business_days += 1
            current_date += timedelta(days=1)
        
        return business_days
    
    def can_start(self):
        """Verifica si la tarea puede iniciar (todas sus dependencias completadas)"""
        return all(dep.is_completed for dep in self.dependencies.all())


class ScheduleComment(models.Model):
    """Modelo para comentarios en las tareas"""
    
    task = models.ForeignKey(
        ScheduleTask,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Tarea'
    )
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    
    comment = models.TextField(
        verbose_name='Comentario'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    class Meta:
        verbose_name = 'Comentario de Tarea'
        verbose_name_plural = 'Comentarios de Tareas'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.user.username} - {self.task.title}"


class TicketApproval(models.Model):
    """Modelo para aprobaciones de clientes en tickets públicos"""
    
    ticket = models.OneToOneField(
        Ticket,
        on_delete=models.CASCADE,
        related_name='client_approval',
        verbose_name='Ticket'
    )
    
    client_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Cliente'
    )
    
    client_email = models.EmailField(
        verbose_name='Correo del Cliente'
    )
    
    approved_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Aprobación'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas Adicionales',
        help_text='Comentarios adicionales del cliente'
    )
    
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP',
        help_text='IP desde donde se realizó la aprobación'
    )
    
    class Meta:
        verbose_name = 'Aprobación de Cliente'
        verbose_name_plural = 'Aprobaciones de Clientes'
        ordering = ['-approved_at']
    
    def __str__(self):
        return f"{self.ticket.ticket_number} - Aprobado por {self.client_name}"


class SatisfactionSurvey(models.Model):
    """Modelo para encuestas de satisfacción de clientes"""
    
    RATING_CHOICES = [
        (1, '1 - Muy Insatisfecho'),
        (2, '2 - Insatisfecho'),
        (3, '3 - Neutral'),
        (4, '4 - Satisfecho'),
        (5, '5 - Muy Satisfecho'),
    ]
    
    RESOLUTION_QUALITY_CHOICES = [
        (1, '1 - Muy Mala'),
        (2, '2 - Mala'),
        (3, '3 - Regular'),
        (4, '4 - Buena'),
        (5, '5 - Excelente'),
    ]
    
    RESPONSE_TIME_CHOICES = [
        (1, '1 - Muy Lento'),
        (2, '2 - Lento'),
        (3, '3 - Aceptable'),
        (4, '4 - Rápido'),
        (5, '5 - Muy Rápido'),
    ]
    
    ticket = models.OneToOneField(
        Ticket,
        on_delete=models.CASCADE,
        related_name='satisfaction_survey',
        verbose_name='Ticket'
    )
    
    client_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Cliente'
    )
    
    client_email = models.EmailField(
        verbose_name='Correo del Cliente'
    )
    
    # Calificaciones
    overall_satisfaction = models.IntegerField(
        choices=RATING_CHOICES,
        verbose_name='Satisfacción General',
        help_text="¿Qué tan satisfecho está con el servicio general?"
    )
    
    resolution_quality = models.IntegerField(
        choices=RESOLUTION_QUALITY_CHOICES,
        verbose_name='Calidad de la Solución',
        help_text="¿Cómo califica la calidad de la solución proporcionada?"
    )
    
    response_time = models.IntegerField(
        choices=RESPONSE_TIME_CHOICES,
        verbose_name='Tiempo de Respuesta',
        help_text="¿Cómo califica nuestro tiempo de respuesta?"
    )
    
    communication = models.IntegerField(
        choices=RATING_CHOICES,
        verbose_name='Calidad de la Comunicación',
        help_text="¿Cómo califica la comunicación durante el proceso?"
    )
    
    # Comentarios
    what_went_well = models.TextField(
        blank=True,
        null=True,
        verbose_name='¿Qué funcionó bien?',
        help_text='Describa los aspectos positivos de nuestro servicio'
    )
    
    what_could_improve = models.TextField(
        blank=True,
        null=True,
        verbose_name='¿Qué se podría mejorar?',
        help_text='Sugerencias para mejorar nuestro servicio'
    )
    
    additional_comments = models.TextField(
        blank=True,
        null=True,
        verbose_name='Comentarios Adicionales',
        help_text='Cualquier otro comentario que desee compartir'
    )
    
    # Recomendación
    would_recommend = models.BooleanField(
        default=True,
        verbose_name='¿Recomendaría nuestros servicios?',
        help_text="¿Recomendaría nuestros servicios a otros?"
    )
    
    recommendation_reason = models.TextField(
        blank=True,
        null=True,
        verbose_name='Razón de recomendación',
        help_text='¿Por qué sí o por qué no recomendaría nuestros servicios?'
    )
    
    # Metadatos
    submitted_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Envío'
    )
    
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP',
        help_text='IP desde donde se realizó la encuesta'
    )
    
    class Meta:
        verbose_name = "Encuesta de Satisfacción"
        verbose_name_plural = "Encuestas de Satisfacción"
        ordering = ['-submitted_at']
    
    def __str__(self):
        return f"Encuesta de {self.client_name} para ticket #{self.ticket.id} - {self.overall_satisfaction}/5"
    
    @property
    def average_rating(self):
        """Calcula el promedio de todas las calificaciones"""
        ratings = [
            self.overall_satisfaction,
            self.resolution_quality,
            self.response_time,
            self.communication
        ]
        return round(sum(ratings) / len(ratings), 1)
    
    @property
    def rating_summary(self):
        """Devuelve un resumen de las calificaciones"""
        avg = self.average_rating
        if avg >= 4.5:
            return "Excelente"
        elif avg >= 3.5:
            return "Bueno"
        elif avg >= 2.5:
            return "Regular"
        else:
            return "Necesita Mejora"
    
    @property
    def rating_color(self):
        """Devuelve el color Bootstrap según la calificación"""
        avg = self.average_rating
        if avg >= 4.5:
            return "success"
        elif avg >= 3.5:
            return "primary"
        elif avg >= 2.5:
            return "warning"
        else:
            return "danger"


# Señales para actualizar automáticamente las fechas del cronograma
@receiver(post_save, sender=ScheduleTask)
def update_schedule_dates_on_task_save(sender, instance, **kwargs):
    """Actualiza las fechas del cronograma cuando se guarda una tarea"""
    instance.schedule.update_auto_dates()


@receiver(models.signals.post_delete, sender=ScheduleTask)
def update_schedule_dates_on_task_delete(sender, instance, **kwargs):
    """Actualiza las fechas del cronograma cuando se elimina una tarea"""
    instance.schedule.update_auto_dates()


class FinancialAction(models.Model):
    """Modelo para gestionar acciones financieras que se muestran en el ticker"""
    
    symbol = models.CharField(
        max_length=20,
        unique=True,
        verbose_name='Símbolo',
        help_text='Símbolo de la acción (ej: EUR/USD, AAPL, GOOGL)'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre',
        help_text='Nombre completo de la acción o par de divisas'
    )
    
    current_price = models.DecimalField(
        max_digits=12,
        decimal_places=4,
        null=True,
        blank=True,
        verbose_name='Precio Actual',
        help_text='Se actualiza automáticamente desde APIs externas'
    )
    
    previous_price = models.DecimalField(
        max_digits=12,
        decimal_places=4,
        null=True,
        blank=True,
        verbose_name='Precio Anterior',
        help_text='Se actualiza automáticamente desde APIs externas'
    )
    
    currency = models.CharField(
        max_length=10,
        default='USD',
        verbose_name='Moneda',
        help_text='Moneda en la que se cotiza (USD, EUR, etc.)'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si está activa se muestra en el ticker'
    )
    
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en el ticker (menor número = primera)'
    )
    
    last_updated = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    class Meta:
        ordering = ['order', 'symbol']
        verbose_name = 'Acción Financiera'
        verbose_name_plural = 'Acciones Financieras'
    
    def __str__(self):
        price_display = self.current_price if self.current_price is not None else "N/A"
        return f"{self.symbol} - {price_display} {self.currency}"
    
    @property
    def price_change(self):
        """Calcula el cambio de precio"""
        if self.previous_price and self.current_price is not None:
            return self.current_price - self.previous_price
        return 0
    
    @property
    def price_change_percent(self):
        """Calcula el porcentaje de cambio"""
        if self.previous_price and self.previous_price != 0 and self.current_price is not None:
            return ((self.current_price - self.previous_price) / self.previous_price) * 100
        return 0
    
    @property
    def is_positive_change(self):
        """Determina si el cambio es positivo"""
        return self.price_change > 0
    
    @property
    def is_negative_change(self):
        """Determina si el cambio es negativo"""
        return self.price_change < 0
    
    @property
    def change_color_class(self):
        """Devuelve la clase CSS según el cambio"""
        if self.is_positive_change:
            return 'text-success'
        elif self.is_negative_change:
            return 'text-danger'
        return 'text-muted'
    
    def get_price_from_yesterday(self):
        """Obtiene el precio del día anterior a la misma hora aproximada"""
        from datetime import datetime, timedelta
        
        # Calcular la fecha y hora del día anterior
        yesterday = timezone.now() - timedelta(days=1)
        
        # Buscar el precio más cercano a la misma hora del día anterior
        # Con una ventana de +/- 2 horas
        start_time = yesterday - timedelta(hours=2)
        end_time = yesterday + timedelta(hours=2)
        
        history_record = self.price_history.filter(
            recorded_at__range=[start_time, end_time]
        ).order_by('recorded_at').first()
        
        if history_record:
            return history_record.price
        
        # Si no hay registro del día anterior, usar el previous_price actual
        return self.previous_price
    
    def save_price_history(self):
        """Guarda el precio actual en el historial"""
        FinancialPriceHistory.objects.create(
            financial_action=self,
            price=self.current_price
        )


class FinancialPriceHistory(models.Model):
    """Modelo para almacenar el historial de precios de las acciones financieras"""
    
    financial_action = models.ForeignKey(
        FinancialAction,
        on_delete=models.CASCADE,
        related_name='price_history',
        verbose_name='Acción Financiera'
    )
    
    price = models.DecimalField(
        max_digits=12,
        decimal_places=4,
        null=True,
        blank=True,
        verbose_name='Precio'
    )
    
    recorded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Registro'
    )
    
    class Meta:
        ordering = ['-recorded_at']
        verbose_name = 'Historial de Precio'
        verbose_name_plural = 'Historiales de Precios'
        unique_together = ['financial_action', 'recorded_at']
    
    def __str__(self):
        return f"{self.financial_action.symbol} - {self.price} - {self.recorded_at}"


class Product(models.Model):
    """
    Modelo para gestionar productos
    """
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Producto'
    )
    
    price = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Precio',
        help_text='Precio del producto'
    )
    
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del producto'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de Actualización'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está marcado, el producto estará disponible'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_products',
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Producto'
        verbose_name_plural = 'Productos'
    
    def __str__(self):
        return self.name
    
    def get_currency_symbol(self):
        """Obtiene el símbolo de la moneda configurada en el sistema"""
        config = SystemConfiguration.objects.first()
        if not config:
            return '€'  # Por defecto EUR
        
        currency_symbols = {
            'EUR': '€',
            'USD': '$',
            'GBP': '£',
            'JPY': '¥',
            'CAD': 'C$',
            'AUD': 'A$',
            'CHF': 'CHF',
            'CNY': '¥',
            'MXN': '$',
            'COP': '$',
        }
        return currency_symbols.get(config.default_currency, '€')
    
    def get_currency_code(self):
        """Obtiene el código de la moneda configurada en el sistema"""
        config = SystemConfiguration.objects.first()
        if not config:
            return 'EUR'
        return config.default_currency
    
    def get_formatted_price(self):
        """Obtiene el precio formateado con la moneda del sistema"""
        symbol = self.get_currency_symbol()
        return f"{symbol}{self.price:,.2f}"
    
    @property
    def price_with_currency(self):
        """Propiedad para obtener el precio con moneda fácilmente"""
        return self.get_formatted_price()


class ClientProjectAccess(models.Model):
    """Modelo para configurar el acceso público a proyectos para imputación de horas por clientes"""
    
    project = models.OneToOneField(
        Project,
        on_delete=models.CASCADE,
        related_name='client_access',
        verbose_name='Proyecto'
    )
    public_token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token Público',
        help_text='Token único para acceso público al formulario de horas'
    )
    dashboard_token = models.CharField(
        max_length=64,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token Dashboard Público',
        help_text='Token único para acceso público al dashboard de estadísticas'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está desactivado, el enlace público no funcionará'
    )
    max_entries_per_day = models.PositiveIntegerField(
        default=50,
        verbose_name='Máximo de entradas por día',
        help_text='Límite de entradas de tiempo que se pueden crear por día'
    )
    requires_phone = models.BooleanField(
        default=True,
        verbose_name='Requiere teléfono',
        help_text='Si es obligatorio que el cliente proporcione su teléfono'
    )
    allowed_categories = models.JSONField(
        default=list,
        verbose_name='Categorías permitidas',
        help_text='Lista de categorías que el cliente puede seleccionar'
    )
    instructions = models.TextField(
        blank=True,
        verbose_name='Instrucciones',
        help_text='Instrucciones que verá el cliente en el formulario'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        verbose_name = 'Acceso Público de Cliente'
        verbose_name_plural = 'Accesos Públicos de Clientes'
    
    def __str__(self):
        return f"Acceso público para {self.project.name}"
    
    def save(self, *args, **kwargs):
        if not self.public_token:
            self.public_token = self.generate_token()
        if not self.dashboard_token:
            self.dashboard_token = self.generate_token()
        if not self.allowed_categories:
            self.allowed_categories = ['capacitacion', 'pruebas', 'uso']
        super().save(*args, **kwargs)
    
    def generate_token(self):
        """Genera un token único"""
        import secrets
        return secrets.token_urlsafe(32)
    
    def get_public_url(self):
        """Obtiene la URL pública para el formulario"""
        from django.urls import reverse
        return reverse('client_time_entry_form', kwargs={'token': self.public_token})
    
    def get_dashboard_url(self):
        """Obtiene la URL pública para el dashboard"""
        from django.urls import reverse
        return reverse('client_dashboard_public', kwargs={'token': self.dashboard_token})
    
    def can_accept_entries_today(self):
        """Verifica si se pueden aceptar más entradas hoy"""
        from django.utils import timezone
        today = timezone.now().date()
        today_entries = ClientTimeEntry.objects.filter(
            project=self.project,
            created_at__date=today
        ).count()
        return today_entries < self.max_entries_per_day


class ClientTimeEntry(models.Model):
    """Modelo para almacenar las entradas de tiempo de clientes"""
    
    CATEGORY_CHOICES = [
        ('capacitacion', 'Capacitación'),
        ('pruebas', 'Pruebas'),
        ('uso', 'Uso'),
    ]
    
    project = models.ForeignKey(
        Project,
        on_delete=models.CASCADE,
        related_name='client_time_entries',
        verbose_name='Proyecto'
    )
    client_name = models.CharField(
        max_length=100,
        verbose_name='Nombre del Cliente',
        help_text='Nombre completo del cliente'
    )
    client_email = models.EmailField(
        verbose_name='Correo del Cliente',
        help_text='Correo electrónico del cliente'
    )
    client_phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Teléfono del Cliente',
        help_text='Número de teléfono del cliente (opcional)'
    )
    hours = models.DecimalField(
        max_digits=5,
        decimal_places=2,
        verbose_name='Horas',
        help_text='Cantidad de horas a imputar (ej: 2.5)'
    )
    category = models.CharField(
        max_length=20,
        choices=CATEGORY_CHOICES,
        verbose_name='Categoría',
        help_text='Categoría de las horas imputadas'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional del trabajo realizado'
    )
    entry_date = models.DateField(
        default=timezone.now,
        verbose_name='Fecha de las horas',
        help_text='Fecha cuando se realizaron las horas de trabajo'
    )
    client_ip = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='IP del Cliente',
        help_text='Dirección IP desde donde se registró la entrada'
    )
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent',
        help_text='Información del navegador del cliente'
    )
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')
    
    class Meta:
        verbose_name = 'Entrada de Tiempo de Cliente'
        verbose_name_plural = 'Entradas de Tiempo de Clientes'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.client_name} ({self.client_email}) - {self.hours}h en {self.project.name} ({self.get_category_display()})"
    
    def get_formatted_hours(self):
        """Retorna las horas formateadas"""
        if self.hours == int(self.hours):
            return f"{int(self.hours)}h"
        return f"{self.hours}h"


class ShortUrl(models.Model):
    """Modelo para el acortador de URLs"""
    
    original_url = models.URLField(
        max_length=2000,
        verbose_name='URL Original',
        help_text='La URL larga que se quiere acortar'
    )
    short_code = models.CharField(
        max_length=10,
        unique=True,
        verbose_name='Código Corto',
        help_text='Código único para la URL corta'
    )
    title = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Título',
        help_text='Título descriptivo opcional para la URL'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional de la URL'
    )
    clicks = models.PositiveIntegerField(
        default=0,
        verbose_name='Clics',
        help_text='Número de veces que se ha accedido a la URL corta'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si la URL corta está activa o deshabilitada'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        help_text='Usuario que creó la URL corta'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración',
        help_text='Fecha opcional cuando expira la URL corta'
    )
    
    class Meta:
        verbose_name = 'URL Corta'
        verbose_name_plural = 'URLs Cortas'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.short_code} → {self.original_url[:50]}..."
    
    def get_short_url(self, request=None):
        """Retorna la URL corta completa"""
        from django.conf import settings
        
        # Si se proporciona request, usar su dominio
        if request:
            domain = request.get_host()
            protocol = 'https' if request.is_secure() else 'http'
        else:
            # Usar configuración del settings
            domain = getattr(settings, 'SITE_DOMAIN', 'ticketproo.com')
            protocol = 'https' if not settings.DEBUG else 'http'
        
        return f"{protocol}://{domain}/s/{self.short_code}"
    
    def is_expired(self):
        """Verifica si la URL ha expirado"""
        if self.expires_at:
            return timezone.now() > self.expires_at
        return False
    
    def increment_clicks(self):
        """Incrementa el contador de clics"""
        self.clicks += 1
        self.save(update_fields=['clicks'])
    
    @staticmethod
    def generate_short_code():
        """Genera un código corto único"""
        import random
        import string
        
        chars = string.ascii_letters + string.digits
        while True:
            code = ''.join(random.choice(chars) for _ in range(6))
            if not ShortUrl.objects.filter(short_code=code).exists():
                return code


class ShortUrlClick(models.Model):
    """Modelo para rastrear cada clic en una URL corta"""
    
    short_url = models.ForeignKey(
        ShortUrl,
        on_delete=models.CASCADE,
        related_name='click_records',
        verbose_name='URL Corta'
    )
    clicked_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha del clic',
        db_index=True
    )
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País',
        db_index=True
    )
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad'
    )
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    referer = models.URLField(
        max_length=500,
        blank=True,
        verbose_name='Referer'
    )
    
    class Meta:
        verbose_name = 'Clic de URL Corta'
        verbose_name_plural = 'Clics de URLs Cortas'
        ordering = ['-clicked_at']
        indexes = [
            models.Index(fields=['short_url', '-clicked_at']),
            models.Index(fields=['country', '-clicked_at']),
        ]
    
    def __str__(self):
        return f"Clic en {self.short_url.short_code} - {self.clicked_at}"


class ClientRequest(models.Model):
    """Modelo para solicitudes al cliente"""
    
    sequence = models.CharField(
        max_length=50,
        unique=True,
        verbose_name='Secuencia',
        help_text='Número de secuencia único de la solicitud'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción de lo que se solicita'
    )
    requested_to = models.CharField(
        max_length=200,
        verbose_name='Solicitado a',
        help_text='Nombre de la persona a quien se solicita'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='client_requests_created',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        verbose_name='Empresa',
        help_text='Empresa asociada a la solicitud'
    )
    status = models.CharField(
        max_length=20,
        choices=[
            ('pending', 'Pendiente'),
            ('in_progress', 'En Progreso'),
            ('completed', 'Completada'),
            ('cancelled', 'Cancelada')
        ],
        default='pending',
        verbose_name='Estado'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    attachment = models.FileField(
        upload_to='client_request_attachments/',
        null=True,
        blank=True,
        verbose_name='Archivo adjunto',
        help_text='Documento o archivo de referencia para el cliente'
    )
    
    class Meta:
        verbose_name = 'Solicitud al Cliente'
        verbose_name_plural = 'Solicitudes a Clientes'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['sequence']),
            models.Index(fields=['status', '-created_at']),
            models.Index(fields=['company', '-created_at']),
        ]
    
    def __str__(self):
        return f"{self.sequence} - {self.title}"
    
    def save(self, *args, **kwargs):
        if not self.sequence and self.company:
            # Generar secuencia automática por empresa: REQ-{company_id}-{número}
            last_request = ClientRequest.objects.filter(company=self.company).order_by('-id').first()
            if last_request and last_request.sequence:
                try:
                    # Extraer el último número de la secuencia REQ-XX-YYYY
                    parts = last_request.sequence.split('-')
                    last_num = int(parts[-1])
                    self.sequence = f"REQ-{self.company.id:02d}-{last_num + 1:04d}"
                except:
                    self.sequence = f"REQ-{self.company.id:02d}-0001"
            else:
                self.sequence = f"REQ-{self.company.id:02d}-0001"
        elif not self.sequence:
            # Fallback si no hay empresa
            self.sequence = f"REQ-00-0001"
        super().save(*args, **kwargs)


class ClientRequestResponse(models.Model):
    """Modelo para respuestas del cliente a las solicitudes"""
    
    request = models.OneToOneField(
        ClientRequest,
        on_delete=models.CASCADE,
        related_name='response',
        verbose_name='Solicitud'
    )
    response_text = models.TextField(
        verbose_name='Respuesta',
        help_text='Texto de la respuesta del cliente'
    )
    attachment = models.FileField(
        upload_to='client_request_attachments/%Y/%m/',
        null=True,
        blank=True,
        verbose_name='Adjunto'
    )
    responded_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Respondido por'
    )
    responded_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de respuesta'
    )
    
    class Meta:
        verbose_name = 'Respuesta de Solicitud'
        verbose_name_plural = 'Respuestas de Solicitudes'
        ordering = ['-responded_at']
    
    def __str__(self):
        return f"Respuesta a {self.request.sequence}"


class ClientRequestTemplate(models.Model):
    """Modelo para plantillas de solicitudes al cliente"""
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Plantilla',
        help_text='Nombre descriptivo para identificar este conjunto de solicitudes'
    )
    description = models.TextField(
        verbose_name='Descripción',
        blank=True,
        help_text='Descripción de para qué sirve esta plantilla'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='client_request_templates',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si está activa, se puede usar para crear solicitudes'
    )
    
    class Meta:
        verbose_name = 'Plantilla de Solicitud'
        verbose_name_plural = 'Plantillas de Solicitudes'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.name


class ClientRequestTemplateItem(models.Model):
    """Modelo para ítems individuales dentro de una plantilla"""
    
    template = models.ForeignKey(
        ClientRequestTemplate,
        on_delete=models.CASCADE,
        related_name='items',
        verbose_name='Plantilla'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    description = models.TextField(
        verbose_name='Descripción'
    )
    requested_to = models.CharField(
        max_length=200,
        verbose_name='Solicitado a',
        help_text='Nombre de la persona a quien se solicita'
    )
    order = models.IntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en la lista'
    )
    attachment = models.FileField(
        upload_to='client_request_template_attachments/',
        null=True,
        blank=True,
        verbose_name='Archivo adjunto',
        help_text='Documento de referencia para esta solicitud'
    )
    
    class Meta:
        verbose_name = 'Item de Plantilla'
        verbose_name_plural = 'Items de Plantilla'
        ordering = ['order', 'id']
    
    def __str__(self):
        return f"{self.template.name} - {self.title}"


class ProductSet(models.Model):
    """Modelo para conjuntos de productos"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre del conjunto de productos'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del conjunto de productos'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Si está marcado, los clientes podrán ver este conjunto'
    )
    public_token = models.CharField(
        max_length=32,
        unique=True,
        blank=True,
        verbose_name='Token Público',
        help_text='Token único para acceso público'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        help_text='Usuario que creó el conjunto'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Campos de estadísticas
    view_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Vistas',
        help_text='Número de veces que se ha visto el conjunto público'
    )
    last_viewed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última vista',
        help_text='Fecha y hora de la última vista pública'
    )
    
    class Meta:
        verbose_name = 'Conjunto de Productos'
        verbose_name_plural = 'Conjuntos de Productos'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        # Generar token público si no existe
        if not self.public_token:
            import uuid
            self.public_token = str(uuid.uuid4()).replace('-', '')[:16]
        super().save(*args, **kwargs)
    
    def get_public_url(self):
        """Retorna la URL pública del conjunto"""
        from django.urls import reverse
        return reverse('product_set_public', kwargs={'token': self.public_token})
    
    def increment_views(self):
        """Incrementa el contador de vistas y actualiza la fecha de última vista"""
        from django.utils import timezone
        self.view_count += 1
        self.last_viewed_at = timezone.now()
        self.save(update_fields=['view_count', 'last_viewed_at'])
    
    @property
    def product_count(self):
        """Retorna el número de productos en el conjunto"""
        return self.products.count()


class ProductItem(models.Model):
    """Modelo para productos individuales dentro de un conjunto"""
    
    product_set = models.ForeignKey(
        ProductSet,
        on_delete=models.CASCADE,
        related_name='products',
        verbose_name='Conjunto de Productos'
    )
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Producto',
        help_text='Nombre descriptivo del producto'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del producto'
    )
    video_link = models.URLField(
        blank=True,
        verbose_name='Link del Video',
        help_text='URL del video demostrativo (opcional)'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en la lista'
    )
    price = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0.00,
        verbose_name='Precio',
        help_text='Precio del producto individual'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Producto'
        verbose_name_plural = 'Productos'
        ordering = ['order', 'name']
    
    def __str__(self):
        return f"{self.product_set.title} - {self.name}"
    
    @property
    def product_code(self):
        """Genera el código único del producto"""
        return f"PRO{str(self.id).zfill(3)}"
    
    def get_formatted_price(self):
        """Retorna el precio formateado con la moneda del sistema"""
        if self.price > 0:
            try:
                config = SystemConfiguration.objects.get(pk=1)
                return config.format_currency(self.price)
            except SystemConfiguration.DoesNotExist:
                return f"€{self.price:,.2f}"  # Fallback
        return None


class Precotizador(models.Model):
    """Modelo para el sistema de precotizador con IA"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Precotizador',
        help_text='Nombre descriptivo del precotizador'
    )
    client_description = models.TextField(
        verbose_name='Descripción del Cliente',
        help_text='Descripción detallada del tipo de cliente y sus necesidades'
    )
    hourly_rate = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Precio por Hora',
        help_text='Tu tarifa por hora para este tipo de trabajo'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        help_text='Usuario que creó el precotizador'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Campos para enlace público de solicitud
    public_request_token = models.CharField(
        max_length=255,
        blank=True,
        null=True,
        unique=True,
        verbose_name='Token de Solicitud Pública',
        help_text='Token único para el enlace público donde clientes pueden solicitar cotizaciones'
    )
    allow_public_requests = models.BooleanField(
        default=False,
        verbose_name='Permitir Solicitudes Públicas',
        help_text='Permite que los clientes soliciten cotizaciones a través de un enlace público'
    )
    
    class Meta:
        verbose_name = 'Precotizador'
        verbose_name_plural = 'Precotizadores'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.title
    
    def get_formatted_hourly_rate(self):
        """Retorna la tarifa por hora formateada con la moneda del sistema"""
        try:
            config = SystemConfiguration.objects.get(pk=1)
            return config.format_currency(self.hourly_rate)
        except SystemConfiguration.DoesNotExist:
            return f"€{self.hourly_rate:,.2f}"  # Fallback
    
    def generate_public_request_token(self):
        """Genera un token único para solicitudes públicas"""
        import uuid
        if not self.public_request_token:
            self.public_request_token = str(uuid.uuid4())
            self.save(update_fields=['public_request_token'])
        return self.public_request_token
    
    def get_public_request_url(self):
        """Retorna la URL pública para solicitar cotizaciones"""
        from django.urls import reverse
        if self.public_request_token:
            return reverse('precotizador_public_request', kwargs={'token': self.public_request_token})
        return None
    
    def toggle_public_requests(self):
        """Activa/desactiva las solicitudes públicas"""
        self.allow_public_requests = not self.allow_public_requests
        if self.allow_public_requests and not self.public_request_token:
            self.generate_public_request_token()
        self.save(update_fields=['allow_public_requests', 'public_request_token'])
        return self.allow_public_requests


class PrecotizadorExample(models.Model):
    """Modelo para los ejemplos de casos del precotizador"""
    
    precotizador = models.ForeignKey(
        Precotizador,
        on_delete=models.CASCADE,
        related_name='examples',
        verbose_name='Precotizador'
    )
    description = models.TextField(
        verbose_name='Descripción del Caso',
        help_text='Descripción detallada del trabajo a realizar'
    )
    estimated_hours = models.DecimalField(
        max_digits=8,
        decimal_places=2,
        verbose_name='Horas Estimadas',
        help_text='Número de horas estimadas para completar el trabajo'
    )
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en la lista'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Ejemplo de Caso'
        verbose_name_plural = 'Ejemplos de Casos'
        ordering = ['order', 'created_at']
    
    def __str__(self):
        return f"{self.precotizador.title} - Caso {self.order + 1}"
    
    def get_estimated_cost(self):
        """Calcula el costo estimado basado en las horas y tarifa del precotizador"""
        return self.estimated_hours * self.precotizador.hourly_rate
    
    def get_formatted_cost(self):
        """Retorna el costo estimado formateado con la moneda del sistema"""
        cost = self.get_estimated_cost()
        try:
            config = SystemConfiguration.objects.get(pk=1)
            return config.format_currency(cost)
        except SystemConfiguration.DoesNotExist:
            return f"€{cost:,.2f}"  # Fallback


class PrecotizadorQuote(models.Model):
    """Modelo para las cotizaciones generadas por IA"""
    
    precotizador = models.ForeignKey(
        Precotizador,
        on_delete=models.CASCADE,
        related_name='quotes',
        verbose_name='Precotizador'
    )
    client_request = models.TextField(
        verbose_name='Solicitud del Cliente',
        help_text='Descripción del trabajo solicitado por el cliente'
    )
    ai_estimated_hours = models.DecimalField(
        max_digits=8,
        decimal_places=2,
        verbose_name='Horas Estimadas por IA',
        help_text='Número de horas estimadas por la IA'
    )
    ai_detailed_description = models.TextField(
        verbose_name='Descripción Detallada por IA',
        help_text='Descripción detallada del trabajo generada por la IA'
    )
    ai_reasoning = models.TextField(
        blank=True,
        verbose_name='Razonamiento de la IA',
        help_text='Explicación del razonamiento de la IA para la estimación'
    )
    ai_full_response = models.TextField(
        blank=True,
        verbose_name='Respuesta Completa de la IA',
        help_text='Respuesta completa de la IA en formato texto plano'
    )
    
    # Campos para compartir públicamente
    public_share_token = models.CharField(
        max_length=100,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token de Compartir Público',
        help_text='Token único para compartir la cotización públicamente'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Es Público',
        help_text='Si la cotización puede ser vista públicamente'
    )
    
    # Estado de la cotización
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('accepted', 'Aceptada por el Cliente'),
        ('rejected', 'Rechazada por el Cliente'),
        ('expired', 'Expirada'),
    ]
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado',
        help_text='Estado actual de la cotización'
    )
    
    # Información del cliente
    client_response_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de Respuesta del Cliente',
        help_text='Fecha en que el cliente respondió'
    )
    client_comments = models.TextField(
        blank=True,
        verbose_name='Comentarios del Cliente',
        help_text='Comentarios adicionales del cliente'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Cotización IA'
        verbose_name_plural = 'Cotizaciones IA'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"Cotización {self.id} - {self.precotizador.title}"
    
    def get_formatted_cost(self):
        """Retorna el costo formateado con la moneda del sistema"""
        from decimal import Decimal
        # Convertir a Decimal para evitar errores de tipos
        hours = Decimal(str(self.ai_estimated_hours))
        total_cost = hours * self.precotizador.hourly_rate
        try:
            config = SystemConfiguration.objects.get(pk=1)
            return config.format_currency(total_cost)
        except SystemConfiguration.DoesNotExist:
            return f"€{total_cost:,.2f}"  # Fallback
    
    def generate_public_token(self):
        """Genera un token único para compartir públicamente"""
        import uuid
        if not self.public_share_token:
            self.public_share_token = str(uuid.uuid4())
            self.save(update_fields=['public_share_token'])
        return self.public_share_token
    
    def get_public_url(self):
        """Retorna la URL pública para compartir"""
        from django.urls import reverse
        if self.public_share_token:
            return reverse('precotizador_quote_public', kwargs={'token': self.public_share_token})
        return None
    
    def can_be_responded(self):
        """Verifica si la cotización puede ser respondida por el cliente"""
        return self.status == 'pending' and self.is_public
    
    def accept(self, client_comments=''):
        """Acepta la cotización"""
        if self.status != 'pending':
            raise ValueError(f"No se puede aceptar una cotización que ya está {self.get_status_display().lower()}")
        
        self.status = 'accepted'
        self.client_comments = client_comments
        self.client_response_date = timezone.now()
        self.save(update_fields=['status', 'client_comments', 'client_response_date'])
    
    def reject(self, client_comments=''):
        """Rechaza la cotización"""
        if self.status != 'pending':
            raise ValueError(f"No se puede rechazar una cotización que ya está {self.get_status_display().lower()}")
        
        self.status = 'rejected'
        self.client_comments = client_comments
        self.client_response_date = timezone.now()
        self.save(update_fields=['status', 'client_comments', 'client_response_date'])
    
    def get_status_color(self):
        """Retorna el color Bootstrap para el estado"""
        colors = {
            'pending': 'warning',
            'accepted': 'success',
            'rejected': 'danger',
            'expired': 'secondary',
        }
        return colors.get(self.status, 'primary')
    
    def get_response_time_info(self):
        """Retorna información sobre el tiempo de respuesta del cliente"""
        if not self.client_response_date:
            return None
        
        from django.utils import timezone
        time_diff = self.client_response_date - self.created_at
        
        if time_diff.days > 0:
            return f"Respondió en {time_diff.days} día{'s' if time_diff.days > 1 else ''}"
        elif time_diff.seconds > 3600:
            hours = time_diff.seconds // 3600
            return f"Respondió en {hours} hora{'s' if hours > 1 else ''}"
        else:
            minutes = time_diff.seconds // 60
            return f"Respondió en {minutes} minuto{'s' if minutes > 1 else ''}"
    
    def is_recently_responded(self):
        """Verifica si el cliente respondió recientemente (últimas 24 horas)"""
        if not self.client_response_date:
            return False
        
        from django.utils import timezone
        from datetime import timedelta
        
        return timezone.now() - self.client_response_date <= timedelta(hours=24)
    
    def get_estimated_cost(self):
        """Calcula el costo estimado basado en las horas y tarifa del precotizador"""
        return self.ai_estimated_hours * self.precotizador.hourly_rate
    
    def get_formatted_cost(self):
        """Retorna el costo estimado formateado con la moneda del sistema"""
        cost = self.get_estimated_cost()
        try:
            config = SystemConfiguration.objects.get(pk=1)
            return config.format_currency(cost)
        except SystemConfiguration.DoesNotExist:
            return f"€{cost:,.2f}"  # Fallback


class CompanyDocumentation(models.Model):
    """Documentación múltiple para empresas con URLs y contraseñas"""
    
    title = models.CharField(max_length=200, verbose_name="Título")
    description = models.TextField(blank=True, verbose_name="Descripción")
    company = models.ForeignKey(
        Company, 
        on_delete=models.CASCADE,
        related_name='documentations',
        verbose_name="Empresa"
    )
    is_public = models.BooleanField(default=False, verbose_name="Es público")
    public_token = models.UUIDField(
        default=uuid.uuid4, 
        editable=False, 
        unique=True,
        verbose_name="Token público"
    )
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE,
        verbose_name="Creado por"
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Fecha de actualización")
    
    class Meta:
        verbose_name = "Documentación de Empresa"
        verbose_name_plural = "Documentaciones de Empresas"
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.title} - {self.company.name}"
    
    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('company_documentation_detail', kwargs={'pk': self.pk})
    
    def get_public_url(self):
        from django.urls import reverse
        return reverse('company_documentation_public', kwargs={'token': self.public_token})


class CompanyDocumentationURL(models.Model):
    """URLs asociadas a documentaciones de empresas"""
    
    documentation = models.ForeignKey(
        CompanyDocumentation,
        on_delete=models.CASCADE,
        related_name='urls',
        verbose_name="Documentación"
    )
    title = models.CharField(max_length=200, verbose_name="Título del enlace")
    url = models.URLField(verbose_name="URL")
    description = models.TextField(blank=True, verbose_name="Descripción")
    username = models.CharField(max_length=100, blank=True, verbose_name="Usuario")
    password = models.CharField(max_length=200, blank=True, verbose_name="Contraseña")
    notes = models.TextField(blank=True, verbose_name="Notas adicionales")
    is_active = models.BooleanField(default=True, verbose_name="Está activo")
    order = models.PositiveIntegerField(default=0, verbose_name="Orden")
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de creación")
    
    class Meta:
        verbose_name = "URL de Documentación"
        verbose_name_plural = "URLs de Documentación"
        ordering = ['order', 'title']
    
    def __str__(self):
        return f"{self.title} - {self.documentation.title}"
    
    def has_credentials(self):
        """Verifica si tiene credenciales (usuario y/o contraseña)"""
        return bool(self.username or self.password)


class TermsOfUse(models.Model):
    """Condiciones de uso de la aplicación"""
    
    title = models.CharField(max_length=200, verbose_name="Título")
    content = models.TextField(verbose_name="Contenido")
    effective_date = models.DateField(verbose_name="Fecha de vigencia")
    is_active = models.BooleanField(default=True, verbose_name="Está activo")
    version = models.CharField(max_length=50, blank=True, verbose_name="Versión")
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='terms_created',
        verbose_name="Creado por"
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Última actualización")
    
    class Meta:
        verbose_name = "Condición de Uso"
        verbose_name_plural = "Condiciones de Uso"
        ordering = ['-effective_date', '-created_at']
    
    def __str__(self):
        return f"{self.title} - {self.effective_date}"
    
    def save(self, *args, **kwargs):
        # Si se marca como activo, desactivar las demás
        if self.is_active:
            TermsOfUse.objects.filter(is_active=True).exclude(pk=self.pk).update(is_active=False)
        super().save(*args, **kwargs)


class ContactGenerator(models.Model):
    """Generador de contactos con URL pública"""
    
    title = models.CharField(max_length=200, verbose_name="Título")
    description = models.TextField(blank=True, verbose_name="Descripción")
    company = models.ForeignKey(
        Company, 
        on_delete=models.CASCADE,
        related_name='contact_generators',
        verbose_name="Empresa"
    )
    is_active = models.BooleanField(default=True, verbose_name="Está activo")
    public_token = models.UUIDField(
        default=uuid.uuid4, 
        editable=False, 
        unique=True,
        verbose_name="Token público"
    )
    
    # Campos personalizables del formulario
    collect_phone = models.BooleanField(default=False, verbose_name="Recopilar teléfono")
    collect_company = models.BooleanField(default=True, verbose_name="Recopilar empresa")
    collect_position = models.BooleanField(default=True, verbose_name="Recopilar cargo")
    collect_notes = models.BooleanField(default=False, verbose_name="Recopilar notas")
    
    # Mensajes personalizables
    welcome_message = models.TextField(
        default="¡Gracias por tu interés! Por favor, completa la siguiente información para que podamos contactarte.",
        verbose_name="Mensaje de bienvenida"
    )
    success_message = models.TextField(
        default="¡Gracias! Hemos recibido tu información y nos pondremos en contacto contigo pronto.",
        verbose_name="Mensaje de éxito"
    )
    
    # Configuración de diseño
    background_color = models.CharField(
        max_length=7, 
        default="#007bff", 
        verbose_name="Color de fondo",
        help_text="Color en formato hexadecimal (ej: #007bff)"
    )
    
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE,
        verbose_name="Creado por"
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Fecha de actualización")
    
    class Meta:
        verbose_name = "Generador de Contactos"
        verbose_name_plural = "Generadores de Contactos"
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.title} - {self.company.name}"
    
    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('contact_generator_detail', kwargs={'pk': self.pk})
    
    def get_public_url(self):
        from django.urls import reverse
        return reverse('contact_generator_public', kwargs={'token': self.public_token})
    
    def get_contacts_count(self):
        """Retorna el número de contactos generados por este generador"""
        return Contact.objects.filter(
            source=f"Generador: {self.title}",
            created_by=self.created_by
        ).count()


class CompanyRequestGenerator(models.Model):
    """Generador de formularios de Solicitudes de Empresas"""

    title = models.CharField(max_length=200, verbose_name="Título")
    description = models.TextField(blank=True, verbose_name="Descripción")
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='company_request_generators',
        verbose_name="Empresa"
    )
    is_active = models.BooleanField(default=True, verbose_name="Está activo")
    public_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name="Token público"
    )

    # Secuencia y prefijo
    sequence_prefix = models.CharField(
        max_length=20,
        blank=True,
        default='',
        verbose_name='Prefijo de Secuencia',
        help_text='Prefijo que se antepone al número de solicitud (ej: SR-)'
    )
    next_sequence = models.PositiveIntegerField(
        default=1,
        verbose_name='Siguiente secuencia',
        help_text='Número secuencial que se incrementa por cada solicitud'
    )

    # Campos personalizables del formulario
    collect_date = models.BooleanField(default=True, verbose_name="Incluir fecha")
    collect_text = models.BooleanField(default=True, verbose_name="Incluir texto de solicitud")
    collect_url = models.BooleanField(default=False, verbose_name="Incluir URL relacionada")

    # Mensajes personalizables
    welcome_message = models.TextField(
        default="Por favor, envía tu solicitud y nuestro equipo la revisará.",
        verbose_name="Mensaje de bienvenida"
    )
    success_message = models.TextField(
        default="Solicitud recibida. Te contactaremos pronto.",
        verbose_name="Mensaje de éxito"
    )

    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name="Creado por"
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Fecha de actualización")

    class Meta:
        verbose_name = "Generador de Solicitudes Empresa"
        verbose_name_plural = "Generadores de Solicitudes Empresas"
        ordering = ['-created_at']

    def __str__(self):
        return f"{self.title} - {self.company.name}"

    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('company_request_generator_detail', kwargs={'pk': self.pk})

    def get_public_url(self):
        from django.urls import reverse
        return reverse('company_request_public', kwargs={'token': self.public_token})

    def generate_sequence(self):
        """Retorna la secuencia completa y la incrementa de forma atómica"""
        from django.db import transaction
        with transaction.atomic():
            seq = self.next_sequence
            full = f"{self.sequence_prefix}{seq:04d}" if self.sequence_prefix else f"{seq:04d}"
            self.next_sequence = seq + 1
            self.save(update_fields=['next_sequence'])
            return full

    def get_requests_count(self):
        return CompanyRequest.objects.filter(generator=self).count()

    def get_recent_requests(self, limit=10):
        return CompanyRequest.objects.filter(generator=self).order_by('-created_at')[:limit]


class CompanyRequest(models.Model):
    """Registro de una solicitud enviada mediante un generador"""

    STATUS_CHOICES = [
        ('new', 'Nueva'),
        ('review', 'En revisión'),
        ('accepted', 'Aceptada'),
        ('executed', 'Ejecutada'),
        ('rejected', 'Rechazada'),
        ('closed', 'Cerrada'),
    ]

    sequence = models.CharField(max_length=50, verbose_name='Número de solicitud', unique=True)
    generator = models.ForeignKey(
        CompanyRequestGenerator,
        on_delete=models.CASCADE,
        related_name='requests',
        verbose_name='Generador'
    )
    request_date = models.DateTimeField(default=timezone.now, verbose_name='Fecha de solicitud')
    text = models.TextField(blank=True, verbose_name='Texto de la solicitud')
    url = models.URLField(blank=True, verbose_name='URL relacionada')
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='new', verbose_name='Estado')
    executed_at = models.DateTimeField(null=True, blank=True, verbose_name='Fecha de ejecución')
    executed_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='executed_company_requests',
        verbose_name='Ejecutado por'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')

    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Solicitud de Empresa'
        verbose_name_plural = 'Solicitudes de Empresas'

    def __str__(self):
        return f"{self.sequence} - {self.generator.title}"

    def get_identifier(self):
        return self.sequence


class CompanyRequestComment(models.Model):
    """Comentarios internos para una CompanyRequest"""

    request = models.ForeignKey(
        CompanyRequest,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Solicitud'
    )
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Usuario')
    content = models.TextField(verbose_name='Comentario')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de creación')

    class Meta:
        ordering = ['created_at']
        verbose_name = 'Comentario de Solicitud'
        verbose_name_plural = 'Comentarios de Solicitudes'

    def __str__(self):
        return f"Comentario de {self.user.username} en {self.request.sequence}"


# ============= SISTEMA DE FORMULARIOS =============

class Form(models.Model):
    """Formulario con múltiples preguntas"""
    
    title = models.CharField(max_length=200, verbose_name='Título del formulario')
    description = models.TextField(blank=True, verbose_name='Descripción')
    theme = models.CharField(
        max_length=20,
        choices=[
            ('blue', 'Azul Clásico'),
            ('green', 'Verde Moderno'),
            ('sunset', 'Sunset'),
            ('ocean', 'Océano'),
            ('purple', 'Púrpura'),
            ('pink', 'Rosa'),
            ('minimal', 'Minimal'),
        ],
        default='blue',
        verbose_name='Tema visual'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='forms',
        verbose_name='Empresa'
    )
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    public_token = models.UUIDField(default=uuid.uuid4, editable=False, unique=True)
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='forms_created',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última actualización')

    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Formulario'
        verbose_name_plural = 'Formularios'

    def __str__(self):
        return f"{self.title} - {self.company.name}"

    def get_public_url(self):
        """Retorna la URL pública del formulario"""
        return f"/form/{self.public_token}/"

    def get_total_responses(self):
        """Retorna el total de respuestas del formulario"""
        return self.responses.count()

    def get_average_score(self):
        """Retorna el promedio de puntuación del formulario"""
        responses = self.responses.all()
        if not responses:
            return 0
        
        total_score = sum(response.total_score for response in responses)
        return round(total_score / len(responses), 2)


class FormQuestion(models.Model):
    """Pregunta de un formulario"""
    
    QUESTION_TYPES = [
        ('number', 'Cantidad (Número)'),
        ('multiple_choice', 'Selección múltiple'),
        ('text', 'Texto libre'),
    ]
    
    form = models.ForeignKey(
        Form,
        on_delete=models.CASCADE,
        related_name='questions',
        verbose_name='Formulario'
    )
    question_text = models.TextField(verbose_name='Texto de la pregunta')
    question_type = models.CharField(
        max_length=20,
        choices=QUESTION_TYPES,
        verbose_name='Tipo de pregunta'
    )
    is_required = models.BooleanField(default=True, verbose_name='Obligatoria')
    order = models.PositiveIntegerField(default=0, verbose_name='Orden')
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')

    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Pregunta'
        verbose_name_plural = 'Preguntas'

    def __str__(self):
        return f"{self.form.title} - {self.question_text[:50]}..."


class FormQuestionOption(models.Model):
    """Opciones para preguntas de selección múltiple"""
    
    question = models.ForeignKey(
        FormQuestion,
        on_delete=models.CASCADE,
        related_name='options',
        verbose_name='Pregunta'
    )
    option_text = models.CharField(max_length=200, verbose_name='Texto de la opción')
    score = models.IntegerField(default=0, verbose_name='Puntuación')
    order = models.PositiveIntegerField(default=0, verbose_name='Orden')

    class Meta:
        ordering = ['order']
        verbose_name = 'Opción de Pregunta'
        verbose_name_plural = 'Opciones de Preguntas'

    def __str__(self):
        return f"{self.option_text} ({self.score} pts)"


class FormResponse(models.Model):
    """Respuesta completa a un formulario"""
    
    form = models.ForeignKey(
        Form,
        on_delete=models.CASCADE,
        related_name='responses',
        verbose_name='Formulario'
    )
    respondent_name = models.CharField(max_length=100, blank=True, verbose_name='Nombre del encuestado')
    respondent_email = models.EmailField(blank=True, verbose_name='Email del encuestado')
    total_score = models.IntegerField(default=0, verbose_name='Puntuación total')
    response_date = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de respuesta')

    class Meta:
        ordering = ['-response_date']
        verbose_name = 'Respuesta al Formulario'
        verbose_name_plural = 'Respuestas a Formularios'

    def __str__(self):
        return f"Respuesta de {self.respondent_name or 'Anónimo'} - {self.form.title}"

    def calculate_total_score(self):
        """Calcula y actualiza la puntuación total"""
        total = 0
        for answer in self.answers.all():
            if answer.selected_option:
                total += answer.selected_option.score
        self.total_score = total
        self.save()
        return total


class FormAnswer(models.Model):
    """Respuesta individual a una pregunta"""
    
    response = models.ForeignKey(
        FormResponse,
        on_delete=models.CASCADE,
        related_name='answers',
        verbose_name='Respuesta al formulario'
    )
    question = models.ForeignKey(
        FormQuestion,
        on_delete=models.CASCADE,
        verbose_name='Pregunta'
    )
    # Para preguntas de texto
    text_answer = models.TextField(blank=True, verbose_name='Respuesta de texto')
    # Para preguntas de número
    number_answer = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Respuesta numérica'
    )
    # Para preguntas de selección múltiple
    selected_option = models.ForeignKey(
        FormQuestionOption,
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        verbose_name='Opción seleccionada'
    )

    class Meta:
        unique_together = ('response', 'question')
        verbose_name = 'Respuesta Individual'
        verbose_name_plural = 'Respuestas Individuales'

    def __str__(self):
        return f"Respuesta a: {self.question.question_text[:30]}..."

    def get_answer_display(self):
        """Retorna la respuesta en formato legible"""
        if self.question.question_type == 'text':
            return self.text_answer
        elif self.question.question_type == 'number':
            return str(self.number_answer) if self.number_answer is not None else ''
        elif self.question.question_type == 'multiple_choice':
            return self.selected_option.option_text if self.selected_option else ''
        return ''


class FormAIAnalysis(models.Model):
    """Análisis de IA para formularios"""
    
    form = models.ForeignKey(
        Form,
        on_delete=models.CASCADE,
        related_name='ai_analyses',
        verbose_name='Formulario'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de análisis'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Analizado por'
    )
    
    # Campos del análisis
    overall_score = models.DecimalField(
        max_digits=3,
        decimal_places=1,
        verbose_name='Puntuación general'
    )
    problems = models.JSONField(
        default=list,
        verbose_name='Problemas identificados',
        help_text='Lista de problemas encontrados en formato JSON'
    )
    improvements = models.JSONField(
        default=list,
        verbose_name='Mejoras sugeridas',
        help_text='Lista de mejoras sugeridas en formato JSON'
    )
    strengths = models.JSONField(
        default=list,
        verbose_name='Fortalezas',
        help_text='Lista de fortalezas identificadas'
    )
    main_recommendation = models.TextField(
        verbose_name='Recomendación principal'
    )
    
    # Metadatos del análisis
    total_responses = models.IntegerField(
        verbose_name='Total de respuestas analizadas'
    )
    analysis_version = models.CharField(
        max_length=10,
        default='1.0',
        verbose_name='Versión del análisis'
    )
    
    class Meta:
        verbose_name = 'Análisis de IA'
        verbose_name_plural = 'Análisis de IA'
        ordering = ['-created_at']
        
    def __str__(self):
        return f"Análisis de {self.form.title} - {self.created_at.strftime('%d/%m/%Y %H:%M')}"
    
    def get_problems_count(self):
        """Retorna el número de problemas identificados"""
        return len(self.problems) if self.problems else 0
    
    def get_improvements_count(self):
        """Retorna el número de mejoras sugeridas"""
        return len(self.improvements) if self.improvements else 0
    
    def get_score_color(self):
        """Retorna el color basado en la puntuación"""
        if self.overall_score >= 8:
            return 'success'
        elif self.overall_score >= 6:
            return 'warning'
        else:
            return 'danger'


# Modelo de Alcance
class Alcance(models.Model):
    titulo = models.CharField(max_length=200, verbose_name='Título')
    categoria = models.CharField(max_length=100, verbose_name='Categoría')
    descripcion = models.TextField(verbose_name='Descripción')
    url = models.URLField(max_length=500, blank=True, null=True, verbose_name='URL del artículo relacionado')
    publico = models.BooleanField(default=True, verbose_name='¿Es público?')
    creado_por = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Creado por'
    )
    creado_en = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    actualizado_en = models.DateTimeField(auto_now=True, verbose_name='Última actualización')

    class Meta:
        verbose_name = 'Alcance'
        verbose_name_plural = 'Alcances'
        ordering = ['-creado_en']

    def __str__(self):
        return self.titulo


# ==================== MODELO DE CONTROL DE LICENCIAS ====================

class License(models.Model):
    """Modelo para gestionar licencias de productos"""
    
    STATUS_CHOICES = [
        ('active', 'Activa'),
        ('expired', 'Expirada'),
        ('suspended', 'Suspendida'),
        ('revoked', 'Revocada'),
    ]
    
    license_key = models.CharField(
        max_length=12,
        unique=True,
        verbose_name='Clave de Licencia',
        help_text='Clave alfanumérica de 12 caracteres en mayúsculas'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name='Empresa',
        related_name='licenses'
    )
    product = models.CharField(
        max_length=200,
        verbose_name='Producto'
    )
    start_date = models.DateField(
        verbose_name='Fecha de Inicio'
    )
    end_date = models.DateField(
        verbose_name='Fecha de Vencimiento'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='active',
        verbose_name='Estado'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre la licencia'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Creado por',
        related_name='created_licenses'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # UUID para acceso público
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='UUID Público'
    )
    
    class Meta:
        verbose_name = 'Licencia'
        verbose_name_plural = 'Licencias'
        ordering = ['-created_at']
        
    def __str__(self):
        return f"{self.license_key} - {self.product} ({self.company.name})"
    
    def save(self, *args, **kwargs):
        # Generar clave automáticamente si no existe
        if not self.license_key:
            self.license_key = self.generate_license_key()
        # Asegurar que la clave esté en mayúsculas
        self.license_key = self.license_key.upper()
        super().save(*args, **kwargs)
    
    @staticmethod
    def generate_license_key():
        """Genera una clave de licencia de 12 caracteres alfanuméricos"""
        import string
        import secrets
        alphabet = string.ascii_uppercase + string.digits
        while True:
            key = ''.join(secrets.choice(alphabet) for _ in range(12))
            # Verificar que no exista ya
            if not License.objects.filter(license_key=key).exists():
                return key
    
    def is_active(self):
        """Verifica si la licencia está activa y no ha expirado"""
        from django.utils import timezone
        today = timezone.now().date()
        return (
            self.status == 'active' and 
            self.start_date <= today <= self.end_date
        )
    
    def is_expired(self):
        """Verifica si la licencia ha expirado"""
        from django.utils import timezone
        today = timezone.now().date()
        return today > self.end_date
    
    def days_until_expiry(self):
        """Calcula los días restantes hasta el vencimiento"""
        from django.utils import timezone
        today = timezone.now().date()
        if today > self.end_date:
            return 0
        return (self.end_date - today).days
    
    def get_status_display_color(self):
        """Retorna el color para mostrar el estado"""
        if self.is_expired():
            return 'danger'
        elif self.status == 'active':
            days_left = self.days_until_expiry()
            if days_left <= 30:
                return 'warning'
            return 'success'
        elif self.status == 'suspended':
            return 'warning'
        else:
            return 'danger'
    
    def get_qr_data(self):
        """Retorna los datos para el código QR"""
        from django.urls import reverse
        from django.conf import settings
        
        # Usar localhost en desarrollo o el dominio configurado en producción
        if settings.DEBUG:
            domain = "http://localhost:8000"
        else:
            # En producción, usar el primer host permitido o localhost como fallback
            domain = "http://localhost:8000"
            if hasattr(settings, 'ALLOWED_HOSTS') and settings.ALLOWED_HOSTS:
                for host in settings.ALLOWED_HOSTS:
                    if host != '*' and host != '' and host != 'localhost':
                        domain = f"https://{host}"
                        break
            
        public_url = f"{domain}{reverse('license_public', kwargs={'uuid': self.public_uuid})}"
        return public_url


# ==================== MODELOS DE WHATSAPP ====================

class WhatsAppConnection(models.Model):
    """Modelo para gestionar la conexión de WhatsApp"""
    STATUS_CHOICES = [
        ('disconnected', 'Desconectado'),
        ('connecting', 'Conectando'),
        ('connected', 'Conectado'),
        ('qr_pending', 'QR Pendiente'),
        ('error', 'Error'),
    ]
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='disconnected',
        verbose_name='Estado'
    )
    qr_code = models.TextField(
        blank=True,
        null=True,
        verbose_name='Código QR'
    )
    phone_number = models.CharField(
        max_length=20,
        blank=True,
        null=True,
        verbose_name='Número de teléfono'
    )
    session_data = models.TextField(
        blank=True,
        null=True,
        verbose_name='Datos de sesión'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    last_connected = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Última conexión'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )

    class Meta:
        verbose_name = 'Conexión WhatsApp'
        verbose_name_plural = 'Conexiones WhatsApp'
        ordering = ['-created_at']

    def __str__(self):
        return f"{self.user.username} - {self.get_status_display()}"


class WhatsAppKeyword(models.Model):
    """Modelo para palabras clave y respuestas automáticas"""
    connection = models.ForeignKey(
        WhatsAppConnection,
        on_delete=models.CASCADE,
        related_name='keywords',
        verbose_name='Conexión'
    )
    keyword = models.CharField(
        max_length=200,
        verbose_name='Palabra clave'
    )
    response = models.TextField(
        verbose_name='Respuesta'
    )
    is_exact_match = models.BooleanField(
        default=False,
        verbose_name='Coincidencia exacta',
        help_text='Si es verdadero, debe coincidir exactamente. Si es falso, busca la palabra en el mensaje.'
    )
    is_case_sensitive = models.BooleanField(
        default=False,
        verbose_name='Sensible a mayúsculas'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    priority = models.IntegerField(
        default=0,
        verbose_name='Prioridad',
        help_text='Mayor número = mayor prioridad'
    )
    use_count = models.IntegerField(
        default=0,
        verbose_name='Veces usado'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )

    class Meta:
        verbose_name = 'Palabra Clave WhatsApp'
        verbose_name_plural = 'Palabras Clave WhatsApp'
        ordering = ['-priority', 'keyword']
        unique_together = ['connection', 'keyword']

    def __str__(self):
        return f"{self.keyword} → {self.response[:50]}..."

    def increment_use_count(self):
        """Incrementa el contador de uso"""
        self.use_count += 1
        self.save(update_fields=['use_count'])


class WhatsAppMessage(models.Model):
    """Modelo para registrar mensajes de WhatsApp"""
    MESSAGE_TYPE_CHOICES = [
        ('received', 'Recibido'),
        ('sent', 'Enviado'),
        ('auto_reply', 'Respuesta Automática'),
    ]
    
    connection = models.ForeignKey(
        WhatsAppConnection,
        on_delete=models.CASCADE,
        related_name='messages',
        verbose_name='Conexión'
    )
    message_type = models.CharField(
        max_length=20,
        choices=MESSAGE_TYPE_CHOICES,
        verbose_name='Tipo de mensaje'
    )
    from_number = models.CharField(
        max_length=50,
        verbose_name='De'
    )
    to_number = models.CharField(
        max_length=50,
        verbose_name='Para'
    )
    message_text = models.TextField(
        verbose_name='Mensaje'
    )
    keyword_matched = models.ForeignKey(
        WhatsAppKeyword,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Palabra clave coincidente'
    )
    timestamp = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha y hora'
    )

    class Meta:
        verbose_name = 'Mensaje WhatsApp'
        verbose_name_plural = 'Mensajes WhatsApp'
        ordering = ['-timestamp']

    def __str__(self):
        return f"{self.get_message_type_display()} - {self.from_number} → {self.to_number}"


# ==================== MODELO IMAGE TO PROMPT ====================

class ImagePrompt(models.Model):
    """Modelo para guardar imágenes y sus prompts generados por IA"""
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='image_prompts',
        verbose_name='Usuario'
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo para identificar esta imagen'
    )
    image = models.ImageField(
        upload_to='image_prompts/%Y/%m/',
        verbose_name='Imagen',
        help_text='Sube una imagen para generar un prompt'
    )
    generated_prompt = models.TextField(
        verbose_name='Prompt Generado',
        blank=True,
        help_text='Prompt generado automáticamente por IA'
    )
    custom_prompt = models.TextField(
        verbose_name='Prompt Personalizado',
        blank=True,
        null=True,
        help_text='Prompt editado manualmente por el usuario'
    )
    tags = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas para organizar'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='¿Es público?',
        help_text='Si está marcado, otros usuarios pueden ver este prompt'
    )
    tokens_used = models.IntegerField(
        default=0,
        verbose_name='Tokens usados',
        help_text='Cantidad de tokens de IA utilizados'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )

    class Meta:
        verbose_name = 'Image to Prompt'
        verbose_name_plural = 'Image to Prompts'
        ordering = ['-created_at']
        
    def __str__(self):
        return f"{self.title} - {self.user.username}"
    
    def get_final_prompt(self):
        """Retorna el prompt personalizado si existe, sino el generado"""
        return self.custom_prompt if self.custom_prompt else self.generated_prompt
    
    def delete(self, *args, **kwargs):
        """Elimina la imagen del sistema de archivos antes de eliminar el objeto"""
        # Eliminar el archivo de imagen si existe
        if self.image:
            if os.path.isfile(self.image.path):
                os.remove(self.image.path)
        super().delete(*args, **kwargs)


class AIManager(models.Model):
    """Modelo para Gerentes de IA por empresa"""
    CATEGORY_CHOICES = [
        ('ventas', 'Ventas'),
        ('compras', 'Compras'),
        ('contabilidad', 'Contabilidad'),
        ('crm', 'CRM'),
        ('inventario', 'Inventario'),
        ('rrhh', 'Recursos Humanos'),
        ('marketing', 'Marketing'),
        ('operaciones', 'Operaciones'),
        ('finanzas', 'Finanzas'),
        ('ti', 'Tecnología'),
        ('otro', 'Otro'),
    ]
    
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='ai_managers',
        verbose_name='Empresa'
    )
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Gerente IA',
        help_text='Ejemplo: Gerente de Ventas IA'
    )
    category = models.CharField(
        max_length=50,
        choices=CATEGORY_CHOICES,
        default='otro',
        verbose_name='Categoría'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Describe el rol y responsabilidades de este gerente IA'
    )
    instructions = models.TextField(
        verbose_name='Instrucciones para la IA',
        help_text='Contexto y personalidad del gerente IA',
        default='Eres un gerente experimentado y empático. Tu objetivo es ayudar al equipo a mejorar.'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    # Integración con Telegram
    telegram_enabled = models.BooleanField(
        default=False,
        verbose_name='Telegram Habilitado',
        help_text='Enviar resúmenes automáticamente por Telegram'
    )
    telegram_bot_token = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Token del Bot de Telegram',
        help_text='Token único del bot de Telegram para este gerente'
    )
    telegram_chat_id = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Chat ID de Telegram',
        help_text='ID del chat o canal donde se enviarán los resúmenes'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='created_ai_managers',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        verbose_name = 'Gerente IA'
        verbose_name_plural = 'Gerentes IA'
        ordering = ['company', 'category', 'name']
        unique_together = ['company', 'name']
    
    def __str__(self):
        return f"{self.name} - {self.company.name}"
    
    def get_total_meetings(self):
        """Retorna el total de reuniones de este gerente"""
        return self.meetings.count()
    
    def get_active_users(self):
        """Retorna usuarios únicos que han tenido reuniones"""
        return User.objects.filter(ai_manager_meetings__ai_manager=self).distinct().count()


class AIManagerMeeting(models.Model):
    """Modelo para reuniones entre usuarios y gerentes IA"""
    INPUT_TYPE_CHOICES = [
        ('text', 'Texto'),
        ('audio', 'Audio'),
    ]
    
    ai_manager = models.ForeignKey(
        AIManager,
        on_delete=models.CASCADE,
        related_name='meetings',
        verbose_name='Gerente IA'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_manager_meetings',
        verbose_name='Usuario'
    )
    title = models.CharField(
        max_length=300,
        verbose_name='Título',
        help_text='Resumen breve de la reunión'
    )
    input_type = models.CharField(
        max_length=10,
        choices=INPUT_TYPE_CHOICES,
        default='text',
        verbose_name='Tipo de entrada'
    )
    user_input = models.TextField(
        verbose_name='Entrada del usuario',
        help_text='Lo que el usuario escribió o transcripción del audio'
    )
    audio_file = models.FileField(
        upload_to='ai_manager_meetings/%Y/%m/',
        blank=True,
        null=True,
        verbose_name='Archivo de audio'
    )
    ai_response = models.TextField(
        verbose_name='Respuesta de la IA',
        help_text='Análisis y recomendaciones del gerente IA'
    )
    ai_summary = models.TextField(
        verbose_name='Resumen de la IA',
        help_text='Resumen corto de la reunión'
    )
    improvement_areas = models.TextField(
        blank=True,
        verbose_name='Áreas de mejora',
        help_text='Áreas identificadas para mejorar'
    )
    tokens_used = models.IntegerField(
        default=0,
        verbose_name='Tokens usados'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de reunión'
    )
    
    class Meta:
        verbose_name = 'Reunión con Gerente IA'
        verbose_name_plural = 'Reuniones con Gerentes IA'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.title} - {self.user.username} con {self.ai_manager.name}"
    
    def delete(self, *args, **kwargs):
        """Elimina el archivo de audio si existe"""
        if self.audio_file:
            if os.path.isfile(self.audio_file.path):
                os.remove(self.audio_file.path)
        super().delete(*args, **kwargs)


class AIManagerMeetingAttachment(models.Model):
    """Modelo para adjuntos de reuniones con gerentes IA"""
    meeting = models.ForeignKey(
        AIManagerMeeting,
        on_delete=models.CASCADE,
        related_name='attachments',
        verbose_name='Reunión'
    )
    file = models.FileField(
        upload_to='ai_manager_meeting_attachments/%Y/%m/',
        verbose_name='Archivo adjunto',
        help_text='Documentos, imágenes, PDFs que el gerente debe considerar'
    )
    file_name = models.CharField(
        max_length=255,
        verbose_name='Nombre del archivo'
    )
    file_type = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Tipo de archivo'
    )
    file_size = models.IntegerField(
        default=0,
        verbose_name='Tamaño del archivo (bytes)'
    )
    uploaded_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de carga'
    )
    
    class Meta:
        verbose_name = 'Adjunto de Reunión IA'
        verbose_name_plural = 'Adjuntos de Reuniones IA'
        ordering = ['uploaded_at']
    
    def __str__(self):
        return f"{self.file_name} - {self.meeting.title}"
    
    def save(self, *args, **kwargs):
        if self.file:
            self.file_name = self.file.name
            self.file_type = self.file.content_type if hasattr(self.file, 'content_type') else ''
            self.file_size = self.file.size
        super().save(*args, **kwargs)
    
    def delete(self, *args, **kwargs):
        # Eliminar el archivo físico al borrar el registro
        if self.file:
            if os.path.isfile(self.file.path):
                os.remove(self.file.path)
        super().delete(*args, **kwargs)
    
    def get_file_size_display(self):
        """Retorna el tamaño en formato legible"""
        size = self.file_size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"


class AIManagerSummary(models.Model):
    """Modelo para resúmenes generales del gerente IA"""
    ai_manager = models.ForeignKey(
        AIManager,
        on_delete=models.CASCADE,
        related_name='summaries',
        verbose_name='Gerente IA'
    )
    period_start = models.DateField(
        verbose_name='Inicio del período'
    )
    period_end = models.DateField(
        verbose_name='Fin del período'
    )
    summary_text = models.TextField(
        verbose_name='Resumen general',
        help_text='Resumen de todas las reuniones del período'
    )
    key_insights = models.TextField(
        verbose_name='Insights clave',
        help_text='Puntos clave identificados'
    )
    recommendations = models.TextField(
        verbose_name='Recomendaciones',
        help_text='Recomendaciones para la empresa'
    )
    total_meetings = models.IntegerField(
        default=0,
        verbose_name='Total de reuniones'
    )
    participants_count = models.IntegerField(
        default=0,
        verbose_name='Participantes únicos'
    )
    tokens_used = models.IntegerField(
        default=0,
        verbose_name='Tokens usados'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Resumen del Gerente IA'
        verbose_name_plural = 'Resúmenes del Gerente IA'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"Resumen {self.ai_manager.name} - {self.period_start} a {self.period_end}"


class CompanyAISummary(models.Model):
    """Modelo para resúmenes empresariales generales (todos los gerentes)"""
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='ai_summaries',
        verbose_name='Empresa'
    )
    period_start = models.DateField(
        verbose_name='Inicio del período'
    )
    period_end = models.DateField(
        verbose_name='Fin del período'
    )
    executive_summary = models.TextField(
        verbose_name='Resumen ejecutivo',
        help_text='Resumen general de toda la empresa'
    )
    key_metrics = models.TextField(
        verbose_name='Métricas clave',
        help_text='KPIs y métricas importantes'
    )
    department_highlights = models.TextField(
        verbose_name='Destacados por departamento',
        help_text='Logros por área/gerente'
    )
    challenges = models.TextField(
        verbose_name='Desafíos identificados',
        help_text='Principales retos de la empresa'
    )
    strategic_recommendations = models.TextField(
        verbose_name='Recomendaciones estratégicas',
        help_text='Acciones recomendadas para la dirección'
    )
    total_managers = models.IntegerField(
        default=0,
        verbose_name='Total de gerentes'
    )
    total_meetings = models.IntegerField(
        default=0,
        verbose_name='Total de reuniones'
    )
    total_participants = models.IntegerField(
        default=0,
        verbose_name='Total de participantes'
    )
    tokens_used = models.IntegerField(
        default=0,
        verbose_name='Tokens usados'
    )
    generated_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Generado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Resumen Empresarial IA'
        verbose_name_plural = 'Resúmenes Empresariales IA'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"Resumen {self.company.name} - {self.period_start} a {self.period_end}"


class UserAIPerformanceEvaluation(models.Model):
    """Modelo para evaluaciones de desempeño de usuarios con IA"""
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_performance_evaluations',
        verbose_name='Usuario'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='user_evaluations',
        verbose_name='Empresa',
        null=True,
        blank=True
    )
    evaluation_date = models.DateField(
        verbose_name='Fecha de evaluación'
    )
    overall_score = models.IntegerField(
        verbose_name='Puntuación general (1-100)',
        help_text='Calificación automática generada por IA'
    )
    productivity_score = models.IntegerField(
        default=0,
        verbose_name='Productividad (1-100)'
    )
    communication_score = models.IntegerField(
        default=0,
        verbose_name='Comunicación (1-100)'
    )
    goal_achievement_score = models.IntegerField(
        default=0,
        verbose_name='Logro de objetivos (1-100)'
    )
    consistency_score = models.IntegerField(
        default=0,
        verbose_name='Consistencia (1-100)'
    )
    improvement_areas = models.TextField(
        verbose_name='Áreas de mejora',
        help_text='Áreas identificadas que necesitan mejora'
    )
    training_recommendations = models.TextField(
        verbose_name='Recomendaciones de capacitación',
        help_text='Cursos o entrenamientos sugeridos'
    )
    strengths = models.TextField(
        verbose_name='Fortalezas identificadas',
        help_text='Puntos fuertes del usuario'
    )
    meetings_analyzed = models.IntegerField(
        default=0,
        verbose_name='Reuniones analizadas'
    )
    period_start = models.DateField(
        verbose_name='Inicio del período',
        null=True,
        blank=True
    )
    period_end = models.DateField(
        verbose_name='Fin del período',
        null=True,
        blank=True
    )
    ai_summary = models.TextField(
        verbose_name='Resumen IA',
        help_text='Resumen narrativo generado por IA',
        blank=True
    )
    tokens_used = models.IntegerField(
        default=0,
        verbose_name='Tokens usados'
    )
    generated_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='generated_evaluations',
        verbose_name='Generado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        verbose_name = 'Evaluación de Desempeño IA'
        verbose_name_plural = 'Evaluaciones de Desempeño IA'
        ordering = ['-evaluation_date', '-created_at']
        unique_together = ['user', 'evaluation_date']
    
    def __str__(self):
        return f"{self.user.get_full_name()} - {self.evaluation_date} - {self.overall_score}/100"
    
    def get_score_color(self):
        """Retorna el color según la puntuación"""
        if self.overall_score >= 90:
            return 'success'
        elif self.overall_score >= 75:
            return 'info'
        elif self.overall_score >= 60:
            return 'warning'
        else:
            return 'danger'
    
    def get_performance_level(self):
        """Retorna el nivel de desempeño"""
        if self.overall_score >= 90:
            return 'Excelente'
        elif self.overall_score >= 75:
            return 'Bueno'
        elif self.overall_score >= 60:
            return 'Regular'
        else:
            return 'Necesita Mejora'


class WebsiteTracker(models.Model):
    """Modelo para rastrear y analizar sitios web/IPs"""
    
    # Información básica
    target = models.CharField(
        max_length=500,
        verbose_name='URL o IP',
        help_text='URL completa (https://ejemplo.com) o dirección IP'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='website_trackers',
        verbose_name='Usuario'
    )
    
    # Información de conectividad
    is_active = models.BooleanField(
        default=False,
        verbose_name='Está activo'
    )
    ping_response_time = models.FloatField(
        null=True,
        blank=True,
        verbose_name='Tiempo de respuesta (ms)',
        help_text='Tiempo de respuesta del ping en milisegundos'
    )
    
    # Información de DNS
    ip_address = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Dirección IP'
    )
    cname_records = models.JSONField(
        default=list,
        blank=True,
        verbose_name='Registros CNAME'
    )
    txt_records = models.JSONField(
        default=list,
        blank=True,
        verbose_name='Registros TXT'
    )
    mx_records = models.JSONField(
        default=list,
        blank=True,
        verbose_name='Registros MX'
    )
    ns_records = models.JSONField(
        default=list,
        blank=True,
        verbose_name='Registros NS'
    )
    
    # Información HTTP
    http_status_code = models.IntegerField(
        null=True,
        blank=True,
        verbose_name='Código de estado HTTP'
    )
    http_headers = models.JSONField(
        default=dict,
        blank=True,
        verbose_name='Headers HTTP'
    )
    redirect_url = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='URL de redirección'
    )
    
    # Información del servidor
    server_software = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Software del servidor'
    )
    technologies = models.JSONField(
        default=list,
        blank=True,
        verbose_name='Tecnologías detectadas',
        help_text='Frameworks, CMS, librerías detectadas'
    )
    
    # Información de seguridad
    ssl_valid = models.BooleanField(
        null=True,
        blank=True,
        verbose_name='SSL válido'
    )
    ssl_issuer = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Emisor del certificado SSL'
    )
    ssl_expiry_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de expiración SSL'
    )
    
    # Información de contenido
    page_title = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Título de la página'
    )
    meta_description = models.TextField(
        blank=True,
        verbose_name='Meta descripción'
    )
    page_size = models.IntegerField(
        null=True,
        blank=True,
        verbose_name='Tamaño de página (bytes)'
    )
    load_time = models.FloatField(
        null=True,
        blank=True,
        verbose_name='Tiempo de carga (s)'
    )
    
    # Información adicional
    whois_data = models.JSONField(
        default=dict,
        blank=True,
        verbose_name='Datos WHOIS'
    )
    error_message = models.TextField(
        blank=True,
        verbose_name='Mensaje de error',
        help_text='Si ocurre algún error durante el rastreo'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de rastreo'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        verbose_name = 'Rastreador Web'
        verbose_name_plural = 'Rastreadores Web'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.target} - {self.created_at.strftime('%d/%m/%Y %H:%M')}"
    
    def get_status_badge(self):
        """Retorna el color del badge según el estado"""
        if self.is_active:
            return 'success'
        elif self.error_message:
            return 'danger'
        else:
            return 'secondary'
    
    def get_status_text(self):
        """Retorna el texto del estado"""
        if self.is_active:
            return 'Activo'
        elif self.error_message:
            return 'Error'
        else:
            return 'Inactivo'


class LegalContract(models.Model):
    """Modelo para contratos legales generados con IA"""
    
    # Información básica
    name = models.CharField(
        max_length=255,
        verbose_name='Nombre del contrato',
        help_text='Ej: Contrato de Soporte de Software'
    )
    contract_type = models.CharField(
        max_length=100,
        verbose_name='Tipo de contrato',
        help_text='Ej: Soporte de software, Desarrollo, Consultoría, etc.'
    )
    
    # Partes del contrato
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='legal_contracts_as_provider',
        verbose_name='Empresa proveedora',
        help_text='Tu empresa que provee el servicio'
    )
    client_company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='legal_contracts_as_client',
        verbose_name='Empresa cliente',
        help_text='La empresa que contrata el servicio'
    )
    
    # Prompt y generación
    objective_prompt = models.TextField(
        verbose_name='Objetivo del contrato',
        help_text='Describe el objetivo y alcance del contrato. Ej: Soporte técnico 24/7 para sistema ERP'
    )
    generated_content = models.TextField(
        blank=True,
        verbose_name='Contenido generado',
        help_text='Contenido del contrato generado por IA'
    )
    
    # Detalles del contrato
    start_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de inicio'
    )
    end_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de finalización'
    )
    amount = models.DecimalField(
        max_digits=12,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Monto del contrato',
        help_text='Valor económico del contrato'
    )
    currency = models.CharField(
        max_length=3,
        default='USD',
        verbose_name='Moneda',
        help_text='Código de moneda (USD, EUR, MXN, etc.)'
    )
    
    # Control de acceso
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='legal_contracts',
        verbose_name='Creado por'
    )
    public_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token público',
        help_text='Token único para compartir el contrato públicamente'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Si está marcado, el contrato se puede ver públicamente con el enlace'
    )
    
    # Estado
    status = models.CharField(
        max_length=20,
        choices=[
            ('draft', 'Borrador'),
            ('generated', 'Generado'),
            ('sent', 'Enviado'),
            ('signed', 'Firmado'),
            ('active', 'Activo'),
            ('completed', 'Completado'),
            ('cancelled', 'Cancelado'),
        ],
        default='draft',
        verbose_name='Estado'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    generated_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de generación',
        help_text='Cuándo se generó el contenido con IA'
    )
    
    # Configuración adicional
    notes = models.TextField(
        blank=True,
        verbose_name='Notas internas',
        help_text='Notas privadas sobre el contrato'
    )
    
    class Meta:
        verbose_name = 'Contrato Legal'
        verbose_name_plural = 'Contratos Legales'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['public_token']),
            models.Index(fields=['status']),
            models.Index(fields=['-created_at']),
        ]
    
    def __str__(self):
        return f"{self.name} - {self.company.name} / {self.client_company.name}"
    
    def get_absolute_url(self):
        from django.urls import reverse
        return reverse('legal_contract_detail', kwargs={'pk': self.pk})
    
    def get_public_url(self):
        from django.urls import reverse
        return reverse('legal_contract_public', kwargs={'token': self.public_token})
    
    def get_status_badge(self):
        """Retorna el color del badge según el estado"""
        status_colors = {
            'draft': 'secondary',
            'generated': 'info',
            'sent': 'warning',
            'signed': 'primary',
            'active': 'success',
            'completed': 'success',
            'cancelled': 'danger',
        }
        return status_colors.get(self.status, 'secondary')
    
    def get_status_display_with_icon(self):
        """Retorna el estado con icono"""
        status_icons = {
            'draft': 'bi-pencil',
            'generated': 'bi-magic',
            'sent': 'bi-send',
            'signed': 'bi-pen',
            'active': 'bi-check-circle',
            'completed': 'bi-check-circle-fill',
            'cancelled': 'bi-x-circle',
        }
        icon = status_icons.get(self.status, 'bi-circle')
        return f'<i class="bi {icon}"></i> {self.get_status_display()}'
    
    def generate_with_ai(self):
        """Genera el contenido del contrato usando OpenAI"""
        from tickets.models import SystemConfiguration
        import re
        
        def remove_markdown_formatting(text):
            """Elimina formato Markdown y devuelve texto plano"""
            # Eliminar negritas **texto** y dejar solo texto
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Eliminar cursivas *texto* y dejar solo texto
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Eliminar títulos ## y # y dejar solo el texto
            text = re.sub(r'^#{1,6}\s*(.*?)$', r'\1', text, flags=re.MULTILINE)
            return text
        
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            raise Exception("OpenAI no está configurado. Por favor configura la API key en Configuración del Sistema.")
        
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            prompt = f"""
            Genera un contrato legal profesional y detallado con los siguientes datos:
            
            TIPO DE CONTRATO: {self.contract_type}
            OBJETIVO: {self.objective_prompt}
            
            EMPRESA PROVEEDORA:
            - Nombre: {self.company.name}
            - Dirección: {self.company.address}
            - Teléfono: {self.company.phone}
            - Email: {self.company.email}
            
            EMPRESA CLIENTE:
            - Nombre: {self.client_company.name}
            - Dirección: {self.client_company.address}
            - Teléfono: {self.client_company.phone}
            - Email: {self.client_company.email}
            
            DETALLES ADICIONALES:
            - Fecha de inicio: {self.start_date if self.start_date else 'No especificada'}
            - Fecha de finalización: {self.end_date if self.end_date else 'No especificada'}
            - Monto: {f'{self.amount} {self.currency}' if self.amount else 'No especificado'}
            
            El contrato debe incluir:
            1. Encabezado con las partes contratantes
            2. Antecedentes y objeto del contrato
            3. Obligaciones de cada parte
            4. Términos y condiciones
            5. Duración y vigencia
            6. Aspectos económicos
            7. Cláusulas de confidencialidad
            8. Causas de rescisión
            9. Solución de controversias
            10. Firmas y datos de contacto
            
            Usa un lenguaje legal formal pero comprensible. Incluye cláusulas estándar para este tipo de contrato.
            NO uses formato Markdown. Usa texto plano con saltos de línea para separar secciones.
            """
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Eres un abogado experto en redacción de contratos legales comerciales. Generas contratos profesionales, detallados y legalmente sólidos. NO uses formato Markdown en tus respuestas, solo texto plano con saltos de línea."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=4000
            )
            
            generated_content = response.choices[0].message.content
            # Limpiar cualquier formato Markdown y devolver texto plano
            self.generated_content = remove_markdown_formatting(generated_content)
            self.generated_at = timezone.now()
            self.status = 'generated'
            self.save()
            
            return True
            
        except Exception as e:
            raise Exception(f"Error al generar contrato: {str(e)}")
    
class SupplierContractReview(models.Model):
    """Modelo para revisar contratos de proveedores con IA"""
    
    REVIEW_STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('reviewed', 'Revisado'),
        ('approved', 'Aprobado'),
        ('rejected', 'Rechazado'),
    ]
    
    RISK_LEVEL_CHOICES = [
        ('low', 'Bajo'),
        ('medium', 'Medio'),
        ('high', 'Alto'),
        ('critical', 'Crítico'),
    ]
    
    # Información básica
    name = models.CharField(max_length=200, verbose_name="Nombre del Contrato")
    supplier_name = models.CharField(max_length=200, verbose_name="Nombre del Proveedor")
    company = models.ForeignKey(Company, on_delete=models.CASCADE, related_name='supplier_contract_reviews', verbose_name="Empresa Revisora")
    
    # Contenido del contrato
    contract_text = models.TextField(verbose_name="Texto del Contrato", help_text="Pegue aquí el texto del contrato o déjelo vacío si subirá un archivo")
    contract_file = models.FileField(upload_to='contract_reviews/%Y/%m/', blank=True, null=True, verbose_name="Archivo del Contrato")
    
    # Análisis de IA
    ai_review = models.TextField(blank=True, verbose_name="Revisión de IA", help_text="Análisis generado por IA")
    risk_assessment = models.TextField(blank=True, verbose_name="Evaluación de Riesgos")
    key_clauses = models.TextField(blank=True, verbose_name="Cláusulas Clave Identificadas")
    recommendations = models.TextField(blank=True, verbose_name="Recomendaciones")
    risk_level = models.CharField(max_length=20, choices=RISK_LEVEL_CHOICES, default='medium', verbose_name="Nivel de Riesgo")
    
    # Estado y seguimiento
    status = models.CharField(max_length=20, choices=REVIEW_STATUS_CHOICES, default='pending', verbose_name="Estado")
    reviewed_at = models.DateTimeField(null=True, blank=True, verbose_name="Fecha de Revisión")
    
    # Notas internas
    internal_notes = models.TextField(blank=True, verbose_name="Notas Internas")
    
    # Metadatos
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Usuario")
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de Creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Última Actualización")
    
    class Meta:
        verbose_name = "Revisión de Contrato de Proveedor"
        verbose_name_plural = "Revisiones de Contratos de Proveedores"
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.name} - {self.supplier_name}"
    
    def get_status_badge(self):
        """Retorna el badge HTML para el estado"""
        badges = {
            'pending': '<span class="badge bg-warning">Pendiente</span>',
            'reviewed': '<span class="badge bg-info">Revisado</span>',
            'approved': '<span class="badge bg-success">Aprobado</span>',
            'rejected': '<span class="badge bg-danger">Rechazado</span>',
        }
        return badges.get(self.status, '<span class="badge bg-secondary">Desconocido</span>')
    
    def get_risk_badge(self):
        """Retorna el badge HTML para el nivel de riesgo"""
        badges = {
            'low': '<span class="badge bg-success"><i class="bi bi-shield-check"></i> Bajo</span>',
            'medium': '<span class="badge bg-warning"><i class="bi bi-shield-exclamation"></i> Medio</span>',
            'high': '<span class="badge bg-danger"><i class="bi bi-shield-x"></i> Alto</span>',
            'critical': '<span class="badge bg-dark"><i class="bi bi-shield-fill-x"></i> Crítico</span>',
        }
        return badges.get(self.risk_level, '<span class="badge bg-secondary">N/A</span>')
    
    def get_contract_content(self):
        """Obtiene el contenido del contrato, ya sea del texto o del archivo"""
        if self.contract_text:
            return self.contract_text
        
        if self.contract_file:
            try:
                file_extension = self.contract_file.name.split('.')[-1].lower()
                
                if file_extension == 'pdf':
                    # Extraer texto de PDF
                    import PyPDF2
                    with self.contract_file.open('rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        text = ""
                        for page in pdf_reader.pages:
                            text += page.extract_text()
                        return text
                
                elif file_extension in ['doc', 'docx']:
                    # Extraer texto de Word
                    import docx
                    doc = docx.Document(self.contract_file.path)
                    text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
                    return text
                
                elif file_extension == 'txt':
                    # Leer archivo de texto
                    with self.contract_file.open('r', encoding='utf-8') as file:
                        return file.read()
                
                else:
                    return "Formato de archivo no soportado. Use PDF, DOCX o TXT."
            
            except Exception as e:
                return f"Error al extraer texto del archivo: {str(e)}"
        
        return "No hay contenido disponible"
    
    def review_with_ai(self):
        """Revisa el contrato usando IA"""
        from django.utils import timezone
        
        # Verificar configuración de OpenAI
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            raise Exception("OpenAI no está configurado. Por favor configura la API key en Configuración del Sistema.")
        
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            # Obtener el contenido del contrato
            contract_content = self.get_contract_content()
            
            if not contract_content or contract_content.startswith("Error") or contract_content == "No hay contenido disponible":
                raise Exception("No se puede obtener el contenido del contrato para revisar")
            
            prompt = f"""
            Eres un abogado experto en revisión de contratos comerciales con proveedores. 
            Analiza el siguiente contrato de proveedor de manera exhaustiva y profesional.
            
            PROVEEDOR: {self.supplier_name}
            EMPRESA REVISORA: {self.company.name}
            
            CONTRATO:
            {contract_content[:15000]}  # Limitar a 15000 caracteres para no exceder tokens
            
            Proporciona un análisis detallado en el siguiente formato:
            
            ## RESUMEN EJECUTIVO
            [Resumen breve del contrato y su propósito]
            
            ## ANÁLISIS DETALLADO
            [Análisis completo del contrato, términos principales, obligaciones de cada parte]
            
            ## CLÁUSULAS CLAVE
            [Lista numerada de las cláusulas más importantes y su implicación]
            
            ## EVALUACIÓN DE RIESGOS
            [Identificación de riesgos potenciales, cláusulas problemáticas, términos desfavorables]
            
            ## NIVEL DE RIESGO GENERAL
            [Indica: BAJO, MEDIO, ALTO o CRÍTICO]
            
            ## RECOMENDACIONES
            [Lista de recomendaciones específicas: qué negociar, qué cláusulas modificar, qué aclaraciones solicitar]
            
            ## PUNTOS A NEGOCIAR
            [Lista de puntos específicos que deberían negociarse antes de firmar]
            
            Sé específico, profesional y enfócate en proteger los intereses de {self.company.name}.
            """
            
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "Eres un abogado corporativo experto en revisión de contratos con proveedores. Tu objetivo es identificar riesgos, proteger los intereses del cliente y proporcionar recomendaciones prácticas."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.5,
                max_tokens=4000
            )
            
            full_review = response.choices[0].message.content
            
            # Extraer información estructurada
            self.ai_review = full_review
            
            # Extraer cláusulas clave
            if "## CLÁUSULAS CLAVE" in full_review:
                clauses_section = full_review.split("## CLÁUSULAS CLAVE")[1].split("##")[0]
                self.key_clauses = clauses_section.strip()
            
            # Extraer evaluación de riesgos
            if "## EVALUACIÓN DE RIESGOS" in full_review:
                risk_section = full_review.split("## EVALUACIÓN DE RIESGOS")[1].split("##")[0]
                self.risk_assessment = risk_section.strip()
            
            # Extraer recomendaciones
            if "## RECOMENDACIONES" in full_review:
                recommendations_section = full_review.split("## RECOMENDACIONES")[1].split("##")[0]
                self.recommendations = recommendations_section.strip()
            
            # Determinar nivel de riesgo
            if "## NIVEL DE RIESGO GENERAL" in full_review:
                risk_level_section = full_review.split("## NIVEL DE RIESGO GENERAL")[1].split("##")[0].strip().lower()
                if "crítico" in risk_level_section or "critical" in risk_level_section:
                    self.risk_level = 'critical'
                elif "alto" in risk_level_section or "high" in risk_level_section:
                    self.risk_level = 'high'
                elif "medio" in risk_level_section or "medium" in risk_level_section:
                    self.risk_level = 'medium'
                else:
                    self.risk_level = 'low'
            
            self.status = 'reviewed'
            self.reviewed_at = timezone.now()
            self.save()
            
            return True
            
        except Exception as e:
            raise Exception(f"Error al revisar contrato: {str(e)}")


class PayPalPaymentLink(models.Model):
    """Modelo para enlaces de pago de PayPal"""
    
    STATUS_CHOICES = [
        ('active', 'Activo'),
        ('paid', 'Pagado'),
        ('expired', 'Expirado'),
        ('cancelled', 'Cancelado'),
    ]
    
    # Información básica
    product_name = models.CharField(max_length=200, verbose_name="Nombre del Producto/Servicio")
    description = models.TextField(blank=True, verbose_name="Descripción")
    amount = models.DecimalField(max_digits=10, decimal_places=2, verbose_name="Importe (EUR)")
    
    # Archivo adjunto
    attachment = models.FileField(
        upload_to='paypal_attachments/%Y/%m/',
        blank=True,
        null=True,
        verbose_name="Archivo Adjunto",
        help_text="Archivo que el cliente podrá descargar después del pago"
    )
    attachment_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name="Nombre del archivo",
        help_text="Nombre descriptivo del archivo adjunto"
    )
    
    # URL y token
    public_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name="Token Público"
    )
    public_url = models.URLField(blank=True, verbose_name="URL Pública")
    
    # Información del pago
    paypal_order_id = models.CharField(max_length=200, blank=True, verbose_name="ID de Orden PayPal")
    paypal_payer_id = models.CharField(max_length=200, blank=True, verbose_name="ID del Pagador")
    paypal_payer_email = models.EmailField(blank=True, verbose_name="Email del Pagador")
    paypal_payer_name = models.CharField(max_length=200, blank=True, verbose_name="Nombre del Pagador")
    paypal_payment_status = models.CharField(max_length=50, blank=True, verbose_name="Estado del Pago PayPal")
    payment_date = models.DateTimeField(null=True, blank=True, verbose_name="Fecha de Pago")
    
    # Estado y seguimiento
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='active',
        verbose_name="Estado"
    )
    download_count = models.IntegerField(default=0, verbose_name="Número de Descargas")
    
    # Token temporal de descarga (válido por 72 horas después del pago)
    download_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name="Token de Descarga Temporal"
    )
    download_token_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name="Fecha de Expiración del Token de Descarga"
    )
    
    # Metadatos
    user = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name="Usuario Creador")
    company = models.ForeignKey(
        Company,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name="Empresa",
        help_text="Empresa asociada (opcional)"
    )
    created_at = models.DateTimeField(auto_now_add=True, verbose_name="Fecha de Creación")
    updated_at = models.DateTimeField(auto_now=True, verbose_name="Última Actualización")
    expires_at = models.DateTimeField(null=True, blank=True, verbose_name="Fecha de Expiración")
    
    # Notas internas
    notes = models.TextField(blank=True, verbose_name="Notas Internas")
    
    class Meta:
        verbose_name = "Enlace de Pago PayPal"
        verbose_name_plural = "Enlaces de Pago PayPal"
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.product_name} - €{self.amount}"
    
    def save(self, *args, **kwargs):
        # Generar URL pública si no existe
        if not self.public_url:
            from django.conf import settings
            base_url = getattr(settings, 'BASE_URL', 'http://localhost:8000')
            self.public_url = f"{base_url}/paypal-payment/{self.public_token}/"
        
        # Establecer nombre del archivo si no existe
        if self.attachment and not self.attachment_name:
            self.attachment_name = self.attachment.name.split('/')[-1]
        
        super().save(*args, **kwargs)
    
    def get_status_badge(self):
        """Retorna el badge HTML para el estado"""
        badges = {
            'active': '<span class="badge bg-primary">Activo</span>',
            'paid': '<span class="badge bg-success"><i class="bi bi-check-circle"></i> Pagado</span>',
            'expired': '<span class="badge bg-secondary">Expirado</span>',
            'cancelled': '<span class="badge bg-danger">Cancelado</span>',
        }
        return badges.get(self.status, '<span class="badge bg-secondary">Desconocido</span>')
    
    def is_paid(self):
        """Verifica si el enlace ha sido pagado"""
        return self.status == 'paid'
    
    def is_active(self):
        """Verifica si el enlace está activo (solo verifica expiración por fecha)"""
        from django.utils import timezone
        # Solo verificar expiración por fecha - los enlaces son reutilizables
        if self.expires_at and self.expires_at < timezone.now():
            return False
        return True
    
    def mark_as_paid(self, order_id, payer_info):
        """Marca el enlace como pagado con información del pagador"""
        from django.utils import timezone
        from datetime import timedelta
        
        self.status = 'paid'
        self.payment_date = timezone.now()
        self.paypal_order_id = order_id
        
        # Establecer token de descarga válido por 72 horas
        self.download_token = uuid.uuid4()
        self.download_token_expires_at = timezone.now() + timedelta(hours=72)
        
        if payer_info:
            self.paypal_payer_id = payer_info.get('payer_id', '')
            self.paypal_payer_email = payer_info.get('email_address', '')
            name_info = payer_info.get('name', {})
            given_name = name_info.get('given_name', '') if isinstance(name_info, dict) else ''
            surname = name_info.get('surname', '') if isinstance(name_info, dict) else ''
            self.paypal_payer_name = f"{given_name} {surname}".strip()
        
        self.save()
    
    def increment_download(self):
        """Incrementa el contador de descargas"""
        self.download_count += 1
        self.save(update_fields=['download_count'])
    
    def is_download_token_valid(self):
        """Verifica si el token de descarga aún es válido (72 horas)"""
        from django.utils import timezone
        if not self.is_paid():
            return False
        if not self.download_token_expires_at:
            return False
        return self.download_token_expires_at > timezone.now()
    
    def get_download_url(self):
        """Retorna la URL de descarga temporal"""
        if self.is_download_token_valid() and self.attachment:
            from django.urls import reverse
            return reverse('paypal_download_with_token', kwargs={'token': self.download_token})
        return None
    
    def get_order_summary_url(self):
        """Retorna la URL del resumen de orden"""
        if self.is_paid():
            from django.urls import reverse
            return reverse('paypal_order_summary', kwargs={'token': self.public_token})
        return None
    
    def get_hours_until_download_expires(self):
        """Retorna las horas restantes hasta que expire el token de descarga"""
        from django.utils import timezone
        if not self.download_token_expires_at:
            return 0
        time_diff = self.download_token_expires_at - timezone.now()
        hours = time_diff.total_seconds() / 3600
        return max(0, round(hours, 1))
    
    def get_total_orders_count(self):
        """Retorna el número total de órdenes/pagos completados para este enlace"""
        return self.orders.filter(status='paid').count()
    
    def get_total_revenue(self):
        """Retorna el total de ingresos generados por este enlace"""
        from django.db.models import Sum
        total = self.orders.filter(status='paid').aggregate(total=Sum('amount'))['total']
        return total or 0


class PayPalOrder(models.Model):
    """Modelo para órdenes individuales de pago (cada compra genera una orden)"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('paid', 'Pagado'),
        ('failed', 'Fallido'),
        ('refunded', 'Reembolsado'),
    ]
    
    # Relación con el enlace de pago original (plantilla)
    payment_link = models.ForeignKey(
        PayPalPaymentLink,
        on_delete=models.CASCADE,
        related_name='orders',
        verbose_name='Enlace de Pago',
        help_text='Enlace de pago desde el cual se generó esta orden'
    )
    
    # Token único para esta orden
    order_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token de Orden',
        help_text='Token único para acceder a esta orden específica'
    )
    
    # Información del producto (copiada del enlace al momento de la compra)
    product_name = models.CharField(max_length=200, verbose_name='Producto')
    description = models.TextField(blank=True, verbose_name='Descripción')
    amount = models.DecimalField(max_digits=10, decimal_places=2, verbose_name='Importe (EUR)')
    
    # Información del pago
    paypal_order_id = models.CharField(
        max_length=200,
        unique=True,
        verbose_name='ID de Orden PayPal'
    )
    paypal_payer_id = models.CharField(max_length=200, blank=True, verbose_name='ID del Pagador')
    paypal_payer_email = models.EmailField(blank=True, verbose_name='Email del Pagador')
    paypal_payer_name = models.CharField(max_length=200, blank=True, verbose_name='Nombre del Pagador')
    payment_date = models.DateTimeField(null=True, blank=True, verbose_name='Fecha de Pago')
    
    # Estado
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    # Token de descarga temporal (válido 72 horas)
    download_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token de Descarga'
    )
    download_token_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Expiración Token de Descarga'
    )
    download_count = models.IntegerField(default=0, verbose_name='Número de Descargas')
    
    # Metadatos
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    
    class Meta:
        verbose_name = 'Orden de Pago PayPal'
        verbose_name_plural = 'Órdenes de Pago PayPal'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['paypal_order_id']),
            models.Index(fields=['order_token']),
            models.Index(fields=['download_token']),
            models.Index(fields=['status', 'created_at']),
        ]
    
    def __str__(self):
        return f"Orden #{self.id:05d} - {self.product_name} - {self.get_status_display()}"
    
    def is_paid(self):
        """Verifica si la orden ha sido pagada"""
        return self.status == 'paid'
    
    def mark_as_paid(self, payer_info):
        """Marca la orden como pagada"""
        from django.utils import timezone
        from datetime import timedelta
        
        self.status = 'paid'
        self.payment_date = timezone.now()
        
        # Establecer token de descarga válido por 72 horas
        self.download_token = uuid.uuid4()
        self.download_token_expires_at = timezone.now() + timedelta(hours=72)
        
        if payer_info:
            self.paypal_payer_id = payer_info.get('payer_id', '')
            self.paypal_payer_email = payer_info.get('email_address', '')
            name_info = payer_info.get('name', {})
            given_name = name_info.get('given_name', '') if isinstance(name_info, dict) else ''
            surname = name_info.get('surname', '') if isinstance(name_info, dict) else ''
            self.paypal_payer_name = f"{given_name} {surname}".strip()
        
        self.save()
    
    def increment_download(self):
        """Incrementa el contador de descargas"""
        self.download_count += 1
        self.save(update_fields=['download_count'])
    
    def is_download_token_valid(self):
        """Verifica si el token de descarga aún es válido"""
        from django.utils import timezone
        if not self.is_paid():
            return False
        if not self.download_token_expires_at:
            return False
        return self.download_token_expires_at > timezone.now()
    
    def get_hours_until_download_expires(self):
        """Retorna las horas restantes hasta que expire el token"""
        from django.utils import timezone
        if not self.download_token_expires_at:
            return 0
        time_diff = self.download_token_expires_at - timezone.now()
        hours = time_diff.total_seconds() / 3600
        return max(0, round(hours, 1))
    
    def get_download_url(self):
        """Retorna la URL de descarga temporal"""
        if self.is_download_token_valid() and self.payment_link.attachment:
            from django.urls import reverse
            return reverse('paypal_order_download', kwargs={'token': self.download_token})
        return None
    
    def get_order_summary_url(self):
        """Retorna la URL del resumen de orden"""
        from django.urls import reverse
        return reverse('paypal_order_detail', kwargs={'token': self.order_token})
    
    def get_status_badge(self):
        """Retorna el badge HTML para el estado"""
        badges = {
            'pending': '<span class="badge bg-warning">Pendiente</span>',
            'paid': '<span class="badge bg-success"><i class="bi bi-check-circle"></i> Pagado</span>',
            'failed': '<span class="badge bg-danger">Fallido</span>',
            'refunded': '<span class="badge bg-secondary">Reembolsado</span>',
        }
        return badges.get(self.status, '<span class="badge bg-secondary">Desconocido</span>')


class TodoItem(models.Model):
    """Modelo para items de TODO list en tickets"""
    ticket = models.ForeignKey(
        Ticket,
        on_delete=models.CASCADE,
        related_name='todo_items',
        verbose_name='Ticket'
    )
    text = models.CharField(
        max_length=500,
        verbose_name='Tarea'
    )
    is_completed = models.BooleanField(
        default=False,
        verbose_name='Completada'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='todo_items_created',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    completed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de completado'
    )
    order = models.IntegerField(
        default=0,
        verbose_name='Orden'
    )
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Item de TODO'
        verbose_name_plural = 'Items de TODO'
    
    def __str__(self):
        status = "✓" if self.is_completed else "○"
        return f"{status} {self.text[:50]}"
    
    def toggle_completed(self):
        """Alterna el estado de completado"""
        self.is_completed = not self.is_completed
        if self.is_completed:
            self.completed_at = timezone.now()
        else:
            self.completed_at = None
        self.save()


class AIBook(models.Model):
    """Modelo para libros generados con IA"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('chapters_proposed', 'Capítulos Propuestos'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completado'),
    ]
    
    title = models.CharField(
        max_length=500,
        verbose_name='Título del Libro'
    )
    author = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Autor',
        help_text='Nombre del autor del libro'
    )
    topic = models.TextField(
        verbose_name='Tema/Descripción',
        help_text='Describe de qué quieres que trate el libro'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_books',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='ai_books',
        verbose_name='Empresa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Libro IA'
        verbose_name_plural = 'Libros IA'
    
    def __str__(self):
        return self.title
    
    def get_status_badge_class(self):
        """Retorna la clase Bootstrap del badge según el estado"""
        badges = {
            'draft': 'bg-secondary',
            'chapters_proposed': 'bg-info',
            'in_progress': 'bg-warning',
            'completed': 'bg-success',
        }
        return badges.get(self.status, 'bg-secondary')
    
    def get_total_chapters(self):
        """Retorna el total de capítulos"""
        return self.chapters.count()
    
    def get_completed_chapters(self):
        """Retorna el número de capítulos completados (con contenido generado)"""
        return self.chapters.filter(status__in=['content_generated', 'completed']).count()
    
    def get_progress_percentage(self):
        """Retorna el porcentaje de progreso"""
        total = self.get_total_chapters()
        if total == 0:
            return 0
        completed = self.get_completed_chapters()
        return int((completed / total) * 100)
    
    def get_total_words(self):
        """Retorna el total de palabras de todos los capítulos"""
        from django.db.models import Sum
        result = self.chapters.aggregate(total=Sum('word_count'))
        return result['total'] or 0


class AIBookChapter(models.Model):
    """Modelo para capítulos de libros generados con IA"""
    
    STATUS_CHOICES = [
        ('proposed', 'Propuesto'),
        ('summary_created', 'Resumen Creado'),
        ('content_generated', 'Contenido Generado'),
        ('completed', 'Completado'),
    ]
    
    book = models.ForeignKey(
        AIBook,
        on_delete=models.CASCADE,
        related_name='chapters',
        verbose_name='Libro'
    )
    title = models.CharField(
        max_length=500,
        verbose_name='Título del Capítulo'
    )
    order = models.IntegerField(
        default=0,
        verbose_name='Orden'
    )
    summary = models.TextField(
        blank=True,
        verbose_name='Resumen del Capítulo',
        help_text='Resumen breve de lo que tratará el capítulo'
    )
    content = models.TextField(
        blank=True,
        verbose_name='Contenido del Capítulo',
        help_text='Contenido completo en texto plano (sin formato Markdown)'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='proposed',
        verbose_name='Estado'
    )
    word_count = models.IntegerField(
        default=0,
        verbose_name='Conteo de Palabras'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['book', 'order']
        verbose_name = 'Capítulo de Libro IA'
        verbose_name_plural = 'Capítulos de Libros IA'
    
    def __str__(self):
        return f"{self.book.title} - Cap. {self.order}: {self.title}"
    
    def get_status_badge_class(self):
        """Retorna la clase Bootstrap del badge según el estado"""
        badges = {
            'proposed': 'bg-secondary',
            'summary_created': 'bg-info',
            'content_generated': 'bg-primary',
            'completed': 'bg-success',
        }
        return badges.get(self.status, 'bg-secondary')
    
    def update_word_count(self):
        """Actualiza el conteo de palabras del contenido"""
        if self.content:
            self.word_count = len(self.content.split())
        else:
            self.word_count = 0
        self.save(update_fields=['word_count'])
    
    def save(self, *args, **kwargs):
        """Override save para actualizar automáticamente el conteo de palabras"""
        if self.content:
            self.word_count = len(self.content.split())
        super().save(*args, **kwargs)


class AIArticleProject(models.Model):
    """Modelo para proyectos de generación de artículos con IA"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('articles_proposed', 'Artículos Propuestos'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completado'),
    ]
    
    FORMAT_CHOICES = [
        ('plain', 'Texto Plano'),
        ('markdown', 'Markdown'),
    ]
    
    title = models.CharField(
        max_length=500,
        verbose_name='Título del Proyecto'
    )
    main_topic = models.TextField(
        verbose_name='Tema Principal',
        help_text='Describe el tema principal sobre el cual quieres generar artículos'
    )
    content_format = models.CharField(
        max_length=20,
        choices=FORMAT_CHOICES,
        default='plain',
        verbose_name='Formato de Contenido',
        help_text='Formato en el que se generará el contenido de los artículos'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_article_projects',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='ai_article_projects',
        verbose_name='Empresa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Proyecto de Artículos IA'
        verbose_name_plural = 'Proyectos de Artículos IA'
    
    def __str__(self):
        return self.title
    
    def get_status_badge_class(self):
        """Retorna la clase Bootstrap del badge según el estado"""
        badges = {
            'draft': 'bg-secondary',
            'articles_proposed': 'bg-info',
            'in_progress': 'bg-warning',
            'completed': 'bg-success',
        }
        return badges.get(self.status, 'bg-secondary')
    
    def get_total_articles(self):
        """Retorna el total de artículos"""
        return self.articles.count()
    
    def get_completed_articles(self):
        """Retorna el número de artículos completados (con contenido generado)"""
        return self.articles.filter(status__in=['content_generated', 'completed']).count()
    
    def get_progress_percentage(self):
        """Retorna el porcentaje de progreso"""
        total = self.get_total_articles()
        if total == 0:
            return 0
        completed = self.get_completed_articles()
        return int((completed / total) * 100)
    
    def get_total_words(self):
        """Retorna el total de palabras de todos los artículos"""
        from django.db.models import Sum
        result = self.articles.aggregate(total=Sum('word_count'))
        return result['total'] or 0


class AIArticle(models.Model):
    """Modelo para artículos individuales generados con IA"""
    
    STATUS_CHOICES = [
        ('proposed', 'Propuesto'),
        ('content_generated', 'Contenido Generado'),
        ('completed', 'Completado'),
    ]
    
    project = models.ForeignKey(
        AIArticleProject,
        on_delete=models.CASCADE,
        related_name='articles',
        verbose_name='Proyecto'
    )
    title = models.CharField(
        max_length=500,
        verbose_name='Título del Artículo'
    )
    keywords = models.CharField(
        max_length=1000,
        blank=True,
        verbose_name='Palabras Clave',
        help_text='Palabras clave separadas por comas'
    )
    content = models.TextField(
        blank=True,
        verbose_name='Contenido del Artículo',
        help_text='Contenido completo en texto plano'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='proposed',
        verbose_name='Estado'
    )
    word_count = models.IntegerField(
        default=0,
        verbose_name='Conteo de Palabras'
    )
    order = models.IntegerField(
        default=0,
        verbose_name='Orden'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['project', 'order']
        verbose_name = 'Artículo IA'
        verbose_name_plural = 'Artículos IA'
    
    def __str__(self):
        return f"{self.project.title} - {self.title}"
    
    def get_status_badge_class(self):
        """Retorna la clase Bootstrap del badge según el estado"""
        badges = {
            'proposed': 'bg-secondary',
            'content_generated': 'bg-primary',
            'completed': 'bg-success',
        }
        return badges.get(self.status, 'bg-secondary')
    
    def get_keywords_list(self):
        """Retorna las palabras clave como lista"""
        if self.keywords:
            return [k.strip() for k in self.keywords.split(',') if k.strip()]
        return []
    
    def save(self, *args, **kwargs):
        """Override save para actualizar automáticamente el conteo de palabras"""
        if self.content:
            self.word_count = len(self.content.split())
        super().save(*args, **kwargs)


class EmployeeRequest(models.Model):
    """Modelo para las solicitudes de empleados"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('approved', 'Aprobada'),
        ('rejected', 'Rechazada'),
    ]
    
    sequence = models.PositiveIntegerField(
        verbose_name='Secuencia',
        help_text='Número de secuencia de la solicitud'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título de la solicitud'
    )
    
    text = models.TextField(
        verbose_name='Texto',
        help_text='Descripción detallada de la solicitud'
    )
    
    date = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha',
        help_text='Fecha de creación de la solicitud'
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado',
        help_text='Estado actual de la solicitud'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='employee_requests',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de Actualización'
    )
    
    class Meta:
        verbose_name = 'Solicitud de Empleado'
        verbose_name_plural = 'Solicitudes de Empleados'
        ordering = ['-sequence', '-created_at']
        unique_together = ['sequence']
    
    def __str__(self):
        return f"Solicitud #{self.sequence} - {self.title}"
    
    def get_status_display_class(self):
        """Retorna la clase CSS para el estado"""
        classes = {
            'draft': 'badge bg-secondary',
            'approved': 'badge bg-success',
            'rejected': 'badge bg-danger',
        }
        return classes.get(self.status, 'badge bg-secondary')
    
    def get_status_icon(self):
        """Retorna el ícono para el estado"""
        icons = {
            'draft': 'bi bi-file-earmark-text',
            'approved': 'bi bi-check-circle',
            'rejected': 'bi bi-x-circle',
        }
        return icons.get(self.status, 'bi bi-file-earmark-text')
    
    def save(self, *args, **kwargs):
        """Override save para generar automáticamente la secuencia"""
        if not self.sequence:
            # Obtener el último número de secuencia y agregar 1
            last_request = EmployeeRequest.objects.order_by('-sequence').first()
            self.sequence = (last_request.sequence + 1) if last_request else 1
        super().save(*args, **kwargs)


class InternalAgreement(models.Model):
    """Modelo para gestionar acuerdos internos de la empresa"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('review', 'En Revisión'),
        ('approved', 'Aprobado'),
        ('rejected', 'Rechazado'),
        ('active', 'Activo'),
        ('expired', 'Expirado'),
        ('cancelled', 'Cancelado'),
    ]
    
    TYPE_CHOICES = [
        ('policy', 'Política'),
        ('procedure', 'Procedimiento'),
        ('contract', 'Contrato'),
        ('guideline', 'Directriz'),
        ('agreement', 'Acuerdo'),
        ('regulation', 'Reglamento'),
    ]
    
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('critical', 'Crítica'),
    ]
    
    sequence = models.PositiveIntegerField(
        unique=True,
        verbose_name='Número de Secuencia',
        help_text='Número único auto-generado'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo del acuerdo interno'
    )
    
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del acuerdo interno'
    )
    
    content = models.TextField(
        verbose_name='Contenido',
        help_text='Contenido completo del acuerdo interno'
    )
    
    agreement_type = models.CharField(
        max_length=20,
        choices=TYPE_CHOICES,
        default='agreement',
        verbose_name='Tipo de Acuerdo'
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    
    priority = models.CharField(
        max_length=20,
        choices=PRIORITY_CHOICES,
        default='medium',
        verbose_name='Prioridad'
    )
    
    effective_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Vigencia',
        help_text='Fecha en que el acuerdo entra en vigencia'
    )
    
    expiration_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Expiración',
        help_text='Fecha en que el acuerdo expira (opcional)'
    )
    
    department = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Departamento',
        help_text='Departamento al que aplica el acuerdo'
    )
    
    applies_to_all = models.BooleanField(
        default=True,
        verbose_name='Aplica a Todos',
        help_text='Si el acuerdo aplica a todos los empleados'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='internal_agreements_created',
        verbose_name='Creado por'
    )
    
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='internal_agreements_approved',
        verbose_name='Aprobado por'
    )
    
    approved_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Aprobación'
    )
    
    # Campos para firmantes del acuerdo
    signer_1 = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='internal_agreements_signed_1',
        verbose_name='Primer Firmante',
        help_text='Usuario que firma el acuerdo como primer firmante'
    )
    
    signer_1_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Firma 1',
        help_text='Fecha en que el primer firmante firmó el acuerdo'
    )
    
    signer_2 = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='internal_agreements_signed_2',
        verbose_name='Segundo Firmante',
        help_text='Usuario que firma el acuerdo como segundo firmante'
    )
    
    signer_2_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Firma 2',
        help_text='Fecha en que el segundo firmante firmó el acuerdo'
    )
    
    tags = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas para facilitar búsquedas'
    )
    
    version = models.CharField(
        max_length=20,
        default='1.0',
        verbose_name='Versión',
        help_text='Versión del acuerdo'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Acuerdo Interno'
        verbose_name_plural = 'Acuerdos Internos'
        indexes = [
            models.Index(fields=['status', 'agreement_type']),
            models.Index(fields=['effective_date', 'expiration_date']),
        ]
    
    def __str__(self):
        return f"#{self.sequence} - {self.title}"
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge del estado"""
        classes = {
            'draft': 'badge bg-secondary',
            'review': 'badge bg-warning',
            'approved': 'badge bg-info',
            'rejected': 'badge bg-danger',
            'active': 'badge bg-success',
            'expired': 'badge bg-dark',
            'cancelled': 'badge bg-light text-dark',
        }
        return classes.get(self.status, 'badge bg-secondary')
    
    def get_status_icon(self):
        """Retorna el ícono para el estado"""
        icons = {
            'draft': 'bi bi-file-earmark-text',
            'review': 'bi bi-eye',
            'approved': 'bi bi-check-circle-fill',
            'rejected': 'bi bi-x-circle-fill',
            'active': 'bi bi-check-circle',
            'expired': 'bi bi-clock-history',
            'cancelled': 'bi bi-x-circle',
        }
        return icons.get(self.status, 'bi bi-file-earmark-text')
    
    def get_type_icon(self):
        """Retorna el ícono para el tipo de acuerdo"""
        icons = {
            'policy': 'bi bi-shield-check',
            'procedure': 'bi bi-list-check',
            'contract': 'bi bi-file-earmark-text',
            'guideline': 'bi bi-signpost',
            'agreement': 'bi bi-handshake',
            'regulation': 'bi bi-bookmark-check',
        }
        return icons.get(self.agreement_type, 'bi bi-file-earmark-check')
    
    def get_priority_badge_class(self):
        """Retorna la clase CSS para el badge de prioridad"""
        classes = {
            'low': 'badge bg-light text-dark',
            'medium': 'badge bg-info',
            'high': 'badge bg-warning text-dark',
            'critical': 'badge bg-danger',
        }
        return classes.get(self.priority, 'badge bg-info')
    
    def is_active(self):
        """Verifica si el acuerdo está activo y vigente"""
        if self.status != 'active':
            return False
        
        today = timezone.now().date()
        
        if self.effective_date and self.effective_date > today:
            return False
        
        if self.expiration_date and self.expiration_date < today:
            return False
        
        return True
    
    def can_edit(self, user):
        """Verifica si el usuario puede editar el acuerdo"""
        from .utils import is_agent
        return is_agent(user) or user.is_staff or user == self.created_by
    
    def can_delete(self, user):
        """Verifica si el usuario puede eliminar el acuerdo"""
        from .utils import is_agent
        return (is_agent(user) or user.is_staff or user == self.created_by) and self.status == 'draft'
    
    def is_signed_by(self, user):
        """Verifica si el usuario ya firmó el acuerdo"""
        return (self.signer_1 == user and self.signer_1_date) or \
               (self.signer_2 == user and self.signer_2_date)
    
    def can_sign(self, user, signer_number=None):
        """Verifica si el usuario puede firmar el acuerdo"""
        # Solo se puede firmar si está activo o aprobado y el usuario no ha firmado ya
        if self.status not in ['active', 'approved']:
            return False
        
        # Si ya firmó, no puede firmar de nuevo
        if self.is_signed_by(user):
            return False
        
        # Si se especifica número de firmante, verificar específicamente
        if signer_number is not None:
            if signer_number == 1:
                return user == self.signer_1 and not self.signer_1_date
            elif signer_number == 2:
                return user == self.signer_2 and not self.signer_2_date
            return False
        
        # Si no hay firmantes asignados, cualquier usuario autorizado puede firmar
        if not self.signer_1 and not self.signer_2:
            from .utils import is_agent
            return is_agent(user) or user.is_staff
        
        # Si hay firmantes asignados, solo ellos pueden firmar
        return user == self.signer_1 or user == self.signer_2
    
    def sign_agreement(self, user, signer_number=None):
        """Firma el acuerdo con el usuario especificado"""
        if signer_number is not None:
            # Firmar como un número específico de firmante
            if not self.can_sign(user, signer_number):
                return False
            
            from django.utils import timezone
            current_time = timezone.now()
            
            if signer_number == 1:
                self.signer_1_date = current_time
            elif signer_number == 2:
                self.signer_2_date = current_time
            else:
                return False
            
            self.save()
            return True
        else:
            # Lógica original para compatibilidad
            if not self.can_sign(user):
                return False
            
            from django.utils import timezone
            current_time = timezone.now()
            
            # Firmar en el primer slot disponible
            if not self.signer_1 or (self.signer_1 == user and not self.signer_1_date):
                self.signer_1 = user
                self.signer_1_date = current_time
            elif not self.signer_2 or (self.signer_2 == user and not self.signer_2_date):
                self.signer_2 = user
                self.signer_2_date = current_time
            else:
                return False
            
            self.save()
            return True
    
    def is_fully_signed(self):
        """Verifica si el acuerdo está completamente firmado"""
        return (self.signer_1 and self.signer_1_date and 
                self.signer_2 and self.signer_2_date)
    
    def get_signature_status(self):
        """Retorna el estado de las firmas"""
        if self.is_fully_signed():
            return "Completamente firmado"
        elif self.signer_1_date or self.signer_2_date:
            return "Parcialmente firmado"
        elif self.signer_1 or self.signer_2:
            return "Pendiente de firma"
        else:
            return "Sin firmantes asignados"
    
    def save(self, *args, **kwargs):
        """Override save para generar automáticamente la secuencia"""
        if not self.sequence:
            # Obtener el último número de secuencia y agregar 1
            last_agreement = InternalAgreement.objects.order_by('-sequence').first()
            self.sequence = (last_agreement.sequence + 1) if last_agreement else 1
        super().save(*args, **kwargs)


class TrainingPlan(models.Model):
    """Plan de Capacitación"""
    sequence = models.PositiveIntegerField(unique=True, verbose_name='Secuencia')
    title = models.CharField(max_length=200, verbose_name='Título del Plan')
    description = models.TextField(blank=True, verbose_name='Descripción')
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Fecha de actualización')
    created_by = models.ForeignKey(User, blank=True, null=True, on_delete=models.SET_NULL, verbose_name='Creado por')

    class Meta:
        verbose_name = 'Plan de Capacitación'
        verbose_name_plural = 'Planes de Capacitación'
        ordering = ['sequence']

    def __str__(self):
        return f"Plan #{self.sequence}: {self.title}"

    def get_total_links(self):
        """Retorna el total de enlaces en este plan"""
        return self.links.count()

    def get_progress_for_user(self, user):
        """Retorna el progreso del usuario en este plan"""
        total_links = self.get_total_links()
        if total_links == 0:
            return 0
        
        # Contar directamente los EmployeeLinkProgress del usuario para este plan
        from .models import EmployeeLinkProgress
        completed_links = EmployeeLinkProgress.objects.filter(
            employee=user,
            link__plan=self
        ).count()
        
        return (completed_links / total_links) * 100

    def is_completed_by_user(self, user):
        """Verifica si el usuario completó todo el plan"""
        try:
            progress = EmployeeTrainingProgress.objects.get(employee=user, plan=self)
            return progress.status == 'completed'
        except EmployeeTrainingProgress.DoesNotExist:
            return False


class TrainingLink(models.Model):
    """Enlace de Capacitación"""
    plan = models.ForeignKey(TrainingPlan, on_delete=models.CASCADE, related_name='links', verbose_name='Plan de Capacitación')
    title = models.CharField(max_length=200, verbose_name='Título del Enlace')
    url = models.URLField(verbose_name='URL')
    description = models.TextField(blank=True, verbose_name='Descripción')
    order = models.PositiveIntegerField(default=0, verbose_name='Orden')
    is_required = models.BooleanField(default=True, verbose_name='Requerido')
    estimated_hours = models.PositiveIntegerField(default=1, verbose_name='Horas estimadas')
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')

    class Meta:
        verbose_name = 'Enlace de Capacitación'
        verbose_name_plural = 'Enlaces de Capacitación'
        ordering = ['order', 'id']

    def __str__(self):
        return f"{self.plan.title} - {self.title}"

    def is_completed_by_user(self, user):
        """Verifica si el usuario completó este enlace"""
        return EmployeeLinkProgress.objects.filter(employee=user, link=self).exists()


class EmployeeTrainingProgress(models.Model):
    """Progreso de Capacitación por Empleado"""
    STATUS_CHOICES = [
        ('not_started', 'No iniciado'),
        ('in_progress', 'En progreso'),
        ('completed', 'Completado'),
    ]
    
    employee = models.ForeignKey(User, on_delete=models.CASCADE, related_name='training_progress', verbose_name='Empleado')
    plan = models.ForeignKey(TrainingPlan, on_delete=models.CASCADE, related_name='employee_progress', verbose_name='Plan de Capacitación')
    status = models.CharField(max_length=20, choices=STATUS_CHOICES, default='not_started', verbose_name='Estado')
    started_at = models.DateTimeField(blank=True, null=True, verbose_name='Iniciado el')
    completed_at = models.DateTimeField(blank=True, null=True, verbose_name='Completado el')
    notes = models.TextField(blank=True, verbose_name='Notas')

    class Meta:
        verbose_name = 'Progreso de Capacitación'
        verbose_name_plural = 'Progreso de Capacitaciones'
        unique_together = ('employee', 'plan')

    def __str__(self):
        return f"{self.employee.username} - {self.plan.title} ({self.get_status_display()})"

    def get_completion_percentage(self):
        """Calcula el porcentaje de completado"""
        total_links = self.plan.get_total_links()
        if total_links == 0:
            return 0
        completed_links = self.completed_links.count()
        return (completed_links / total_links) * 100

    def update_status(self):
        """Actualiza el estado basado en los enlaces completados"""
        total_links = self.plan.get_total_links()
        completed_links = self.completed_links.count()
        
        if completed_links == 0:
            self.status = 'not_started'
            self.started_at = None
            self.completed_at = None
        elif completed_links == total_links:
            self.status = 'completed'
            if not self.completed_at:
                from django.utils import timezone
                self.completed_at = timezone.now()
        else:
            self.status = 'in_progress'
            if not self.started_at:
                from django.utils import timezone
                self.started_at = timezone.now()
            self.completed_at = None
        
        self.save()


class EmployeeLinkProgress(models.Model):
    """Progreso individual de cada enlace por empleado"""
    employee = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Empleado')
    link = models.ForeignKey(TrainingLink, on_delete=models.CASCADE, related_name='employee_progress', verbose_name='Enlace')
    training_progress = models.ForeignKey(EmployeeTrainingProgress, on_delete=models.CASCADE, related_name='completed_links', verbose_name='Progreso de Capacitación')
    completed_at = models.DateTimeField(auto_now_add=True, verbose_name='Completado el')
    time_spent_hours = models.DecimalField(max_digits=5, decimal_places=2, blank=True, null=True, verbose_name='Tiempo invertido (horas)')
    notes = models.TextField(blank=True, verbose_name='Notas')

    class Meta:
        verbose_name = 'Progreso de Enlace'
        verbose_name_plural = 'Progreso de Enlaces'
        unique_together = ('employee', 'link')

    def __str__(self):
        return f"{self.employee.username} - {self.link.title}"

    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        # Actualizar el estado del progreso general del entrenamiento
        self.training_progress.update_status()


# ========== MODELOS DE RECOMENDACIONES IA ==========

class AIRecommendation(models.Model):
    """Modelo para gestionar recomendaciones de IA para respuestas políticamente correctas"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo de la recomendación'
    )
    context_text = models.TextField(
        verbose_name='Texto de Contexto',
        help_text='Contexto base que usará la IA para generar respuestas'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si está activa, podrá ser usada para generar respuestas'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_recommendations',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Recomendación IA'
        verbose_name_plural = 'Recomendaciones IA'
    
    def __str__(self):
        return self.title
    
    def get_usage_count(self):
        """Retorna el número de veces que se ha usado esta recomendación"""
        return self.usage_logs.count()


class AIRecommendationUsage(models.Model):
    """Modelo para registrar el uso de recomendaciones IA"""
    
    recommendation = models.ForeignKey(
        AIRecommendation,
        on_delete=models.CASCADE,
        related_name='usage_logs',
        verbose_name='Recomendación'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_recommendation_usage',
        verbose_name='Usuario'
    )
    client_case = models.TextField(
        verbose_name='Caso del Cliente',
        help_text='Caso específico o situación del cliente enviada para análisis'
    )
    ai_response = models.TextField(
        verbose_name='Respuesta de IA',
        help_text='Respuesta generada por la IA basada en el contexto y caso'
    )
    is_helpful = models.BooleanField(
        default=True,
        verbose_name='¿Fue útil?',
        help_text='Calificación de utilidad de la respuesta'
    )
    feedback = models.TextField(
        blank=True,
        verbose_name='Comentarios',
        help_text='Comentarios adicionales del usuario sobre la respuesta'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de uso'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Uso de Recomendación IA'
        verbose_name_plural = 'Usos de Recomendaciones IA'
    
    def __str__(self):
        return f"{self.recommendation.title} - {self.user.username} ({self.created_at.strftime('%d/%m/%Y')})"


# ========== MODELOS DE AUSENCIAS DE EMPLEADOS ==========

class AbsenceType(models.Model):
    """Tipos de ausencias disponibles"""
    
    name = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Nombre del Tipo',
        help_text='Ej: Vacaciones, Enfermedad, Permiso Personal'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción detallada del tipo de ausencia'
    )
    requires_documentation = models.BooleanField(
        default=False,
        verbose_name='Requiere Documentación',
        help_text='Si se requiere documentación de respaldo (ej: certificado médico)'
    )
    is_paid = models.BooleanField(
        default=True,
        verbose_name='Es Remunerado',
        help_text='Si la ausencia es con goce de sueldo'
    )
    max_days_per_year = models.PositiveIntegerField(
        blank=True,
        null=True,
        verbose_name='Máximo Días por Año',
        help_text='Límite anual de días para este tipo de ausencia (opcional)'
    )
    color = models.CharField(
        max_length=7,
        default='#007bff',
        verbose_name='Color',
        help_text='Color para mostrar en el calendario (formato hex: #RRGGBB)'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si este tipo de ausencia está disponible para uso'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['name']
        verbose_name = 'Tipo de Ausencia'
        verbose_name_plural = 'Tipos de Ausencias'
    
    def __str__(self):
        return self.name
    
    def get_usage_count(self):
        """Retorna el número de ausencias registradas de este tipo"""
        return self.absences.count()


class EmployeeAbsence(models.Model):
    """Registro de ausencias de empleados"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('approved', 'Aprobada'),
        ('rejected', 'Rechazada'),
        ('cancelled', 'Cancelada'),
    ]
    
    employee = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='absences',
        verbose_name='Empleado'
    )
    absence_type = models.ForeignKey(
        AbsenceType,
        on_delete=models.CASCADE,
        related_name='absences',
        verbose_name='Tipo de Ausencia'
    )
    start_date = models.DateField(
        verbose_name='Fecha de Inicio',
        help_text='Primer día de ausencia'
    )
    end_date = models.DateField(
        verbose_name='Fecha de Fin',
        help_text='Último día de ausencia'
    )
    lost_hours = models.DecimalField(
        max_digits=8,
        decimal_places=2,
        default=0.00,
        verbose_name='Horas Perdidas',
        help_text='Total de horas laborales perdidas por la ausencia'
    )
    reason = models.TextField(
        verbose_name='Motivo',
        help_text='Explicación detallada del motivo de la ausencia'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    documentation_file = models.FileField(
        upload_to='absence_documentation/%Y/%m/',
        blank=True,
        null=True,
        verbose_name='Documentación',
        help_text='Archivo de respaldo (certificado médico, etc.)'
    )
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='approved_absences',
        verbose_name='Aprobado por'
    )
    approved_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de aprobación'
    )
    rejection_reason = models.TextField(
        blank=True,
        verbose_name='Motivo de rechazo',
        help_text='Razón por la cual se rechazó la ausencia'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas adicionales',
        help_text='Comentarios internos sobre la ausencia'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de solicitud'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-start_date', '-created_at']
        verbose_name = 'Ausencia de Empleado'
        verbose_name_plural = 'Ausencias de Empleados'
    
    def __str__(self):
        return f"{self.employee.get_full_name() or self.employee.username} - {self.absence_type.name} ({self.start_date} a {self.end_date})"
    
    def clean(self):
        """Validación personalizada"""
        from django.core.exceptions import ValidationError
        
        if self.start_date and self.end_date:
            if self.start_date > self.end_date:
                raise ValidationError('La fecha de inicio no puede ser posterior a la fecha de fin.')
    
    def save(self, *args, **kwargs):
        self.full_clean()
        super().save(*args, **kwargs)
    
    def get_duration_days(self):
        """Retorna la duración en días de la ausencia"""
        if self.start_date and self.end_date:
            return (self.end_date - self.start_date).days + 1
        return 0
    
    def calculate_lost_hours(self, hours_per_day=8):
        """Calcula las horas perdidas basado en los días de ausencia"""
        duration_days = self.get_duration_days()
        return duration_days * hours_per_day
    
    def get_lost_hours_display(self):
        """Retorna las horas perdidas formateadas"""
        if self.lost_hours:
            return f"{self.lost_hours} horas"
        return "0 horas"
    
    def get_average_hours_per_day(self):
        """Retorna el promedio de horas perdidas por día"""
        duration_days = self.get_duration_days()
        if duration_days > 0 and self.lost_hours:
            return round(float(self.lost_hours) / duration_days, 1)
        return 0.0
    
    def get_status_display_class(self):
        """Retorna la clase CSS para el estado"""
        status_classes = {
            'pending': 'warning',
            'approved': 'success',
            'rejected': 'danger',
            'cancelled': 'secondary',
        }
        return status_classes.get(self.status, 'secondary')
    
    def is_current(self):
        """Verifica si la ausencia está actualmente en curso"""
        from django.utils import timezone
        today = timezone.now().date()
        return self.start_date <= today <= self.end_date and self.status == 'approved'
    
    def approve(self, approved_by_user):
        """Aprueba la ausencia"""
        self.status = 'approved'
        self.approved_by = approved_by_user
        self.approved_at = timezone.now()
        self.save()
    
    def reject(self, reason, rejected_by_user):
        """Rechaza la ausencia"""
        self.status = 'rejected'
        self.rejection_reason = reason
        self.approved_by = rejected_by_user
        self.approved_at = timezone.now()
        self.save()


class CompanyProtocol(models.Model):
    """Modelo para gestionar protocolos de empresa"""
    
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='protocols',
        verbose_name='Empresa',
        help_text='Empresa a la que pertenece este protocolo',
        null=True,
        blank=True
    )
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Protocolo',
        help_text='Título descriptivo del protocolo empresarial'
    )
    content = models.TextField(
        verbose_name='Contenido del Protocolo',
        help_text='Contenido completo del protocolo. Puedes usar formato básico.'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_protocols',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Indica si el protocolo está activo y disponible'
    )
    version = models.CharField(
        max_length=10,
        default='1.0',
        verbose_name='Versión',
        help_text='Versión del protocolo (ej: 1.0, 1.1, 2.0)'
    )
    is_public = models.BooleanField(
        default=False,
        verbose_name='Compartir públicamente',
        help_text='Permite que este protocolo sea accesible públicamente sin autenticación'
    )
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        verbose_name='UUID público',
        help_text='Identificador único para acceso público'
    )
    public_shared_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Compartido públicamente en',
        help_text='Fecha cuando se compartió públicamente por primera vez'
    )
    public_views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Vistas públicas',
        help_text='Número de veces que se ha visto públicamente este protocolo'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Protocolo de Empresa'
        verbose_name_plural = 'Protocolos de Empresa'
    
    def __str__(self):
        return f"{self.title} (v{self.version})"
    
    def get_short_content(self, max_length=100):
        """Retorna una versión truncada del contenido"""
        if len(self.content) <= max_length:
            return self.content
        return self.content[:max_length] + '...'
    
    def get_word_count(self):
        """Retorna el número de palabras en el contenido"""
        return len(self.content.split())
    
    def increment_version(self):
        """Incrementa automáticamente la versión del protocolo"""
        try:
            version_parts = self.version.split('.')
            major = int(version_parts[0])
            minor = int(version_parts[1]) if len(version_parts) > 1 else 0
            minor += 1
            if minor >= 10:
                major += 1
                minor = 0
            self.version = f"{major}.{minor}"
        except (ValueError, IndexError):
            self.version = "1.0"
    
    def make_public(self):
        """Hace el protocolo público y registra la fecha"""
        if not self.is_public:
            self.is_public = True
            self.public_shared_at = timezone.now()
            self.save()
    
    def make_private(self):
        """Hace el protocolo privado"""
        self.is_public = False
        self.save()
    
    def get_public_url(self):
        """Retorna la URL pública del protocolo"""
        if self.is_public:
            return f"/public/protocols/{self.public_uuid}/"
        return None
    
    def increment_public_views(self):
        """Incrementa el contador de vistas públicas"""
        self.public_views_count += 1
        self.save(update_fields=['public_views_count'])
    
    def get_public_stats(self):
        """Retorna estadísticas de visualización pública"""
        return {
            'total_views': self.public_views_count,
            'is_public': self.is_public,
            'shared_date': self.public_shared_at,
            'public_url': self.get_public_url()
        }
    
    def generate_ai_content(self, content_type="comprehensive"):
        """Genera contenido automático del protocolo usando IA"""
        from tickets.models import SystemConfiguration
        import re
        
        def remove_markdown_formatting(text):
            """Elimina formato Markdown y devuelve texto plano"""
            # Eliminar negritas **texto** y dejar solo texto
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Eliminar cursivas *texto* y dejar solo texto
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Eliminar títulos ## y # y dejar solo el texto
            text = re.sub(r'^#{1,6}\s*(.*?)$', r'\1', text, flags=re.MULTILINE)
            # Eliminar listas numeradas 1. y dejar solo el texto
            text = re.sub(r'^\d+\.\s*', '', text, flags=re.MULTILINE)
            # Eliminar listas con viñetas - y * y dejar solo el texto
            text = re.sub(r'^[-*]\s*', '', text, flags=re.MULTILINE)
            return text
            text = re.sub(r'\n(\d+)\.\s+', r'</p><ol><li>', text)
            # Convertir listas con viñetas
            text = re.sub(r'\n[-*]\s+', r'</p><ul><li>', text)
            # Limpiar párrafos vacíos
            text = re.sub(r'<p>\s*</p>', '', text)
            return text
        
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            return {"success": False, "error": "OpenAI no está configurado en el sistema"}
        
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            company_context = f" para la empresa {self.company.name}" if self.company else ""
            
            prompts = {
                "comprehensive": f"""
                Genera un protocolo empresarial completo y profesional con el título: "{self.title}"{company_context}.
                
                El protocolo debe incluir:
                1. Introducción y propósito
                2. Alcance y aplicabilidad
                3. Responsabilidades
                4. Procedimientos detallados paso a paso
                5. Requisitos y estándares
                6. Medidas de seguridad (si aplica)
                7. Documentación y registros
                8. Revisión y actualización
                
                Escribe en español, usando un lenguaje profesional y claro. El protocolo debe ser práctico y actionable.
                NO uses formato Markdown (**, *, etc.). Usa texto plano con saltos de línea para separar secciones.
                """,
                "summary": f"""
                Basándote en el título "{self.title}"{company_context}, genera un resumen ejecutivo y puntos clave 
                para un protocolo empresarial. Incluye objetivos principales, beneficios esperados y requisitos básicos.
                Máximo 500 palabras.
                NO uses formato Markdown. Usa texto plano con saltos de línea.
                """,
                "checklist": f"""
                Crea una lista de verificación (checklist) detallada para el protocolo "{self.title}"{company_context}.
                Incluye todos los pasos necesarios, verificaciones de seguridad, y puntos de control de calidad.
                Formato de lista numerada y clara.
                NO uses formato Markdown. Usa texto plano con numeración simple.
                """
            }
            
            response = client.chat.completions.create(
                model=config.openai_model or "gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un experto consultor empresarial especializado en crear protocolos y procedimientos corporativos de alta calidad. NO uses formato Markdown en tus respuestas, solo texto plano con saltos de línea."},
                    {"role": "user", "content": prompts.get(content_type, prompts["comprehensive"])}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            generated_content = response.choices[0].message.content.strip()
            # Limpiar cualquier formato Markdown y devolver texto plano
            generated_content = remove_markdown_formatting(generated_content)
            
            return {
                "success": True,
                "content": generated_content,
                "word_count": len(generated_content.split()),
                "content_type": content_type
            }
            
        except Exception as e:
            return {"success": False, "error": f"Error generando contenido: {str(e)}"}
    
    def analyze_readability(self):
        """Analiza la legibilidad del protocolo"""
        import re
        
        try:
            text = self.content
            if not text:
                return {"success": False, "error": "No hay contenido para analizar"}
            
            # Métricas básicas
            word_count = len(text.split())
            sentence_count = len(re.split(r'[.!?]+', text))
            paragraph_count = len([p for p in text.split('\n\n') if p.strip()])
            
            # Promedio de palabras por oración
            avg_words_per_sentence = word_count / max(sentence_count, 1)
            
            # Promedio de oraciones por párrafo
            avg_sentences_per_paragraph = sentence_count / max(paragraph_count, 1)
            
            # Análisis de complejidad
            complex_words = len([word for word in text.split() if len(word) > 6])
            complexity_ratio = complex_words / max(word_count, 1) * 100
            
            # Puntuación de legibilidad (escala 1-10)
            readability_score = 10
            
            if avg_words_per_sentence > 20:
                readability_score -= 2
            elif avg_words_per_sentence > 15:
                readability_score -= 1
                
            if complexity_ratio > 30:
                readability_score -= 2
            elif complexity_ratio > 20:
                readability_score -= 1
                
            if avg_sentences_per_paragraph > 8:
                readability_score -= 1
            
            readability_score = max(1, readability_score)
            
            # Nivel de legibilidad
            if readability_score >= 8:
                level = "Excelente"
                color = "success"
            elif readability_score >= 6:
                level = "Bueno"
                color = "info"
            elif readability_score >= 4:
                level = "Regular"
                color = "warning"
            else:
                level = "Difícil"
                color = "danger"
            
            return {
                "success": True,
                "score": readability_score,
                "level": level,
                "color": color,
                "metrics": {
                    "word_count": word_count,
                    "sentence_count": sentence_count,
                    "paragraph_count": paragraph_count,
                    "avg_words_per_sentence": round(avg_words_per_sentence, 1),
                    "avg_sentences_per_paragraph": round(avg_sentences_per_paragraph, 1),
                    "complexity_ratio": round(complexity_ratio, 1)
                }
            }
            
        except Exception as e:
            return {"success": False, "error": f"Error analizando legibilidad: {str(e)}"}
    
    def get_ai_suggestions(self):
        """Genera sugerencias de mejora usando IA"""
        from tickets.models import SystemConfiguration
        import re
        
        def remove_markdown_formatting(text):
            """Elimina formato Markdown y devuelve texto plano"""
            # Eliminar negritas **texto** y dejar solo texto
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Eliminar cursivas *texto* y dejar solo texto
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Eliminar títulos ## y # y dejar solo el texto
            text = re.sub(r'^#{1,6}\s*(.*?)$', r'\1', text, flags=re.MULTILINE)
            # Eliminar listas numeradas 1. y dejar solo el texto
            text = re.sub(r'^\d+\.\s*', '', text, flags=re.MULTILINE)
            # Eliminar listas con viñetas - y * y dejar solo el texto
            text = re.sub(r'^[-*]\s*', '', text, flags=re.MULTILINE)
            return text
            return text
        
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            return {"success": False, "error": "OpenAI no está configurado en el sistema"}
        
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            # Análisis de legibilidad primero
            readability = self.analyze_readability()
            
            prompt = f"""
            Analiza el siguiente protocolo empresarial y proporciona sugerencias específicas de mejora:
            
            TÍTULO: {self.title}
            EMPRESA: {self.company.name if self.company else 'No especificada'}
            
            CONTENIDO:
            {self.content[:1500]}...
            
            Proporciona sugerencias en las siguientes categorías:
            1. CLARIDAD: Cómo mejorar la claridad del lenguaje
            2. ESTRUCTURA: Sugerencias de organización del contenido
            3. COMPLETITUD: Qué elementos podrían estar faltando
            4. PROFESIONALISMO: Cómo hacer el protocolo más profesional
            5. ACCIONABILIDAD: Cómo hacer las instrucciones más específicas
            
            Proporciona máximo 3 sugerencias por categoría, sean concretas y actionables.
            NO uses formato Markdown. Usa texto plano con saltos de línea.
            """
            
            response = client.chat.completions.create(
                model=config.openai_model or "gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un consultor experto en documentación empresarial y mejora de procesos corporativos. NO uses formato Markdown en tus respuestas, solo texto plano con saltos de línea."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=1000,
                temperature=0.6
            )
            
            suggestions = response.choices[0].message.content.strip()
            # Limpiar cualquier formato Markdown y devolver texto plano
            suggestions = remove_markdown_formatting(suggestions)
            
            return {
                "success": True,
                "suggestions": suggestions,
                "readability_info": readability if readability["success"] else None
            }
            
        except Exception as e:
            return {"success": False, "error": f"Error generando sugerencias: {str(e)}"}
    
    def generate_executive_summary(self):
        """Genera un resumen ejecutivo del protocolo"""
        from tickets.models import SystemConfiguration
        import re
        
        def remove_markdown_formatting(text):
            """Elimina formato Markdown y devuelve texto plano"""
            # Eliminar negritas **texto** y dejar solo texto
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
            # Eliminar cursivas *texto* y dejar solo texto
            text = re.sub(r'\*(.*?)\*', r'\1', text)
            # Eliminar títulos ## y # y dejar solo el texto
            text = re.sub(r'^#{1,6}\s*(.*?)$', r'\1', text, flags=re.MULTILINE)
            # Eliminar listas numeradas 1. y dejar solo el texto
            text = re.sub(r'^\d+\.\s*', '', text, flags=re.MULTILINE)
            # Eliminar listas con viñetas - y * y dejar solo el texto
            text = re.sub(r'^[-*]\s*', '', text, flags=re.MULTILINE)
            return text
        
        config = SystemConfiguration.objects.first()
        if not config or not config.ai_chat_enabled or not config.openai_api_key:
            return {"success": False, "error": "OpenAI no está configurado en el sistema"}
        
        try:
            import openai
            client = openai.OpenAI(api_key=config.openai_api_key)
            
            prompt = f"""
            Crea un resumen ejecutivo profesional del siguiente protocolo:
            
            TÍTULO: {self.title}
            EMPRESA: {self.company.name if self.company else 'No especificada'}
            
            CONTENIDO:
            {self.content}
            
            El resumen ejecutivo debe incluir:
            1. Propósito y objetivos principales
            2. Beneficios clave
            3. Impacto en la organización
            4. Recursos necesarios
            5. Cronograma de implementación (si aplica)
            
            Máximo 300 palabras, lenguaje ejecutivo y persuasivo.
            NO uses formato Markdown. Usa texto plano con saltos de línea.
            """
            
            response = client.chat.completions.create(
                model=config.openai_model or "gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "Eres un director ejecutivo escribiendo resúmenes para la junta directiva. NO uses formato Markdown en tus respuestas, solo texto plano con saltos de línea."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.5
            )
            
            summary = response.choices[0].message.content.strip()
            # Limpiar cualquier formato Markdown y devolver texto plano
            summary = remove_markdown_formatting(summary)
            
            return {
                "success": True,
                "summary": summary,
                "word_count": len(summary.split())
            }
            
        except Exception as e:
            return {"success": False, "error": f"Error generando resumen: {str(e)}"}


class QAComplaint(models.Model):
    """Modelo para gestionar quejas y sugerencias de QA asociadas a empresas"""
    
    TYPE_CHOICES = [
        ('complaint', 'Queja'),
        ('suggestion', 'Sugerencia'),
        ('feedback', 'Retroalimentación'),
        ('improvement', 'Mejora'),
        ('quality_issue', 'Problema de Calidad'),
    ]
    
    PRIORITY_CHOICES = [
        ('low', 'Baja'),
        ('medium', 'Media'),
        ('high', 'Alta'),
        ('critical', 'Crítica'),
    ]
    
    STATUS_CHOICES = [
        ('open', 'Abierto'),
        ('in_progress', 'En Progreso'),
        ('pending_review', 'Pendiente de Revisión'),
        ('resolved', 'Resuelto'),
        ('closed', 'Cerrado'),
        ('rejected', 'Rechazado'),
    ]
    
    # Campos principales
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo de la queja o sugerencia'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del problema o sugerencia'
    )
    type = models.CharField(
        max_length=20,
        choices=TYPE_CHOICES,
        default='complaint',
        verbose_name='Tipo',
        help_text='Tipo de reporte: queja, sugerencia, etc.'
    )
    priority = models.CharField(
        max_length=10,
        choices=PRIORITY_CHOICES,
        default='medium',
        verbose_name='Prioridad',
        help_text='Nivel de prioridad del reporte'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='open',
        verbose_name='Estado',
        help_text='Estado actual del reporte'
    )
    
    # Relaciones
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='qa_complaints',
        verbose_name='Empresa',
        help_text='Empresa asociada a esta queja o sugerencia'
    )
    reported_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='reported_complaints',
        verbose_name='Reportado por',
        help_text='Usuario que reportó la queja o sugerencia'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='assigned_complaints',
        verbose_name='Asignado a',
        help_text='Usuario responsable de resolver la queja'
    )
    
    # Campos adicionales
    contact_email = models.EmailField(
        blank=True,
        verbose_name='Email de contacto',
        help_text='Email del cliente o usuario que reporta (opcional)'
    )
    contact_phone = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Teléfono de contacto',
        help_text='Teléfono del cliente o usuario que reporta (opcional)'
    )
    
    # Campos de seguimiento
    resolution = models.TextField(
        blank=True,
        verbose_name='Resolución',
        help_text='Descripción de cómo se resolvió la queja o implementó la sugerencia'
    )
    resolution_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de resolución'
    )
    resolved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='resolved_complaints',
        verbose_name='Resuelto por',
        help_text='Usuario que resolvió la queja'
    )
    
    # Campos de auditoría
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Campos de archivo
    attachment = models.FileField(
        upload_to='qa_complaints/%Y/%m/',
        blank=True,
        null=True,
        verbose_name='Archivo adjunto',
        help_text='Imagen, documento o archivo relacionado con la queja'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Queja/Sugerencia QA'
        verbose_name_plural = 'Quejas/Sugerencias QA'
    
    def __str__(self):
        return f"{self.get_type_display()}: {self.title} ({self.company.name})"
    
    def get_status_color(self):
        """Retorna el color Bootstrap para el estado"""
        colors = {
            'open': 'danger',
            'in_progress': 'warning',
            'pending_review': 'info',
            'resolved': 'success',
            'closed': 'secondary',
            'rejected': 'dark',
        }
        return colors.get(self.status, 'secondary')
    
    def get_priority_color(self):
        """Retorna el color Bootstrap para la prioridad"""
        colors = {
            'low': 'success',
            'medium': 'warning',
            'high': 'danger',
            'critical': 'dark',
        }
        return colors.get(self.priority, 'secondary')
    
    def get_type_icon(self):
        """Retorna el ícono Bootstrap para el tipo"""
        icons = {
            'complaint': 'exclamation-triangle',
            'suggestion': 'lightbulb',
            'feedback': 'chat-dots',
            'improvement': 'arrow-up-circle',
            'quality_issue': 'bug',
        }
        return icons.get(self.type, 'file-text')
    
    def is_overdue(self):
        """Verifica si la queja está vencida (más de 7 días sin resolver)"""
        if self.status in ['resolved', 'closed']:
            return False
        
        from datetime import timedelta
        overdue_date = self.created_at + timedelta(days=7)
        return timezone.now() > overdue_date
    
    def get_age_days(self):
        """Retorna la antigüedad en días"""
        return (timezone.now() - self.created_at).days
    
    def can_edit(self, user):
        """Verifica si un usuario puede editar esta queja"""
        from .utils import is_agent
        return (
            user.is_staff or 
            user.is_superuser or 
            user == self.reported_by or 
            user == self.assigned_to or
            is_agent(user)
        )
    
    def can_resolve(self, user):
        """Verifica si un usuario puede resolver esta queja"""
        from .utils import is_agent
        return (
            user.is_staff or 
            user.is_superuser or 
            user == self.assigned_to or
            is_agent(user)
        )
    
    def resolve(self, resolution_text, resolved_by):
        """Marca la queja como resuelta"""
        self.status = 'resolved'
        self.resolution = resolution_text
        self.resolution_date = timezone.now()
        self.resolved_by = resolved_by
        if not self.assigned_to:
            self.assigned_to = resolved_by
        self.save()
    
    def assign_to(self, user):
        """Asigna la queja a un usuario"""
        self.assigned_to = user
        if self.status == 'open':
            self.status = 'in_progress'
        self.save()


class QAComplaintComment(models.Model):
    """Modelo para comentarios en quejas y sugerencias QA"""
    
    complaint = models.ForeignKey(
        QAComplaint,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Queja/Sugerencia'
    )
    author = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Autor'
    )
    content = models.TextField(
        verbose_name='Comentario',
        help_text='Contenido del comentario o nota'
    )
    is_internal = models.BooleanField(
        default=False,
        verbose_name='Comentario interno',
        help_text='Marca si es un comentario interno (no visible para el cliente)'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Comentario QA'
        verbose_name_plural = 'Comentarios QA'
    
    def __str__(self):
        return f"Comentario de {self.author.username} en {self.complaint.title[:50]}"


# ==================== MODELOS DE CONTROL DE ACTIVOS ====================

class Asset(models.Model):
    """Modelo para gestionar activos de la empresa"""
    STATUS_CHOICES = [
        ('available', 'Disponible'),
        ('assigned', 'Asignado'),
        ('maintenance', 'En Mantenimiento'),
        ('retired', 'Dado de Baja'),
        ('lost', 'Perdido'),
        ('damaged', 'Dañado'),
    ]
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Activo',
        help_text='Nombre descriptivo del activo'
    )
    manufacturer = models.CharField(
        max_length=100,
        verbose_name='Fabricante',
        help_text='Marca o fabricante del activo'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción detallada del activo'
    )
    serial_number = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Número de Serie',
        help_text='Número de serie único del activo'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='available',
        verbose_name='Estado'
    )
    assigned_to = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='assigned_assets',
        verbose_name='Asignado a'
    )
    purchase_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Compra'
    )
    purchase_price = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Precio de Compra'
    )
    warranty_expiry = models.DateField(
        null=True,
        blank=True,
        verbose_name='Vencimiento de Garantía'
    )
    location = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Ubicación',
        help_text='Ubicación física del activo'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el activo'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='created_assets',
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Activo'
        verbose_name_plural = 'Activos'
        indexes = [
            models.Index(fields=['serial_number']),
            models.Index(fields=['status']),
            models.Index(fields=['assigned_to']),
        ]
    
    def __str__(self):
        return f"{self.name} ({self.serial_number})"
    
    def get_status_display_color(self):
        """Retorna el color CSS para el estado"""
        colors = {
            'available': 'success',
            'assigned': 'primary',
            'maintenance': 'warning',
            'retired': 'secondary',
            'lost': 'danger',
            'damaged': 'danger',
        }
        return colors.get(self.status, 'secondary')
    
    def is_available(self):
        """Verifica si el activo está disponible para asignación"""
        return self.status == 'available'
    
    def can_be_assigned(self):
        """Verifica si el activo puede ser asignado"""
        return self.status in ['available', 'assigned']
    
    def save(self, *args, **kwargs):
        """Override save para actualizar estado automáticamente"""
        # Si se está asignando a un usuario y el estado es 'available'
        if self.assigned_to and self.status == 'available':
            self.status = 'assigned'
        # Si se desasigna (assigned_to = None) y el estado es 'assigned'
        elif not self.assigned_to and self.status == 'assigned':
            self.status = 'available'
        
        super().save(*args, **kwargs)


class AssetHistory(models.Model):
    """Modelo para el historial de asignaciones de activos"""
    ACTION_CHOICES = [
        ('created', 'Creado'),
        ('assigned', 'Asignado'),
        ('unassigned', 'Desasignado'),
        ('transferred', 'Transferido'),
        ('maintenance', 'Enviado a Mantenimiento'),
        ('returned', 'Devuelto de Mantenimiento'),
        ('retired', 'Dado de Baja'),
        ('status_changed', 'Cambio de Estado'),
    ]
    
    asset = models.ForeignKey(
        Asset,
        on_delete=models.CASCADE,
        related_name='history',
        verbose_name='Activo'
    )
    action = models.CharField(
        max_length=20,
        choices=ACTION_CHOICES,
        verbose_name='Acción'
    )
    previous_employee = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='asset_history_previous',
        verbose_name='Usuario Anterior'
    )
    new_employee = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='asset_history_new',
        verbose_name='Nuevo Usuario'
    )
    previous_status = models.CharField(
        max_length=20,
        choices=Asset.STATUS_CHOICES,
        null=True,
        blank=True,
        verbose_name='Estado Anterior'
    )
    new_status = models.CharField(
        max_length=20,
        choices=Asset.STATUS_CHOICES,
        null=True,
        blank=True,
        verbose_name='Nuevo Estado'
    )
    reason = models.TextField(
        blank=True,
        verbose_name='Motivo',
        help_text='Motivo del cambio o acción'
    )
    date = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha'
    )
    performed_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='asset_actions',
        verbose_name='Realizado por'
    )
    
    class Meta:
        ordering = ['-date']
        verbose_name = 'Historial de Activo'
        verbose_name_plural = 'Historiales de Activos'
        indexes = [
            models.Index(fields=['asset', '-date']),
            models.Index(fields=['action']),
        ]
    
    def __str__(self):
        return f"{self.asset.name} - {self.get_action_display()} ({self.date.strftime('%d/%m/%Y')})"
    
    def get_action_display_color(self):
        """Retorna el color CSS para la acción"""
        colors = {
            'created': 'success',
            'assigned': 'primary',
            'unassigned': 'warning',
            'transferred': 'info',
            'maintenance': 'warning',
            'returned': 'success',
            'retired': 'secondary',
            'status_changed': 'info',
        }
        return colors.get(self.action, 'secondary')


# ==========================================
# MODELOS PARA TUTOR IA
# ==========================================

class AITutor(models.Model):
    """Modelo para el sistema de Tutor IA"""
    
    STATUS_CHOICES = [
        ('active', 'Activo'),
        ('completed', 'Completado'),
        ('paused', 'Pausado'),
        ('cancelled', 'Cancelado'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Objetivo',
        help_text='Título descriptivo de lo que quieres lograr'
    )
    
    objective = models.TextField(
        verbose_name='Objetivo',
        help_text='Describe claramente qué quieres lograr'
    )
    
    current_state = models.TextField(
        verbose_name='Estado Actual',
        help_text='Describe dónde estás ahora en relación a tu objetivo'
    )
    
    context = models.TextField(
        verbose_name='Contexto',
        help_text='Proporciona contexto adicional relevante (experiencia, recursos, limitaciones, etc.)'
    )
    
    tutor_persona = models.TextField(
        verbose_name='Personalidad del Tutor',
        help_text='Describe qué tipo de tutor quieres (ej: un experimentado profesor de Django, un mentor de carrera, etc.)',
        default='Un experimentado mentor y profesor'
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='active',
        verbose_name='Estado'
    )
    
    is_private = models.BooleanField(
        default=False,
        verbose_name='Privado',
        help_text='Si está marcado, solo tú podrás ver este tutor'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='ai_tutors',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última Actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Tutor IA'
        verbose_name_plural = 'Tutores IA'
        indexes = [
            models.Index(fields=['created_by']),
            models.Index(fields=['status']),
            models.Index(fields=['is_private']),
        ]
    
    def __str__(self):
        return f"{self.title} - {self.created_by.username}"
    
    def get_status_display_color(self):
        """Retorna el color CSS para el estado"""
        colors = {
            'active': 'success',
            'completed': 'primary',
            'paused': 'warning',
            'cancelled': 'danger',
        }
        return colors.get(self.status, 'secondary')
    
    def get_progress_count(self):
        """Retorna el número de reportes de progreso"""
        return self.progress_reports.count()


class AITutorProgressReport(models.Model):
    """Modelo para reportar avances al Tutor IA"""
    
    tutor = models.ForeignKey(
        AITutor,
        on_delete=models.CASCADE,
        related_name='progress_reports',
        verbose_name='Tutor IA'
    )
    
    progress_description = models.TextField(
        verbose_name='Descripción del Avance',
        help_text='Describe qué has logrado desde tu último reporte'
    )
    
    challenges = models.TextField(
        blank=True,
        verbose_name='Desafíos Encontrados',
        help_text='Describe los obstáculos o dificultades que has enfrentado (opcional)'
    )
    
    ai_feedback = models.TextField(
        blank=True,
        verbose_name='Retroalimentación de la IA',
        help_text='Recomendaciones y consejos generados por la IA'
    )
    
    ai_recommendations = models.TextField(
        blank=True,
        verbose_name='Recomendaciones de Mejora',
        help_text='Sugerencias específicas para mejorar'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='tutor_progress_reports',
        verbose_name='Reportado por'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha del Reporte'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Reporte de Progreso'
        verbose_name_plural = 'Reportes de Progreso'
        indexes = [
            models.Index(fields=['tutor']),
            models.Index(fields=['created_by']),
        ]
    
    def __str__(self):
        return f"Reporte {self.pk} - {self.tutor.title}"


class AITutorAttachment(models.Model):
    """Modelo para adjuntos del Tutor IA"""
    
    tutor = models.ForeignKey(
        AITutor,
        on_delete=models.CASCADE,
        related_name='attachments',
        verbose_name='Tutor IA',
        null=True,
        blank=True
    )
    
    progress_report = models.ForeignKey(
        AITutorProgressReport,
        on_delete=models.CASCADE,
        related_name='attachments',
        verbose_name='Reporte de Progreso',
        null=True,
        blank=True
    )
    
    file = models.FileField(
        upload_to='tutor_attachments/%Y/%m/',
        verbose_name='Archivo'
    )
    
    original_filename = models.CharField(
        max_length=255,
        verbose_name='Nombre Original'
    )
    
    file_size = models.PositiveIntegerField(
        verbose_name='Tamaño del Archivo',
        help_text='Tamaño en bytes'
    )
    
    file_type = models.CharField(
        max_length=100,
        verbose_name='Tipo de Archivo'
    )
    
    description = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional del archivo'
    )
    
    ai_analysis = models.TextField(
        blank=True,
        verbose_name='Análisis de la IA',
        help_text='Análisis del archivo generado por la IA'
    )
    
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='tutor_attachments',
        verbose_name='Subido por'
    )
    
    uploaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Subida'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Adjunto del Tutor'
        verbose_name_plural = 'Adjuntos del Tutor'
        indexes = [
            models.Index(fields=['tutor']),
            models.Index(fields=['progress_report']),
            models.Index(fields=['uploaded_by']),
        ]
    
    def __str__(self):
        return f"{self.original_filename}"
    
    def get_file_size_display(self):
        """Formato legible del tamaño del archivo"""
        size_bytes = self.file_size
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024**2:
            return f"{size_bytes/1024:.1f} KB"
        elif size_bytes < 1024**3:
            return f"{size_bytes/(1024**2):.1f} MB"
        else:
            return f"{size_bytes/(1024**3):.1f} GB"
    
    def save(self, *args, **kwargs):
        if self.file:
            self.original_filename = self.file.name
            self.file_size = self.file.size
            self.file_type = self.file.name.split('.')[-1].lower() if '.' in self.file.name else 'unknown'
        super().save(*args, **kwargs)


class ExpenseReport(models.Model):
    """Modelo para rendición de gastos de empleados"""
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('submitted', 'Enviado'),
        ('approved', 'Aprobado'),
        ('rejected', 'Rechazado'),
        ('paid', 'Pagado'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la rendición',
        help_text='Título descriptivo de la rendición de gastos'
    )
    description = models.TextField(
        blank=True,
        null=True,
        verbose_name='Descripción',
        help_text='Descripción general de los gastos'
    )
    employee = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='expense_reports',
        verbose_name='Empleado'
    )
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='expense_reports',
        verbose_name='Empresa',
        blank=True,
        null=True
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    start_date = models.DateField(
        verbose_name='Fecha inicio',
        help_text='Fecha de inicio del período de gastos'
    )
    end_date = models.DateField(
        verbose_name='Fecha fin',
        help_text='Fecha de fin del período de gastos'
    )
    total_amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0,
        verbose_name='Monto total'
    )
    currency = models.CharField(
        max_length=3,
        default='USD',
        verbose_name='Moneda'
    )
    submitted_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de envío'
    )
    approved_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        related_name='approved_expense_reports',
        verbose_name='Aprobado por'
    )
    approved_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de aprobación'
    )
    rejection_reason = models.TextField(
        blank=True,
        null=True,
        verbose_name='Motivo de rechazo'
    )
    payment_reference = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Referencia de pago'
    )
    paid_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de pago'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Rendición de Gastos'
        verbose_name_plural = 'Rendiciones de Gastos'
        indexes = [
            models.Index(fields=['employee', 'status']),
            models.Index(fields=['company', 'status']),
            models.Index(fields=['status', 'created_at']),
        ]
    
    def __str__(self):
        return f"{self.title} - {self.employee.get_full_name() or self.employee.username}"
    
    def can_edit(self):
        """Verifica si la rendición puede ser editada"""
        return self.status in ['draft', 'rejected']
    
    def can_submit(self):
        """Verifica si la rendición puede ser enviada"""
        return self.status == 'draft' and self.expense_items.exists()
    
    def can_approve(self):
        """Verifica si la rendición puede ser aprobada"""
        return self.status == 'submitted'
    
    def calculate_total(self):
        """Calcula el total de todos los gastos"""
        total = self.expense_items.aggregate(
            total=models.Sum('amount')
        )['total'] or 0
        return total
    
    def save(self, *args, **kwargs):
        # Actualizar total automáticamente
        super().save(*args, **kwargs)
        self.total_amount = self.calculate_total()
        if self.total_amount != 0:
            super().save(update_fields=['total_amount'])


class ExpenseItem(models.Model):
    """Modelo para items individuales de gastos"""
    CATEGORY_CHOICES = [
        ('food', 'Comida'),
        ('transport', 'Transporte'),
        ('hotel', 'Hotel/Alojamiento'),
        ('fuel', 'Combustible'),
        ('materials', 'Materiales'),
        ('entertainment', 'Entretenimiento'),
        ('communication', 'Comunicaciones'),
        ('office', 'Oficina/Suministros'),
        ('medical', 'Gastos médicos'),
        ('training', 'Capacitación'),
        ('other', 'Otros'),
    ]
    
    expense_report = models.ForeignKey(
        ExpenseReport,
        on_delete=models.CASCADE,
        related_name='expense_items',
        verbose_name='Rendición de gastos'
    )
    date = models.DateField(
        verbose_name='Fecha del gasto'
    )
    category = models.CharField(
        max_length=20,
        choices=CATEGORY_CHOICES,
        verbose_name='Categoría'
    )
    description = models.CharField(
        max_length=200,
        verbose_name='Descripción del gasto'
    )
    amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Monto'
    )
    currency = models.CharField(
        max_length=3,
        default='USD',
        verbose_name='Moneda'
    )
    vendor = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Proveedor/Comercio'
    )
    location = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Ubicación'
    )
    notes = models.TextField(
        blank=True,
        null=True,
        verbose_name='Notas adicionales'
    )
    receipt_file = models.FileField(
        upload_to='expense_receipts/',
        blank=True,
        null=True,
        verbose_name='Comprobante/Recibo'
    )
    is_reimbursable = models.BooleanField(
        default=True,
        verbose_name='Reembolsable',
        help_text='Marcar si este gasto debe ser reembolsado'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-date', '-created_at']
        verbose_name = 'Item de Gasto'
        verbose_name_plural = 'Items de Gastos'
        indexes = [
            models.Index(fields=['expense_report', 'category']),
            models.Index(fields=['date', 'category']),
        ]
    
    def __str__(self):
        return f"{self.description} - {self.amount} {self.currency}"
    
    def get_category_display_icon(self):
        """Retorna el icono correspondiente a la categoría"""
        icons = {
            'food': 'bi-cup-hot',
            'transport': 'bi-bus-front',
            'hotel': 'bi-building',
            'fuel': 'bi-fuel-pump',
            'materials': 'bi-tools',
            'entertainment': 'bi-music-note',
            'communication': 'bi-telephone',
            'office': 'bi-file-text',
            'medical': 'bi-heart-pulse',
            'training': 'bi-book',
            'other': 'bi-three-dots',
        }
        return icons.get(self.category, 'bi-receipt')
    
    def save(self, *args, **kwargs):
        super().save(*args, **kwargs)
        # Actualizar el total de la rendición
        if self.expense_report:
            self.expense_report.save()


class ExpenseComment(models.Model):
    """Modelo para comentarios en rendiciones de gastos"""
    expense_report = models.ForeignKey(
        ExpenseReport,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Rendición de gastos'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    comment = models.TextField(
        verbose_name='Comentario'
    )
    is_internal = models.BooleanField(
        default=False,
        verbose_name='Comentario interno',
        help_text='Solo visible para administradores'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Comentario de Gasto'
        verbose_name_plural = 'Comentarios de Gastos'
    
    def __str__(self):
        return f"Comentario de {self.user.get_full_name() or self.user.username} - {self.created_at.strftime('%d/%m/%Y')}"


class VideoMeeting(models.Model):
    """Modelo para gestionar reuniones de video y transcripciones"""
    MEETING_STATUS_CHOICES = [
        ('scheduled', 'Programada'),
        ('in_progress', 'En progreso'),
        ('completed', 'Completada'),
        ('cancelled', 'Cancelada'),
    ]
    
    TRANSCRIPTION_STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('processing', 'Procesando'),
        ('completed', 'Completada'),
        ('failed', 'Falló'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título de la reunión'
    )
    description = models.TextField(
        blank=True,
        null=True,
        verbose_name='Descripción'
    )
    organizer = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='organized_video_meetings',
        verbose_name='Organizador'
    )
    participants = models.ManyToManyField(
        User,
        related_name='attended_meetings',
        blank=True,
        verbose_name='Participantes'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Empresa'
    )
    scheduled_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha programada'
    )
    duration_minutes = models.PositiveIntegerField(
        blank=True,
        null=True,
        verbose_name='Duración en minutos'
    )
    status = models.CharField(
        max_length=20,
        choices=MEETING_STATUS_CHOICES,
        default='scheduled',
        verbose_name='Estado'
    )
    meeting_url = models.URLField(
        blank=True,
        null=True,
        verbose_name='URL de la reunión'
    )
    
    # Archivo de video/audio
    recording_file = models.FileField(
        upload_to='meeting_recordings/',
        blank=True,
        null=True,
        verbose_name='Archivo de grabación',
        help_text='Formatos soportados: MP4, AVI, MOV, MP3, WAV, M4A, WebM'
    )
    
    # Transcripción
    transcription_status = models.CharField(
        max_length=20,
        choices=TRANSCRIPTION_STATUS_CHOICES,
        default='pending',
        verbose_name='Estado de transcripción'
    )
    transcription_text = models.TextField(
        blank=True,
        null=True,
        verbose_name='Transcripción'
    )
    transcription_confidence = models.FloatField(
        blank=True,
        null=True,
        verbose_name='Confianza de transcripción (%)'
    )
    transcription_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de transcripción'
    )
    
    # Resumen automático generado por IA
    transcription_summary = models.TextField(
        blank=True,
        null=True,
        verbose_name='Resumen de transcripción'
    )
    
    # Resumen automático generado por IA
    ai_summary = models.TextField(
        blank=True,
        null=True,
        verbose_name='Resumen generado por IA'
    )
    
    # Puntos clave extraídos por IA
    key_points = models.TextField(
        blank=True,
        null=True,
        verbose_name='Puntos clave'
    )
    
    # Acciones identificadas por IA
    action_items = models.TextField(
        blank=True,
        null=True,
        verbose_name='Elementos de acción'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    
    class Meta:
        ordering = ['-scheduled_date', '-created_at']
        verbose_name = 'Reunión de Video'
        verbose_name_plural = 'Reuniones de Video'
    
    def __str__(self):
        return f"{self.title} - {self.organizer.get_full_name() or self.organizer.username}"
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado"""
        status_classes = {
            'scheduled': 'bg-primary',
            'in_progress': 'bg-warning',
            'completed': 'bg-success',
            'cancelled': 'bg-danger',
        }
        return status_classes.get(self.status, 'bg-secondary')
    
    def get_status_color(self):
        """Retorna el color del badge para Bootstrap"""
        color_map = {
            'scheduled': 'primary',
            'in_progress': 'warning',
            'completed': 'success',
            'cancelled': 'danger',
        }
        return color_map.get(self.status, 'secondary')
    
    def get_transcription_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado de transcripción"""
        status_classes = {
            'pending': 'bg-secondary',
            'processing': 'bg-warning',
            'completed': 'bg-success',
            'failed': 'bg-danger',
        }
        return status_classes.get(self.transcription_status, 'bg-secondary')
    
    def can_edit(self):
        """Verifica si la reunión puede ser editada"""
        return self.status in ['scheduled', 'in_progress']
    
    def get_file_size_display(self):
        """Retorna el tamaño del archivo en formato legible"""
        if self.recording_file:
            try:
                size = self.recording_file.size
                if size < 1024:
                    return f"{size} B"
                elif size < 1024 * 1024:
                    return f"{size / 1024:.1f} KB"
                elif size < 1024 * 1024 * 1024:
                    return f"{size / (1024 * 1024):.1f} MB"
                else:
                    return f"{size / (1024 * 1024 * 1024):.1f} GB"
            except:
                return "Tamaño desconocido"
        return None
    
    def get_duration_display(self):
        """Retorna la duración en formato legible"""
        if self.duration_minutes:
            hours = self.duration_minutes // 60
            minutes = self.duration_minutes % 60
            if hours > 0:
                return f"{hours}h {minutes}m"
            else:
                return f"{minutes}m"
        return "Sin duración especificada"


class MeetingNote(models.Model):
    """Modelo para notas adicionales de reuniones"""
    meeting = models.ForeignKey(
        VideoMeeting,
        on_delete=models.CASCADE,
        related_name='notes',
        verbose_name='Reunión'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario'
    )
    content = models.TextField(
        verbose_name='Contenido de la nota'
    )
    is_private = models.BooleanField(
        default=False,
        verbose_name='Nota privada',
        help_text='Solo visible para el autor'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Nota de Reunión'
        verbose_name_plural = 'Notas de Reunión'
    
    def __str__(self):
        return f"Nota de {self.user.get_full_name() or self.user.username} - {self.meeting.title}"


class MeetingAttachment(models.Model):
    """Modelo para archivos adjuntos de reuniones"""
    meeting = models.ForeignKey(
        VideoMeeting,
        on_delete=models.CASCADE,
        related_name='attachments',
        verbose_name='Reunión'
    )
    file = models.FileField(
        upload_to='meeting_recordings/',
        verbose_name='Archivo adjunto'
    )
    filename = models.CharField(
        max_length=255,
        verbose_name='Nombre del archivo'
    )
    file_size = models.BigIntegerField(
        default=0,
        verbose_name='Tamaño del archivo (bytes)'
    )
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Subido por'
    )
    uploaded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de subida'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Adjunto de Reunión'
        verbose_name_plural = 'Adjuntos de Reunión'
    
    def __str__(self):
        return f"{self.filename} - {self.meeting.title}"
    
    def get_file_size_display(self):
        """Retorna el tamaño del archivo en formato legible"""
        size = self.file_size
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size < 1024.0:
                return f"{size:.1f} {unit}"
            size /= 1024.0
        return f"{size:.1f} TB"


class QuoteGenerator(models.Model):
    """Modelo para generar citas temáticas con IA"""
    title = models.CharField(
        max_length=200,
        verbose_name='Título del generador'
    )
    topic = models.CharField(
        max_length=500,
        verbose_name='Tema de las citas',
        help_text='Ej: "citas de Einstein sobre el amor", "frases motivacionales de líderes", etc.'
    )
    generated_quotes = models.TextField(
        blank=True,
        verbose_name='Citas generadas (JSON)',
        help_text='Almacena las citas en formato JSON'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Empresa'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está activo, las citas aparecerán en el dashboard'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Generador de Citas'
        verbose_name_plural = 'Generadores de Citas'
    
    def __str__(self):
        return f"{self.title} - {self.topic[:50]}"
    
    def get_quotes_list(self):
        """Obtiene las citas como lista de Python"""
        if self.generated_quotes:
            try:
                import json
                return json.loads(self.generated_quotes)
            except:
                return []
        return []
    
    def set_quotes_list(self, quotes_list):
        """Guarda las citas como JSON"""
        import json
        self.generated_quotes = json.dumps(quotes_list, ensure_ascii=False)


class CountdownTimer(models.Model):
    """Modelo para gestionar cuentas regresivas"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre descriptivo para la cuenta regresiva'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional del evento'
    )
    target_date = models.DateTimeField(
        verbose_name='Fecha objetivo',
        help_text='Fecha y hora del evento'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Empresa'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está activo, la cuenta regresiva será visible'
    )
    is_private = models.BooleanField(
        default=False,
        verbose_name='Privada',
        help_text='Si está marcada como privada, solo el creador podrá verla'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    
    class Meta:
        ordering = ['target_date']
        verbose_name = 'Cuenta Regresiva'
        verbose_name_plural = 'Cuentas Regresivas'
    
    def __str__(self):
        return f"{self.title} - {self.target_date.strftime('%d/%m/%Y')}"
    
    def days_remaining(self):
        """Calcula los días restantes hasta la fecha objetivo"""
        from datetime import datetime
        now = timezone.now()
        if self.target_date > now:
            delta = self.target_date - now
            return delta.days
        else:
            return 0
    
    def is_expired(self):
        """Verifica si la fecha objetivo ya pasó"""
        return timezone.now() > self.target_date
    
    def time_remaining(self):
        """Retorna un diccionario con días, horas, minutos y segundos restantes"""
        from datetime import datetime
        now = timezone.now()
        if self.target_date > now:
            delta = self.target_date - now
            days = delta.days
            hours, remainder = divmod(delta.seconds, 3600)
            minutes, seconds = divmod(remainder, 60)
            return {
                'days': days,
                'hours': hours,
                'minutes': minutes,
                'seconds': seconds,
                'total_seconds': delta.total_seconds()
            }
        else:
            return {
                'days': 0,
                'hours': 0,
                'minutes': 0,
                'seconds': 0,
                'total_seconds': 0
            }


class Procedure(models.Model):
    """Modelo para gestionar procedimientos/documentos de procesos"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título descriptivo del procedimiento'
    )
    url = models.URLField(
        blank=True,
        verbose_name='URL',
        help_text='Enlace externo relacionado al procedimiento (opcional)'
    )
    attachment = models.FileField(
        upload_to='procedures/',
        blank=True,
        verbose_name='Adjunto',
        help_text='Archivo del procedimiento (PDF, DOC, etc.) - Opcional'
    )
    sequence = models.PositiveIntegerField(
        default=1,
        verbose_name='Secuencia',
        help_text='Orden de visualización del procedimiento'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción detallada del procedimiento'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        blank=True,
        null=True,
        verbose_name='Empresa',
        help_text='Empresa a la que pertenece el procedimiento (opcional)',
        related_name='procedures'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si está marcado, el procedimiento será visible'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        related_name='created_procedures'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Fecha de actualización'
    )
    
    class Meta:
        ordering = ['sequence', 'title']
        verbose_name = 'Procedimiento'
        verbose_name_plural = 'Procedimientos'
    
    def __str__(self):
        return f"{self.sequence}. {self.title}"
    
    def get_file_extension(self):
        """Retorna la extensión del archivo adjunto"""
        if self.attachment:
            return os.path.splitext(self.attachment.name)[1].lower()
        return ''
    
    def get_file_icon(self):
        """Retorna el icono apropiado según el tipo de archivo"""
        if not self.attachment:
            return 'bi-file-earmark'
        
        extension = self.get_file_extension()
        icons = {
            '.pdf': 'bi-file-earmark-pdf',
            '.doc': 'bi-file-earmark-word',
            '.docx': 'bi-file-earmark-word',
            '.xls': 'bi-file-earmark-excel',
            '.xlsx': 'bi-file-earmark-excel',
            '.ppt': 'bi-file-earmark-ppt',
            '.pptx': 'bi-file-earmark-ppt',
            '.txt': 'bi-file-earmark-text',
            '.jpg': 'bi-file-earmark-image',
            '.jpeg': 'bi-file-earmark-image',
            '.png': 'bi-file-earmark-image',
            '.gif': 'bi-file-earmark-image',
        }
        return icons.get(extension, 'bi-file-earmark')
    
    def get_file_size(self):
        """Retorna el tamaño del archivo en formato legible"""
        if not self.attachment:
            return "Sin archivo"
        
        if hasattr(self.attachment, 'size'):
            size = self.attachment.size
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size < 1024.0:
                    return f"{size:.1f} {unit}"
                size /= 1024.0
            return f"{size:.1f} TB"
        return "N/A"


class MonthlyCumplimiento(models.Model):
    """Modelo para gestionar metas de cumplimiento mensual"""
    FREQUENCY_CHOICES = [
        ('daily', 'Todos los días'),
        ('specific', 'Número específico de días'),
    ]
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la meta',
        help_text='Ejemplo: Ejercicio diario, Lectura, Meditación'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario',
        help_text='Usuario que debe cumplir esta meta',
        related_name='monthly_cumplimientos'
    )
    frequency_type = models.CharField(
        max_length=20,
        choices=FREQUENCY_CHOICES,
        default='daily',
        verbose_name='Tipo de frecuencia'
    )
    target_days = models.PositiveIntegerField(
        default=30,
        verbose_name='Días objetivo',
        help_text='Número de días que debe cumplir en el mes (solo para tipo específico)'
    )
    description = models.TextField(
        blank=True,
        null=True,
        verbose_name='Descripción',
        help_text='Descripción detallada de la meta'
    )
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        verbose_name='ID público',
        help_text='ID único para acceso público'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Cumplimiento Mensual'
        verbose_name_plural = 'Cumplimientos Mensuales'
        indexes = [
            models.Index(fields=['public_uuid']),
            models.Index(fields=['created_by', 'is_active']),
        ]
    
    def __str__(self):
        return f"{self.name} - {self.user}"
    
    def get_public_url(self):
        """Retorna la URL pública para este cumplimiento"""
        from django.urls import reverse
        return reverse('monthly_cumplimiento_public', kwargs={'uuid': self.public_uuid})
    
    def get_current_month_progress(self):
        """Retorna el progreso del mes actual"""
        from datetime import date, timedelta
        today = date.today()
        current_month_start = today.replace(day=1)
        
        # Obtener cumplimientos del mes actual
        cumplimientos = self.daily_cumplimientos.filter(
            date__gte=current_month_start,
            date__year=today.year,
            date__month=today.month
        )
        
        completed_days = cumplimientos.filter(completed=True).count()
        
        # Calcular días transcurridos del mes (hasta hoy)
        days_elapsed = today.day
        
        # Total de días en el mes
        if today.month == 12:
            next_month = today.replace(year=today.year + 1, month=1, day=1)
        else:
            next_month = today.replace(month=today.month + 1, day=1)
        
        total_days_in_month = (next_month - timedelta(days=1)).day
        
        # Meta de días (depende del tipo de frecuencia)
        if self.frequency_type == 'daily':
            target_days = days_elapsed  # Días transcurridos hasta hoy
            progress_base = days_elapsed  # Progreso basado en días transcurridos
        else:
            target_days = self.target_days  # Número específico de días
            progress_base = min(target_days, days_elapsed)  # Lo que menor sea
        
        # Calcular porcentaje de progreso
        if progress_base > 0:
            percentage = (completed_days / progress_base) * 100
        else:
            percentage = 0
        
        # Limitar el porcentaje al 100%
        percentage = min(percentage, 100)
        
        return {
            'completed_days': completed_days,
            'total_days_in_month': total_days_in_month,
            'days_elapsed': days_elapsed,
            'target_days': target_days,
            'remaining_days': max(target_days - completed_days, 0),
            'percentage': percentage
        }


class DailyCumplimiento(models.Model):
    """Modelo para registrar el cumplimiento diario"""
    monthly_cumplimiento = models.ForeignKey(
        MonthlyCumplimiento,
        on_delete=models.CASCADE,
        related_name='daily_cumplimientos',
        verbose_name='Cumplimiento mensual'
    )
    date = models.DateField(
        verbose_name='Fecha'
    )
    completed = models.BooleanField(
        default=False,
        verbose_name='Cumplido'
    )
    notes = models.TextField(
        blank=True,
        null=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el cumplimiento de este día'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Registrado el'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Actualizado el'
    )
    
    class Meta:
        ordering = ['-date']
        verbose_name = 'Cumplimiento Diario'
        verbose_name_plural = 'Cumplimientos Diarios'
        unique_together = ['monthly_cumplimiento', 'date']
        indexes = [
            models.Index(fields=['monthly_cumplimiento', 'date']),
            models.Index(fields=['date', 'completed']),
        ]


class QRCode(models.Model):
    """Modelo para generar y almacenar códigos QR"""
    title = models.CharField(max_length=200, verbose_name='Título')
    content = models.TextField(verbose_name='Contenido', help_text='URL, texto o datos para generar el QR')
    description = models.TextField(blank=True, null=True, verbose_name='Descripción')
    qr_type = models.CharField(
        max_length=20,
        choices=[
            ('url', 'URL/Enlace'),
            ('text', 'Texto'),
            ('email', 'Email'),
            ('phone', 'Teléfono'),
            ('sms', 'SMS'),
            ('wifi', 'WiFi'),
            ('vcard', 'Tarjeta de Contacto'),
        ],
        default='text',
        verbose_name='Tipo de QR'
    )
    size = models.IntegerField(default=256, verbose_name='Tamaño (px)')
    created_by = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Creado por')
    created_at = models.DateTimeField(auto_now_add=True, verbose_name='Fecha de creación')
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    
    # Campos para URL pública
    is_public = models.BooleanField(default=False, verbose_name='Público')
    public_uuid = models.UUIDField(default=uuid.uuid4, null=True, blank=True, verbose_name='UUID público')
    public_token = models.CharField(max_length=32, null=True, blank=True, unique=True, verbose_name='Token público')
    public_views = models.PositiveIntegerField(default=0, verbose_name='Vistas públicas')

    def __str__(self):
        return f"{self.title} ({self.qr_type})"

    def get_formatted_content(self):
        """Retorna el contenido formateado según el tipo"""
        if self.qr_type == 'email':
            return f"mailto:{self.content}"
        elif self.qr_type == 'phone':
            return f"tel:{self.content}"
        elif self.qr_type == 'sms':
            return f"sms:{self.content}"
        elif self.qr_type == 'wifi':
            # Formato: WIFI:T:WPA;S:SSID;P:password;;
            return self.content
        else:
            return self.content
    
    def save(self, *args, **kwargs):
        # Generar token si se activa como público y no tiene token
        if self.is_public and not self.public_token:
            self.generate_public_token()
        # Limpiar token si se desactiva como público
        elif not self.is_public and self.public_token:
            self.public_token = None
        super().save(*args, **kwargs)
    
    def generate_public_token(self):
        """Genera un token único para acceso público"""
        import secrets
        import string
        
        while True:
            # Generar token de 16 caracteres alfanuméricos
            token = ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(16))
            # Verificar que no existe otro token igual
            if not QRCode.objects.filter(public_token=token).exists():
                self.public_token = token
                break
        return token
    
    def get_public_url(self):
        """Retorna la URL pública para mostrar este QR"""
        from django.urls import reverse
        if self.is_public and self.public_token:
            return reverse('qr_public', kwargs={'token': self.public_token})
        return None

    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Código QR'
        verbose_name_plural = 'Códigos QR'


class QuickTodo(models.Model):
    """Modelo para tareas rápidas (TODO list simple)"""
    text = models.CharField(
        max_length=500,
        verbose_name='Tarea'
    )
    completed = models.BooleanField(
        default=False,
        verbose_name='Completada'
    )
    created_by = models.ForeignKey(
        User, 
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['completed', '-created_at']
        verbose_name = 'Tarea Rápida'
        verbose_name_plural = 'Tareas Rápidas'
        indexes = [
            models.Index(fields=['created_by', 'completed']),
            models.Index(fields=['created_at']),
        ]
    
    def __str__(self):
        status = "✓" if self.completed else "○"
        return f"{status} {self.text[:50]}..."


class Quotation(models.Model):
    """Modelo para gestionar cotizaciones"""
    
    STATUS_CHOICES = [
        ('draft', 'Borrador'),
        ('sent', 'Enviada'),
        ('approved', 'Aprobada'),
        ('rejected', 'Rechazada'),
        ('expired', 'Expirada'),
    ]
    
    # Secuencia automática
    sequence = models.CharField(
        max_length=20, 
        unique=True, 
        blank=True,
        verbose_name='Secuencia'
    )
    
    # Empresa cliente
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        verbose_name='Empresa',
        related_name='quotations'
    )
    
    # Vendedor (usuario)
    salesperson = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Vendedor',
        related_name='quotations'
    )
    
    # Fecha de cotización
    date = models.DateField(
        default=timezone.now,
        verbose_name='Fecha'
    )
    
    # Fecha de vencimiento (por defecto 30 días después de la fecha de cotización)
    expiry_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Fecha de Vencimiento',
        help_text='Fecha límite para aceptar esta cotización'
    )
    
    # Estado
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='draft',
        verbose_name='Estado'
    )
    
    # Campos adicionales
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    
    # Fechas del servicio (opcional para servicios con duración específica)
    service_start_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Fecha de Inicio del Servicio',
        help_text='Fecha en que comenzará el servicio (opcional)'
    )
    
    service_end_date = models.DateField(
        blank=True,
        null=True,
        verbose_name='Fecha de Fin del Servicio', 
        help_text='Fecha estimada de finalización del servicio (opcional)'
    )
    
    # Token para acceso público
    public_token = models.CharField(
        max_length=50,
        unique=True,
        blank=True,
        null=True,
        verbose_name='Token Público',
        help_text='Token único para compartir la cotización públicamente'
    )
    
    # Estado del cliente (respuesta pública)
    CLIENT_STATUS_CHOICES = [
        ('pending', 'Pendiente de Respuesta'),
        ('accepted', 'Aceptada por Cliente'),
        ('rejected', 'Rechazada por Cliente'),
    ]
    
    client_status = models.CharField(
        max_length=20,
        choices=CLIENT_STATUS_CHOICES,
        default='pending',
        verbose_name='Estado del Cliente',
        help_text='Estado de la respuesta del cliente a la cotización'
    )
    
    # Información de respuesta del cliente
    client_response_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre del Cliente',
        help_text='Nombre de quien respondió a la cotización'
    )
    
    client_response_email = models.EmailField(
        blank=True,
        verbose_name='Email del Cliente',
        help_text='Email de quien respondió a la cotización'
    )
    
    client_response_comments = models.TextField(
        blank=True,
        verbose_name='Comentarios del Cliente',
        help_text='Comentarios adicionales del cliente'
    )
    
    client_response_date = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de Respuesta del Cliente',
        help_text='Fecha y hora en que el cliente respondió'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Cotización'
        verbose_name_plural = 'Cotizaciones'
        indexes = [
            models.Index(fields=['sequence']),
            models.Index(fields=['company']),
            models.Index(fields=['salesperson']),
            models.Index(fields=['date']),
            models.Index(fields=['status']),
            models.Index(fields=['created_at']),
        ]
    
    def save(self, *args, **kwargs):
        if not self.sequence:
            # Generar secuencia automática
            current_year = timezone.now().year
            last_quotation = Quotation.objects.filter(
                sequence__startswith=f'COT-{current_year}'
            ).order_by('-sequence').first()
            
            if last_quotation:
                last_number = int(last_quotation.sequence.split('-')[-1])
                new_number = last_number + 1
            else:
                new_number = 1
            
            self.sequence = f'COT-{current_year}-{new_number:04d}'
        
        # Calcular fecha de vencimiento si no está establecida
        if not self.expiry_date and self.date:
            from datetime import timedelta, datetime
            # Convertir self.date a objeto date si es string
            if isinstance(self.date, str):
                date_obj = datetime.strptime(self.date, '%Y-%m-%d').date()
            else:
                date_obj = self.date
            self.expiry_date = date_obj + timedelta(days=30)
        
        super().save(*args, **kwargs)
    
    def __str__(self):
        return f"{self.sequence} - {self.company.name}"
    
    def get_status_display_badge(self):
        """Retorna el estado con el badge HTML correspondiente"""
        badge_colors = {
            'draft': 'secondary',
            'sent': 'primary', 
            'approved': 'success',
            'rejected': 'danger',
            'expired': 'warning'
        }
        color = badge_colors.get(self.status, 'secondary')
        return f'<span class="badge bg-{color}">{self.get_status_display()}</span>'
    
    def get_total_amount(self):
        """Calcula el total de la cotización sumando todas las líneas"""
        total = self.lines.aggregate(
            total=models.Sum(models.F('quantity') * models.F('unit_price'))
        )['total']
        return total or 0
    
    def get_lines_count(self):
        """Retorna el número de líneas de la cotización"""
        return self.lines.count()
    
    def generate_public_token(self):
        """Genera un token único para acceso público"""
        import uuid
        import hashlib
        
        if not self.public_token:
            # Generar token único basado en UUID y timestamp
            unique_string = f"{self.id}-{self.sequence}-{timezone.now().timestamp()}"
            token = hashlib.md5(unique_string.encode()).hexdigest()[:32]
            self.public_token = token
            self.save(update_fields=['public_token'])
        return self.public_token
    
    def get_public_url(self, request=None):
        """Retorna la URL pública de la cotización"""
        if not self.public_token:
            self.generate_public_token()
        
        if request:
            from django.urls import reverse
            return request.build_absolute_uri(
                reverse('quotation_public_view', kwargs={'token': self.public_token})
            )
        return f"/quotations/public/{self.public_token}/"
    
    def can_be_responded(self):
        """Verifica si la cotización puede ser respondida (aceptada/rechazada)"""
        # Si el estado del cliente es 'pending', siempre se permite responder
        # independientemente del estado interno
        return self.client_status == 'pending' and not self.is_expired()
    
    def is_expired(self):
        """Verifica si la cotización está vencida"""
        if not self.expiry_date:
            return False
        from datetime import date
        return date.today() > self.expiry_date
    
    def get_days_until_expiry(self):
        """Retorna los días restantes hasta el vencimiento"""
        if not self.expiry_date:
            return None
        from datetime import date
        days_left = (self.expiry_date - date.today()).days
        return days_left if days_left >= 0 else 0
    
    def get_expiry_status_display(self):
        """Retorna el estado de vencimiento con formato"""
        if not self.expiry_date:
            return "Sin fecha de vencimiento"
        
        days_left = self.get_days_until_expiry()
        if days_left is None:
            return "Sin fecha de vencimiento"
        elif days_left == 0:
            return "Vence hoy"
        elif days_left < 0:
            return f"Vencida hace {abs(days_left)} días"
        elif days_left <= 3:
            return f"Vence en {days_left} días (¡Urgente!)"
        elif days_left <= 7:
            return f"Vence en {days_left} días"
        else:
            return f"Vence en {days_left} días"
    
    def get_client_status_display_badge(self):
        """Retorna el estado del cliente con el badge HTML correspondiente"""
        badge_colors = {
            'pending': 'warning',
            'accepted': 'success',
            'rejected': 'danger'
        }
        color = badge_colors.get(self.client_status, 'secondary')
        return f'<span class="badge bg-{color}">{self.get_client_status_display()}</span>'


class QuotationView(models.Model):
    """Modelo para registrar las visitas a las cotizaciones publicas"""
    
    quotation = models.ForeignKey(
        'Quotation',
        on_delete=models.CASCADE,
        verbose_name='Cotizacion',
        related_name='views'
    )
    
    viewed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Visita'
    )
    
    ip_address = models.GenericIPAddressField(
        verbose_name='Direccion IP',
        null=True,
        blank=True
    )
    
    country = models.CharField(
        max_length=100,
        verbose_name='Pais',
        blank=True,
        help_text='Pais desde donde se accedio'
    )
    
    country_code = models.CharField(
        max_length=2,
        verbose_name='Codigo de Pais',
        blank=True,
        help_text='Codigo ISO del pais (ej: US, MX, ES)'
    )
    
    user_agent = models.TextField(
        verbose_name='Navegador/Dispositivo',
        blank=True
    )
    
    class Meta:
        ordering = ['-viewed_at']
        verbose_name = 'Visita a Cotizacion'
        verbose_name_plural = 'Visitas a Cotizaciones'
        indexes = [
            models.Index(fields=['quotation']),
            models.Index(fields=['viewed_at']),
            models.Index(fields=['country']),
        ]
    
    def __str__(self):
        return f"{self.quotation.sequence} - {self.viewed_at.strftime('%Y-%m-%d %H:%M')} - {self.country or 'Desconocido'}"


class QuotationLine(models.Model):
    """Modelo para las líneas de productos en una cotización"""
    
    quotation = models.ForeignKey(
        'Quotation',
        on_delete=models.CASCADE,
        verbose_name='Cotización',
        related_name='lines'
    )
    
    product = models.ForeignKey(
        'Product',
        on_delete=models.CASCADE,
        verbose_name='Producto'
    )
    
    quantity = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=1,
        verbose_name='Cantidad'
    )
    
    unit_price = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Precio Unitario'
    )
    
    discount_percentage = models.DecimalField(
        max_digits=5,
        decimal_places=2,
        default=0,
        verbose_name='Descuento (%)',
        help_text='Porcentaje de descuento sobre el precio unitario'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción adicional o personalizada del producto'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['id']
        verbose_name = 'Línea de Cotización'
        verbose_name_plural = 'Líneas de Cotización'
        indexes = [
            models.Index(fields=['quotation']),
            models.Index(fields=['product']),
        ]
    
    def save(self, *args, **kwargs):
        # Si no se ha establecido el precio unitario, usar el precio del producto
        if not self.unit_price:
            self.unit_price = self.product.price
        super().save(*args, **kwargs)
    
    def get_subtotal(self):
        """Calcula el subtotal de la línea (cantidad × precio unitario)"""
        return self.quantity * self.unit_price
    
    def get_discount_amount(self):
        """Calcula el monto del descuento"""
        return self.get_subtotal() * (self.discount_percentage / 100)
    
    def get_total(self):
        """Calcula el total de la línea (subtotal - descuento)"""
        return self.get_subtotal() - self.get_discount_amount()
    
    def __str__(self):
        return f"{self.quotation.sequence} - {self.product.name} (x{self.quantity})"


class QuotationTemplate(models.Model):
    """Modelo para plantillas de cotización que se pueden compartir públicamente"""
    
    # Cotización base que sirve como plantilla
    template_quotation = models.ForeignKey(
        'Quotation',
        on_delete=models.CASCADE,
        verbose_name='Cotización Plantilla',
        related_name='public_templates'
    )
    
    # Token público único para acceder al formulario
    public_token = models.CharField(
        max_length=50,
        unique=True,
        verbose_name='Token Público',
        help_text='Token único para el enlace público'
    )
    
    # Información del enlace
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Enlace',
        help_text='Título descriptivo para el cliente'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción que verá el cliente al acceder al enlace'
    )
    
    # Estado
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo',
        help_text='Si el enlace está activo y acepta nuevas solicitudes'
    )
    
    # Estadísticas
    views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de Vistas',
        help_text='Cuántas veces se ha visitado el enlace'
    )
    
    conversions_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Conversiones',
        help_text='Cuántas cotizaciones se han generado desde este enlace'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        related_name='quotation_templates'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Plantilla de Cotización Pública'
        verbose_name_plural = 'Plantillas de Cotización Públicas'
        indexes = [
            models.Index(fields=['public_token']),
            models.Index(fields=['is_active']),
            models.Index(fields=['created_at']),
        ]
    
    def save(self, *args, **kwargs):
        if not self.public_token:
            # Generar token único
            import uuid
            import hashlib
            unique_string = f"{self.template_quotation.id}-{self.title}-{timezone.now().timestamp()}"
            token = hashlib.md5(unique_string.encode()).hexdigest()[:32]
            self.public_token = token
        super().save(*args, **kwargs)
    
    def __str__(self):
        return f"{self.title} - {self.template_quotation.sequence}"
    
    def get_public_url(self, request=None):
        """Retorna la URL pública del formulario"""
        if request:
            from django.urls import reverse
            return request.build_absolute_uri(
                reverse('quotation_template_public_form', kwargs={'token': self.public_token})
            )
        return f"/quotations/template/{self.public_token}/"
    
    def increment_views(self):
        """Incrementa el contador de vistas"""
        self.views_count += 1
        self.save(update_fields=['views_count'])
    
    def increment_conversions(self):
        """Incrementa el contador de conversiones"""
        self.conversions_count += 1
        self.save(update_fields=['conversions_count'])


class QuotationTemplateRequest(models.Model):
    """Solicitudes generadas desde plantillas públicas"""
    
    # Plantilla que generó esta solicitud
    template = models.ForeignKey(
        'QuotationTemplate',
        on_delete=models.CASCADE,
        verbose_name='Plantilla',
        related_name='requests'
    )
    
    # Datos del cliente
    client_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Cliente'
    )
    
    client_email = models.EmailField(
        verbose_name='Email del Cliente'
    )
    
    client_phone = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Teléfono del Cliente'
    )
    
    client_company = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Empresa del Cliente'
    )
    
    # Comentarios del cliente
    client_comments = models.TextField(
        blank=True,
        verbose_name='Comentarios del Cliente'
    )
    
    # Cotización generada
    generated_quotation = models.ForeignKey(
        'Quotation',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Cotización Generada',
        related_name='template_requests'
    )
    
    # Estado
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('processed', 'Procesada'),
        ('quoted', 'Cotizada'),
    ]
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    # Metadatos
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de solicitud'
    )
    
    processed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de procesamiento'
    )
    
    processed_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        verbose_name='Procesado por',
        related_name='processed_template_requests'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Solicitud de Plantilla'
        verbose_name_plural = 'Solicitudes de Plantillas'
        indexes = [
            models.Index(fields=['template']),
            models.Index(fields=['status']),
            models.Index(fields=['created_at']),
        ]
    
    def __str__(self):
        return f"{self.client_name} - {self.template.title}"
    
    def process_request(self, processed_by_user):
        """Procesa la solicitud creando una cotización personalizada"""
        try:
            # Crear la empresa del cliente si no existe, o actualizar contacto si existe
            from .models import Company
            
            company, created = Company.objects.get_or_create(
                email=self.client_email,
                defaults={
                    'name': self.client_company or f"Empresa de {self.client_name}",
                    'phone': self.client_phone or '',
                    'address': '',
                    'description': f'Cliente generado automáticamente desde plantilla: {self.template.title}. Contacto: {self.client_name}',
                    'is_active': True,
                }
            )
            
            # Si la empresa ya existía, actualizar datos de contacto
            if not created:
                if self.client_phone:
                    company.phone = self.client_phone
                if self.client_company:
                    company.name = self.client_company
                # Actualizar descripción para incluir el nuevo contacto
                company.description = f'Cliente generado automáticamente desde plantilla: {self.template.title}. Contacto: {self.client_name}'
                company.save()
            
            # Crear la nueva cotización basada en la plantilla
            template_quotation = self.template.template_quotation
            new_quotation = Quotation.objects.create(
                company=company,
                salesperson=processed_by_user,
                date=timezone.now().date(),
                status='sent',  # La enviamos directamente
                client_status='pending',  # Pendiente de respuesta del cliente
                description=f"Cotización generada automáticamente para {self.client_name}\n\nComentarios del cliente:\n{self.client_comments}" if self.client_comments else f"Cotización generada automáticamente para {self.client_name}",
            )
            
            # Copiar las líneas de la plantilla
            for line in template_quotation.lines.all():
                new_quotation.lines.create(
                    product=line.product,
                    quantity=line.quantity,
                    unit_price=line.unit_price,
                    discount_percentage=line.discount_percentage,
                    description=line.description
                )
            
            # Generar token público para la nueva cotización
            new_quotation.generate_public_token()
            
            # Actualizar esta solicitud
            self.generated_quotation = new_quotation
            self.status = 'quoted'
            self.processed_at = timezone.now()
            self.processed_by = processed_by_user
            self.save()
            
            # Incrementar contador de conversiones
            self.template.increment_conversions()
            
            return new_quotation
            
        except Exception as e:
            # Marcar como procesada con error
            self.status = 'processed'
            self.processed_at = timezone.now()
            self.processed_by = processed_by_user
            self.save()
            raise e


class CrmQuestion(models.Model):
    """Modelo para gestionar preguntas del CRM"""
    
    sequence = models.PositiveIntegerField(
        verbose_name='Secuencia',
        help_text='Número de secuencia de la pregunta'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.CASCADE,
        related_name='crm_questions',
        verbose_name='Empresa',
        help_text='Empresa asociada a la pregunta'
    )
    question = models.TextField(
        verbose_name='Pregunta',
        help_text='Contenido de la pregunta'
    )
    answer = models.TextField(
        blank=True,
        null=True,
        verbose_name='Respuesta',
        help_text='Respuesta a la pregunta'
    )
    person_name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la persona',
        help_text='Nombre completo de quien hace la pregunta'
    )
    person_email = models.EmailField(
        verbose_name='Email de la persona',
        help_text='Correo electrónico de quien hace la pregunta'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        related_name='crm_questions_created',
        verbose_name='Creado por',
        help_text='Usuario que creó la pregunta (NULL para preguntas públicas)'
    )
    answered_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='crm_questions_answered',
        verbose_name='Respondido por'
    )
    answered_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de respuesta'
    )
    
    class Meta:
        ordering = ['sequence', '-created_at']
        verbose_name = 'Pregunta CRM'
        verbose_name_plural = 'Preguntas CRM'
        unique_together = [['company', 'sequence']]
    
    def __str__(self):
        return f"#{self.sequence} - {self.person_name} ({self.company.name})"
    
    def save(self, *args, **kwargs):
        # Auto-asignar secuencia si no está definida
        if not self.sequence:
            last_question = CrmQuestion.objects.filter(company=self.company).order_by('-sequence').first()
            self.sequence = (last_question.sequence + 1) if last_question else 1
        
        # Manejar el estado de la respuesta automáticamente
        if self.answer and self.answer.strip():
            # Si hay respuesta y no estaba respondida antes
            if not self.answered_at:
                self.answered_at = timezone.now()
            # Si no hay usuario que respondió asignado, mantener el que ya estaba
            # (esto se manejará en las vistas para asignar el usuario correcto)
        else:
            # Si se borra la respuesta, limpiar campos relacionados
            self.answered_at = None
            self.answered_by = None
        
        super().save(*args, **kwargs)
    
    @property
    def is_answered(self):
        """Retorna True si la pregunta ya fue respondida"""
        return bool(self.answer and self.answer.strip())
    
    @property
    def is_public_question(self):
        """Retorna True si es una pregunta pública (sin usuario creador)"""
        return self.created_by is None
    
    def get_creator_display(self):
        """Retorna información sobre quién creó la pregunta"""
        if self.is_public_question:
            return "Pregunta Pública"
        else:
            if self.created_by:
                return self.created_by.get_full_name() or self.created_by.username
            return "Usuario no disponible"
    
    def get_origin_badge_class(self):
        """Retorna la clase CSS para el badge de origen"""
        return "bg-success" if self.is_public_question else "bg-primary"
    
    def get_origin_icon(self):
        """Retorna el icono para el origen de la pregunta"""
        return "fas fa-globe" if self.is_public_question else "fas fa-user-tie"
    
    def get_answered_by_display(self):
        """Retorna información sobre quién respondió la pregunta"""
        if not self.answered_by:
            return "Usuario no disponible"
        return self.answered_by.get_full_name() or self.answered_by.username
    
    def mark_as_answered(self, user, answer_text):
        """Marca la pregunta como respondida con el texto y usuario especificado"""
        self.answer = answer_text
        self.answered_by = user
        self.answered_at = timezone.now()
        self.save()
    
    def mark_as_pending(self):
        """Marca la pregunta como pendiente (sin respuesta)"""
        self.answer = None
        self.answered_by = None
        self.answered_at = None
        self.save()
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge de estado"""
        return "bg-success" if self.is_answered else "bg-warning"
    
    def get_status_icon(self):
        """Retorna el icono FontAwesome para el estado"""
        return "fas fa-check" if self.is_answered else "fas fa-clock"
    
    @property
    def status_display(self):
        """Retorna el estado de la pregunta para mostrar"""
        return "Respondida" if self.is_answered else "Pendiente"
    
    @property
    def status_color(self):
        """Retorna el color del estado para la UI"""
        return "success" if self.is_answered else "warning"


class SupportMeeting(models.Model):
    """Modelo para gestionar reuniones de soporte"""
    
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        verbose_name='Empresa',
        help_text='Empresa para la cual se realiza la reunión'
    )
    sequence = models.PositiveIntegerField(
        verbose_name='Secuencia',
        help_text='Número secuencial de la reunión'
    )
    date = models.DateField(
        default=timezone.now,
        verbose_name='Fecha',
        help_text='Fecha de la reunión'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por',
        help_text='Usuario que creó la reunión'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción de lo que se vio en la reunión'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['company', '-sequence']
        verbose_name = 'Reunión de Soporte'
        verbose_name_plural = 'Reuniones de Soporte'
        unique_together = [['company', 'sequence']]
    
    def __str__(self):
        return f"#{self.sequence} - {self.company.name} ({self.date})"
    
    def save(self, *args, **kwargs):
        # Auto-asignar secuencia si no está definida
        if not self.sequence:
            last_meeting = SupportMeeting.objects.filter(company=self.company).order_by('-sequence').first()
            self.sequence = (last_meeting.sequence + 1) if last_meeting else 1
        
        super().save(*args, **kwargs)
    
    @property
    def points_count(self):
        """Retorna el número total de puntos de la reunión"""
        return self.support_meeting_points.count()
    
    @property
    def selected_points_count(self):
        """Retorna el número de puntos seleccionados"""
        return self.support_meeting_points.filter(is_selected=True).count()


class SupportMeetingPoint(models.Model):
    """Modelo para los puntos de las reuniones de soporte"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('completed', 'Completado'),
        ('cancelled', 'Cancelado'),
    ]
    
    meeting = models.ForeignKey(
        SupportMeeting,
        on_delete=models.CASCADE,
        related_name='support_meeting_points',
        verbose_name='Reunión'
    )
    description = models.TextField(
        verbose_name='Descripción del punto',
        help_text='Descripción detallada del punto'
    )
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado',
        help_text='Estado actual del punto'
    )
    is_selected = models.BooleanField(
        default=False,
        verbose_name='Seleccionado',
        help_text='Marcar para crear como ticket'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Punto de Reunión'
        verbose_name_plural = 'Puntos de Reunión'
    
    def __str__(self):
        return f"{self.meeting} - {self.description[:50]}..."
    
    def get_status_badge_class(self):
        """Retorna la clase CSS para el badge del estado"""
        status_classes = {
            'pending': 'bg-secondary',
            'in_progress': 'bg-warning',
            'completed': 'bg-success',
            'cancelled': 'bg-danger',
        }
        return status_classes.get(self.status, 'bg-secondary')


class SupportMeetingPublicLink(models.Model):
    """Modelo para enlaces públicos de reuniones de soporte"""
    
    meeting = models.OneToOneField(
        SupportMeeting,
        on_delete=models.CASCADE,
        related_name='public_link',
        verbose_name='Reunión'
    )
    token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token de acceso'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    show_company_meetings = models.BooleanField(
        default=True,
        verbose_name='Mostrar otras reuniones de la empresa',
        help_text='Si está marcado, se mostrarán todas las reuniones públicas de la empresa'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Creado por'
    )
    last_accessed = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Último acceso'
    )
    access_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Número de accesos'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Enlace Público de Reunión'
        verbose_name_plural = 'Enlaces Públicos de Reuniones'
    
    def __str__(self):
        return f"Enlace público - {self.meeting}"
    
    def save(self, *args, **kwargs):
        if not self.token:
            import secrets
            self.token = secrets.token_urlsafe(32)
        super().save(*args, **kwargs)
    
    def get_public_url(self, request=None):
        """Obtener URL pública completa"""
        from django.urls import reverse
        from django.conf import settings
        
        if request:
            # Si tenemos request, usar el host actual
            base_url = f"{request.scheme}://{request.get_host()}"
        else:
            # Fallback para cuando no hay request
            base_url = getattr(settings, 'BASE_URL', 'http://127.0.0.1:8000')
        
        path = reverse('support_meeting_public', kwargs={'token': self.token})
        return f"{base_url}{path}"
    
    def register_access(self):
        """Registrar un acceso al enlace"""
        from django.utils import timezone
        self.last_accessed = timezone.now()
        self.access_count += 1
        self.save(update_fields=['last_accessed', 'access_count'])


class ScheduledTask(models.Model):
    """Modelo para tareas programadas"""
    
    FREQUENCY_UNIT_CHOICES = [
        ('minute', 'Minuto'),
        ('hour', 'Hora'),  
        ('day', 'Día'),
        ('month', 'Mes'),
    ]
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Tarea'
    )
    frequency = models.PositiveIntegerField(
        verbose_name='Frecuencia',
        help_text='Cada cuántas unidades de tiempo se ejecuta'
    )
    frequency_unit = models.CharField(
        max_length=10,
        choices=FREQUENCY_UNIT_CHOICES,
        default='hour',
        verbose_name='Unidad de Frecuencia'
    )
    code = models.TextField(
        verbose_name='Código a Ejecutar',
        help_text='Código Python que se ejecutará en cada iteración'
    )
    is_active = models.BooleanField(
        default=False,
        verbose_name='Activo'
    )
    last_execution = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última Ejecución'
    )
    last_result = models.TextField(
        blank=True,
        verbose_name='Último Resultado',
        help_text='Resultado de la última ejecución'
    )
    success_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Ejecuciones Exitosas'
    )
    error_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Ejecuciones con Error'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='created_scheduled_tasks',
        verbose_name='Creado por'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Tarea Programada'
        verbose_name_plural = 'Tareas Programadas'
    
    def __str__(self):
        return f"{self.name} (cada {self.frequency} {self.get_frequency_unit_display().lower()})"
    
    def get_next_execution(self):
        """Calcula cuándo debe ejecutarse la próxima vez"""
        if not self.last_execution:
            return timezone.now()
        
        if self.frequency_unit == 'minute':
            delta = timedelta(minutes=self.frequency)
        elif self.frequency_unit == 'hour':
            delta = timedelta(hours=self.frequency)
        elif self.frequency_unit == 'day':
            delta = timedelta(days=self.frequency)
        elif self.frequency_unit == 'month':
            delta = timedelta(days=self.frequency * 30)  # Aproximado
        else:
            delta = timedelta(hours=1)
            
        return self.last_execution + delta
    
    def should_execute(self):
        """Verifica si la tarea debe ejecutarse"""
        if not self.is_active:
            return False
        return timezone.now() >= self.get_next_execution()


class ScheduledTaskExecution(models.Model):
    """Historial de ejecuciones de tareas programadas"""
    
    STATUS_CHOICES = [
        ('success', 'Exitosa'),
        ('error', 'Error'),
        ('timeout', 'Timeout'),
    ]
    
    task = models.ForeignKey(
        ScheduledTask,
        on_delete=models.CASCADE,
        related_name='executions',
        verbose_name='Tarea'
    )
    executed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Ejecutado en'
    )
    status = models.CharField(
        max_length=10,
        choices=STATUS_CHOICES,
        verbose_name='Estado'
    )
    result = models.TextField(
        blank=True,
        verbose_name='Resultado'
    )
    execution_time = models.FloatField(
        null=True,
        blank=True,
        verbose_name='Tiempo de Ejecución (segundos)'
    )
    
    class Meta:
        ordering = ['-executed_at']
        verbose_name = 'Ejecución de Tarea'
        verbose_name_plural = 'Ejecuciones de Tareas'
    
    def __str__(self):
        return f"{self.task.name} - {self.executed_at.strftime('%d/%m/%Y %H:%M')} ({self.get_status_display()})"


class QARating(models.Model):
    """Modelo para calificaciones de QA con caras (triste, neutro, feliz)"""
    
    RATING_CHOICES = [
        ('sad', '😞 Triste'),
        ('neutral', '😐 Neutro'),
        ('happy', '😊 Feliz'),
    ]
    
    # Calificación principal
    rating = models.CharField(
        max_length=10,
        choices=RATING_CHOICES,
        verbose_name='Calificación',
        help_text='Calificación: triste, neutro o feliz'
    )
    
    # Opinión
    opinion = models.TextField(
        verbose_name='Opinión',
        help_text='Tu opinión sobre el servicio de QA'
    )
    
    # Información opcional del usuario
    name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre',
        help_text='Tu nombre (opcional)'
    )
    email = models.EmailField(
        blank=True,
        verbose_name='Correo electrónico',
        help_text='Tu correo electrónico (opcional)'
    )
    
    # Relación opcional con empresa
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='qa_ratings',
        verbose_name='Empresa',
        help_text='Empresa asociada a esta calificación (opcional)'
    )
    
    # Usuario registrado (si aplica)
    user = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='qa_ratings',
        verbose_name='Usuario',
        help_text='Usuario registrado que dejó la calificación'
    )
    
    # Metadatos
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP',
        help_text='IP desde donde se realizó la calificación'
    )
    user_agent = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='User Agent',
        help_text='Información del navegador'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    is_public = models.BooleanField(
        default=True,
        verbose_name='Visible públicamente',
        help_text='Si la calificación puede mostrarse públicamente'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Calificación QA'
        verbose_name_plural = 'Calificaciones QA'
    
    def __str__(self):
        name_display = self.name or self.email or 'Anónimo'
        return f"{name_display} - {self.get_rating_display()} - {self.created_at.strftime('%d/%m/%Y')}"
    
    def get_rating_emoji(self):
        """Devuelve el emoji correspondiente a la calificación"""
        emojis = {
            'sad': '😞',
            'neutral': '😐',
            'happy': '😊',
        }
        return emojis.get(self.rating, '😐')
    
    def get_rating_color(self):
        """Devuelve el color correspondiente a la calificación"""
        colors = {
            'sad': '#dc3545',      # Rojo
            'neutral': '#ffc107',  # Amarillo
            'happy': '#28a745',    # Verde
        }
        return colors.get(self.rating, '#6c757d')


class GameCounter(models.Model):
    """Modelo para contador de juegos de pádel"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('finished', 'Finalizado'),
    ]
    
    # Información del juego
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Juego',
        help_text='Nombre descriptivo del juego (ej: Final Torneo 2025)'
    )
    player1_name = models.CharField(
        max_length=100,
        verbose_name='Jugador 1',
        help_text='Nombre del primer jugador'
    )
    player2_name = models.CharField(
        max_length=100,
        verbose_name='Jugador 2',
        help_text='Nombre del segundo jugador'
    )
    
    # Puntuación
    player1_score = models.IntegerField(
        default=0,
        verbose_name='Puntos Jugador 1'
    )
    player2_score = models.IntegerField(
        default=0,
        verbose_name='Puntos Jugador 2'
    )
    
    # Estado
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado del Juego'
    )
    
    # URL pública
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='UUID Público',
        help_text='Identificador único para acceso público'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_games',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='games',
        verbose_name='Empresa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    started_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de inicio'
    )
    finished_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de finalización'
    )
    
    # Configuración opcional
    max_points = models.IntegerField(
        default=0,
        verbose_name='Puntos máximos',
        help_text='0 = sin límite'
    )
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el juego'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Contador de Juego'
        verbose_name_plural = 'Contadores de Juegos'
    
    def __str__(self):
        return f"{self.name} - {self.player1_name} vs {self.player2_name}"
    
    def get_public_url(self):
        """Devuelve la URL pública del juego"""
        from django.urls import reverse
        return reverse('game_counter_public', kwargs={'uuid': self.public_uuid})
    
    def start_game(self):
        """Inicia el juego"""
        if self.status == 'pending':
            self.status = 'in_progress'
            self.started_at = timezone.now()
            self.save()
    
    def finish_game(self):
        """Finaliza el juego"""
        if self.status == 'in_progress':
            self.status = 'finished'
            self.finished_at = timezone.now()
            self.save()
    
    def add_point_player1(self):
        """Añade un punto al jugador 1"""
        self.player1_score += 1
        self.save()
        
        # Auto-finalizar si alcanza max_points
        if self.max_points > 0 and self.player1_score >= self.max_points:
            self.finish_game()
    
    def add_point_player2(self):
        """Añade un punto al jugador 2"""
        self.player2_score += 1
        self.save()
        
        # Auto-finalizar si alcanza max_points
        if self.max_points > 0 and self.player2_score >= self.max_points:
            self.finish_game()
    
    def get_winner(self):
        """Devuelve el nombre del ganador si el juego ha terminado"""
        if self.status == 'finished':
            if self.player1_score > self.player2_score:
                return self.player1_name
            elif self.player2_score > self.player1_score:
                return self.player2_name
            else:
                return 'Empate'
        return None
    
    def get_duration(self):
        """Devuelve la duración del juego si ha terminado"""
        if self.started_at and self.finished_at:
            duration = self.finished_at - self.started_at
            hours = duration.seconds // 3600
            minutes = (duration.seconds % 3600) // 60
            if hours > 0:
                return f"{hours}h {minutes}m"
            return f"{minutes}m"
        return None


class ExerciseCounter(models.Model):
    """Modelo para contador de ejercicios deportivos"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('in_progress', 'En Progreso'),
        ('paused', 'Pausado'),
        ('finished', 'Finalizado'),
    ]
    
    # Información del ejercicio
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre del ejercicio (ej: Rutina de Flexiones)'
    )
    
    # Contadores
    sets_target = models.IntegerField(
        verbose_name='Tandas Objetivo',
        help_text='Número de tandas objetivo'
    )
    reps_target = models.IntegerField(
        verbose_name='Repeticiones Objetivo',
        help_text='Número de repeticiones objetivo por tanda'
    )
    
    current_sets = models.IntegerField(
        default=0,
        verbose_name='Tandas Actuales'
    )
    current_reps = models.IntegerField(
        default=0,
        verbose_name='Repeticiones Actuales'
    )
    
    # Estado
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    # URL pública
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='UUID Público',
        help_text='Identificador único para acceso público'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_exercises',
        verbose_name='Creado por'
    )
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='exercises',
        verbose_name='Empresa'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    started_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de inicio'
    )
    finished_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de finalización'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el ejercicio'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Contador de Ejercicio'
        verbose_name_plural = 'Contadores de Ejercicios'
    
    def __str__(self):
        return f"{self.title} - {self.current_sets}/{self.sets_target} tandas"
    
    def get_public_url(self):
        """Devuelve la URL pública del ejercicio"""
        from django.urls import reverse
        return reverse('exercise_counter_public', kwargs={'uuid': self.public_uuid})
    
    def start_exercise(self):
        """Inicia el ejercicio"""
        if self.status == 'pending':
            self.status = 'in_progress'
            self.started_at = timezone.now()
            self.save()
    
    def pause_exercise(self):
        """Pausa el ejercicio"""
        if self.status == 'in_progress':
            self.status = 'paused'
            self.save()
    
    def resume_exercise(self):
        """Reanuda el ejercicio"""
        if self.status == 'paused':
            self.status = 'in_progress'
            self.save()
    
    def finish_exercise(self):
        """Finaliza el ejercicio"""
        if self.status in ['in_progress', 'paused']:
            self.status = 'finished'
            self.finished_at = timezone.now()
            self.save()
    
    def add_set(self):
        """Añade una tanda"""
        self.current_sets += 1
        self.current_reps = 0  # Resetear repeticiones al completar una tanda
        self.save()
        
        # Auto-finalizar si alcanza las tandas objetivo
        if self.current_sets >= self.sets_target:
            self.finish_exercise()
    
    def add_rep(self):
        """Añade una repetición"""
        self.current_reps += 1
        self.save()
        
        # Auto-pasar a siguiente tanda si alcanza repeticiones objetivo
        if self.current_reps >= self.reps_target:
            self.add_set()
    
    def get_progress(self):
        """Devuelve el progreso en porcentaje"""
        total_reps_target = self.sets_target * self.reps_target
        total_reps_current = (self.current_sets * self.reps_target) + self.current_reps
        if total_reps_target > 0:
            return int((total_reps_current / total_reps_target) * 100)
        return 0
    
    def get_duration(self):
        """Devuelve la duración del ejercicio si ha terminado"""
        if self.started_at and self.finished_at:
            duration = self.finished_at - self.started_at
            hours = duration.seconds // 3600
            minutes = (duration.seconds % 3600) // 60
            seconds = duration.seconds % 60
            if hours > 0:
                return f"{hours}h {minutes}m {seconds}s"
            elif minutes > 0:
                return f"{minutes}m {seconds}s"
            return f"{seconds}s"
        return None


class SportGoal(models.Model):
    """Modelo para objetivos deportivos con seguimiento de tiempo y distancia"""
    
    # Información del objetivo
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Objetivo',
        help_text='Nombre del objetivo deportivo (ej: Correr 5K en menos de 30 minutos)'
    )
    
    player_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Jugador',
        help_text='Nombre de la persona que realizará el objetivo'
    )
    
    # Objetivos
    target_time = models.IntegerField(
        verbose_name='Tiempo Objetivo (segundos)',
        help_text='Tiempo objetivo en segundos para completar el objetivo',
        null=True,
        blank=True
    )
    
    target_distance = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Distancia Objetivo (km)',
        help_text='Distancia objetivo en kilómetros',
        null=True,
        blank=True
    )
    
    # URL pública
    public_uuid = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='UUID Público',
        help_text='Identificador único para acceso público'
    )
    
    # Metadatos
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='created_sport_goals',
        verbose_name='Creado por'
    )
    
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='sport_goals',
        verbose_name='Empresa'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el objetivo'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Objetivo Deportivo'
        verbose_name_plural = 'Objetivos Deportivos'
    
    def __str__(self):
        return f"{self.title} - {self.player_name}"
    
    def get_records(self):
        """Retorna todos los registros del objetivo"""
        return self.records.all().order_by('recorded_at')
    
    def get_best_time(self):
        """Retorna el mejor tiempo registrado"""
        best = self.records.filter(actual_time__isnull=False).order_by('actual_time').first()
        return best.actual_time if best else None
    
    def get_best_distance(self):
        """Retorna la mejor distancia registrada"""
        best = self.records.filter(actual_distance__isnull=False).order_by('-actual_distance').first()
        return best.actual_distance if best else None
    
    def get_progress_percentage(self):
        """Calcula el porcentaje de progreso basado en el mejor registro"""
        if self.target_time and self.get_best_time():
            # Para tiempo: menor es mejor, entonces invertimos el cálculo
            if self.get_best_time() <= self.target_time:
                return 100
            else:
                return min(100, (self.target_time / self.get_best_time()) * 100)
        
        if self.target_distance and self.get_best_distance():
            # Para distancia: mayor es mejor
            return min(100, (float(self.get_best_distance()) / float(self.target_distance)) * 100)
        
        return 0
    
    def format_target_time(self):
        """Formatea el tiempo objetivo"""
        if not self.target_time:
            return "N/A"
        hours = self.target_time // 3600
        minutes = (self.target_time % 3600) // 60
        seconds = self.target_time % 60
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        return f"{minutes:02d}:{seconds:02d}"


class SportGoalRecord(models.Model):
    """Modelo para registros de intentos de objetivos deportivos"""
    
    sport_goal = models.ForeignKey(
        SportGoal,
        on_delete=models.CASCADE,
        related_name='records',
        verbose_name='Objetivo Deportivo'
    )
    
    # Registros reales
    actual_time = models.IntegerField(
        verbose_name='Tiempo Real (segundos)',
        help_text='Tiempo real en segundos que tomó completar',
        null=True,
        blank=True
    )
    
    actual_distance = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Distancia Real (km)',
        help_text='Distancia real recorrida en kilómetros',
        null=True,
        blank=True
    )
    
    # Metadatos
    recorded_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de registro'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas sobre este intento específico'
    )
    
    class Meta:
        ordering = ['-recorded_at']
        verbose_name = 'Registro de Objetivo'
        verbose_name_plural = 'Registros de Objetivos'
    
    def __str__(self):
        return f"{self.sport_goal.title} - {self.recorded_at.strftime('%Y-%m-%d %H:%M')}"
    
    def format_actual_time(self):
        """Formatea el tiempo real"""
        if not self.actual_time:
            return "N/A"
        hours = self.actual_time // 3600
        minutes = (self.actual_time % 3600) // 60
        seconds = self.actual_time % 60
        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
        return f"{minutes:02d}:{seconds:02d}"
    
    def is_goal_achieved(self):
        """Verifica si se alcanzó el objetivo en este registro"""
        goal = self.sport_goal
        
        # Verificar tiempo
        if goal.target_time and self.actual_time:
            if self.actual_time > goal.target_time:
                return False
        
        # Verificar distancia
        if goal.target_distance and self.actual_distance:
            if self.actual_distance < goal.target_distance:
                return False
        
        return True


class Event(models.Model):
    """Modelo para gestionar eventos del calendario"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título del evento'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción detallada del evento'
    )
    
    event_date = models.DateTimeField(
        verbose_name='Fecha y Hora',
        help_text='Fecha y hora del evento'
    )
    
    location = models.CharField(
        max_length=300,
        blank=True,
        verbose_name='Ubicación',
        help_text='Lugar donde se realizará el evento'
    )
    
    color = models.CharField(
        max_length=7,
        default='#007bff',
        verbose_name='Color',
        help_text='Color del evento en el calendario (formato hexadecimal)'
    )
    
    is_all_day = models.BooleanField(
        default=False,
        verbose_name='Todo el día',
        help_text='Marcar si el evento dura todo el día'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='created_events',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['event_date']
        verbose_name = 'Evento'
        verbose_name_plural = 'Eventos'
    
    def __str__(self):
        return f"{self.title} - {self.event_date.strftime('%Y-%m-%d %H:%M')}"


class Trip(models.Model):
    """Modelo para gestionar viajes"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Viaje'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    
    destination = models.CharField(
        max_length=200,
        verbose_name='Destino Principal'
    )
    
    start_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Inicio'
    )
    
    end_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Fin'
    )
    
    is_public = models.BooleanField(
        default=False,
        verbose_name='Público',
        help_text='Permitir que el viaje sea visible públicamente'
    )
    
    public_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        verbose_name='Token Público',
        help_text='Token para compartir el viaje públicamente'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='trips',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Viaje'
        verbose_name_plural = 'Viajes'
    
    def __str__(self):
        return f"{self.title} - {self.destination}"
    
    def get_public_url(self):
        """Obtener URL pública del viaje"""
        from django.urls import reverse
        return reverse('trip_public', kwargs={'token': self.public_token})


class TripStop(models.Model):
    """Modelo para puntos de interés en un viaje"""
    
    trip = models.ForeignKey(
        Trip,
        on_delete=models.CASCADE,
        related_name='stops',
        verbose_name='Viaje'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Lugar'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción'
    )
    
    address = models.CharField(
        max_length=300,
        blank=True,
        verbose_name='Dirección'
    )
    
    latitude = models.DecimalField(
        max_digits=9,
        decimal_places=6,
        null=True,
        blank=True,
        verbose_name='Latitud'
    )
    
    longitude = models.DecimalField(
        max_digits=9,
        decimal_places=6,
        null=True,
        blank=True,
        verbose_name='Longitud'
    )
    
    order = models.PositiveIntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de visita en el itinerario'
    )
    
    visit_date = models.DateField(
        null=True,
        blank=True,
        verbose_name='Fecha de Visita'
    )
    
    visit_time = models.TimeField(
        null=True,
        blank=True,
        verbose_name='Hora de Visita'
    )
    
    duration_minutes = models.PositiveIntegerField(
        null=True,
        blank=True,
        verbose_name='Duración (minutos)',
        help_text='Tiempo estimado de visita en minutos'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas',
        help_text='Notas adicionales sobre el lugar'
    )
    
    historical_info = models.TextField(
        blank=True,
        verbose_name='Información Histórica',
        help_text='Historia e información interesante sobre el lugar (aprox. 200 palabras)'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['trip', 'order']
        verbose_name = 'Parada de Viaje'
        verbose_name_plural = 'Paradas de Viaje'
    
    def __str__(self):
        return f"{self.order}. {self.name} - {self.trip.title}"
    
    def get_google_maps_url(self):
        """Generar URL de Google Maps"""
        if self.latitude and self.longitude:
            return f"https://www.google.com/maps/search/?api=1&query={self.latitude},{self.longitude}"
        elif self.address:
            import urllib.parse
            query = urllib.parse.quote(self.address)
            return f"https://www.google.com/maps/search/?api=1&query={query}"
        return None


class WebCounter(models.Model):
    """Contador web para rastrear visitas en sitios externos"""
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        verbose_name='Usuario',
        related_name='web_counters'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Contador'
    )
    
    token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token',
        help_text='Token único para identificar este contador'
    )
    
    domain = models.CharField(
        max_length=255,
        blank=True,
        verbose_name='Dominio',
        help_text='Dominio donde se instalará el contador (opcional)'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    total_visits = models.PositiveIntegerField(
        default=0,
        verbose_name='Total de Visitas'
    )
    
    total_page_views = models.PositiveIntegerField(
        default=0,
        verbose_name='Total de Páginas Vistas'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Contador Web'
        verbose_name_plural = 'Contadores Web'
    
    def __str__(self):
        return f"{self.title} ({self.token})"
    
    def save(self, *args, **kwargs):
        if not self.token:
            import secrets
            self.token = secrets.token_urlsafe(32)
        super().save(*args, **kwargs)
    
    def get_script_html(self, request=None):
        """Generar el script HTML para insertar en sitios web"""
        from django.conf import settings
        
        # Usar configuración del settings o detectar automáticamente
        if request:
            base_url = request.build_absolute_uri('/').rstrip('/')
        else:
            # En producción usar el dominio configurado en settings
            base_url = getattr(settings, 'WEB_COUNTER_BASE_URL', 'https://ticketproo.com')
        
        script = f'''<!-- TicketProo Counter -->
<script type="text/javascript">
!function(){{var t="{self.token}",e="{base_url}/api/web-counter/track/";function n(){{fetch(e,{{method:"POST",headers:{{"Content-Type":"application/json"}},body:JSON.stringify({{token:t,url:window.location.href,referrer:document.referrer,user_agent:navigator.userAgent,screen_resolution:window.screen.width+"x"+window.screen.height,language:navigator.language}})}}).catch(function(){{}});}}"loading"===document.readyState?document.addEventListener("DOMContentLoaded",n):n();}}();
</script>'''
        return script


class WebCounterVisit(models.Model):
    """Registro de visitas del contador web"""
    counter = models.ForeignKey(
        WebCounter,
        on_delete=models.CASCADE,
        related_name='visits',
        verbose_name='Contador'
    )
    
    url = models.CharField(
        max_length=2048,
        verbose_name='URL Visitada'
    )
    
    referrer = models.CharField(
        max_length=2048,
        blank=True,
        verbose_name='Referrer'
    )
    
    ip_address = models.GenericIPAddressField(
        verbose_name='IP'
    )
    
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País'
    )
    
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad'
    )
    
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    browser = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Navegador'
    )
    
    os = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Sistema Operativo'
    )
    
    device_type = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Tipo de Dispositivo',
        help_text='mobile, tablet, desktop'
    )
    
    screen_resolution = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Resolución de Pantalla'
    )
    
    language = models.CharField(
        max_length=10,
        blank=True,
        verbose_name='Idioma'
    )
    
    session_id = models.CharField(
        max_length=64,
        blank=True,
        verbose_name='ID de Sesión'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Visita',
        db_index=True
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Visita de Contador Web'
        verbose_name_plural = 'Visitas de Contadores Web'
        indexes = [
            models.Index(fields=['-created_at', 'counter']),
            models.Index(fields=['counter', 'session_id']),
        ]
    
    def __str__(self):
        return f"{self.counter.title} - {self.url} ({self.created_at})"


class QuickQuote(models.Model):
    """Modelo para cotizaciones rápidas"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('accepted', 'Aceptado'),
        ('rejected', 'Rechazado'),
        ('expired', 'Expirado'),
    ]
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='quick_quotes',
        verbose_name='Usuario'
    )
    
    sequence_number = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Número de Secuencia',
        help_text='Formato: CR-YYYY-XX'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    
    hours = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Cantidad de Horas'
    )
    
    hourly_rate = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Costo por Hora'
    )
    
    total_amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Importe Total',
        editable=False
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    valid_until = models.DateTimeField(
        verbose_name='Válido Hasta'
    )
    
    public_token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token Público',
        help_text='Token único para acceso público'
    )
    
    company = models.ForeignKey(
        'Company',
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='quick_quotes',
        verbose_name='Empresa'
    )
    
    client_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre del Cliente'
    )
    
    client_email = models.EmailField(
        blank=True,
        verbose_name='Email del Cliente'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas'
    )
    
    response_date = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Respuesta'
    )
    
    response_notes = models.TextField(
        blank=True,
        verbose_name='Notas de Respuesta del Cliente'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Cotización Rápida'
        verbose_name_plural = 'Cotizaciones Rápidas'
    
    def save(self, *args, **kwargs):
        # Calcular el total automáticamente
        from decimal import Decimal
        self.total_amount = Decimal(str(self.hours)) * Decimal(str(self.hourly_rate))
        
        # Generar token público si no existe
        if not self.public_token:
            import secrets
            self.public_token = secrets.token_urlsafe(32)
        
        # Establecer fecha de vencimiento si no existe
        if not self.valid_until:
            from datetime import timedelta
            self.valid_until = timezone.now() + timedelta(days=7)
        
        # Generar número de secuencia si no existe
        if not self.sequence_number:
            year = timezone.now().year
            # Obtener el último número del año actual
            last_quote = QuickQuote.objects.filter(
                sequence_number__startswith=f'CR-{year}-'
            ).order_by('-sequence_number').first()
            
            if last_quote and last_quote.sequence_number:
                try:
                    last_num = int(last_quote.sequence_number.split('-')[-1])
                    new_num = last_num + 1
                except (ValueError, IndexError):
                    new_num = 1
            else:
                new_num = 1
            
            self.sequence_number = f'CR-{year}-{new_num:02d}'
        
        super().save(*args, **kwargs)
    
    def is_expired(self):
        """Verifica si la cotización ha expirado"""
        return timezone.now() > self.valid_until and self.status == 'pending'
    
    def get_public_url(self, request=None):
        """Retorna la URL pública para que el cliente responda"""
        from django.urls import reverse
        path = reverse('quick_quote_public', kwargs={'token': self.public_token})
        if request:
            return request.build_absolute_uri(path)
        from django.conf import settings
        base_url = getattr(settings, 'SITE_DOMAIN', 'localhost:8000')
        protocol = 'https' if 'localhost' not in base_url else 'http'
        return f"{protocol}://{base_url}{path}"
    
    def __str__(self):
        return f"{self.title} - {self.get_status_display()} - ${self.total_amount}"


class QuickQuoteComment(models.Model):
    """Modelo para comentarios públicos en cotizaciones"""
    
    quote = models.ForeignKey(
        QuickQuote,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Cotización'
    )
    
    author_name = models.CharField(
        max_length=200,
        verbose_name='Nombre'
    )
    
    author_email = models.EmailField(
        verbose_name='Email'
    )
    
    comment = models.TextField(
        verbose_name='Comentario'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    
    class Meta:
        verbose_name = 'Comentario de Cotización'
        verbose_name_plural = 'Comentarios de Cotizaciones'
        ordering = ['-created_at']
    
    def __str__(self):
        return f"{self.author_name} - {self.quote.sequence_number}"


class QuickQuoteView(models.Model):
    """Modelo para registrar visualizaciones de cotizaciones públicas"""
    
    quote = models.ForeignKey(
        QuickQuote,
        on_delete=models.CASCADE,
        related_name='views',
        verbose_name='Cotización'
    )
    
    viewed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Visualización'
    )
    
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    
    user_agent = models.TextField(
        blank=True,
        verbose_name='User Agent'
    )
    
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País'
    )
    
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad'
    )
    
    browser = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Navegador'
    )
    
    device_type = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Tipo de Dispositivo'
    )
    
    class Meta:
        ordering = ['-viewed_at']
        verbose_name = 'Visualización de Cotización'
        verbose_name_plural = 'Visualizaciones de Cotizaciones'
    
    def __str__(self):
        return f"Vista de {self.quote.title} - {self.viewed_at.strftime('%d/%m/%Y %H:%M')}"


class MultiMeasurement(models.Model):
    """Modelo para mediciones múltiples"""
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='multi_measurements',
        verbose_name='Usuario'
    )
    
    sequence_number = models.CharField(
        max_length=20,
        blank=True,
        verbose_name='Número de Secuencia',
        help_text='Formato: MM-YYYY-XX'
    )
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título'
    )
    
    public_token = models.CharField(
        max_length=64,
        unique=True,
        verbose_name='Token Público',
        help_text='Token único para acceso público'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    # Etiquetas para las 5 mediciones
    label_1 = models.CharField(
        max_length=100,
        default='Medición 1',
        verbose_name='Etiqueta Medición 1'
    )
    
    label_2 = models.CharField(
        max_length=100,
        default='Medición 2',
        verbose_name='Etiqueta Medición 2'
    )
    
    label_3 = models.CharField(
        max_length=100,
        default='Medición 3',
        verbose_name='Etiqueta Medición 3'
    )
    
    label_4 = models.CharField(
        max_length=100,
        default='Medición 4',
        verbose_name='Etiqueta Medición 4'
    )
    
    label_5 = models.CharField(
        max_length=100,
        default='Medición 5',
        verbose_name='Etiqueta Medición 5'
    )
    
    class Meta:
        verbose_name = 'Medición Múltiple'
        verbose_name_plural = 'Mediciones Múltiples'
        ordering = ['-created_at']
    
    def save(self, *args, **kwargs):
        # Generar token público si no existe
        if not self.public_token:
            import secrets
            self.public_token = secrets.token_urlsafe(32)
        
        # Generar número de secuencia si no existe
        if not self.sequence_number:
            year = timezone.now().year
            # Obtener el último número del año actual
            last_measurement = MultiMeasurement.objects.filter(
                sequence_number__startswith=f'MM-{year}-'
            ).order_by('-sequence_number').first()
            
            if last_measurement and last_measurement.sequence_number:
                try:
                    last_num = int(last_measurement.sequence_number.split('-')[-1])
                    new_num = last_num + 1
                except (ValueError, IndexError):
                    new_num = 1
            else:
                new_num = 1
            
            self.sequence_number = f'MM-{year}-{new_num:02d}'
        
        super().save(*args, **kwargs)
    
    def get_public_url(self, request=None):
        """Retorna la URL pública para cargar mediciones"""
        from django.urls import reverse
        path = reverse('multi_measurement_public', kwargs={'token': self.public_token})
        if request:
            return request.build_absolute_uri(path)
        from django.conf import settings
        base_url = getattr(settings, 'SITE_DOMAIN', 'localhost:8000')
        protocol = 'https' if 'localhost' not in base_url else 'http'
        return f"{protocol}://{base_url}{path}"
    
    def __str__(self):
        return f"{self.sequence_number} - {self.title}"


class MultiMeasurementRecord(models.Model):
    """Modelo para registros diarios de mediciones"""
    
    measurement = models.ForeignKey(
        MultiMeasurement,
        on_delete=models.CASCADE,
        related_name='records',
        verbose_name='Medición'
    )
    
    date = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha y Hora'
    )
    
    value_1 = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Valor 1'
    )
    
    value_2 = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Valor 2'
    )
    
    value_3 = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Valor 3'
    )
    
    value_4 = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Valor 4'
    )
    
    value_5 = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        null=True,
        blank=True,
        verbose_name='Valor 5'
    )
    
    notes = models.TextField(
        blank=True,
        verbose_name='Notas'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    ip_address = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='Dirección IP'
    )
    
    class Meta:
        verbose_name = 'Registro de Medición'
        verbose_name_plural = 'Registros de Mediciones'
        ordering = ['-date']
    
    def __str__(self):
        return f"{self.measurement.sequence_number} - {self.date.strftime('%d/%m/%Y %H:%M')}"


class PersonalBudget(models.Model):
    """Modelo para presupuestos personales anuales"""
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='personal_budgets',
        verbose_name='Usuario'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Presupuesto'
    )
    
    year = models.IntegerField(
        verbose_name='Año'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Creación'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    is_private = models.BooleanField(
        default=True,
        verbose_name='Privado',
        help_text='Si está marcado, solo el creador puede ver este presupuesto'
    )
    
    class Meta:
        verbose_name = 'Presupuesto Personal'
        verbose_name_plural = 'Presupuestos Personales'
        ordering = ['-year', 'name']
        unique_together = ['user', 'name', 'year']
    
    def __str__(self):
        return f"{self.name} - {self.year}"
    
    def get_month_display(self):
        """Retorna el nombre del mes"""
        months = ['', 'Enero', 'Febrero', 'Marzo', 'Abril', 'Mayo', 'Junio', 
                  'Julio', 'Agosto', 'Septiembre', 'Octubre', 'Noviembre', 'Diciembre']
        return months[1]  # Placeholder, no se usa más
    
    def get_total_income_budget(self, month=None):
        """Total presupuestado de ingresos (mensual o anual)"""
        from decimal import Decimal
        total_monthly = self.income_items.aggregate(total=models.Sum('budgeted_amount'))['total'] or Decimal('0')
        # Si se pide vista mensual, retornar el total mensual; si es anual, multiplicar por 12
        return total_monthly if month else total_monthly * Decimal('12')
    
    def get_total_expense_budget(self, month=None):
        """Total presupuestado de egresos (mensual o anual)"""
        from decimal import Decimal
        total_monthly = self.expense_items.aggregate(total=models.Sum('budgeted_amount'))['total'] or Decimal('0')
        # Si se pide vista mensual, retornar el total mensual; si es anual, multiplicar por 12
        return total_monthly if month else total_monthly * Decimal('12')
    
    def get_total_income_actual(self, month=None):
        """Total real de ingresos (filtrado por mes si se especifica)"""
        from decimal import Decimal
        qs = self.transactions.filter(transaction_type='income')
        if month:
            qs = qs.filter(month=month, year=self.year)
        return qs.aggregate(total=models.Sum('amount'))['total'] or Decimal('0')
    
    def get_total_expense_actual(self, month=None):
        """Total real de egresos (filtrado por mes si se especifica)"""
        from decimal import Decimal
        qs = self.transactions.filter(transaction_type='expense')
        if month:
            qs = qs.filter(month=month, year=self.year)
        return qs.aggregate(total=models.Sum('amount'))['total'] or Decimal('0')


class BudgetIncomeItem(models.Model):
    """Partidas de ingresos del presupuesto"""
    
    budget = models.ForeignKey(
        PersonalBudget,
        on_delete=models.CASCADE,
        related_name='income_items',
        verbose_name='Presupuesto'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Partida'
    )
    
    budgeted_amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0,
        verbose_name='Monto Presupuestado Mensual (€)',
        help_text='Monto mensual en euros para esta partida'
    )
    
    order = models.IntegerField(
        default=0,
        verbose_name='Orden'
    )
    
    class Meta:
        verbose_name = 'Partida de Ingreso'
        verbose_name_plural = 'Partidas de Ingresos'
        ordering = ['order', 'name']
    
    def __str__(self):
        return f"{self.budget} - {self.name}"
    
    def get_actual_amount(self, month=None, year=None):
        """Monto real de transacciones de ingreso (filtrado por mes/año si se especifica)"""
        from decimal import Decimal
        qs = self.transactions.all()
        if month and year:
            qs = qs.filter(month=month, year=year)
        return qs.aggregate(total=models.Sum('amount'))['total'] or Decimal('0')
    
    def get_percentage_budget(self):
        """Porcentaje respecto al presupuesto total de ingresos"""
        total = self.budget.get_total_income_budget()
        if total > 0:
            return (float(self.budgeted_amount) / float(total)) * 100
        return 0
    
    def get_percentage_actual(self):
        """Porcentaje real respecto al total real de ingresos"""
        total = self.budget.get_total_income_actual()
        if total > 0:
            return (float(self.get_actual_amount()) / float(total)) * 100
        return 0


class BudgetExpenseItem(models.Model):
    """Partidas de egresos del presupuesto"""
    
    budget = models.ForeignKey(
        PersonalBudget,
        on_delete=models.CASCADE,
        related_name='expense_items',
        verbose_name='Presupuesto'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Partida'
    )
    
    budgeted_amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        default=0,
        verbose_name='Monto Presupuestado Mensual (€)',
        help_text='Monto mensual en euros para esta partida'
    )
    
    order = models.IntegerField(
        default=0,
        verbose_name='Orden'
    )
    
    class Meta:
        verbose_name = 'Partida de Egreso'
        verbose_name_plural = 'Partidas de Egresos'
        ordering = ['order', 'name']
    
    def __str__(self):
        return f"{self.budget} - {self.name}"
    
    def get_actual_amount(self, month=None, year=None):
        """Monto real de transacciones de egreso (filtrado por mes/año si se especifica)"""
        from decimal import Decimal
        qs = self.transactions.all()
        if month and year:
            qs = qs.filter(month=month, year=year)
        return qs.aggregate(total=models.Sum('amount'))['total'] or Decimal('0')
    
    def get_percentage_budget(self):
        """Porcentaje respecto al presupuesto total de egresos"""
        total = self.budget.get_total_expense_budget()
        if total > 0:
            return (float(self.budgeted_amount) / float(total)) * 100
        return 0
    
    def get_percentage_actual(self):
        """Porcentaje real respecto al total real de egresos"""
        total = self.budget.get_total_expense_actual()
        if total > 0:
            return (float(self.get_actual_amount()) / float(total)) * 100
        return 0


class BudgetTransaction(models.Model):
    """Transacciones (ingresos o egresos) del presupuesto"""
    
    TRANSACTION_TYPE_CHOICES = [
        ('income', 'Ingreso'),
        ('expense', 'Egreso'),
    ]
    
    budget = models.ForeignKey(
        PersonalBudget,
        on_delete=models.CASCADE,
        related_name='transactions',
        verbose_name='Presupuesto'
    )
    
    transaction_type = models.CharField(
        max_length=10,
        choices=TRANSACTION_TYPE_CHOICES,
        verbose_name='Tipo'
    )
    
    income_item = models.ForeignKey(
        BudgetIncomeItem,
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        related_name='transactions',
        verbose_name='Partida de Ingreso'
    )
    
    expense_item = models.ForeignKey(
        BudgetExpenseItem,
        on_delete=models.CASCADE,
        null=True,
        blank=True,
        related_name='transactions',
        verbose_name='Partida de Egreso'
    )
    
    amount = models.DecimalField(
        max_digits=10,
        decimal_places=2,
        verbose_name='Monto'
    )
    
    description = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Descripción'
    )
    
    transaction_date = models.DateField(
        verbose_name='Fecha de Transacción'
    )
    
    month = models.IntegerField(
        default=1,
        verbose_name='Mes',
        help_text='Mes de la transacción (1-12)'
    )
    
    year = models.IntegerField(
        default=2025,
        verbose_name='Año',
        help_text='Año de la transacción'
    )
    
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de Registro'
    )
    
    class Meta:
        verbose_name = 'Transacción'
        verbose_name_plural = 'Transacciones'
        ordering = ['-transaction_date', '-created_at']
    
    def __str__(self):
        return f"{self.get_transaction_type_display()} - ${self.amount} - {self.transaction_date}"
    
    def save(self, *args, **kwargs):
        # Extraer mes y año de la fecha de transacción
        if self.transaction_date:
            self.month = self.transaction_date.month
            self.year = self.transaction_date.year
        super().save(*args, **kwargs)


# Añadir campos Many-to-Many al modelo PersonalBudget
PersonalBudget.add_to_class('income_transactions', 
    models.ManyToManyField(BudgetTransaction, related_name='income_budgets', blank=True))
PersonalBudget.add_to_class('expense_transactions', 
    models.ManyToManyField(BudgetTransaction, related_name='expense_budgets', blank=True))


# ============= MODELOS PARA TABLAS DINÁMICAS API =============

class DynamicTable(models.Model):
    """Modelo para definir tablas dinámicas con API CRUD"""
    
    name = models.CharField(
        max_length=100,
        unique=True,
        verbose_name='Nombre de la tabla',
        help_text='Nombre único de la tabla (ej: usuarios, comentarios)'
    )
    display_name = models.CharField(
        max_length=200,
        verbose_name='Nombre de visualización',
        help_text='Nombre amigable para mostrar'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción del propósito de la tabla'
    )
    api_token = models.UUIDField(
        default=uuid.uuid4,
        unique=True,
        editable=False,
        verbose_name='Token de API',
        help_text='Token de seguridad para acceder a la API de esta tabla'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Si la tabla y su API están activas'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='dynamic_tables_created',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Configuración de permisos
    allow_public_read = models.BooleanField(
        default=False,
        verbose_name='Permitir lectura pública',
        help_text='Permitir GET sin autenticación'
    )
    allow_public_create = models.BooleanField(
        default=False,
        verbose_name='Permitir creación pública',
        help_text='Permitir POST sin autenticación'
    )
    
    # Configuración de formulario público móvil
    public_form_enabled = models.BooleanField(
        default=False,
        verbose_name='Formulario público habilitado',
        help_text='Habilitar formulario web público para crear registros'
    )
    public_form_token = models.UUIDField(
        default=uuid.uuid4,
        null=True,
        blank=True,
        unique=True,
        verbose_name='Token del formulario público',
        help_text='Token único para acceder al formulario público'
    )
    public_form_password = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Contraseña del formulario',
        help_text='Contraseña opcional para proteger el formulario público'
    )
    public_form_title = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Título del formulario',
        help_text='Título personalizado para el formulario público'
    )
    public_form_description = models.TextField(
        blank=True,
        verbose_name='Descripción del formulario',
        help_text='Texto descriptivo que se muestra en el formulario'
    )
    public_form_success_message = models.CharField(
        max_length=500,
        default='¡Gracias! Tu información ha sido registrada correctamente.',
        verbose_name='Mensaje de éxito',
        help_text='Mensaje que se muestra después de enviar el formulario'
    )
    
    class Meta:
        verbose_name = 'Tabla Dinámica'
        verbose_name_plural = 'Tablas Dinámicas'
        ordering = ['-created_at']
    
    def __str__(self):
        return self.display_name
    
    def get_api_url(self):
        """Retorna la URL base de la API para esta tabla"""
        return f'/api/dynamic-tables/{self.name}/'
    
    def regenerate_token(self):
        """Regenera el token de API"""
        self.api_token = uuid.uuid4()
        self.save()
        return self.api_token
    
    def get_public_form_url(self):
        """Retorna la URL del formulario público"""
        return f'/public-form/{self.public_form_token}/'
    
    def regenerate_public_form_token(self):
        """Regenera el token del formulario público"""
        self.public_form_token = uuid.uuid4()
        self.save()
        return self.public_form_token


class DynamicTableField(models.Model):
    """Modelo para definir campos de una tabla dinámica"""
    
    FIELD_TYPES = [
        ('text', 'Texto corto'),
        ('textarea', 'Texto largo'),
        ('number', 'Número'),
        ('decimal', 'Decimal'),
        ('boolean', 'Verdadero/Falso'),
        ('date', 'Fecha'),
        ('datetime', 'Fecha y hora'),
        ('email', 'Email'),
        ('url', 'URL'),
        ('json', 'JSON'),
    ]
    
    table = models.ForeignKey(
        DynamicTable,
        on_delete=models.CASCADE,
        related_name='fields',
        verbose_name='Tabla'
    )
    name = models.CharField(
        max_length=100,
        verbose_name='Nombre del campo',
        help_text='Nombre del campo (ej: titulo, email, contenido)'
    )
    display_name = models.CharField(
        max_length=200,
        verbose_name='Nombre de visualización',
        help_text='Nombre amigable para mostrar'
    )
    field_type = models.CharField(
        max_length=20,
        choices=FIELD_TYPES,
        default='text',
        verbose_name='Tipo de campo'
    )
    is_required = models.BooleanField(
        default=False,
        verbose_name='Requerido',
        help_text='Si el campo es obligatorio'
    )
    default_value = models.TextField(
        blank=True,
        null=True,
        verbose_name='Valor por defecto',
        help_text='Valor por defecto del campo'
    )
    max_length = models.IntegerField(
        null=True,
        blank=True,
        verbose_name='Longitud máxima',
        help_text='Longitud máxima para campos de texto'
    )
    help_text = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Texto de ayuda',
        help_text='Descripción o ayuda del campo'
    )
    order = models.IntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de visualización del campo'
    )
    is_unique = models.BooleanField(
        default=False,
        verbose_name='Único',
        help_text='Si el valor debe ser único en la tabla'
    )
    is_indexed = models.BooleanField(
        default=False,
        verbose_name='Indexado',
        help_text='Si el campo debe ser indexado para búsquedas rápidas'
    )
    
    class Meta:
        verbose_name = 'Campo de Tabla'
        verbose_name_plural = 'Campos de Tablas'
        ordering = ['table', 'order', 'id']
        unique_together = [['table', 'name']]
    
    def __str__(self):
        return f'{self.table.name}.{self.name}'


class DynamicTableRecord(models.Model):
    """Modelo para almacenar registros de tablas dinámicas"""
    
    table = models.ForeignKey(
        DynamicTable,
        on_delete=models.CASCADE,
        related_name='records',
        verbose_name='Tabla'
    )
    data = models.JSONField(
        verbose_name='Datos',
        help_text='Datos del registro en formato JSON'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    created_by_ip = models.GenericIPAddressField(
        null=True,
        blank=True,
        verbose_name='IP de creación',
        help_text='Dirección IP desde donde se creó el registro'
    )
    
    class Meta:
        verbose_name = 'Registro de Tabla'
        verbose_name_plural = 'Registros de Tablas'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['table', '-created_at']),
        ]
    
    def __str__(self):
        return f'{self.table.name} #{self.id}'


class SavedApiRequest(models.Model):
    """Modelo para guardar configuraciones de peticiones API"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Nombre descriptivo para identificar esta petición'
    )
    url = models.URLField(
        max_length=500,
        verbose_name='URL',
        help_text='URL completa del endpoint'
    )
    method = models.CharField(
        max_length=10,
        choices=[
            ('GET', 'GET'),
            ('POST', 'POST'),
            ('PUT', 'PUT'),
            ('PATCH', 'PATCH'),
            ('DELETE', 'DELETE'),
        ],
        default='GET',
        verbose_name='Método HTTP'
    )
    headers = models.TextField(
        blank=True,
        verbose_name='Headers',
        help_text='Headers personalizados (formato: Name: Value, uno por línea)'
    )
    body = models.TextField(
        blank=True,
        verbose_name='Body',
        help_text='Cuerpo de la petición (JSON)'
    )
    token = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Token',
        help_text='Token de autenticación (Bearer, API Key, etc.)'
    )
    api_type = models.CharField(
        max_length=20,
        choices=[
            ('internal', 'API Interna'),
            ('external', 'API Externa'),
        ],
        default='external',
        verbose_name='Tipo de API'
    )
    dynamic_table = models.ForeignKey(
        DynamicTable,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='saved_requests',
        verbose_name='Tabla Dinámica',
        help_text='Solo si es API interna'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Notas sobre esta petición'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='saved_api_requests',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    is_favorite = models.BooleanField(
        default=False,
        verbose_name='Favorito',
        help_text='Marcar como favorito para acceso rápido'
    )
    last_used = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Último uso',
        help_text='Última vez que se ejecutó esta petición'
    )
    
    class Meta:
        verbose_name = 'API Guardada'
        verbose_name_plural = 'APIs Guardadas'
        ordering = ['-is_favorite', '-last_used', '-created_at']
        indexes = [
            models.Index(fields=['created_by', '-created_at']),
            models.Index(fields=['created_by', 'is_favorite']),
        ]
    
    def __str__(self):
        return self.title
    
    def mark_as_used(self):
        """Actualiza el timestamp de último uso"""
        self.last_used = timezone.now()
        self.save(update_fields=['last_used'])


class SocialPost(models.Model):
    """Publicación en la red social interna"""
    
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='social_posts',
        verbose_name='Usuario'
    )
    content = models.TextField(
        blank=True,
        verbose_name='Contenido',
        help_text='Texto de la publicación'
    )
    image = models.ImageField(
        upload_to='social_posts/%Y/%m/',
        blank=True,
        null=True,
        verbose_name='Imagen',
        help_text='Imagen de la publicación'
    )
    youtube_url = models.URLField(
        blank=True,
        null=True,
        verbose_name='URL de YouTube',
        help_text='URL del video de YouTube'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    is_private = models.BooleanField(
        default=False,
        verbose_name='Privado',
        help_text='Solo visible para el creador'
    )
    share_token = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        unique=True,
        verbose_name='Token de compartir',
        help_text='Token único para compartir públicamente'
    )
    company_only = models.BooleanField(
        default=True,
        verbose_name='Solo mi empresa',
        help_text='Si está marcado, solo los usuarios de la misma empresa pueden ver esta publicación'
    )
    public_views_count = models.PositiveIntegerField(
        default=0,
        verbose_name='Vistas públicas',
        help_text='Número de veces que se ha visto la publicación pública'
    )
    
    class Meta:
        verbose_name = 'Publicación Social'
        verbose_name_plural = 'Publicaciones Sociales'
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['-created_at']),
            models.Index(fields=['user', '-created_at']),
            models.Index(fields=['share_token']),
            models.Index(fields=['company_only', '-created_at']),
        ]
    
    def __str__(self):
        content_preview = self.content[:50] if self.content else 'Sin texto'
        return f'{self.user.username} - {content_preview}'
    
    def get_likes_count(self):
        """Retorna el número de me gusta"""
        return self.likes.filter(reaction_type='like').count()
    
    def get_loves_count(self):
        """Retorna el número de me encanta"""
        return self.likes.filter(reaction_type='love').count()
    
    def get_dislikes_count(self):
        """Retorna el número de no me gusta"""
        return self.likes.filter(reaction_type='dislike').count()
    
    def get_comments_count(self):
        """Retorna el número de comentarios"""
        return self.comments.filter(is_active=True).count()
    
    def user_has_liked(self, user):
        """Verifica si un usuario ha dado like"""
        if not user.is_authenticated:
            return False
        return self.likes.filter(user=user).exists()
    
    def get_youtube_embed_url(self):
        """Convierte URL de YouTube a URL embebida"""
        if not self.youtube_url:
            return None
        
        import re
        # Patrones para diferentes formatos de URL de YouTube
        patterns = [
            r'(?:https?://)?(?:www\.)?youtube\.com/watch\?v=([\w-]+)',
            r'(?:https?://)?(?:www\.)?youtu\.be/([\w-]+)',
            r'(?:https?://)?(?:www\.)?youtube\.com/embed/([\w-]+)'
        ]
        
        for pattern in patterns:
            match = re.search(pattern, self.youtube_url)
            if match:
                video_id = match.group(1)
                return f'https://www.youtube.com/embed/{video_id}'
        
        return None
    
    def generate_share_token(self):
        """Genera un token único para compartir"""
        import uuid
        if not self.share_token:
            self.share_token = str(uuid.uuid4())
            self.save()
        return self.share_token
    
    def get_public_share_url(self):
        """Obtiene la URL pública para compartir"""
        if not self.share_token:
            return None
        from django.urls import reverse
        return reverse('social_post_public', kwargs={'token': self.share_token})


class SocialPostLike(models.Model):
    """Reacción en una publicación"""
    
    REACTION_CHOICES = [
        ('like', 'Me gusta'),
        ('love', 'Me encanta'),
        ('dislike', 'No me gusta'),
    ]
    
    post = models.ForeignKey(
        SocialPost,
        on_delete=models.CASCADE,
        related_name='likes',
        verbose_name='Publicación'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='social_likes',
        verbose_name='Usuario'
    )
    reaction_type = models.CharField(
        max_length=10,
        choices=REACTION_CHOICES,
        default='like',
        verbose_name='Tipo de reacción'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha'
    )
    
    class Meta:
        verbose_name = 'Like'
        verbose_name_plural = 'Likes'
        unique_together = ('post', 'user')
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['post', 'user']),
        ]
    
    def __str__(self):
        return f'{self.user.username} - Post #{self.post.id}'


class SocialPostComment(models.Model):
    """Comentario en una publicación"""
    
    post = models.ForeignKey(
        SocialPost,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Publicación'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='social_comments',
        verbose_name='Usuario',
        blank=True,
        null=True
    )
    # Campos para usuarios anónimos (comentarios públicos)
    guest_name = models.CharField(
        max_length=100,
        blank=True,
        null=True,
        verbose_name='Nombre del invitado'
    )
    guest_email = models.EmailField(
        blank=True,
        null=True,
        verbose_name='Email del invitado'
    )
    content = models.TextField(
        verbose_name='Comentario'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    class Meta:
        verbose_name = 'Comentario'
        verbose_name_plural = 'Comentarios'
        ordering = ['created_at']
        indexes = [
            models.Index(fields=['post', 'created_at']),
        ]
    
    def __str__(self):
        content_preview = self.content[:50]
        username = self.user.username if self.user else self.guest_name or 'Anónimo'
        return f'{username} - {content_preview}'
    
    def get_display_name(self):
        """Retorna el nombre a mostrar (usuario o invitado)"""
        if self.user:
            return self.user.get_full_name() or self.user.username
        return self.guest_name or 'Anónimo'
    
    def get_likes_count(self):
        """Retorna el número de me gusta en el comentario"""
        return self.comment_reactions.filter(reaction_type='like').count()
    
    def get_loves_count(self):
        """Retorna el número de me encanta en el comentario"""
        return self.comment_reactions.filter(reaction_type='love').count()
    
    def get_dislikes_count(self):
        """Retorna el número de no me gusta en el comentario"""
        return self.comment_reactions.filter(reaction_type='dislike').count()


class SocialCommentReaction(models.Model):
    """Reacción en un comentario"""
    
    REACTION_CHOICES = [
        ('like', 'Me gusta'),
        ('love', 'Me encanta'),
        ('dislike', 'No me gusta'),
    ]
    
    comment = models.ForeignKey(
        SocialPostComment,
        on_delete=models.CASCADE,
        related_name='comment_reactions',
        verbose_name='Comentario'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='social_comment_reactions',
        verbose_name='Usuario'
    )
    reaction_type = models.CharField(
        max_length=10,
        choices=REACTION_CHOICES,
        default='like',
        verbose_name='Tipo de reacción'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha'
    )
    
    class Meta:
        verbose_name = 'Reacción de Comentario'
        verbose_name_plural = 'Reacciones de Comentarios'
        unique_together = ('comment', 'user')
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['comment', 'user']),
        ]
    
    def __str__(self):
        return f'{self.user.username} - {self.reaction_type} - Comment #{self.comment.id}'


class SocialPostFavorite(models.Model):
    """Publicación marcada como favorita por un usuario"""
    
    post = models.ForeignKey(
        SocialPost,
        on_delete=models.CASCADE,
        related_name='favorites',
        verbose_name='Publicación'
    )
    user = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='favorite_posts',
        verbose_name='Usuario'
    )
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha'
    )
    
    class Meta:
        verbose_name = 'Favorito'
        verbose_name_plural = 'Favoritos'
        unique_together = ('post', 'user')
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['user', '-created_at']),
            models.Index(fields=['post', 'user']),
        ]
    
    def __str__(self):
        return f'{self.user.username} - Post #{self.post.id}'


class FunctionalRequirementDocument(models.Model):
    """Documento de Requerimientos Funcionales (DRF)"""
    
    company = models.ForeignKey(
        Company,
        on_delete=models.CASCADE,
        related_name='functional_requirement_documents',
        verbose_name='Empresa',
        null=True,
        blank=True
    )
    company_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre de Empresa',
        help_text='Nombre libre de la empresa si no está en el listado'
    )
    title = models.CharField(
        max_length=300,
        verbose_name='Título'
    )
    description = models.TextField(
        blank=True,
        verbose_name='Descripción General',
        help_text='Descripción general del proyecto o alcance'
    )
    sequence = models.CharField(
        max_length=50,
        verbose_name='Secuencia',
        help_text='Se genera automáticamente',
        unique=True,
        editable=False
    )
    public_share_token = models.UUIDField(
        default=uuid.uuid4,
        editable=False,
        unique=True,
        verbose_name='Token público'
    )
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='created_frd_documents',
        verbose_name='Creado por'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    # Campos de aceptación/rechazo del documento completo
    document_status = models.CharField(
        max_length=20,
        choices=[
            ('pending', 'Pendiente'),
            ('accepted', 'Aceptado'),
            ('rejected', 'Rechazado')
        ],
        default='pending',
        verbose_name='Estado del Documento'
    )
    document_reviewed_by_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Revisado por'
    )
    document_reviewed_by_email = models.EmailField(
        blank=True,
        verbose_name='Email del revisor'
    )
    document_reviewed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de revisión'
    )
    document_review_comments = models.TextField(
        blank=True,
        verbose_name='Comentarios de revisión'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Documento de Requerimientos Funcionales'
        verbose_name_plural = 'Documentos de Requerimientos Funcionales'
    
    def __str__(self):
        return f"{self.sequence} - {self.title}"
    
    def get_company_display(self):
        """Retorna el nombre de la empresa (FK o libre)"""
        if self.company:
            return self.company.name
        return self.company_name or 'Sin empresa'
    
    @staticmethod
    def generate_sequence():
        """Genera la secuencia automática en formato DRF-YYYY-NNN"""
        from datetime import datetime
        year = datetime.now().year
        prefix = f"DRF-{year}-"
        
        # Buscar el último documento del año
        last_doc = FunctionalRequirementDocument.objects.filter(
            sequence__startswith=prefix
        ).order_by('-sequence').first()
        
        if last_doc:
            # Extraer el número y incrementar
            last_number = int(last_doc.sequence.split('-')[-1])
            new_number = last_number + 1
        else:
            new_number = 1
        
        return f"{prefix}{new_number:03d}"
    
    def get_public_url(self):
        """Retorna la URL pública del documento"""
        from django.urls import reverse
        return reverse('public_frd', kwargs={'token': self.public_share_token})
    
    def get_approved_count(self):
        """Retorna el número de requerimientos aprobados"""
        return self.requirements.filter(is_approved=True).count()
    
    def get_total_count(self):
        """Retorna el número total de requerimientos"""
        return self.requirements.count()
    
    def get_approval_percentage(self):
        """Retorna el porcentaje de aprobación"""
        total = self.get_total_count()
        if total == 0:
            return 0
        return round((self.get_approved_count() / total) * 100, 1)
    
    def get_pending_count(self):
        """Retorna el número de requerimientos pendientes"""
        return self.get_total_count() - self.get_approved_count()
    
    def get_total_views(self):
        """Retorna el número total de visitas a la página pública"""
        return self.views.count()
    
    def get_unique_visitors(self):
        """Retorna el número de visitantes únicos por IP"""
        return self.views.values('ip_address').distinct().count()
    
    def get_recent_views(self, limit=10):
        """Retorna las últimas visitas"""
        return self.views.order_by('-viewed_at')[:limit]
    
    def get_views_by_date(self):
        """Retorna las visitas agrupadas por fecha"""
        from django.db.models import Count
        from django.db.models.functions import TruncDate
        return self.views.annotate(
            date=TruncDate('viewed_at')
        ).values('date').annotate(
            count=Count('id')
        ).order_by('-date')
    
    def accept_document(self, name, email, comments=''):
        """Acepta el documento completo"""
        self.document_status = 'accepted'
        self.document_reviewed_by_name = name
        self.document_reviewed_by_email = email
        self.document_review_comments = comments
        self.document_reviewed_at = timezone.now()
        self.save()
    
    def reject_document(self, name, email, comments=''):
        """Rechaza el documento completo"""
        self.document_status = 'rejected'
        self.document_reviewed_by_name = name
        self.document_reviewed_by_email = email
        self.document_review_comments = comments
        self.document_reviewed_at = timezone.now()
        self.save()
    
    def is_document_accepted(self):
        """Retorna True si el documento está aceptado"""
        return self.document_status == 'accepted'
    
    def is_document_rejected(self):
        """Retorna True si el documento está rechazado"""
        return self.document_status == 'rejected'
    
    def is_document_pending(self):
        """Retorna True si el documento está pendiente"""
        return self.document_status == 'pending'


class FunctionalRequirement(models.Model):
    """Requerimiento individual dentro de un DRF"""
    
    document = models.ForeignKey(
        FunctionalRequirementDocument,
        on_delete=models.CASCADE,
        related_name='requirements',
        verbose_name='Documento'
    )
    number = models.PositiveIntegerField(
        verbose_name='Número',
        help_text='Número secuencial del requerimiento'
    )
    title = models.CharField(
        max_length=300,
        blank=True,
        verbose_name='Título',
        help_text='Título breve del requerimiento'
    )
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Descripción del alcance o requerimiento'
    )
    is_approved = models.BooleanField(
        default=False,
        verbose_name='Aprobado'
    )
    approved_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de aprobación'
    )
    approved_by_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Aprobado por (nombre)'
    )
    approved_by_email = models.EmailField(
        blank=True,
        verbose_name='Aprobado por (email)'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    class Meta:
        ordering = ['document', 'number']
        verbose_name = 'Requerimiento Funcional'
        verbose_name_plural = 'Requerimientos Funcionales'
        unique_together = ['document', 'number']
    
    def __str__(self):
        return f"{self.document.sequence} - Req #{self.number}"
    
    def approve(self, name, email):
        """Aprueba el requerimiento"""
        self.is_approved = True
        self.approved_at = timezone.now()
        self.approved_by_name = name
        self.approved_by_email = email
        self.save()


class FunctionalRequirementDocumentView(models.Model):
    """Registra las visitas a la página pública de un DRF"""
    
    document = models.ForeignKey(
        FunctionalRequirementDocument,
        on_delete=models.CASCADE,
        related_name='views',
        verbose_name='Documento'
    )
    ip_address = models.GenericIPAddressField(
        verbose_name='Dirección IP',
        null=True,
        blank=True
    )
    user_agent = models.TextField(
        verbose_name='User Agent',
        blank=True
    )
    referrer = models.URLField(
        verbose_name='Referrer',
        blank=True,
        max_length=500
    )
    viewed_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de visita'
    )
    
    class Meta:
        ordering = ['-viewed_at']
        verbose_name = 'Vista de DRF'
        verbose_name_plural = 'Vistas de DRF'
    
    def __str__(self):
        return f"{self.document.sequence} - {self.viewed_at.strftime('%d/%m/%Y %H:%M')}"


class FunctionalRequirementComment(models.Model):
    """Comentarios en los requerimientos funcionales"""
    
    requirement = models.ForeignKey(
        FunctionalRequirement,
        on_delete=models.CASCADE,
        related_name='comments',
        verbose_name='Requerimiento'
    )
    author_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del autor'
    )
    author_email = models.EmailField(
        verbose_name='Email del autor'
    )
    comment = models.TextField(
        verbose_name='Comentario'
    )
    created_at = models.DateTimeField(
        default=timezone.now,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['created_at']
        verbose_name = 'Comentario de Requerimiento'
        verbose_name_plural = 'Comentarios de Requerimientos'
    
    def __str__(self):
        return f"Comentario de {self.author_name} en Req #{self.requirement.number}"


class AccessGroup(models.Model):
    """Modelo para grupos de accesos rápidos"""
    title = models.CharField(max_length=200, verbose_name='Título del Grupo')
    description = models.TextField(blank=True, verbose_name='Descripción')
    is_public = models.BooleanField(default=False, verbose_name='Público')
    share_token = models.CharField(max_length=50, unique=True, blank=True, verbose_name='Token de Compartir')
    created_by = models.ForeignKey(User, on_delete=models.CASCADE, related_name='access_groups', verbose_name='Creado por')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    order = models.IntegerField(default=0, verbose_name='Orden')
    
    class Meta:
        ordering = ['order', '-created_at']
        verbose_name = 'Grupo de Accesos'
        verbose_name_plural = 'Grupos de Accesos'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        if not self.share_token:
            self.share_token = str(uuid.uuid4())[:12]
        super().save(*args, **kwargs)
    
    def get_public_url(self):
        """Retorna la URL pública del grupo"""
        if self.is_public and self.share_token:
            from django.urls import reverse
            return reverse('access_group_public', kwargs={'token': self.share_token})
        return None


class AccessLink(models.Model):
    """Modelo para enlaces dentro de un grupo de accesos"""
    group = models.ForeignKey(AccessGroup, on_delete=models.CASCADE, related_name='links', verbose_name='Grupo')
    title = models.CharField(max_length=200, verbose_name='Título del Enlace')
    url = models.URLField(max_length=500, verbose_name='URL')
    description = models.TextField(blank=True, verbose_name='Descripción')
    icon = models.CharField(
        max_length=50, 
        blank=True, 
        verbose_name='Icono',
        help_text='Clase de icono Bootstrap (ej: bi-link-45deg)'
    )
    color = models.CharField(
        max_length=20,
        default='primary',
        verbose_name='Color',
        help_text='Color del botón (primary, success, info, warning, danger, etc.)'
    )
    order = models.IntegerField(default=0, verbose_name='Orden')
    is_active = models.BooleanField(default=True, verbose_name='Activo')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Enlace de Acceso'
        verbose_name_plural = 'Enlaces de Acceso'
    
    def __str__(self):
        return f"{self.title} ({self.group.title})"


# ============= MODELOS DE PLANIFICACIÓN DE TAREAS =============

class TaskPlan(models.Model):
    """Modelo para plan de tareas con días"""
    title = models.CharField(max_length=200, verbose_name='Título del Plan')
    description = models.TextField(blank=True, verbose_name='Descripción')
    created_by = models.ForeignKey(User, on_delete=models.CASCADE, related_name='task_plans', verbose_name='Creado por')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    is_public = models.BooleanField(default=False, verbose_name='Vista Pública Activada')
    share_token = models.CharField(max_length=50, unique=True, blank=True, verbose_name='Token de Compartir')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Plan de Tareas'
        verbose_name_plural = 'Planes de Tareas'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        if not self.share_token:
            self.share_token = str(uuid.uuid4())[:12]
        super().save(*args, **kwargs)
    
    def get_public_url(self):
        """Retorna la URL pública del plan"""
        if self.is_public and self.share_token:
            from django.urls import reverse
            return reverse('task_plan_public', kwargs={'token': self.share_token})
        return None


class TaskPlanDay(models.Model):
    """Modelo para días dentro de un plan de tareas"""
    plan = models.ForeignKey(TaskPlan, on_delete=models.CASCADE, related_name='days', verbose_name='Plan')
    title = models.CharField(max_length=100, verbose_name='Título del Día', help_text='Ej: Día 1, Lunes, Semana 1')
    description = models.TextField(blank=True, verbose_name='Descripción del Día')
    order = models.IntegerField(default=0, verbose_name='Orden')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Día del Plan'
        verbose_name_plural = 'Días del Plan'
        unique_together = ['plan', 'order']
    
    def __str__(self):
        return f"{self.title} - {self.plan.title}"


class TaskPlanItem(models.Model):
    """Modelo para tareas individuales dentro de un día"""
    day = models.ForeignKey(TaskPlanDay, on_delete=models.CASCADE, related_name='tasks', verbose_name='Día')
    title = models.CharField(max_length=300, verbose_name='Título de la Tarea')
    description = models.TextField(blank=True, verbose_name='Descripción')
    order = models.IntegerField(default=0, verbose_name='Orden')
    is_completed = models.BooleanField(default=False, verbose_name='Completada')
    completed_at = models.DateTimeField(null=True, blank=True, verbose_name='Completada el')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    
    # Campos opcionales para mayor detalle
    estimated_hours = models.DecimalField(
        max_digits=5, 
        decimal_places=2, 
        null=True, 
        blank=True, 
        verbose_name='Horas Estimadas'
    )
    priority = models.CharField(
        max_length=20,
        choices=[
            ('low', 'Baja'),
            ('medium', 'Media'),
            ('high', 'Alta'),
            ('critical', 'Crítica')
        ],
        default='medium',
        verbose_name='Prioridad'
    )
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Tarea del Plan'
        verbose_name_plural = 'Tareas del Plan'
    
    def __str__(self):
        return f"{self.title} ({self.day.title})"


# ===================================
# MODELOS DE CHECKLIST
# ===================================

class Checklist(models.Model):
    """Modelo para checklist con tareas"""
    title = models.CharField(max_length=300, verbose_name='Título del Checklist')
    company = models.CharField(max_length=200, blank=True, verbose_name='Empresa')
    description = models.TextField(blank=True, verbose_name='Descripción')
    is_public = models.BooleanField(default=False, verbose_name='Público')
    share_token = models.CharField(max_length=100, unique=True, blank=True, verbose_name='Token de Compartir')
    views_count = models.IntegerField(default=0, verbose_name='Visualizaciones')
    created_by = models.ForeignKey(User, on_delete=models.CASCADE, verbose_name='Creado por')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Checklist'
        verbose_name_plural = 'Checklists'
    
    def __str__(self):
        return f"{self.title} - {self.company}" if self.company else self.title
    
    def save(self, *args, **kwargs):
        if not self.share_token:
            self.share_token = str(uuid.uuid4())[:12]
        super().save(*args, **kwargs)
    
    def get_public_url(self):
        return f"/checklist/{self.share_token}/"
    
    def get_completion_percentage(self):
        """Calcular porcentaje de completado"""
        total = self.items.count()
        if total == 0:
            return 0
        completed = self.items.filter(is_completed=True).count()
        return int((completed / total) * 100)


class ChecklistItem(models.Model):
    """Modelo para items individuales del checklist"""
    checklist = models.ForeignKey(Checklist, on_delete=models.CASCADE, related_name='items', verbose_name='Checklist')
    title = models.CharField(max_length=500, verbose_name='Título de la Tarea')
    description = models.TextField(blank=True, verbose_name='Descripción')
    cost = models.DecimalField(max_digits=10, decimal_places=2, null=True, blank=True, verbose_name='Costo')
    order = models.IntegerField(default=0, verbose_name='Orden')
    is_completed = models.BooleanField(default=False, verbose_name='Completada')
    completed_at = models.DateTimeField(null=True, blank=True, verbose_name='Completada el')
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    
    class Meta:
        ordering = ['order', 'created_at']
        verbose_name = 'Item del Checklist'
        verbose_name_plural = 'Items del Checklist'
    
    def __str__(self):
        return f"{self.title} ({self.checklist.title})"


class Transaction(models.Model):
    """Modelo para transacciones rápidas - accesos directos a URLs del sistema"""
    
    code = models.CharField(
        max_length=50,
        unique=True,
        verbose_name='Código',
        help_text='Código único de la transacción (ej: T001, TICKET, REPORTE)'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre',
        help_text='Nombre descriptivo de la transacción'
    )
    
    url = models.CharField(
        max_length=500,
        verbose_name='URL',
        help_text='URL de acceso (puede ser relativa o absoluta)'
    )
    
    visible_for_all = models.BooleanField(
        default=True,
        verbose_name='Visible para Todos',
        help_text='Si está marcado, todos pueden ver esta transacción. Si no, solo agentes.'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción opcional de la transacción'
    )
    
    icon = models.CharField(
        max_length=50,
        blank=True,
        default='bi-link-45deg',
        verbose_name='Ícono Bootstrap',
        help_text='Clase del ícono Bootstrap (ej: bi-link-45deg, bi-file-text)'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa',
        help_text='Solo las transacciones activas aparecen en la búsqueda'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='created_transactions',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    
    class Meta:
        ordering = ['code']
        verbose_name = 'Transacción'
        verbose_name_plural = 'Transacciones'
        indexes = [
            models.Index(fields=['code']),
            models.Index(fields=['name']),
            models.Index(fields=['is_active']),
        ]
    
    def __str__(self):
        return f"{self.code} - {self.name}"


class KnowledgeBase(models.Model):
    """Modelo para Base de Conocimientos - conceptos técnicos y código fuente"""
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título',
        help_text='Título del concepto o tema'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción detallada del concepto'
    )
    
    source_code = models.TextField(
        blank=True,
        verbose_name='Código Fuente',
        help_text='Código fuente de ejemplo o snippet'
    )
    
    tags = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Etiquetas',
        help_text='Etiquetas separadas por comas (ej: python, django, api)'
    )
    
    category = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Categoría',
        help_text='Categoría del conocimiento (ej: Backend, Frontend, Base de Datos)'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='knowledge_bases_created',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(default=timezone.now, verbose_name='Fecha de Creación')
    updated_at = models.DateTimeField(auto_now=True, verbose_name='Última Actualización')
    
    # Campos para compartir públicamente
    public_share_enabled = models.BooleanField(default=False, verbose_name='Compartir Públicamente')
    public_share_token = models.CharField(max_length=64, unique=True, null=True, blank=True, verbose_name='Token de Compartir')
    public_share_password = models.CharField(max_length=128, blank=True, verbose_name='Contraseña de Acceso')
    public_share_expires_at = models.DateTimeField(null=True, blank=True, verbose_name='Fecha de Expiración')
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Base de Conocimiento'
        verbose_name_plural = 'Base de Conocimientos'
        indexes = [
            models.Index(fields=['title']),
            models.Index(fields=['category']),
            models.Index(fields=['is_active']),
        ]
    
    def __str__(self):
        return self.title
    
    def generate_share_token(self):
        """Genera un token único para compartir"""
        import secrets
        self.public_share_token = secrets.token_urlsafe(32)
        self.save()
        return self.public_share_token
    
    def get_share_url(self, request=None):
        """Obtiene la URL pública de compartir"""
        if not self.public_share_token:
            return None
        
        from django.urls import reverse
        try:
            path = reverse('knowledge_base_public', kwargs={'token': self.public_share_token})
            if request:
                return request.build_absolute_uri(path)
            return path
        except:
            return None
    
    def is_share_valid(self):
        """Verifica si el compartir está activo y no ha expirado"""
        if not self.public_share_enabled:
            return False
        
        if self.public_share_expires_at:
            from django.utils import timezone
            if timezone.now() > self.public_share_expires_at:
                return False
        
        return True


class Translation(models.Model):
    """Modelo para almacenar traducciones realizadas con IA"""
    
    LANGUAGES = [
        ('es', 'Español'),
        ('en', 'Inglés'),
        ('fr', 'Francés'),
        ('de', 'Alemán'),
        ('it', 'Italiano'),
        ('pt', 'Portugués'),
        ('ru', 'Ruso'),
        ('zh', 'Chino'),
        ('ja', 'Japonés'),
        ('ko', 'Coreano'),
        ('ar', 'Árabe'),
        ('hi', 'Hindi'),
        ('tr', 'Turco'),
        ('nl', 'Neerlandés'),
        ('pl', 'Polaco'),
        ('sv', 'Sueco'),
        ('no', 'Noruego'),
        ('da', 'Danés'),
        ('fi', 'Finlandés'),
        ('cs', 'Checo'),
        ('ca', 'Catalán'),
        ('val', 'Valenciano'),
    ]
    
    source_language = models.CharField(
        max_length=10,
        choices=LANGUAGES,
        default='en',
        verbose_name='Idioma Origen'
    )
    
    target_language = models.CharField(
        max_length=10,
        choices=LANGUAGES,
        default='es',
        verbose_name='Idioma Destino'
    )
    
    source_text = models.TextField(
        verbose_name='Texto Original',
        help_text='Texto a traducir'
    )
    
    translated_text = models.TextField(
        verbose_name='Texto Traducido',
        help_text='Resultado de la traducción'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    class Meta:
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['source_language']),
            models.Index(fields=['target_language']),
            models.Index(fields=['created_at']),
        ]
        verbose_name = 'Traducción'
        verbose_name_plural = 'Traducciones'
    
    def __str__(self):
        return f"{self.get_source_language_display()} → {self.get_target_language_display()} ({self.created_at.strftime('%Y-%m-%d')})"


class SQLQuery(models.Model):
    """Modelo para almacenar consultas SQL y sus resultados"""
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Consulta',
        help_text='Nombre descriptivo para identificar la consulta'
    )
    
    query = models.TextField(
        verbose_name='Consulta SQL',
        help_text='Solo se permiten consultas SELECT'
    )
    
    result = models.TextField(
        blank=True,
        null=True,
        verbose_name='Resultado',
        help_text='Resultado de la última ejecución'
    )
    
    row_count = models.IntegerField(
        default=0,
        verbose_name='Cantidad de Filas',
        help_text='Número de filas retornadas'
    )
    
    execution_time = models.FloatField(
        default=0.0,
        verbose_name='Tiempo de Ejecución',
        help_text='Tiempo de ejecución en segundos'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='sql_queries',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    last_executed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última ejecución'
    )
    
    is_favorite = models.BooleanField(
        default=False,
        verbose_name='Favorito'
    )
    
    class Meta:
        ordering = ['-created_at']
        indexes = [
            models.Index(fields=['created_by']),
            models.Index(fields=['created_at']),
            models.Index(fields=['is_favorite']),
        ]
        verbose_name = 'Consulta SQL'
        verbose_name_plural = 'Consultas SQL'
    
    def __str__(self):
        return f"{self.name} ({self.created_at.strftime('%Y-%m-%d')})"


class OdooConnection(models.Model):
    """Modelo para almacenar configuraciones de conexión a Odoo"""
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre de la Conexión',
        help_text='Nombre descriptivo para identificar esta conexión'
    )
    
    url = models.URLField(
        verbose_name='URL de Odoo',
        help_text='URL completa del servidor Odoo (ej: https://miodoo.com)'
    )
    
    port = models.IntegerField(
        default=8069,
        verbose_name='Puerto',
        help_text='Puerto del servidor Odoo'
    )
    
    database = models.CharField(
        max_length=200,
        verbose_name='Base de Datos',
        help_text='Nombre de la base de datos en Odoo'
    )
    
    username = models.CharField(
        max_length=200,
        verbose_name='Usuario',
        help_text='Usuario de Odoo'
    )
    
    password = models.CharField(
        max_length=200,
        verbose_name='Contraseña',
        help_text='Contraseña del usuario'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    is_private = models.BooleanField(
        default=False,
        verbose_name='Privada',
        help_text='Si está marcada como privada, solo el usuario creador puede verla y usarla'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='odoo_connections',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    last_tested_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Última prueba de conexión'
    )
    
    connection_status = models.CharField(
        max_length=50,
        default='not_tested',
        choices=[
            ('not_tested', 'No probada'),
            ('success', 'Exitosa'),
            ('failed', 'Fallida'),
        ],
        verbose_name='Estado de Conexión'
    )
    
    # Portal público para todas las tablas
    public_portal_enabled = models.BooleanField(
        default=False,
        verbose_name='Portal Público Habilitado',
        help_text='Permite acceso público a todas las tablas de esta conexión'
    )
    
    public_portal_token = models.CharField(
        max_length=64,
        blank=True,
        unique=True,
        null=True,
        verbose_name='Token del Portal Público',
        help_text='Token único para acceso al portal público'
    )
    
    public_portal_password = models.CharField(
        max_length=128,
        blank=True,
        verbose_name='Contraseña del Portal',
        help_text='Dejar en blanco para permitir acceso sin contraseña'
    )
    
    public_portal_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Expiración del Portal',
        help_text='Fecha en que expira el acceso al portal (opcional)'
    )
    
    public_portal_allow_import = models.BooleanField(
        default=True,
        verbose_name='Permitir Importación',
        help_text='Si está activo, los usuarios pueden importar datos. Si no, solo pueden descargar plantillas'
    )
    
    # Portal público de operaciones RPC
    public_operations_enabled = models.BooleanField(
        default=False,
        verbose_name='Portal de Operaciones Público Habilitado',
        help_text='Permite acceso público al portal de operaciones RPC'
    )
    
    public_operations_token = models.CharField(
        max_length=64,
        blank=True,
        unique=True,
        null=True,
        verbose_name='Token del Portal de Operaciones',
        help_text='Token único para acceso al portal de operaciones RPC'
    )
    
    public_operations_password = models.CharField(
        max_length=128,
        blank=True,
        verbose_name='Contraseña del Portal de Operaciones',
        help_text='Dejar en blanco para permitir acceso sin contraseña'
    )
    
    public_operations_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Expiración del Portal de Operaciones',
        help_text='Fecha en que expira el acceso al portal de operaciones (opcional)'
    )
    
    public_operations_read_only = models.BooleanField(
        default=False,
        verbose_name='Portal de Operaciones Solo Lectura',
        help_text='Si está activo, solo permite operaciones de búsqueda (no crear, actualizar o eliminar)'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Conexión Odoo'
        verbose_name_plural = 'Conexiones Odoo'
    
    def __str__(self):
        return f"{self.name} ({self.database})"
    
    def generate_portal_token(self):
        """Generar token único para portal público"""
        import secrets
        self.public_portal_token = secrets.token_urlsafe(32)
        self.save()
        return self.public_portal_token
    
    def get_portal_url(self, request=None):
        """Obtener URL del portal público"""
        if not self.public_portal_enabled or not self.public_portal_token:
            return None
        if request:
            from django.urls import reverse
            return request.build_absolute_uri(
                reverse('odoo_rpc_public_portal', kwargs={'token': self.public_portal_token})
            )
        return f"/public/odoo-portal/{self.public_portal_token}/"
    
    def is_portal_valid(self):
        """Verificar si el portal público sigue siendo válido"""
        if not self.public_portal_enabled:
            return False
        if self.public_portal_expires_at:
            from django.utils import timezone
            return timezone.now() < self.public_portal_expires_at
        return True
    
    def generate_operations_token(self):
        """Generar token único para portal de operaciones RPC"""
        import secrets
        self.public_operations_token = secrets.token_urlsafe(32)
        self.save()
        return self.public_operations_token
    
    def get_operations_url(self, request=None):
        """Obtener URL del portal de operaciones público"""
        if not self.public_operations_enabled or not self.public_operations_token:
            return None
        if request:
            from django.urls import reverse
            try:
                return request.build_absolute_uri(
                    reverse('odoo_rpc_public_operations', kwargs={'token': self.public_operations_token})
                )
            except:
                # Fallback si la ruta no existe aún
                return request.build_absolute_uri(f"/public/odoo-operations/{self.public_operations_token}/")
        return f"/public/odoo-operations/{self.public_operations_token}/"
    
    def is_operations_valid(self):
        """Verificar si el portal de operaciones público sigue siendo válido"""
        if not self.public_operations_enabled:
            return False
        if self.public_operations_expires_at:
            from django.utils import timezone
            return timezone.now() < self.public_operations_expires_at
        return True


class OdooRPCTable(models.Model):
    """Modelo para definir tablas de Odoo a sincronizar"""
    
    connection = models.ForeignKey(
        OdooConnection,
        on_delete=models.CASCADE,
        related_name='rpc_tables',
        verbose_name='Conexión'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre Descriptivo',
        help_text='Nombre descriptivo de la tabla'
    )
    
    odoo_model = models.CharField(
        max_length=200,
        verbose_name='Modelo Odoo',
        help_text='Nombre del modelo en Odoo (ej: product.product, res.partner)'
    )
    
    description = models.TextField(
        blank=True,
        verbose_name='Descripción',
        help_text='Descripción de la tabla y su propósito'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='odoo_rpc_tables',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    # Campos para URL pública
    public_url_enabled = models.BooleanField(
        default=False,
        verbose_name='URL Pública Habilitada',
        help_text='Permite acceso público para importar datos'
    )
    
    public_url_token = models.CharField(
        max_length=64,
        blank=True,
        unique=True,
        null=True,
        verbose_name='Token de URL Pública',
        help_text='Token único para acceso público'
    )
    
    public_url_password = models.CharField(
        max_length=128,
        blank=True,
        verbose_name='Contraseña de URL Pública',
        help_text='Dejar en blanco para permitir acceso sin contraseña'
    )
    
    public_url_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Expiración',
        help_text='Fecha en que expira el acceso público (opcional)'
    )
    
    # Campos para formulario público de creación
    public_form_enabled = models.BooleanField(
        default=False,
        verbose_name='Formulario Público Habilitado',
        help_text='Permite crear registros vía formulario público'
    )
    
    public_form_token = models.CharField(
        max_length=64,
        blank=True,
        unique=True,
        null=True,
        verbose_name='Token de Formulario Público',
        help_text='Token único para formulario público'
    )
    
    public_form_password = models.CharField(
        max_length=128,
        blank=True,
        verbose_name='Contraseña de Formulario Público',
        help_text='Dejar en blanco para permitir acceso sin contraseña'
    )
    
    public_form_expires_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de Expiración del Formulario',
        help_text='Fecha en que expira el formulario público (opcional)'
    )
    
    class Meta:
        ordering = ['name']
        unique_together = ['connection', 'odoo_model']
        verbose_name = 'Tabla RPC Odoo'
        verbose_name_plural = 'Tablas RPC Odoo'
    
    def __str__(self):
        return f"{self.name} ({self.odoo_model})"
    
    def generate_public_token(self):
        """Generar token único para URL pública"""
        import secrets
        self.public_url_token = secrets.token_urlsafe(32)
        self.save()
        return self.public_url_token
    
    def get_public_url(self, request=None):
        """Obtener URL pública para importación"""
        if not self.public_url_enabled or not self.public_url_token:
            return None
        if request:
            from django.urls import reverse
            return request.build_absolute_uri(
                reverse('odoo_rpc_public_import', kwargs={'token': self.public_url_token})
            )
        return f"/public/odoo-import/{self.public_url_token}/"
    
    def is_public_url_valid(self):
        """Verificar si la URL pública sigue siendo válida"""
        if not self.public_url_enabled:
            return False
        if self.public_url_expires_at:
            from django.utils import timezone
            return timezone.now() < self.public_url_expires_at
        return True
    
    def generate_form_token(self):
        """Generar token único para formulario público"""
        import secrets
        self.public_form_token = secrets.token_urlsafe(32)
        self.save()
        return self.public_form_token
    
    def get_form_url(self, request=None):
        """Obtener URL pública para formulario de creación"""
        if not self.public_form_enabled or not self.public_form_token:
            return None
        if request:
            from django.urls import reverse
            try:
                return request.build_absolute_uri(
                    reverse('odoo_rpc_public_form', kwargs={'token': self.public_form_token})
                )
            except:
                return request.build_absolute_uri(f"/public/odoo-form/{self.public_form_token}/")
        return f"/public/odoo-form/{self.public_form_token}/"
    
    def is_form_valid(self):
        """Verificar si el formulario público sigue siendo válido"""
        if not self.public_form_enabled:
            return False
        if self.public_form_expires_at:
            from django.utils import timezone
            return timezone.now() < self.public_form_expires_at
        return True


class OdooRPCField(models.Model):
    """Modelo para definir campos de las tablas Odoo"""
    
    FIELD_TYPES = [
        ('char', 'Texto'),
        ('text', 'Texto Largo'),
        ('integer', 'Entero'),
        ('float', 'Decimal'),
        ('boolean', 'Booleano'),
        ('date', 'Fecha'),
        ('datetime', 'Fecha y Hora'),
        ('selection', 'Selección'),
        ('many2one', 'Relación Many2One'),
        ('one2many', 'Relación One2Many'),
        ('many2many', 'Relación Many2Many'),
    ]
    
    table = models.ForeignKey(
        OdooRPCTable,
        on_delete=models.CASCADE,
        related_name='rpc_fields',
        verbose_name='Tabla'
    )
    
    name = models.CharField(
        max_length=200,
        verbose_name='Nombre Descriptivo',
        help_text='Nombre descriptivo del campo'
    )
    
    odoo_field_name = models.CharField(
        max_length=200,
        verbose_name='Nombre del Campo en Odoo',
        help_text='Nombre técnico del campo en Odoo (ej: name, default_code)'
    )
    
    field_type = models.CharField(
        max_length=50,
        choices=FIELD_TYPES,
        default='char',
        verbose_name='Tipo de Campo'
    )
    
    is_required = models.BooleanField(
        default=False,
        verbose_name='Requerido',
        help_text='Si el campo es obligatorio'
    )
    
    order = models.IntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en Excel'
    )
    
    help_text = models.TextField(
        blank=True,
        verbose_name='Texto de Ayuda',
        help_text='Instrucciones para llenar este campo'
    )
    
    class Meta:
        ordering = ['table', 'order', 'name']
        unique_together = ['table', 'odoo_field_name']
        verbose_name = 'Campo RPC Odoo'
        verbose_name_plural = 'Campos RPC Odoo'
    
    def __str__(self):
        return f"{self.name} ({self.odoo_field_name})"


class OdooRPCData(models.Model):
    """Modelo para almacenar datos temporales antes de enviar a Odoo"""
    
    STATUS_CHOICES = [
        ('pending', 'Pendiente'),
        ('processing', 'Procesando'),
        ('success', 'Exitoso'),
        ('failed', 'Fallido'),
        ('fetched', 'Obtenido'),
    ]
    
    table = models.ForeignKey(
        OdooRPCTable,
        on_delete=models.CASCADE,
        related_name='rpc_data',
        verbose_name='Tabla'
    )
    
    data = models.JSONField(
        verbose_name='Datos',
        help_text='Datos en formato JSON para enviar a Odoo'
    )
    
    status = models.CharField(
        max_length=20,
        choices=STATUS_CHOICES,
        default='pending',
        verbose_name='Estado'
    )
    
    odoo_id = models.IntegerField(
        null=True,
        blank=True,
        verbose_name='ID en Odoo',
        help_text='ID del registro creado en Odoo'
    )
    
    error_message = models.TextField(
        blank=True,
        verbose_name='Mensaje de Error',
        help_text='Mensaje de error si la sincronización falló'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='odoo_rpc_data',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    processed_at = models.DateTimeField(
        null=True,
        blank=True,
        verbose_name='Fecha de procesamiento'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Dato RPC Odoo'
        verbose_name_plural = 'Datos RPC Odoo'
    
    def __str__(self):
        return f"{self.table.name} - {self.status} ({self.created_at.strftime('%Y-%m-%d')})"


class OdooRPCImportFile(models.Model):
    """Modelo para almacenar archivos de importación"""
    
    table = models.ForeignKey(
        OdooRPCTable,
        on_delete=models.CASCADE,
        related_name='import_files',
        verbose_name='Tabla'
    )
    
    file = models.FileField(
        upload_to='odoo_rpc_imports/%Y/%m/',
        verbose_name='Archivo'
    )
    
    original_filename = models.CharField(
        max_length=255,
        verbose_name='Nombre Original'
    )
    
    uploaded_by = models.ForeignKey(
        User,
        on_delete=models.SET_NULL,
        null=True,
        related_name='odoo_rpc_imports',
        verbose_name='Subido por'
    )
    
    uploaded_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de subida'
    )
    
    records_imported = models.IntegerField(
        default=0,
        verbose_name='Registros importados'
    )
    
    records_success = models.IntegerField(
        default=0,
        verbose_name='Registros exitosos'
    )
    
    records_failed = models.IntegerField(
        default=0,
        verbose_name='Registros fallidos'
    )
    
    class Meta:
        ordering = ['-uploaded_at']
        verbose_name = 'Archivo de Importación RPC'
        verbose_name_plural = 'Archivos de Importación RPC'
    
    def __str__(self):
        return f"{self.original_filename} - {self.uploaded_at.strftime('%Y-%m-%d %H:%M')}"


# ==================== CHATBOT ====================

class Chatbot(models.Model):
    """Modelo para chatbots configurables"""
    
    TYPE_CHOICES = [
        ('internal', 'Interno (Solo esta web)'),
        ('external', 'Externo (Script integrable)'),
    ]
    
    title = models.CharField(
        max_length=200,
        verbose_name='Título del Chatbot'
    )
    
    type = models.CharField(
        max_length=20,
        choices=TYPE_CHOICES,
        default='internal',
        verbose_name='Tipo de Chatbot'
    )
    
    description = models.TextField(
        verbose_name='Descripción',
        help_text='Describe el propósito del chatbot (ej: Agente de ventas que se dedica a captar leads)'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activo'
    )
    
    use_ai = models.BooleanField(
        default=True,
        verbose_name='Usar IA',
        help_text='Habilitar asistencia de IA para respuestas no configuradas'
    )
    
    ai_context = models.TextField(
        blank=True,
        verbose_name='Contexto para IA',
        help_text='Información adicional para que la IA entienda el contexto del negocio'
    )
    
    welcome_message = models.TextField(
        default='¡Hola! ¿En qué puedo ayudarte?',
        verbose_name='Mensaje de bienvenida'
    )
    
    # Personalización visual
    primary_color = models.CharField(
        max_length=7,
        default='#007bff',
        verbose_name='Color primario',
        help_text='Color del botón y encabezado (formato: #RRGGBB)'
    )
    
    secondary_color = models.CharField(
        max_length=7,
        default='#0056b3',
        verbose_name='Color secundario',
        help_text='Color secundario para efectos hover (formato: #RRGGBB)'
    )
    
    bot_message_color = models.CharField(
        max_length=7,
        default='#f0f0f0',
        verbose_name='Color mensajes del bot',
        help_text='Color de fondo para los mensajes del bot (formato: #RRGGBB)'
    )
    
    user_message_color = models.CharField(
        max_length=7,
        default='#007bff',
        verbose_name='Color mensajes del usuario',
        help_text='Color de fondo para los mensajes del usuario (formato: #RRGGBB)'
    )
    
    icon_choice = models.CharField(
        max_length=20,
        choices=[
            ('chat', 'Chat (por defecto)'),
            ('robot', 'Robot'),
            ('headset', 'Auriculares'),
            ('person', 'Persona'),
        ],
        default='chat',
        verbose_name='Icono del botón',
        help_text='Icono que se muestra en el botón flotante'
    )
    
    position = models.CharField(
        max_length=20,
        choices=[
            ('bottom-right', 'Inferior Derecha'),
            ('bottom-left', 'Inferior Izquierda'),
        ],
        default='bottom-right',
        verbose_name='Posición en pantalla'
    )
    
    # Integración externa
    script_token = models.CharField(
        max_length=64,
        unique=True,
        blank=True,
        verbose_name='Token de script',
        help_text='Token único para integración externa'
    )
    
    allowed_domains = models.TextField(
        blank=True,
        verbose_name='Dominios permitidos',
        help_text='Dominios donde se puede integrar (uno por línea). Dejar vacío para permitir todos.'
    )
    
    created_by = models.ForeignKey(
        User,
        on_delete=models.CASCADE,
        related_name='chatbots_created',
        verbose_name='Creado por'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Comportamiento
    response_delay = models.IntegerField(
        default=0,
        verbose_name='Tiempo de espera (segundos)',
        help_text='Segundos que espera el bot antes de responder (0 = inmediato)'
    )
    
    # Estadísticas
    total_conversations = models.IntegerField(
        default=0,
        verbose_name='Total de conversaciones'
    )
    
    total_messages = models.IntegerField(
        default=0,
        verbose_name='Total de mensajes'
    )
    
    class Meta:
        ordering = ['-created_at']
        verbose_name = 'Chatbot'
        verbose_name_plural = 'Chatbots'
    
    def __str__(self):
        return self.title
    
    def save(self, *args, **kwargs):
        # Generar token si no tiene (para todos los chatbots)
        if not self.script_token:
            import secrets
            self.script_token = secrets.token_urlsafe(32)
        super().save(*args, **kwargs)
    
    def get_script_url(self):
        """Retorna la URL del script para integración"""
        if self.type == 'external' and self.script_token:
            from django.urls import reverse
            return reverse('chatbot_script', kwargs={'token': self.script_token})
        return None


class ChatbotQuestion(models.Model):
    """Preguntas y respuestas configuradas para el chatbot"""
    
    chatbot = models.ForeignKey(
        Chatbot,
        on_delete=models.CASCADE,
        related_name='questions',
        verbose_name='Chatbot'
    )
    
    question = models.CharField(
        max_length=500,
        verbose_name='Pregunta',
        help_text='Pregunta que el usuario puede hacer'
    )
    
    keywords = models.CharField(
        max_length=500,
        blank=True,
        verbose_name='Palabras clave',
        help_text='Palabras clave separadas por comas para detectar esta pregunta'
    )
    
    answer = models.TextField(
        verbose_name='Respuesta',
        help_text='Respuesta que el chatbot dará'
    )
    
    order = models.IntegerField(
        default=0,
        verbose_name='Orden',
        help_text='Orden de aparición en sugerencias'
    )
    
    is_active = models.BooleanField(
        default=True,
        verbose_name='Activa'
    )
    
    use_ai = models.BooleanField(
        default=False,
        verbose_name='Mejorar con IA',
        help_text='La IA puede expandir o mejorar esta respuesta'
    )
    
    created_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Fecha de creación'
    )
    
    updated_at = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actualización'
    )
    
    # Estadísticas
    times_used = models.IntegerField(
        default=0,
        verbose_name='Veces utilizada'
    )
    
    class Meta:
        ordering = ['order', '-created_at']
        verbose_name = 'Pregunta del Chatbot'
        verbose_name_plural = 'Preguntas del Chatbot'
    
    def __str__(self):
        return f"{self.question[:50]}..."
    
    def get_keywords_list(self):
        """Retorna lista de keywords"""
        if self.keywords:
            return [k.strip().lower() for k in self.keywords.split(',')]
        return []


class ChatbotConversation(models.Model):
    """Conversación de un usuario con el chatbot"""
    
    chatbot = models.ForeignKey(
        Chatbot,
        on_delete=models.CASCADE,
        related_name='conversations',
        verbose_name='Chatbot'
    )
    
    session_id = models.CharField(
        max_length=100,
        verbose_name='ID de sesión'
    )
    
    user_email = models.EmailField(
        blank=True,
        verbose_name='Email del usuario',
        help_text='Email capturado durante la conversación'
    )
    
    user_name = models.CharField(
        max_length=200,
        blank=True,
        verbose_name='Nombre del usuario'
    )
    
    user_phone = models.CharField(
        max_length=50,
        blank=True,
        verbose_name='Teléfono del usuario'
    )
    
    ip_address = models.GenericIPAddressField(
        blank=True,
        null=True,
        verbose_name='Dirección IP',
        help_text='IP del usuario en la sesión'
    )
    
    country = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='País',
        help_text='País detectado del usuario'
    )
    
    city = models.CharField(
        max_length=100,
        blank=True,
        verbose_name='Ciudad',
        help_text='Ciudad detectada del usuario'
    )
    
    started_at = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Inicio de conversación'
    )
    
    last_activity = models.DateTimeField(
        auto_now=True,
        verbose_name='Última actividad'
    )
    
    is_lead = models.BooleanField(
        default=False,
        verbose_name='Es lead',
        help_text='Marcado si se capturó información de contacto'
    )
    
    lead_quality = models.CharField(
        max_length=20,
        choices=[
            ('hot', 'Caliente'),
            ('warm', 'Tibio'),
            ('cold', 'Frío'),
        ],
        blank=True,
        verbose_name='Calidad del lead'
    )
    
    hide_chat = models.BooleanField(
        default=False,
        verbose_name='Ocultar chat',
        help_text='Si está marcado, el chat no se mostrará a este usuario'
    )
    
    reviewed = models.BooleanField(
        default=False,
        verbose_name='Revisada',
        help_text='Marca si la conversación ya fue revisada por un administrador'
    )
    
    reviewed_at = models.DateTimeField(
        blank=True,
        null=True,
        verbose_name='Fecha de revisión'
    )
    
    reviewed_by = models.ForeignKey(
        'auth.User',
        on_delete=models.SET_NULL,
        blank=True,
        null=True,
        verbose_name='Revisada por'
    )
    
    class Meta:
        ordering = ['-started_at']
        verbose_name = 'Conversación del Chatbot'
        verbose_name_plural = 'Conversaciones del Chatbot'
    
    def __str__(self):
        return f"Conversación {self.session_id[:8]} - {self.started_at.strftime('%Y-%m-%d %H:%M')}"


class ChatbotMessage(models.Model):
    """Mensajes individuales en una conversación"""
    
    conversation = models.ForeignKey(
        ChatbotConversation,
        on_delete=models.CASCADE,
        related_name='messages',
        verbose_name='Conversación'
    )
    
    is_bot = models.BooleanField(
        verbose_name='Es del bot',
        help_text='True si es mensaje del bot, False si es del usuario'
    )
    
    message = models.TextField(
        verbose_name='Mensaje'
    )
    
    matched_question = models.ForeignKey(
        ChatbotQuestion,
        on_delete=models.SET_NULL,
        null=True,
        blank=True,
        related_name='matched_messages',
        verbose_name='Pregunta coincidente'
    )
    
    used_ai = models.BooleanField(
        default=False,
        verbose_name='Usó IA',
        help_text='True si la respuesta fue generada por IA'
    )
    
    timestamp = models.DateTimeField(
        auto_now_add=True,
        verbose_name='Hora del mensaje'
    )
    
    class Meta:
        ordering = ['timestamp']
        verbose_name = 'Mensaje del Chatbot'
        verbose_name_plural = 'Mensajes del Chatbot'
    
    def __str__(self):
        sender = "Bot" if self.is_bot else "Usuario"
        return f"{sender}: {self.message[:50]}..."





